from pypdf import PdfReader, PdfWriter
from django.core.files.base import ContentFile
from docx import Document
from openpyxl import load_workbook
from PIL import Image
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER
import io
import os
import tempfile


def is_pdf(file):
    """
    Check if a file is a PDF
    
    Args:
        file: Django FileField or UploadedFile object
    
    Returns:
        Boolean indicating if file is PDF
    """
    if hasattr(file, 'name'):
        file_extension = os.path.splitext(file.name)[1].lower()
        return file_extension == '.pdf'
    return False


def convert_image_to_pdf(image_file):
    """
    Convert image (JPG, PNG) to PDF
    
    Args:
        image_file: Django UploadedFile object
    
    Returns:
        ContentFile object containing the PDF
    """
    output = io.BytesIO()
    
    # Open the image
    img = Image.open(image_file)
    
    # Convert RGBA to RGB if necessary
    if img.mode == 'RGBA':
        # Create a white background
        background = Image.new('RGB', img.size, (255, 255, 255))
        background.paste(img, mask=img.split()[3])  # Use alpha channel as mask
        img = background
    elif img.mode != 'RGB':
        img = img.convert('RGB')
    
    # Get image dimensions
    img_width, img_height = img.size
    
    # Determine page size based on image orientation
    if img_width > img_height:
        # Landscape
        page_width, page_height = letter[1], letter[0]  # Swap for landscape
    else:
        # Portrait
        page_width, page_height = letter
    
    # Calculate scaling to fit image on page with margins
    margin = 0.5 * inch
    available_width = page_width - (2 * margin)
    available_height = page_height - (2 * margin)
    
    # Calculate scale factor
    scale = min(available_width / img_width, available_height / img_height)
    scaled_width = img_width * scale
    scaled_height = img_height * scale
    
    # Center the image
    x = (page_width - scaled_width) / 2
    y = (page_height - scaled_height) / 2
    
    # Create PDF with ReportLab
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    c = canvas.Canvas(output, pagesize=(page_width, page_height))
    
    # Convert PIL Image to ImageReader for ReportLab
    img_reader = ImageReader(img)
    
    # Draw image on canvas
    c.drawImage(img_reader, x, y, scaled_width, scaled_height)
    c.save()
    
    output.seek(0)
    return ContentFile(output.read())


def convert_docx_to_pdf(docx_file):
    """
    Convert Word document to PDF
    
    Args:
        docx_file: Django UploadedFile object
    
    Returns:
        ContentFile object containing the PDF
    """
    output = io.BytesIO()
    
    # Load the Word document
    doc = Document(docx_file)
    
    # Create PDF
    pdf = SimpleDocTemplate(output, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Add a title style
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#333333'),
        spaceAfter=12,
        alignment=TA_CENTER
    )
    
    # Add a normal style
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#000000'),
        spaceAfter=12,
        alignment=TA_LEFT
    )
    
    # Process document content
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            # Detect if it's likely a heading (bold or all caps)
            is_heading = False
            if paragraph.runs:
                is_heading = paragraph.runs[0].bold or paragraph.text.isupper()
            
            if is_heading and len(paragraph.text) < 100:
                p = Paragraph(paragraph.text, title_style)
            else:
                p = Paragraph(paragraph.text, normal_style)
            story.append(p)
    
    # Process tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        
        if table_data:
            t = Table(table_data)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(t)
            story.append(Spacer(1, 12))
    
    # Build PDF
    pdf.build(story)
    
    output.seek(0)
    return ContentFile(output.read())


def convert_excel_to_pdf(excel_file):
    """
    Convert Excel spreadsheet to PDF
    
    Args:
        excel_file: Django UploadedFile object
    
    Returns:
        ContentFile object containing the PDF
    """
    output = io.BytesIO()
    
    # Load the Excel workbook
    wb = load_workbook(excel_file, data_only=True)
    
    # Create PDF
    pdf = SimpleDocTemplate(output, pagesize=A4)
    story = []
    styles = getSampleStyleSheet()
    
    # Title style
    title_style = ParagraphStyle(
        'SheetTitle',
        parent=styles['Heading1'],
        fontSize=14,
        textColor=colors.HexColor('#333333'),
        spaceAfter=12,
        alignment=TA_CENTER
    )
    
    # Process each sheet
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        
        # Add sheet name as title
        story.append(Paragraph(f"Sheet: {sheet_name}", title_style))
        story.append(Spacer(1, 12))
        
        # Extract data from sheet
        table_data = []
        for row in ws.iter_rows(values_only=True):
            # Convert None to empty string and all values to strings
            row_data = [str(cell) if cell is not None else '' for cell in row]
            # Only add non-empty rows
            if any(cell.strip() for cell in row_data):
                table_data.append(row_data)
        
        if table_data:
            # Create table
            t = Table(table_data)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(t)
            story.append(Spacer(1, 24))
    
    # Build PDF
    pdf.build(story)
    
    output.seek(0)
    return ContentFile(output.read())


def convert_to_pdf(uploaded_file):
    """
    Convert any supported file type to PDF
    
    Supported formats:
    - PDF: Returns as-is
    - Images (JPG, JPEG, PNG): Converts to PDF
    - Word (DOCX): Converts to PDF
    - Excel (XLSX, XLS): Converts to PDF
    
    Args:
        uploaded_file: Django UploadedFile object
    
    Returns:
        Tuple of (ContentFile, filename)
    
    Raises:
        ValueError: If file type is not supported
    """
    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
    
    # If already PDF, return as-is
    if file_extension == '.pdf':
        uploaded_file.seek(0)
        return ContentFile(uploaded_file.read()), uploaded_file.name
    
    # Generate new filename with .pdf extension
    original_name = os.path.splitext(uploaded_file.name)[0]
    new_filename = f"{original_name}.pdf"
    
    # Convert based on file type
    if file_extension in ['.jpg', '.jpeg', '.png']:
        pdf_content = convert_image_to_pdf(uploaded_file)
        return pdf_content, new_filename
    
    elif file_extension in ['.docx', '.doc']:
        pdf_content = convert_docx_to_pdf(uploaded_file)
        return pdf_content, new_filename
    
    elif file_extension in ['.xlsx', '.xls']:
        pdf_content = convert_excel_to_pdf(uploaded_file)
        return pdf_content, new_filename
    
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")


def merge_pdfs(existing_file, new_file):
    """
    Merge two PDF files and return a ContentFile object
    
    Args:
        existing_file: Django FileField object (existing document)
        new_file: Django UploadedFile, BytesIO, or ContentFile object (new document)
    
    Returns:
        ContentFile object containing the merged PDF
    
    Raises:
        ValueError: If either file is not a PDF
    """
    # Validate existing file is PDF
    if not is_pdf(existing_file):
        raise ValueError("Existing document is not a PDF file")
    
    # Create a PDF writer object
    pdf_writer = PdfWriter()
    
    # Read the existing PDF
    existing_pdf = PdfReader(existing_file)
    for page in existing_pdf.pages:
        pdf_writer.add_page(page)
    
    # Handle different types of new_file
    if isinstance(new_file, (io.BytesIO, ContentFile)):
        new_file.seek(0)
        new_pdf = PdfReader(new_file)
    else:
        new_pdf = PdfReader(new_file)
    
    for page in new_pdf.pages:
        pdf_writer.add_page(page)
    
    # Write to a BytesIO object
    output = io.BytesIO()
    pdf_writer.write(output)
    output.seek(0)
    
    return ContentFile(output.read())


def merge_pdfs_from_bytes(existing_bytes, new_pdf_content):
    """
    Merge two PDFs where the existing file is supplied as raw bytes.
    Used in edit_recipe to avoid Windows file lock (WinError 32) — the existing
    file is read into memory and closed before deletion, then merged from bytes.

    Args:
        existing_bytes: Raw bytes of the existing PDF (already read from disk)
        new_pdf_content: ContentFile or BytesIO of the new PDF (output of convert_to_pdf)

    Returns:
        ContentFile object containing the merged PDF
    """
    pdf_writer = PdfWriter()

    # Read existing PDF from bytes — no file handle needed
    reader1 = PdfReader(io.BytesIO(existing_bytes))
    for page in reader1.pages:
        pdf_writer.add_page(page)

    # Read new PDF
    if hasattr(new_pdf_content, 'read'):
        new_pdf_content.seek(0)
        new_bytes = new_pdf_content.read()
    else:
        new_bytes = bytes(new_pdf_content)
    reader2 = PdfReader(io.BytesIO(new_bytes))
    for page in reader2.pages:
        pdf_writer.add_page(page)

    output = io.BytesIO()
    pdf_writer.write(output)
    output.seek(0)
    return ContentFile(output.read())
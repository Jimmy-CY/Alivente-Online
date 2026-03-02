from django.conf import settings
from django.contrib.auth import authenticate, login, logout, update_session_auth_hash
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from django.contrib.auth.forms import UserCreationForm, UserChangeForm, PasswordChangeForm
from django.core.exceptions import ValidationError
from django.core.files.base import ContentFile
from django.core.files.storage import FileSystemStorage, default_storage
from django.core.mail import EmailMultiAlternatives
from django.core.paginator import Paginator
from django.core.serializers import serialize
from django.db import connection, transaction, models
from django.db.models import Q, Prefetch, Subquery, OuterRef, Sum, F, Count, Max, Min
from django.db.models.functions import Coalesce
from django.http import HttpResponse, HttpResponseServerError, FileResponse, Http404, JsonResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.template.loader import render_to_string, get_template
from django.templatetags.static import static
from django.urls import reverse
from django.utils import timezone
from django.utils.dateparse import parse_date
from django.utils.html import strip_tags
from django.utils.safestring import mark_safe
from django.utils.timezone import now
from django.views.decorators.cache import never_cache
from django.views.decorators.csrf import csrf_exempt, csrf_protect
from django.views.decorators.http import require_POST, require_http_methods
from django.views.static import serve
from docxtpl import DocxTemplate


#from .translation_service import ensure_project_translations, get_translated_text
# Temporarily disabled translation
def ensure_project_translations(request):
    pass

def get_translated_text(text, target_language='en'):
    return text

from . import forms
from .forms import PropForm, TenantForm, PettyForm, InvoicesForm, IssuesForm, DetailsForm, SupplierForm, ValuesForm, RevenueTypesForm, RevenueLineForm, RevenueForm, ExpenseTypesForm, ExpenseLineForm, ExpenseForm, ActExpenseForm 
from .models import (
    props,
    petty,
    issues,
    issues_details, 
    tenant, 
    invoices,
    supplier,
    prop_values,
    revenue_types,
    revenue_line_types,
    revenue,
    expense_types,
    expense_line_types,
    expense,
    act_expense,
    Project, 
    ProjectTask,
    ProjectDocument,
    Passport,
    Recipe,
    RecipeIngredient,
    RecipeIngredientText,
    RecipeInstruction,
    RecipeCourse,
    RecipeCategory,
    Ingredient,
    MeasurementUnit,
    IngredientCategory,
    CustomProtein,
    PreparationMethod,
    MealPlan,
    MealPlanDay,
    MealPlanRecipe,
    UnitConversion,
    VacancyPeriod,
    Contact,
    CelebrationEvent,
    EventNotification,
    NotificationRecipient,

    )
from decimal import Decimal
from fractions import Fraction
from calendar import monthrange, monthcalendar, month_name
from collections import defaultdict
from datetime import date, datetime, timedelta
from pages.management.commands.email_utils import get_email_recipients, format_email_recipients_for_header
from spellchecker import SpellChecker
from urllib.parse import urlparse, parse_qs
from xhtml2pdf import pisa
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from .utils import merge_pdfs, is_pdf, convert_to_pdf
from PIL import Image
from docx import Document
from io import BytesIO
import decimal
import calendar
import mysql.connector
import smtplib
import io
import os
import re
import uuid
import logging
import json
import tempfile
import base64
import anthropic
import PyPDF2
import string

logger = logging.getLogger(__name__)

### PASSPORT MANAGEMENT ###
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from datetime import date, timedelta
from .models import Passport  # Adjust import based on your app structure

@login_required
def passport_management(request):
    if not request.user.is_superuser:
        messages.error(request, "You don't have permission to access this page.")
        return redirect('home')
    
    if request.method == 'POST':
        action = request.POST.get('action')
        
        # ADD NEW PASSPORT/ID
        if action == 'add':
            holder_name = request.POST.get('holder_name')
            document_type = request.POST.get('document_type')
            document_number = request.POST.get('document_number')
            country_of_issue = request.POST.get('country_of_issue')
            date_of_issue = request.POST.get('date_of_issue') or None
            expiry_date = request.POST.get('expiry_date') or None
            status = request.POST.get('status')
            document_file = request.FILES.get('document_file')
            
            try:
                Passport.objects.create(
                    holder_name=holder_name,
                    document_type=document_type,
                    document_number=document_number,
                    country_of_issue=country_of_issue,
                    date_of_issue=date_of_issue,
                    expiry_date=expiry_date,
                    status=status,
                    document_file=document_file
                )
                messages.success(request, f'Passport/ID for {holder_name} added successfully!')
            except Exception as e:
                messages.error(request, f'Error adding passport/ID: {str(e)}')
        
        # EDIT PASSPORT/ID
        elif action == 'edit':
            passport_id = request.POST.get('passport_id')
            passport = get_object_or_404(Passport, id=passport_id)
            
            passport.holder_name = request.POST.get('holder_name')
            passport.document_type = request.POST.get('document_type')
            passport.document_number = request.POST.get('document_number')
            passport.country_of_issue = request.POST.get('country_of_issue')
            passport.date_of_issue = request.POST.get('date_of_issue') or None
            passport.expiry_date = request.POST.get('expiry_date') or None
            passport.status = request.POST.get('status')
            
            # Update file if a new one is uploaded
            if request.FILES.get('document_file'):
                passport.document_file = request.FILES.get('document_file')
            
            try:
                passport.save()
                messages.success(request, f'Passport/ID for {passport.holder_name} updated successfully!')
            except Exception as e:
                messages.error(request, f'Error updating passport/ID: {str(e)}')
        
        # UPLOAD DOCUMENT
        elif action == 'upload':
            passport_id = request.POST.get('passport_id')
            passport = get_object_or_404(Passport, id=passport_id)
            document_file = request.FILES.get('document_file')
            
            if document_file:
                passport.document_file = document_file
                try:
                    passport.save()
                    messages.success(request, f'Document uploaded successfully for {passport.holder_name}!')
                except Exception as e:
                    messages.error(request, f'Error uploading document: {str(e)}')
            else:
                messages.error(request, 'No file selected.')
        
        # DELETE PASSPORT/ID
        elif action == 'delete':
            passport_id = request.POST.get('passport_id')
            passport = get_object_or_404(Passport, id=passport_id)
            holder_name = passport.holder_name
            
            try:
                # Delete the file from storage
                if passport.document_file:
                    passport.document_file.delete()
                passport.delete()
                messages.success(request, f'Passport/ID for {holder_name} deleted successfully!')
            except Exception as e:
                messages.error(request, f'Error deleting passport/ID: {str(e)}')
        
        return redirect('passport_management')
    
    # GET request - display all passports with filters
    passports = Passport.objects.all()
    
    # Get filter parameters from request
    selected_holder = request.GET.get('holder', '')
    selected_doc_type = request.GET.get('doc_type', '')
    selected_country = request.GET.get('country', '')
    selected_status = request.GET.get('status', '')
    
    # Apply holder filter
    if selected_holder:
        passports = passports.filter(holder_name=selected_holder)
    
    # Apply document type filter
    if selected_doc_type:
        passports = passports.filter(document_type=selected_doc_type)
    
    # Apply country filter
    if selected_country:
        passports = passports.filter(country_of_issue=selected_country)
    
    # Apply status filter
    if selected_status:
        passports = passports.filter(status=selected_status)
    
    # Order by creation date (newest first)
    passports = passports.order_by('-created_at')
    
    # Add expiry warning flag to each passport (for 6-month warning)
    today = date.today()
    six_months_from_now = today + timedelta(days=180)
    
    for passport in passports:
        if passport.expiry_date:
            # Flag if expiry date is within 6 months or already expired
            passport.expiring_soon = passport.expiry_date <= six_months_from_now
        else:
            passport.expiring_soon = False
    
    context = {
        'passports': passports,
        'selected_holder': selected_holder,
        'selected_doc_type': selected_doc_type,
        'selected_country': selected_country,
        'selected_status': selected_status,
    }
    
    return render(request, 'passport_management.html', context)

### LEASE TEMPLATE GENERATOR ###
import re

@login_required
@require_POST
def spell_check_instructions(request):
    """Spell check recipe instructions"""
    try:
        data = json.loads(request.body)
        instructions = data.get('instructions', [])
        
        if not instructions:
            return JsonResponse({'success': False, 'error': 'No instructions provided'})
        
        # Initialize spell checker
        spell = SpellChecker()
        
        # Common cooking terms to ignore
        cooking_terms = {
            'tsp', 'tbsp', 'mins', 'hrs', 'preheat', 'saute', 'sauteed', 
            'broil', 'simmer', 'whisk', 'preheated',
            'mins', 'secs', 'ml', 'oz', 'fahrenheit', 'celsius'
        }
        spell.word_frequency.load_words(cooking_terms)
        
        errors = []
        
        for idx, instruction in enumerate(instructions):
            if not instruction.strip():
                continue
                
            # Remove common cooking abbreviations and numbers
            text = instruction.lower()
            
            # Split into words, removing punctuation
            words = re.findall(r'\b[a-z]+\b', text)
            
            # Find misspelled words
            misspelled = spell.unknown(words)
            
            if misspelled:
                for word in misspelled:
                    # Get suggestions - handle None case
                    candidates = spell.candidates(word)
                    
                    # Convert to list and handle None
                    if candidates is None:
                        suggestions = []
                    else:
                        suggestions = list(candidates)[:5]
                    
                    errors.append({
                        'step': idx + 1,
                        'word': word,
                        'suggestions': suggestions,  # Will be empty list if no suggestions
                        'context': instruction
                    })
        
        return JsonResponse({
            'success': True,
            'errors': errors,
            'total_errors': len(errors)
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()  # Print full error to console for debugging
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

def is_superuser(user):
    """Check if user is superuser"""
    return user.is_superuser

@login_required
@user_passes_test(is_superuser)
def generate_lease_agreement_view(request):
    """
    View to display the lease agreement generation form and handle document generation
    """
    # Get properties and tenants from your actual models
    properties = props.objects.filter(prop_available_for_rent='Yes').order_by('prop_name')
    tenants = tenant.objects.select_related('prop').all().order_by('tenant_name')
    
    context = {
        'properties': properties,
        'tenants': tenants,
    }
    
    if request.method == 'POST':
        try:
            # Extract form data
            country = request.POST.get('country')
            language = request.POST.get('language')
            property_id = request.POST.get('property')
            tenant_id = request.POST.get('tenant')
            
            # Get additional data from modal (JSON format)
            additional_data_json = request.POST.get('additional_data', '{}')
            
            # Validate additional_data is valid JSON
            try:
                additional_data = json.loads(additional_data_json)
            except json.JSONDecodeError:
                additional_data = {}
                logger.warning(f"Invalid JSON in additional_data: {additional_data_json}")
            
            # Validate required fields
            if not all([country, language, property_id, tenant_id]):
                messages.error(request, 'Please fill in all required fields.')
                return render(request, 'generate_lease_agreement.html', context)
            
            # Validate additional data was provided
            if not additional_data:
                messages.error(request, 'Please complete the additional data by clicking "Review & Complete Data".')
                return render(request, 'generate_lease_agreement.html', context)
            
            # Get property object from database
            try:
                property_obj = get_object_or_404(props, prop_id=property_id)
            except:
                messages.error(request, 'Selected property not found.')
                return render(request, 'generate_lease_agreement.html', context)
            
            # Handle tenant (existing or new)
            if tenant_id == 'new_tenant':
                # Create a mock tenant object with new tenant data
                new_tenant_data = additional_data.get('new_tenant_data', {})
                if not new_tenant_data.get('tenant_name'):
                    messages.error(request, 'Tenant name is required for new tenant.')
                    return render(request, 'generate_lease_agreement.html', context)
                
                tenant_obj = type('MockTenant', (), {
                    'tenant_name': new_tenant_data.get('tenant_name', ''),
                    'tenant_type': new_tenant_data.get('tenant_type', 'Individual'),
                    'tenant_passport_id': new_tenant_data.get('tenant_passport_id', ''),
                    'tenant_passport_country': new_tenant_data.get('tenant_passport_country', ''),
                    'tenant_email': new_tenant_data.get('tenant_email', ''),
                    'tenant_contact_number': new_tenant_data.get('tenant_contact_number', ''),
                    'tenant_address': new_tenant_data.get('tenant_address', ''),
                    'tenant_rent': None,
                    'tenant_deposit': None,
                    'tenant_lease_start_date': None,
                    'tenant_lease_end_date': None,
                    'tenant_levies': None,
                    'tenant_payment_terms': None,
                    'tenant_rental_type': None,
                    'tenant_renewal': None,
                    'tenant_renewal_period': None,
                })()
            else:
                # Get existing tenant
                try:
                    tenant_obj = get_object_or_404(tenant, tenant_id=tenant_id)
                except:
                    messages.error(request, 'Selected tenant not found.')
                    return render(request, 'generate_lease_agreement.html', context)
            
            # Generate the lease agreement
            document_path = generate_lease_document(
                country=country,
                language=language,
                property_obj=property_obj,
                tenant_obj=tenant_obj,
                additional_data=additional_data
            )
            
            if document_path:
                # Create filename
                filename = f'lease_agreement_{property_obj.prop_name}_{tenant_obj.tenant_name}_{country}_{language}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
                filename = filename.replace(' ', '_').replace('/', '_')  # Clean filename
                
                # Return the generated document
                try:
                    response = FileResponse(
                        open(document_path, 'rb'),
                        as_attachment=True,
                        filename=filename
                    )
                    # messages.success(request, 'Lease agreement generated successfully!')
                    return response
                except Exception as e:
                    logger.error(f"Error returning file: {str(e)}")
                    messages.error(request, 'Error downloading the generated document.')
            else:
                messages.error(request, 'Error generating lease agreement. Please check the template file exists.')
                
        except Exception as e:
            logger.error(f"Error generating lease agreement: {str(e)}")
            messages.error(request, f'An error occurred while generating the lease agreement: {str(e)}')
    
    return render(request, 'generate_lease_agreement.html', context)

@csrf_exempt
@login_required
@user_passes_test(is_superuser)
@require_http_methods(["POST"])
def get_property_tenant_data(request):
    """
    AJAX endpoint to get property and tenant data for the modal
    """
    try:
        data = json.loads(request.body)
        property_id = data.get('property_id')
        tenant_id = data.get('tenant_id')
        
        response_data = {}
        
        # Get property data
        if property_id:
            try:
                property_obj = props.objects.get(prop_id=property_id)
                
                # Build full address
                address_parts = [
                    property_obj.prop_address1,
                    property_obj.prop_address2,
                    property_obj.prop_suburb,
                    property_obj.prop_city,
                    property_obj.prop_province,
                    property_obj.prop_country,
                    property_obj.prop_pcode
                ]
                full_address = ', '.join([part for part in address_parts if part and part.strip()])
                
                response_data['property'] = {
                    'prop_id': property_obj.prop_id,
                    'prop_name': property_obj.prop_name or '',
                    'prop_address1': property_obj.prop_address1 or '',
                    'prop_address2': property_obj.prop_address2 or '',
                    'prop_suburb': property_obj.prop_suburb or '',
                    'prop_city': property_obj.prop_city or '',
                    'prop_province': property_obj.prop_province or '',
                    'prop_country': property_obj.prop_country or '',
                    'prop_pcode': property_obj.prop_pcode or '',
                    'prop_floor_area': property_obj.prop_floor_area or '',
                    'prop_year_built': property_obj.prop_year_built or '',
                    'full_address': full_address,
                    'prop_electricity': property_obj.prop_electricity or '',
                    'prop_water': property_obj.prop_water or '',
                    'prop_refuse': property_obj.prop_refuse or '',
                    'prop_sewerage': property_obj.prop_sewerage or '',
                    'prop_insurance': property_obj.prop_insurance or '',
                }
            except props.DoesNotExist:
                response_data['property_error'] = 'Property not found'
                logger.warning(f"Property with ID {property_id} not found")
        
        # Get tenant data
        if tenant_id:
            try:
                tenant_obj = tenant.objects.get(tenant_id=tenant_id)
                response_data['tenant'] = {
                    'tenant_id': tenant_obj.tenant_id,
                    'tenant_name': tenant_obj.tenant_name or '',
                    'tenant_type': tenant_obj.tenant_type or '',
                    'tenant_contact_person': tenant_obj.tenant_contact_person or '',
                    'tenant_contact_number': tenant_obj.tenant_contact_number or '',
                    'tenant_email': tenant_obj.tenant_email or '',
                    'tenant_lease_start_date': tenant_obj.tenant_lease_start_date.strftime('%Y-%m-%d') if tenant_obj.tenant_lease_start_date else '',
                    'tenant_lease_end_date': tenant_obj.tenant_lease_end_date.strftime('%Y-%m-%d') if tenant_obj.tenant_lease_end_date else '',
                    'tenant_rent': float(tenant_obj.tenant_rent) if tenant_obj.tenant_rent else 0,
                    'tenant_deposit': float(tenant_obj.tenant_deposit) if tenant_obj.tenant_deposit else 0,
                    'tenant_levies': float(tenant_obj.tenant_levies) if tenant_obj.tenant_levies else 0,
                    'tenant_payment_terms': tenant_obj.tenant_payment_terms or '',
                    'tenant_rental_type': tenant_obj.tenant_rental_type or '',
                    'tenant_renewal': tenant_obj.tenant_renewal or '',
                    'tenant_renewal_period': tenant_obj.tenant_renewal_period or '',
                }
            except tenant.DoesNotExist:
                response_data['tenant_error'] = 'Tenant not found'
                logger.warning(f"Tenant with ID {tenant_id} not found")
        elif tenant_id is None:
            # Handle case where no tenant_id is provided (new tenant case)
            response_data['tenant'] = None
        
        return JsonResponse(response_data)
        
    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error in get_property_tenant_data: {str(e)}")
        return JsonResponse({'error': 'Invalid JSON data'}, status=400)
    except Exception as e:
        logger.error(f"Error in get_property_tenant_data: {str(e)}")
        return JsonResponse({'error': 'Server error occurred'}, status=500)

def get_ordinal_day(day):
    """Convert day number to ordinal (1st, 2nd, 3rd, etc.)"""
    day = int(day)
    if 10 <= day % 100 <= 20:
        suffix = 'th'
    else:
        suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    return f"{day}{suffix}"

def number_to_words_greek(num):
    """Convert number to Greek words with proper capitalization"""
    if num == 0:
        return "Μηδέν"
    
    ones = ["", "Ένα", "Δύο", "Τρία", "Τέσσερα", "Πέντε", "Έξι", "Επτά", "Οκτώ", "Εννέα",
            "Δέκα", "Έντεκα", "Δώδεκα", "Δεκατρία", "Δεκατέσσερα", "Δεκαπέντε", 
            "Δεκαέξι", "Δεκαεπτά", "Δεκαοκτώ", "Δεκαεννέα"]
    
    tens = ["", "", "Είκοσι", "Τριάντα", "Σαράντα", "Πενήντα", "Εξήντα", "Εβδομήντα", "Ογδόντα", "Ενενήντα"]
    
    hundreds = ["", "Εκατό", "Διακόσια", "Τριακόσια", "Τετρακόσια", "Πεντακόσια", 
                "Εξακόσια", "Επτακόσια", "Οκτακόσια", "Εννιακόσια"]
    
    num = int(num)  # Remove decimals
    
    if num < 20:
        return ones[num]
    elif num < 100:
        return tens[num // 10] + ("" if num % 10 == 0 else " " + ones[num % 10])
    elif num < 1000:
        return hundreds[num // 100] + ("" if num % 100 == 0 else " " + number_to_words_greek(num % 100))
    elif num < 1000000:
        thousands = num // 1000
        remainder = num % 1000
        
        if thousands == 1:
            thousands_text = "Χίλια"
        elif thousands < 20:
            thousands_text = ones[thousands] + " Χιλιάδες"
        else:
            thousands_text = number_to_words_greek(thousands) + " Χιλιάδες"
            
        return thousands_text + ("" if remainder == 0 else " " + number_to_words_greek(remainder))
    else:
        return str(int(num))  # Fallback for very large numbers

def format_date_greek(date_obj):
    """Convert date object to Greek format with ordinal day and Greek month name"""
    if not date_obj:
        return 'N/A'
    
    # Greek month names (genitive case for dates)
    greek_months = {
        1: "Ιανουαρίου", 2: "Φεβρουαρίου", 3: "Μαρτίου", 4: "Απριλίου",
        5: "Μαΐου", 6: "Ιουνίου", 7: "Ιουλίου", 8: "Αυγούστου",
        9: "Σεπτεμβρίου", 10: "Οκτωβρίου", 11: "Νοεμβρίου", 12: "Δεκεμβρίου"
    }
    
    day = date_obj.day
    month = greek_months[date_obj.month]
    year = date_obj.year
    
    # Add Greek ordinal suffix to day
    ordinal_day = f"{day}η"
    
    return f"{ordinal_day} {month} {year}"

def translate_to_greek(text):
    """Translate common property and name terms to Greek"""
    if not text or text.strip() == '':
        return text
    
    # Complete translations dictionary
    translations = {
        # Common street types
        'Street': 'Οδός',
        'St': 'Οδός',
        'Avenue': 'Λεωφόρος',
        'Ave': 'Λεωφόρος',
        'Road': 'Δρόμος',
        'Rd': 'Δρόμος',
        'Lane': 'Δρομάκι',
        'Square': 'Πλατεία',
        'Plaza': 'Πλατεία',
        
        # Specific street names and areas
        'Eleftheroupoleos': 'Ελευθερουπόλεως',
        'Agias': 'Αγίας',
        'Annas': 'Άννας',
        'Dikaiosynis': 'Δικαιοσύνης',
        'Ionion': 'Ιόνιον',
        'Aristoteli': 'Αριστοτέλη',
        'Valaoriti': 'Βαλαωρίτη',
        'Pindarou': 'Πινδάρου',
        'Evagora': 'Ευαγόρα',
        'Palikaridi': 'Παλικαρίδη',
        'Agios': 'Άγιος',
        'Dometios': 'Δομέτιος',
        'Aristoteli Valaoriti': 'Αριστοτέλη Βαλαωρίτη',
        
        # Common area names in Cyprus
        'Nicosia': 'Λευκωσία',
        'Limassol': 'Λεμεσός',
        'Larnaca': 'Λάρνακα',
        'Paphos': 'Πάφος',
        'Famagusta': 'Αμμόχωστος',
        'Kyrenia': 'Κερύνεια',
        'Engomi': 'Έγκωμη',
        'Strovolos': 'Στρόβολος',
        'Lakatamia': 'Λακαταμια',
        'Latsia': 'Λατσιά',
        'Aglandjia': 'Αγλαντζιά',
        'Agia Thekla': 'Αγία Θέκλα',
        'Sotira': 'Σωτήρα',
        'Cyprus': 'Κύπρος',
        'Agios Dometios': 'Άγιος Δομέτιος',
        
        # Common building types
        'Flat': 'Διαμέρισμα',
        'Apartment': 'Διαμέρισμα',
        'House': 'Σπίτι',
        'Villa': 'Βίλα',
        'Villas': 'Βίλες',
        'Office': 'Γραφείο',
        'Building': 'Κτίριο',
        'Complex': 'Συγκρότημα',
        'Tower': 'Πύργος',
        'Center': 'Κέντρο',
        'Centre': 'Κέντρο',
        
        # FURNITURE AND APPLIANCES
        'Television': 'Τηλεόραση',
        'Televisions': 'Τηλεοράσεις',
        'TV': 'Τηλεόραση',
        'TVs': 'Τηλεοράσεις',
        'Washing Machine': 'Πλυντήριο',
        'Washing Machines': 'Πλυντήρια',
        'Dishwasher': 'Πλυντήριο Πιάτων',
        'Dishwashers': 'Πλυντήρια Πιάτων',
        'Oven': 'Φούρνος',
        'Ovens': 'Φούρνοι',
        'Stove': 'Ηλεκτρικές εστίες',
        'Stoves': 'Ηλεκτρικές εστίες',
        'Extractor Fan': 'Απορροφητήρας',
        'Extractor Fans': 'Απορροφητήρες',
        'Fridge Freezer': 'Ψυγειοκαταψύκτης',
        'Fridge Freezers': 'Ψυγειοκαταψύκτες',
        'Airconditioner': 'Κλιματιστικό',
        'Airconditioners': 'Κλιματιστικά',
        'Air Conditioner': 'Κλιματιστικό',
        'Air Conditioners': 'Κλιματιστικά',
        'Microwave': 'Φούρνος Μικροκυμάτων',
        'Microwaves': 'Φούρνοι Μικροκυμάτων',
        'Three Seater Couch': 'Τριθέσιος Καναπές',
        'Three Seater Couches': 'Τριθέσιοι Καναπέδες',
        'Two Seater Couch': 'Διθέσιος Καναπές',
        'Two Seater Couches': 'Διθέσιοι Καναπέδες',
        'Lounge Chair': 'Πολυθρόνα',
        'Lounge Chairs': 'Πολυθρόνες',
        'Coffee Tables': 'Τραπεζάκια Σαλονιού',
        'Coffee Table': 'Τραπεζάκι Σαλονιού',
        'Corner Tables': 'Γωνιακά Τραπέζια',
        'Corner Table': 'Γωνιακό Τραπέζι',
        'Standing Lamps': 'Φωτιστικά Δαπέδου',
        'Standing Lamp': 'Φωτιστικό Δαπέδου',
        'Table Lamps': 'Φωτιστικά Τραπεζιού',
        'Table Lamp': 'Φωτιστικό Τραπεζιού',
        'Dining Tables': 'Τραπεζαρίες',
        'Dining Table': 'Τραπεζαρία',
        'Dining Chairs': 'Καρέκλες Τραπεζαρίας',
        'Dining Chair': 'Καρέκλα Τραπεζαρίας',
        'Single Beds': 'Μονά Κρεβάτια',
        'Single Bed': 'Μονό Κρεβάτι',
        'Queen Beds': 'Διπλά Κρεβάτια',
        'Queen Bed': 'Διπλό Κρεβάτι',
        'Double Bed': 'Διπλό Κρεβάτι',
        'Double Beds': 'Διπλά Κρεβάτια',
        'Outside Tables': 'Εξωτερικά Τραπέζια',
        'Outside Table': 'Εξωτερικό Τραπέζι',
        'Outside Chairs': 'Εξωτερικές Καρέκλες',
        'Outside Chair': 'Εξωτερική Καρέκλα',
        'TV Table': 'Τραπέζι Tηλεόρασης',
        'TV Tables': 'Τραπέζια Tηλεόρασης',
        'Bed': 'Κρεβάτι',
        'Beds': 'Κρεβάτια',
        'Sofa': 'Καναπές',
        'Sofas': 'Καναπέδες',
        'Table': 'Τραπέζι',
        'Tables': 'Τραπέζια',
        'Chair': 'Καρέκλα',
        'Chairs': 'Καρέκλες',
        'Lamp': 'Λάμπα',
        'Lamps': 'Λάμπες',
        'Desk': 'Γραφείο',
        'Desks': 'Γραφεία',
        'Wardrobe': 'Ντουλάπα',
        'Wardrobes': 'Ντουλάπες',
        'Refrigerator': 'Ψυγείο',
        'Refrigerators': 'Ψυγεία',
        'Fridge': 'Ψυγείο',
        'Fridges': 'Ψυγεία',
        'Heater': 'Θερμάστρα',
        'Heaters': 'Θερμάστρες',
        'Curtain': 'Κουρτίνα',
        'Curtains': 'Κουρτίνες',
        'Mirror': 'Καθρέφτης',
        'Mirrors': 'Καθρέφτες',
        'Bookshelf': 'Βιβλιοθήκη',
        'Bookshelves': 'Βιβλιοθήκες',
        'AC': 'Κλιματιστικό',
        'ACs': 'Κλιματιστικά',
        'Bedside Table': 'Κομοδίνο',
        'Bedside Tables': 'Κομοδίνα',
        'Roller Blind': 'Ρολό',
        'Roller Blinds': 'Ρολά',
        'Stool': 'Σκαμπό',
        'Stools': 'Σκαμπό',
        
        # Keys - Complete singular and plural forms
        'Key': 'Κλειδί',
        'Keys': 'Κλειδιά',
        'Front Door': 'Κεντρική Πόρτα',
        'Front Door Key': 'Κλειδί Κεντρικής Πόρτας',
        'Front Door Keys': 'Κλειδιά Κεντρικής Πόρτας',
        'Main Door': 'Κεντρική Πόρτα',
        'Main Door Key': 'Κλειδί Κεντρικής Πόρτας',
        'Main Door Keys': 'Κλειδιά Κεντρικής Πόρτας',
        'Building Key': 'Κλειδί Κτιρίου',
        'Building Keys': 'Κλειδιά Κτιρίου',
        'Entrance': 'Είσοδος',
        'Entrance Key': 'Κλειδί Εισόδου',
        'Entrance Keys': 'Κλειδιά Εισόδου',
        'Mailbox': 'Γραμματοκιβώτιο',
        'Mailbox Key': 'Κλειδί Γραμματοκιβωτίου',
        'Mailbox Keys': 'Κλειδιά Γραμματοκιβωτίου',
        'Post Box Key': 'Κλειδί Γραμματοκιβωτίου',
        'Post Box Keys': 'Κλειδιά Γραμματοκιβωτίου',
        'Storeroom Key': 'Κλειδί Αποθήκης',
        'Storeroom Keys': 'Κλειδιά Αποθήκης',
        'Storage Key': 'Κλειδί Αποθήκης',
        'Storage Keys': 'Κλειδιά Αποθήκης',
        'Garage Key': 'Κλειδί Γκαράζ',
        'Garage Keys': 'Κλειδιά Γκαράζ',
        'Garage': 'Γκαράζ',
        'Storage': 'Αποθήκη',
        'Balcony Key': 'Κλειδί Μπαλκονιού',
        'Balcony Keys': 'Κλειδιά Μπαλκονιού',
        'Balcony': 'Μπαλκόνι',
        'Terrace Key': 'Κλειδί Βεράντας',
        'Terrace Keys': 'Κλειδιά Βεράντας',
        'Terrace': 'Βεράντα',
        
        # Common names - CORRECTED
        'Demetri': 'Δημήτρης',
        'Manias': 'Μανιάς',
        'Foti': 'Φώτι',  # CORRECTED - was Φώτης
        'Pitta': 'Πίττα',
        'George': 'Γεώργιος',
        'Maria': 'Μαρία',
        'Andreas': 'Ανδρέας',
        'Christina': 'Χριστίνα',
        'Kostas': 'Κώστας',
        'Nikos': 'Νίκος',
        
        # Company terms
        'Limited': 'Περιορισμένη',
        'Ltd': 'Περιορισμένη',
        'Company': 'Εταιρεία',
        'Co': 'Εταιρεία',
    }
    
    # Simple direct matching - no regex complications
    if text in translations:
        return translations[text]
    
    # Case insensitive fallback
    for eng, greek in translations.items():
        if text.lower() == eng.lower():
            return greek
    
    # Handle compound names and addresses by translating individual parts
    if ' ' in text:
        words = text.split()
        translated_words = []
        for word in words:
            # Clean punctuation
            clean_word = word.strip('.,;:!?()[]{}')
            punctuation = word[len(clean_word):]
            
            if clean_word in translations:
                translated_words.append(translations[clean_word] + punctuation)
            else:
                # Case insensitive check
                found = False
                for eng, greek in translations.items():
                    if clean_word.lower() == eng.lower():
                        translated_words.append(greek + punctuation)
                        found = True
                        break
                if not found:
                    translated_words.append(word)
        
        result = ' '.join(translated_words)
        if result != text:
            return result
    
    return text

def generate_lease_document(country, language, property_obj, tenant_obj, additional_data):
    """
    Generate the actual lease agreement document using Word templates
    """
    try:
        # Define template paths based on country and language
        template_mapping = {
            'cyprus': {
                'english': 'lease_templates/cyprus_english_lease_template.docx',
                'greek': 'lease_templates/cyprus_greek_lease_template.docx'
            },
            'greece': {
                'greek': 'lease_templates/greece_greek_lease_template.docx'
            },
            'spain': {
                'spanish': 'lease_templates/spain_spanish_lease_template.docx'
            }
        }
        
        # Get template path
        template_filename = template_mapping.get(country, {}).get(language)
        if not template_filename:
            logger.error(f"No template mapping found for country: {country}, language: {language}")
            return None
            
        template_path = os.path.join(settings.BASE_DIR, 'pages', 'templates', template_filename)

        # Check if template exists
        if not os.path.exists(template_path):
            logger.error(f"Template file not found: {template_path}")
            # Try fallback - create basic document
            return create_basic_lease_document(country, language, property_obj, tenant_obj, additional_data)
        
        # Load template
        doc = DocxTemplate(template_path)
        
        # Prepare data for template
        template_data = prepare_lease_template_data(
            country, language, property_obj, tenant_obj, additional_data
        )
        
        # Render document
        doc.render(template_data)
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        
        logger.info(f"Lease document generated successfully: {temp_file.name}")
        return temp_file.name
        
    except Exception as e:
        logger.error(f"Error in generate_lease_document: {str(e)}")
        return None

def prepare_lease_template_data(country, language, property_obj, tenant_obj, additional_data):
    """
    Prepare data dictionary for the lease agreement template using your models
    """
    # Get current date components
    current_date = datetime.now()
    
    # Helper function to convert numbers to words with proper capitalization
    def number_to_words(num):
        """Convert number to English words with proper capitalization"""
        if num == 0:
            return "Zero"
        
        ones = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine", 
                "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen", 
                "Seventeen", "Eighteen", "Nineteen"]
        
        tens = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
        
        num = int(num)  # Remove decimals
        
        if num < 20:
            return ones[num]
        elif num < 100:
            return tens[num // 10] + ("" if num % 10 == 0 else " " + ones[num % 10])
        elif num < 1000:
            return ones[num // 100] + " Hundred" + ("" if num % 100 == 0 else " " + number_to_words(num % 100))
        elif num < 1000000:
            return number_to_words(num // 1000) + " Thousand" + ("" if num % 1000 == 0 else " " + number_to_words(num % 1000))
        else:
            return str(int(num))  # Fallback for very large numbers
    
    # Helper function to pluralize items
    def pluralize_item(item_name, count):
        """Add 's' to item names when count > 1"""
        if count > 1:
            # Handle special cases - check the last word only for compound items
            words = item_name.split()
            last_word = words[-1]
            
            if last_word.endswith('y') and len(last_word) > 1 and last_word[-2] not in 'aeiou':
                # Only change y to ies for words ending in consonant + y (like "Key" but not "Boy")
                words[-1] = last_word[:-1] + 'ies'
                return ' '.join(words)
            elif last_word.endswith(('s', 'sh', 'ch', 'x', 'z')):
                words[-1] = last_word + 'es'
                return ' '.join(words)
            else:
                words[-1] = last_word + 's'
                return ' '.join(words)
        return item_name
    
    # Get dates from additional_data or tenant object
    start_date_str = additional_data.get('start_date') or (
        tenant_obj.tenant_lease_start_date.strftime('%Y-%m-%d') 
        if hasattr(tenant_obj, 'tenant_lease_start_date') and tenant_obj.tenant_lease_start_date else ''
    )
    end_date_str = additional_data.get('end_date') or (
        tenant_obj.tenant_lease_end_date.strftime('%Y-%m-%d') 
        if hasattr(tenant_obj, 'tenant_lease_end_date') and tenant_obj.tenant_lease_end_date else ''
    )
    
    # Convert date strings to datetime objects
    start_date_obj = None
    end_date_obj = None
    duration_days = 0
    duration_months = 0
    
    try:
        if start_date_str:
            start_date_obj = datetime.strptime(start_date_str, '%Y-%m-%d').date()
        if end_date_str:
            end_date_obj = datetime.strptime(end_date_str, '%Y-%m-%d').date()
        
        if start_date_obj and end_date_obj:
            duration_days = (end_date_obj - start_date_obj).days
            duration_months = max(1, round(duration_days / 30.44))  # More accurate month calculation
    except ValueError as e:
        logger.warning(f"Date parsing error: {str(e)}")
    
    # Get financial information - prefer additional_data over database (remove decimals)
    monthly_rent = int(float(additional_data.get('monthly_rent') or (
        tenant_obj.tenant_rent if hasattr(tenant_obj, 'tenant_rent') and tenant_obj.tenant_rent else 0
    )))
    security_deposit = int(float(additional_data.get('security_deposit') or (
        tenant_obj.tenant_deposit if hasattr(tenant_obj, 'tenant_deposit') and tenant_obj.tenant_deposit else 0
    )))
    communal_expenses = int(float(additional_data.get('communal_expenses', 0)))
    
    # Get extension information
    extension_period = int(additional_data.get('extension_period', 0))
    extension_date_str = additional_data.get('extension_date', '')
    extension_date_obj = None
    if extension_date_str:
        try:
            extension_date_obj = datetime.strptime(extension_date_str, '%Y-%m-%d').date()
        except ValueError:
            pass
    
    # Handle rental increase percentage (remove decimals if whole number)
    rental_increase = float(additional_data.get('rental_increase', 0))
    rental_increase_formatted = f"{int(rental_increase)}" if rental_increase == int(rental_increase) else f"{rental_increase}"
    
    # Get currency symbol based on country
    currency_symbols = {
        'cyprus': '€',
        'greece': '€',
        'spain': '€'
    }
    currency = currency_symbols.get(country, '€')
    
    # Build full property address
    address_parts = [
        property_obj.prop_address1,
        property_obj.prop_address2,
        property_obj.prop_suburb,
        property_obj.prop_city,
        property_obj.prop_province,
        property_obj.prop_country,
        property_obj.prop_pcode
    ]
    full_address = ', '.join([part for part in address_parts if part and part.strip()])
    
    # Get tenant information (handle both existing and new tenants)
    if additional_data.get('is_new_tenant') and additional_data.get('new_tenant_data'):
        new_tenant_data = additional_data['new_tenant_data']
        tenant_name = new_tenant_data.get('tenant_name', 'N/A')
        tenant_address = new_tenant_data.get('tenant_address', 'N/A')
        tenant_phone = new_tenant_data.get('tenant_contact_number', 'N/A')
        tenant_email = new_tenant_data.get('tenant_email', 'N/A')
        tenant_type = new_tenant_data.get('tenant_type', 'Individual')
        
        if tenant_type == 'Company':
            tenant_passport_id = 'N/A'
            tenant_passport_country = 'N/A'
            tenant_registration_number = new_tenant_data.get('tenant_registration_number', 'N/A')
            tenant_contact_person = new_tenant_data.get('tenant_contact_person', 'N/A')
        else:
            tenant_passport_id = new_tenant_data.get('tenant_passport_id', 'N/A')
            tenant_passport_country = new_tenant_data.get('tenant_passport_country', 'N/A')
            tenant_registration_number = 'N/A'
            tenant_contact_person = 'N/A'
    else:
        tenant_name = tenant_obj.tenant_name or 'N/A'
        # Use property address as tenant address for existing tenants
        tenant_address = full_address or 'N/A'
        tenant_phone = getattr(tenant_obj, 'tenant_contact_number', 'N/A') or 'N/A'
        tenant_email = getattr(tenant_obj, 'tenant_email', 'N/A') or 'N/A'
        tenant_type = getattr(tenant_obj, 'tenant_type', 'Individual') or 'Individual'
        tenant_contact_person = getattr(tenant_obj, 'tenant_contact_person', 'N/A') or 'N/A'

        # Get additional information for existing tenants from form data
        existing_tenant_data = additional_data.get('existing_tenant_data', {})
        if tenant_type == 'Company':
            tenant_passport_id = 'N/A'
            tenant_passport_country = 'N/A'
            tenant_registration_number = existing_tenant_data.get('tenant_registration_number', 'N/A')
        else:
            tenant_passport_id = existing_tenant_data.get('tenant_passport_id', 'N/A')
            tenant_passport_country = existing_tenant_data.get('tenant_passport_country', 'N/A')
            tenant_registration_number = 'N/A'
    
    # Second tenant information
    has_second_tenant = additional_data.get('has_second_tenant', False)
    second_tenant_data = additional_data.get('second_tenant_data', {})
    
    # Second tenant variables for template
    if has_second_tenant and second_tenant_data:
        second_tenant_name = second_tenant_data.get('tenant_name', 'N/A')
        second_tenant_passport_id = second_tenant_data.get('tenant_passport_id', 'N/A')
        second_tenant_passport_country = second_tenant_data.get('tenant_passport_country', 'N/A')
        second_tenant_address = second_tenant_data.get('tenant_address', 'N/A')
        second_tenant_contact_number = second_tenant_data.get('tenant_contact_number', 'N/A')
        second_tenant_email = second_tenant_data.get('tenant_email', 'N/A')
    else:
        second_tenant_name = ''
        second_tenant_passport_id = ''
        second_tenant_passport_country = ''
        second_tenant_address = ''
        second_tenant_contact_number = ''
        second_tenant_email = ''
    
    # Process furniture and appliances
    furniture_items = {}
    furniture_list = []
    furniture_list_greek = []
    for key, value in additional_data.items():
        if key.startswith('furniture_') and int(value or 0) > 0:
            count = int(value)
            item_name = key.replace('furniture_', '').replace('_', ' ').title()
            pluralized_name = pluralize_item(item_name, count)
            furniture_items[pluralized_name] = count
            furniture_list.append(f"{count} x {pluralized_name}")
            
            # Create Greek version - translate the pluralized name
            item_name_greek = translate_to_greek(pluralized_name)
            furniture_list_greek.append(f"{count} x {item_name_greek}")
            
    # Get fully furnished checkbox value
    fully_furnished = additional_data.get('fully_furnished') == 'true'
    
    # Get parking and storeroom checkbox values
    has_parking = additional_data.get('amenity_parking') == True  # Boolean, not string
    has_storeroom = additional_data.get('amenity_storeroom') == True  # Boolean, not string

    # Process keys
    keys_items = {}
    keys_list = []
    keys_list_greek = []
    for key, value in additional_data.items():
        if key.startswith('keys_') and int(value or 0) > 0:
            count = int(value)
            key_name = key.replace('keys_', '').replace('_', ' ').title() + ' Key'
            pluralized_key = pluralize_item(key_name, count)
            keys_items[pluralized_key] = count
            keys_list.append(f"{count} x {pluralized_key}")
            
            # Create Greek version - translate the pluralized key name
            key_name_greek = translate_to_greek(pluralized_key)
            keys_list_greek.append(f"{count} x {key_name_greek}")
            
    # Prepare template data
    template_data = {
        # Current date components with ordinal day
        'today_day': get_ordinal_day(current_date.day),
        'today_month': current_date.strftime('%B'),
        'today_year': current_date.strftime('%Y'),
        'today_date': current_date.strftime('%B %d, %Y'),
        
        # Greek date formats
        'today_date_greek': format_date_greek(current_date.date()),
        'lease_start_date_greek': format_date_greek(start_date_obj) if start_date_obj else 'N/A',
        'lease_end_date_greek': format_date_greek(end_date_obj) if end_date_obj else 'N/A',
        
        # Document information
        'document_date': current_date.strftime('%B %d, %Y'),
        'document_id': f'LA-{current_date.strftime("%Y%m%d-%H%M%S")}',
        'country': country.title(),
        'language': language.title(),
        
        # Landlord information
        'landlord_name': additional_data.get('landlord_name', 'Alivente Limited'),
        'landlord_contact_person': additional_data.get('landlord_contact_person', 'Demetri Manias'),
        'landlord_contact_person_greek': translate_to_greek(additional_data.get('landlord_contact_person', 'Demetri Manias')),
        'landlord_registration_number': additional_data.get('landlord_registration_number', 'HE123456'),
        'landlord_phone': additional_data.get('landlord_phone', '+357-96668557'),
        'landlord_phone_number': additional_data.get('landlord_phone', '+357-96668557'),
        'landlord_email': additional_data.get('landlord_email', 'demetri.manias@alivente.com'),
        'landlord_address': additional_data.get('landlord_address', 'Alivente House, Dikaiosynis 13A, Engomi, Nicosia, Cyprus, 2412'),
        
        # Tenant information
        'tenant_name': tenant_name,
        'tenant_full_name': tenant_name,
        'tenant_passport_id': tenant_passport_id,
        'tenant_passport_country': tenant_passport_country,
        'tenant_address': tenant_address,
        'tenant_phone': tenant_phone,
        'tenant_email': tenant_email,
        'tenant_type': tenant_type,
        'tenant_registration_number': tenant_registration_number,
        'tenant_contact_person': tenant_contact_person,
        
        # Second tenant information
        'has_second_tenant': has_second_tenant,
        'second_tenant_name': second_tenant_name,
        'second_tenant_passport_id': second_tenant_passport_id,
        'second_tenant_passport_country': second_tenant_passport_country,
        'second_tenant_address': second_tenant_address,
        'second_tenant_contact_number': second_tenant_contact_number,
        'second_tenant_email': second_tenant_email,
        
        # Property information
        'property_name': property_obj.prop_name or 'N/A',
        'property_address': full_address or 'N/A',
        'property_address1': property_obj.prop_address1 or '',
        'property_address2': property_obj.prop_address2 or '',
        'property_suburb': property_obj.prop_suburb or '',
        'property_city': property_obj.prop_city or '',
        'property_province': property_obj.prop_province or '',
        'property_country': property_obj.prop_country or '',
        'property_pcode': property_obj.prop_pcode or '',
        'property_floor_area': property_obj.prop_floor_area or 'N/A',
        'property_year_built': property_obj.prop_year_built or 'N/A',
        
        # Property information - Greek translations
        'property_address1_greek': translate_to_greek(property_obj.prop_address1 or ''),
        'property_address2_greek': translate_to_greek(property_obj.prop_address2 or ''),
        'property_suburb_greek': translate_to_greek(property_obj.prop_suburb or ''),
        'property_city_greek': translate_to_greek(property_obj.prop_city or ''),
        'property_province_greek': translate_to_greek(property_obj.prop_province or ''),
        
        # Lease terms
        'lease_start_date': start_date_obj.strftime('%B %d, %Y') if start_date_obj else 'N/A',
        'lease_end_date': end_date_obj.strftime('%B %d, %Y') if end_date_obj else 'N/A',
        'lease_start_date_short': start_date_obj.strftime('%d/%m/%Y') if start_date_obj else 'N/A',
        'lease_end_date_short': end_date_obj.strftime('%d/%m/%Y') if end_date_obj else 'N/A',
        'duration_months': duration_months,
        'duration_months_words': number_to_words(duration_months),
        'duration_days': duration_days,
        
        # Financial terms - numbers (without decimals)
        'rental_amount': monthly_rent,
        'communal_expenses': communal_expenses,
        'security_deposit': security_deposit,
        'deposit_amount': security_deposit,
        
        # Financial terms - words (with proper capitalization)
        'rental_amount_words': number_to_words(monthly_rent),
        'communal_expenses_words': number_to_words(communal_expenses),
        'deposit_amount_words': number_to_words(security_deposit),
        
        # Greek financial terms in words
        'rental_amount_words_greek': number_to_words_greek(monthly_rent),
        'deposit_amount_words_greek': number_to_words_greek(security_deposit),
        'communal_expenses_words_greek': number_to_words_greek(communal_expenses),
        
        # Financial terms - with thousand separators (the template uses the raw numbers)
        'rental_amount': f'{monthly_rent:,}',
        'communal_expenses': f'{communal_expenses:,}',
        'security_deposit': f'{security_deposit:,}',
        
        # Additional formatted versions for compatibility
        'rental_amount_formatted': f'{currency}{monthly_rent:,}',
        'communal_expenses_formatted': f'{currency}{communal_expenses:,}',
        'deposit_amount_formatted': f'{currency}{security_deposit:,}',
        'security_deposit_formatted': f'{currency}{security_deposit:,}',
        'total_rent_formatted': f'{currency}{(monthly_rent * duration_months):,}' if duration_months else f'{currency}0',
        'monthly_rent_formatted': f'{currency}{monthly_rent:,}',
        
        # Extension and increase terms
        'lease_extension_period': extension_period,
        'lease_extension_period_words': number_to_words(extension_period),
        'lease_extension_period_words_greek': number_to_words_greek(extension_period),
        'lease_extension_date': extension_date_obj.strftime('%B %d, %Y') if extension_date_obj else 'N/A',
        'lease_extension_date_greek': format_date_greek(extension_date_obj) if extension_date_obj else 'N/A',
        'percentage_rental_increase': rental_increase_formatted,
        
        # Currency
        'currency': currency,
        
        # Furniture and appliances (with proper pluralization and comma-separated lists)
        'furniture_items': furniture_items,
        'furniture_list': ', '.join(furniture_list) if furniture_list else 'None',
        'furniture_list_greek': ', '.join(furniture_list_greek) if furniture_list_greek else 'Κανένα',
        'fully_furnished': fully_furnished,
        
        # Parking and Storeroom
        'has_parking': has_parking,
        'has_storeroom': has_storeroom,
        
        # Keys (with proper pluralization and comma-separated lists)
        'keys_items': keys_items,
        'keys_list': ', '.join(keys_list) if keys_list else 'None',
        'keys_list_greek': ', '.join(keys_list_greek) if keys_list_greek else 'Κανένα',
        
        # Legacy fields for compatibility
        'start_date': start_date_obj.strftime('%B %d, %Y') if start_date_obj else 'N/A',
        'end_date': end_date_obj.strftime('%B %d, %Y') if end_date_obj else 'N/A',
        'monthly_rent': monthly_rent,
        
        # Generation metadata
        'generation_timestamp': current_date.strftime('%Y-%m-%d %H:%M:%S'),
    }
    
    return template_data

def create_basic_lease_document(country, language, property_obj, tenant_obj, additional_data):
    """
    Create a basic lease document when no template is available
    This is a fallback method for demonstration purposes
    """
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Helper function to pluralize items (for fallback function)
        def pluralize_item(item_name, count):
            """Add 's' to item names when count > 1"""
            if count > 1:
                # Handle special cases
                if item_name.endswith('y'):
                    return item_name[:-1] + 'ies'  # e.g., "Key" -> "Keys"
                elif item_name.endswith(('s', 'sh', 'ch', 'x', 'z')):
                    return item_name + 'es'
                else:
                    return item_name + 's'
            return item_name
        
        # Create a new document
        document = Document()
        
        # Add title
        title = document.add_heading('LEASE AGREEMENT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = document.add_paragraph(f'{country.title()} - {language.title()}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        document.add_paragraph('')  # Empty line
        
        # Document info
        document.add_paragraph(f'Document ID: LA-{datetime.now().strftime("%Y%m%d-%H%M%S")}')
        document.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
        document.add_paragraph('')
        
        # Property information
        document.add_heading('Property Information', level=1)
        document.add_paragraph(f'Property Name: {property_obj.prop_name or "N/A"}')
        
        address_parts = [
            property_obj.prop_address1,
            property_obj.prop_address2,
            property_obj.prop_suburb,
            property_obj.prop_city,
            property_obj.prop_province,
            property_obj.prop_country,
            property_obj.prop_pcode
        ]
        full_address = ', '.join([part for part in address_parts if part])
        document.add_paragraph(f'Address: {full_address or "N/A"}')
        
        if property_obj.prop_floor_area:
            document.add_paragraph(f'Floor Area: {property_obj.prop_floor_area} m²')
        if property_obj.prop_year_built:
            document.add_paragraph(f'Year Built: {property_obj.prop_year_built}')
        
        document.add_paragraph('')
        
        # Tenant information
        document.add_heading('Tenant Information', level=1)
        if additional_data.get('is_new_tenant') and additional_data.get('new_tenant_data'):
            new_tenant_data = additional_data['new_tenant_data']
            document.add_paragraph(f'Tenant Name: {new_tenant_data.get("tenant_name", "N/A")}')
            document.add_paragraph(f'Passport/ID: {new_tenant_data.get("tenant_passport_id", "N/A")}')
            document.add_paragraph(f'Email: {new_tenant_data.get("tenant_email", "N/A")}')
            document.add_paragraph(f'Phone: {new_tenant_data.get("tenant_contact_number", "N/A")}')
            document.add_paragraph(f'Address: {new_tenant_data.get("tenant_address", "N/A")}')
        else:
            document.add_paragraph(f'Tenant Name: {tenant_obj.tenant_name or "N/A"}')
            document.add_paragraph(f'Email: {getattr(tenant_obj, "tenant_email", "N/A") or "N/A"}')
            document.add_paragraph(f'Phone: {getattr(tenant_obj, "tenant_contact_number", "N/A") or "N/A"}')
        
        document.add_paragraph('')
        
        # Lease terms
        document.add_heading('Lease Terms', level=1)
        
        # Get dates
        start_date = additional_data.get('start_date') or (
            tenant_obj.tenant_lease_start_date.strftime('%Y-%m-%d') 
            if hasattr(tenant_obj, 'tenant_lease_start_date') and tenant_obj.tenant_lease_start_date else 'N/A'
        )
        end_date = additional_data.get('end_date') or (
            tenant_obj.tenant_lease_end_date.strftime('%Y-%m-%d') 
            if hasattr(tenant_obj, 'tenant_lease_end_date') and tenant_obj.tenant_lease_end_date else 'N/A'
        )
        
        document.add_paragraph(f'Start Date: {start_date}')
        document.add_paragraph(f'End Date: {end_date}')
        
        # Financial terms (without decimals, with thousand separators)
        monthly_rent = int(float(additional_data.get('monthly_rent') or (
            tenant_obj.tenant_rent if hasattr(tenant_obj, 'tenant_rent') and tenant_obj.tenant_rent else 0
        )))
        security_deposit = int(float(additional_data.get('security_deposit') or (
            tenant_obj.tenant_deposit if hasattr(tenant_obj, 'tenant_deposit') and tenant_obj.tenant_deposit else 0
        )))
        communal_expenses = int(float(additional_data.get('communal_expenses', 0)))
        
        document.add_paragraph(f'Monthly Rent: €{monthly_rent:,}')
        document.add_paragraph(f'Security Deposit: €{security_deposit:,}')
        if communal_expenses > 0:
            document.add_paragraph(f'Communal Expenses: €{communal_expenses:,}')
        
        if hasattr(tenant_obj, 'tenant_levies') and tenant_obj.tenant_levies:
            document.add_paragraph(f'Levies: €{int(float(tenant_obj.tenant_levies)):,}')
        
        # Furniture and appliances (with pluralization)
        furniture_list = []
        for key, value in additional_data.items():
            if key.startswith('furniture_') and int(value or 0) > 0:
                count = int(value)
                item_name = key.replace('furniture_', '').replace('_', ' ').title()
                pluralized_name = pluralize_item(item_name, count)
                furniture_list.append(f"{count} x {pluralized_name}")
        
        if furniture_list:
            document.add_paragraph('')
            document.add_heading('Furniture and Appliances', level=1)
            document.add_paragraph(', '.join(furniture_list))
        
        # Keys (with pluralization)
        keys_list = []
        for key, value in additional_data.items():
            if key.startswith('keys_') and int(value or 0) > 0:
                count = int(value)
                key_name = key.replace('keys_', '').replace('_', ' ').title() + ' Key'
                pluralized_key = pluralize_item(key_name, count)
                keys_list.append(f"{count} x {pluralized_key}")
        
        if keys_list:
            document.add_paragraph('')
            document.add_heading('Keys Provided', level=1)
            document.add_paragraph(', '.join(keys_list))
        
        # Signatures
        document.add_paragraph('')
        document.add_heading('Signatures', level=1)
        document.add_paragraph('Landlord: _________________________    Date: _________')
        document.add_paragraph('Tenant: _________________________     Date: _________')
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        document.save(temp_file.name)
        
        logger.info(f"Basic lease document created: {temp_file.name}")
        return temp_file.name
        
    except Exception as e:
        logger.error(f"Error creating basic lease document: {str(e)}")
        return None

### CASH FLOW ###
@login_required
def cashflow_forecast(request):
    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        horizon_months = int(request.GET.get("horizon", 12))
        today = date.today()
        
        # Calculate horizon cutoff date
        horizon_year = today.year + (today.month + horizon_months - 1) // 12
        horizon_month = (today.month + horizon_months - 1) % 12 + 1
        horizon_last_day = monthrange(horizon_year, horizon_month)[1]
        cutoff_date = date(horizon_year, horizon_month, horizon_last_day)
        
        expenses = []
        
        # Handle budget expenses - they repeat every year
        for exp in expense.objects.select_related("prop", "expense_line_types").all():
            month_fields = [
                ("expense_jan", 1), ("expense_feb", 2), ("expense_mar", 3),
                ("expense_apr", 4), ("expense_may", 5), ("expense_jun", 6),
                ("expense_jul", 7), ("expense_aug", 8), ("expense_sep", 9),
                ("expense_oct", 10), ("expense_nov", 11), ("expense_dec", 12),
            ]
            
            for field, month_idx in month_fields:
                amount = getattr(exp, field)
                if amount and amount > decimal.Decimal("0.00"):
                    # Generate expenses for each year within the horizon
                    current_year = today.year
                    while True:
                        last_day = monthrange(current_year, month_idx)[1]
                        due_date = date(current_year, month_idx, last_day)
                        
                        # Break if beyond our cutoff date
                        if due_date > cutoff_date:
                            break
                        
                        # Only include if the date is today or in the future
                        if due_date >= today:
                            days_ahead = (due_date - today).days
                            if days_ahead <= 30:
                                color = "red"
                            elif days_ahead <= 90:
                                color = "orange"
                            else:
                                color = "green"
                                
                            expenses.append({
                                "id": f"{exp.expense_id}-{current_year}-{month_idx}",
                                "property": exp.prop.prop_name,
                                "amount": float(amount),
                                "due_date": due_date.strftime("%Y-%m-%d"),
                                "description": exp.expense_line_types.expense_line_types_name,
                                "line_type": exp.expense_line_types.expense_line_types_name,
                                "color": color,
                            })
                        
                        current_year += 1
        
        expenses.sort(key=lambda x: x["due_date"])
        return JsonResponse({"expenses": expenses})
    
    return render(request, "finance/cashflow_forecast.html")

### NOTIFICATIONS ###
@login_required
def notifications_dashboard(request):
    """
    Notifications Dashboard view - shows property management alerts and status
    """
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        # AJAX request - return JSON data
        try:
            notification_data = get_notification_data()
            return JsonResponse(notification_data)
        except Exception as e:
            return JsonResponse({
                'error': f'Error loading notification data: {str(e)}'
            })
    else:
        # Regular page load - return template
        return render(request, 'notifications.html')

def get_notification_data():
    """
    Get notification data by running similar queries to the management command
    """
    import logging
    from django.db import connection as django_connection
    
    logger = logging.getLogger(__name__)
    mydb = None
    my_cursor = None
    
    try:
        mydb = mysql.connector.connect(
            host=settings.DATABASES['default']['HOST'],
            port=settings.DATABASES['default']['PORT'],
            user=settings.DATABASES['default']['USER'],
            password=settings.DATABASES['default']['PASSWORD'],
            database=settings.DATABASES['default']['NAME'],
        )
        
        my_cursor = mydb.cursor()
        today = date.today()
        
        # Get vacant properties
        vacant_properties = get_vacant_properties(my_cursor)
        
        # Get expiring leases (pending renewals)
        expiring_leases = get_expiring_leases(my_cursor, today)
        
        # Get declined renewals
        declined_renewals = get_declined_renewals(my_cursor, today)
        
        # Get overdue invoices
        overdue_invoices = get_overdue_invoices(my_cursor, today)
        
        # Get expenses waiting for approval
        expenses_waiting_approval = get_expenses_waiting_approval(my_cursor)
        
        # Get expenses waiting for payment
        expenses_waiting_payment = get_expenses_waiting_payment(my_cursor)
        
        return {
            'summary': {
                'vacantProperties': len(vacant_properties),
                'expiringLeases': len(expiring_leases),
                'declinedRenewals': len(declined_renewals),
                'overdueInvoices': len(overdue_invoices),
                'expensesWaitingApproval': len(expenses_waiting_approval),
                'expensesWaitingPayment': len(expenses_waiting_payment)
            },
            'vacantProperties': vacant_properties,
            'expiringLeases': expiring_leases,
            'declinedRenewals': declined_renewals,
            'overdueInvoices': overdue_invoices,
            'expensesWaitingApproval': expenses_waiting_approval,
            'expensesWaitingPayment': expenses_waiting_payment,
            'lastUpdated': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
        
    except mysql.connector.Error as e:
        logger.error(f"MySQL connection error in get_notification_data: {e}")
        # Return empty data structure on error
        return {
            'summary': {
                'vacantProperties': 0,
                'expiringLeases': 0,
                'declinedRenewals': 0,
                'overdueInvoices': 0,
                'expensesWaitingApproval': 0,
                'expensesWaitingPayment': 0
            },
            'vacantProperties': [],
            'expiringLeases': [],
            'declinedRenewals': [],
            'overdueInvoices': [],
            'expensesWaitingApproval': [],
            'expensesWaitingPayment': [],
            'lastUpdated': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
        
    except Exception as e:
        logger.error(f"Unexpected error in get_notification_data: {e}")
        # Return empty data structure on error
        return {
            'summary': {
                'vacantProperties': 0,
                'expiringLeases': 0,
                'declinedRenewals': 0,
                'overdueInvoices': 0,
                'expensesWaitingApproval': 0,
                'expensesWaitingPayment': 0
            },
            'vacantProperties': [],
            'expiringLeases': [],
            'declinedRenewals': [],
            'overdueInvoices': [],
            'expensesWaitingApproval': [],
            'expensesWaitingPayment': [],
            'lastUpdated': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
        
    finally:
        # Always close cursor and connection
        if my_cursor:
            try:
                my_cursor.close()
            except:
                pass
        if mydb and mydb.is_connected():
            try:
                mydb.close()
            except:
                pass
        # Also close Django's connection
        django_connection.close()

def get_expenses_waiting_approval(cursor):
    """Get expenses that require approval (status = 'require_approval')"""
    cursor.execute("""
        SELECT ae.act_expense_id, ae.act_expense_date, ae.act_expense_description,
               ae.act_expense_amount, p.prop_name
        FROM railway.act_expense ae
        JOIN railway.prop p ON ae.prop_id = p.prop_id
        WHERE ae.act_expense_approved = 'No' 
        AND ae.act_expense_paid = 'No'
        ORDER BY ae.act_expense_date ASC
    """)
    
    expenses_data = cursor.fetchall()
    expenses_waiting_approval = []
    
    for row in expenses_data:
        expenses_waiting_approval.append({
            'expense_id': row[0],
            'expense_date': row[1].strftime('%Y-%m-%d') if row[1] else '',
            'description': row[2] or '',
            'amount': float(row[3]) if row[3] else 0.0,
            'property_name': row[4] or '',
            'expense_type': 'General Expense',  # You might want to add this field to your database
            'submitted_date': row[1].strftime('%Y-%m-%d') if row[1] else ''
        })
    
    return expenses_waiting_approval

def get_expenses_waiting_payment(cursor):
    """Get expenses that are approved but not yet paid"""
    cursor.execute("""
        SELECT ae.act_expense_id, ae.act_expense_date, ae.act_expense_description,
               ae.act_expense_amount, p.prop_name
        FROM railway.act_expense ae
        JOIN railway.prop p ON ae.prop_id = p.prop_id
        WHERE ae.act_expense_approved = 'Yes' 
        AND ae.act_expense_paid = 'No'
        ORDER BY ae.act_expense_date ASC
    """)
    
    expenses_data = cursor.fetchall()
    expenses_waiting_payment = []
    
    for row in expenses_data:
        expenses_waiting_payment.append({
            'expense_id': row[0],
            'expense_date': row[1].strftime('%Y-%m-%d') if row[1] else '',
            'description': row[2] or '',
            'amount': float(row[3]) if row[3] else 0.0,
            'property_name': row[4] or '',
            'expense_type': 'General Expense',  # You might want to add this field to your database
            'approved_date': row[1].strftime('%Y-%m-%d') if row[1] else ''  # Using expense date as approximation
        })
    
    return expenses_waiting_payment

def get_vacant_properties(cursor):
    """Get properties that are active and available but have no current tenant"""
    # Get properties with current tenants
    cursor.execute("""
        SELECT prop.prop_name
        FROM railway.tenant
        JOIN railway.prop ON prop.prop_id = tenant.prop_id
        WHERE tenant.tenant_current = 'Yes'
    """)
    prop_active_tenant = [row[0] for row in cursor.fetchall()]
    
    # Get all active properties available for rent
    cursor.execute("""
        SELECT prop.prop_name, prop.prop_country
        FROM railway.prop
        WHERE prop.prop_status = 'Active'
        AND prop.prop_available_for_rent = 'Yes'
        ORDER BY prop.prop_country ASC, prop.prop_name ASC
    """)
    active_properties_data = cursor.fetchall()
    
    # Find vacant properties
    vacant_properties = []
    for prop_data in active_properties_data:
        prop_name = prop_data[0]
        prop_country = prop_data[1]
        
        if prop_name not in prop_active_tenant:
            vacant_properties.append({
                'prop_name': prop_name,
                'prop_country': prop_country
            })
    
    return vacant_properties

def get_expiring_leases(cursor, today):
    """Get leases that are expiring and pending renewal"""
    cursor.execute("""
        SELECT prop.prop_name, prop.prop_country, tenant.tenant_name, 
               tenant.tenant_lease_end_date, tenant.tenant_renewal_period,
               tenant.tenant_renewal_status
        FROM railway.tenant
        JOIN railway.prop ON prop.prop_id = tenant.prop_id
        WHERE tenant.tenant_current = 'Yes'
        ORDER BY prop.prop_country ASC, prop.prop_name ASC
    """)
    tenant_rows = cursor.fetchall()
    
    expiring_leases = []
    
    for row in tenant_rows:
        prop_name = row[0]
        prop_country = row[1]
        tenant_name = row[2]
        lease_end_date = row[3]
        renewal_period = int(row[4])
        renewal_status = row[5] if row[5] else 'pending'
        
        renewal_date = lease_end_date - timedelta(days=renewal_period)
#        warning_date = renewal_date - timedelta(days=30)
        warning_date = renewal_date
        
        if today >= warning_date and renewal_status == 'pending':
            expiring_leases.append({
                'prop_name': prop_name,
                'prop_country': prop_country,
                'tenant_name': tenant_name,
                'lease_end_date': lease_end_date.strftime('%Y-%m-%d'),
                'renewal_date': renewal_date.strftime('%Y-%m-%d')
            })
    
    return expiring_leases

def get_declined_renewals(cursor, today):
    """Get renewals that have been declined"""
    cursor.execute("""
        SELECT prop.prop_name, prop.prop_country, tenant.tenant_name, 
               tenant.tenant_lease_end_date, tenant.tenant_renewal_period,
               tenant.tenant_renewal_status
        FROM railway.tenant
        JOIN railway.prop ON prop.prop_id = tenant.prop_id
        WHERE tenant.tenant_current = 'Yes'
        ORDER BY prop.prop_country ASC, prop.prop_name ASC
    """)
    tenant_rows = cursor.fetchall()
    
    declined_renewals = []
    
    for row in tenant_rows:
        prop_name = row[0]
        prop_country = row[1]
        tenant_name = row[2]
        lease_end_date = row[3]
        renewal_period = int(row[4])
        renewal_status = row[5] if row[5] else 'pending'
        
        renewal_date = lease_end_date - timedelta(days=renewal_period)
        warning_date = renewal_date - timedelta(days=30)
        
        if today >= warning_date and renewal_status == 'declined':
            declined_renewals.append({
                'prop_name': prop_name,
                'prop_country': prop_country,
                'tenant_name': tenant_name,
                'lease_end_date': lease_end_date.strftime('%Y-%m-%d'),
                'message': 'CURRENT TENANT NOT RENEWING LEASE - NEED NEW TENANT'
            })
    
    return declined_renewals

def get_overdue_invoices(cursor, today):
    """Get properties with overdue invoices"""
    cursor.execute("""
        SELECT prop.prop_name, prop.prop_country, tenant.tenant_name, 
               tenant.tenant_payment_terms, tenant.tenant_rent,
               invoice.invoice_date, invoice.invoice_id
        FROM railway.invoice
        JOIN railway.tenant ON invoice.tenant_id = tenant.tenant_id
        JOIN railway.prop ON tenant.prop_id = prop.prop_id
        WHERE invoice.invoice_paid = 'No'
        AND tenant.tenant_current = 'Yes'
        ORDER BY prop.prop_country ASC, prop.prop_name ASC, invoice.invoice_date ASC
    """)
    
    invoice_data = cursor.fetchall()
    
    # Create a flat list of overdue invoices instead of grouping
    overdue_invoices_list = []
    
    for row in invoice_data:
        prop_name = row[0]
        prop_country = row[1]
        tenant_name = row[2]
        payment_terms = int(row[3]) if row[3] else 0
        tenant_rent = row[4]
        invoice_date = row[5]
        invoice_id = row[6]
        
        # Calculate due date based on invoice date and payment terms
        due_date = invoice_date + timedelta(days=payment_terms)
        
        # Only include if invoice is overdue
        if due_date < today:
            # Calculate days overdue
            days_overdue = (today - due_date).days
            
            overdue_invoices_list.append({
                'prop_name': prop_name,
                'prop_country': prop_country,
                'tenant_name': tenant_name,
                'tenant_rent': tenant_rent,
                'invoice_date': invoice_date.strftime('%Y-%m-%d'),
                'due_date': due_date.strftime('%Y-%m-%d'),
                'days_overdue': days_overdue,
                'invoice_id': invoice_id
            })
    
    return overdue_invoices_list

### FINANCIAL DASHBOARD ###
def calculate_occupancy_metrics(property, start_date=None, end_date=None):
    """
    Calculate occupancy rate and average days to fill for a property
    FIXED: Only count days within actual lease period, ignore tenant_current status
    """
    if not end_date:
        end_date = timezone.now().date()
    
    if not start_date:
        first_tenant = tenant.objects.filter(prop=property).order_by('tenant_lease_start_date').first()
        if first_tenant and first_tenant.tenant_lease_start_date:
            start_date = first_tenant.tenant_lease_start_date
        else:
            from datetime import timedelta
            start_date = end_date - timedelta(days=365)
    
    total_days = (end_date - start_date).days
    if total_days <= 0:
        return {
            'occupancy_rate': 0,
            'avg_days_to_fill': 0,
            'current_vacancy_days': 0,
            'is_currently_vacant': False
        }
    
    # Calculate occupied days - DATE BASED ONLY
    occupied_days = 0
    property_tenants = tenant.objects.filter(
        prop=property,
        tenant_lease_start_date__lte=end_date
    )
    
    for t in property_tenants:
        if not t.tenant_lease_start_date:
            continue
            
        lease_start = max(t.tenant_lease_start_date, start_date)
        
        # FIXED: Only count up to lease end date
        if t.tenant_lease_end_date:
            lease_end = min(t.tenant_lease_end_date, end_date)
        else:
            lease_end = end_date
        
        if lease_end >= lease_start:
            occupied_days += (lease_end - lease_start).days + 1
    
    occupied_days = min(occupied_days, total_days)
    occupancy_rate = (occupied_days / total_days * 100) if total_days > 0 else 0
    
    # Calculate average days to fill
    vacancies = VacancyPeriod.objects.filter(
        prop=property,
        status='FILLED',
        reason='BETWEEN_TENANTS',
        start_date__gte=start_date,
        end_date__lte=end_date
    )
    
    if vacancies.exists():
        from django.db.models import Avg
        avg_days = vacancies.aggregate(Avg('days_vacant'))['days_vacant__avg']
        avg_days_to_fill = round(avg_days, 0) if avg_days else 0
    else:
        avg_days_to_fill = 0
    
    # Auto-detect if currently vacant
    is_currently_vacant = True
    for t in property_tenants:
        if t.tenant_lease_start_date and t.tenant_lease_start_date <= end_date:
            if t.tenant_lease_end_date:
                if t.tenant_lease_end_date >= end_date:
                    is_currently_vacant = False
                    break
            else:
                is_currently_vacant = False
                break
    
    # Calculate current vacancy days
    current_vacancy_days = 0
    if is_currently_vacant:
        last_tenant = property_tenants.filter(
            tenant_lease_end_date__isnull=False
        ).order_by('-tenant_lease_end_date').first()
        
        if last_tenant and last_tenant.tenant_lease_end_date:
            current_vacancy_days = (end_date - last_tenant.tenant_lease_end_date).days
    
    return {
        'occupancy_rate': round(occupancy_rate, 1),
        'avg_days_to_fill': int(avg_days_to_fill),
        'current_vacancy_days': current_vacancy_days,
        'is_currently_vacant': is_currently_vacant
    }

def calculate_occupancy_metrics_optimized(property, start_date=None, end_date=None):
    """
    Calculate occupancy metrics using PREFETCHED data
    FIXED: Only count days within actual lease period
    """
    if not end_date:
        end_date = timezone.now().date()
    
    property_tenants = list(property.tenant_set.all())
    
    if not start_date:
        if property_tenants:
            tenants_with_dates = [t for t in property_tenants if t.tenant_lease_start_date]
            if tenants_with_dates:
                first_tenant = min(tenants_with_dates, key=lambda t: t.tenant_lease_start_date)
                start_date = first_tenant.tenant_lease_start_date
            else:
                from datetime import timedelta
                start_date = end_date - timedelta(days=365)
        else:
            from datetime import timedelta
            start_date = end_date - timedelta(days=365)
    
    total_days = (end_date - start_date).days
    if total_days <= 0:
        return {
            'occupancy_rate': 0,
            'avg_days_to_fill': 0,
            'current_vacancy_days': 0,
            'is_currently_vacant': False
        }
    
    # Calculate occupied days - DATE BASED ONLY
    occupied_days = 0
    for t in property_tenants:
        if not t.tenant_lease_start_date or t.tenant_lease_start_date > end_date:
            continue
            
        lease_start = max(t.tenant_lease_start_date, start_date)
        
        # FIXED: Only count up to lease end date
        if t.tenant_lease_end_date:
            lease_end = min(t.tenant_lease_end_date, end_date)
        else:
            lease_end = end_date
        
        if lease_end >= lease_start:
            occupied_days += (lease_end - lease_start).days + 1
    
    occupied_days = min(occupied_days, total_days)
    occupancy_rate = (occupied_days / total_days * 100) if total_days > 0 else 0
    
    # Use prefetched vacancy data
    vacancies = [v for v in property.vacancy_periods.all() 
                 if v.status == 'FILLED' 
                 and v.reason == 'BETWEEN_TENANTS'
                 and v.start_date >= start_date
                 and v.end_date and v.end_date <= end_date]
    
    if vacancies:
        avg_days = sum(v.days_vacant for v in vacancies) / len(vacancies)
        avg_days_to_fill = round(avg_days, 0)
    else:
        avg_days_to_fill = 0
    
    # Auto-detect if currently vacant
    is_currently_vacant = True
    for t in property_tenants:
        if t.tenant_lease_start_date and t.tenant_lease_start_date <= end_date:
            if t.tenant_lease_end_date:
                if t.tenant_lease_end_date >= end_date:
                    is_currently_vacant = False
                    break
            else:
                is_currently_vacant = False
                break
    
    # Calculate current vacancy days
    current_vacancy_days = 0
    if is_currently_vacant:
        tenants_with_end_dates = [t for t in property_tenants if t.tenant_lease_end_date]
        if tenants_with_end_dates:
            last_tenant = max(tenants_with_end_dates, key=lambda t: t.tenant_lease_end_date)
            current_vacancy_days = (end_date - last_tenant.tenant_lease_end_date).days
    
    return {
        'occupancy_rate': round(occupancy_rate, 1),
        'avg_days_to_fill': int(avg_days_to_fill),
        'current_vacancy_days': current_vacancy_days,
        'is_currently_vacant': is_currently_vacant
    }

def calculate_portfolio_occupancy_metrics_optimized(properties, start_date=None, end_date=None):
    """
    Calculate portfolio metrics using SIMPLE AVERAGE method
    Each property judged on its own history
    """
    if not end_date:
        end_date = timezone.now().date()
    
    # Calculate occupancy for each property using its OWN start date
    property_occupancy_rates = []
    
    for prop in properties:
        metrics = calculate_occupancy_metrics_optimized(prop, start_date=None, end_date=end_date)
        property_occupancy_rates.append(metrics['occupancy_rate'])
    
    # Simple average
    if property_occupancy_rates:
        portfolio_occupancy = sum(property_occupancy_rates) / len(property_occupancy_rates)
    else:
        portfolio_occupancy = 0
    
    # For avg_days_to_fill
    if not start_date:
        all_tenants = []
        for prop in properties:
            all_tenants.extend([t for t in prop.tenant_set.all() if t.tenant_lease_start_date])
        
        if all_tenants:
            first_tenant = min(all_tenants, key=lambda t: t.tenant_lease_start_date)
            start_date = first_tenant.tenant_lease_start_date
        else:
            from datetime import timedelta
            start_date = end_date - timedelta(days=365)
    
    # Get all vacancies
    all_vacancies = []
    for prop in properties:
        vacancies = [v for v in prop.vacancy_periods.all()
                    if v.status == 'FILLED'
                    and v.reason == 'BETWEEN_TENANTS'
                    and v.start_date >= start_date
                    and v.end_date and v.end_date <= end_date]
        all_vacancies.extend(vacancies)
    
    # Weighted average
    if all_vacancies and len(list(properties)) > 0:
        total_vacancy_days = sum(v.days_vacant for v in all_vacancies)
        portfolio_avg_days = round(total_vacancy_days / len(list(properties)), 1)
    else:
        portfolio_avg_days = 0
    
    return {
        'occupancy_rate': round(portfolio_occupancy, 1),
        'avg_days_to_fill': portfolio_avg_days
    }

def get_property_first_tenant_date_optimized(prop):
    """
    Get the date when the property became operational using PREFETCHED data.
    Returns None if property has no tenants.
    """
    # Use prefetched tenant_set instead of querying
    property_tenants = [t for t in prop.tenant_set.all() 
                       if t.tenant_lease_start_date is not None]
    
    if not property_tenants:
        return None
    
    # Find first tenant by lease start date
    first_tenant = min(property_tenants, key=lambda t: t.tenant_lease_start_date)
    return first_tenant.tenant_lease_start_date

def calculate_effective_period(first_tenant_date, period_start, period_end):
    """
    Calculate the effective period for a property based on when it became operational.
    Never count days before the first tenant.
    
    Returns: (effective_start, effective_end) or (None, None) if no valid period
    """
    if not first_tenant_date:
        return None, None
    
    # Property wasn't operational before first tenant
    effective_start = max(first_tenant_date, period_start)
    effective_end = period_end
    
    # If effective start is after period end, no valid period
    if effective_start > effective_end:
        return None, None
    
    return effective_start, effective_end

def calculate_occupancy_metrics_with_period(prop, period_start, period_end):
    """
    Calculate occupancy metrics for a specific time period.
    Only counts operational days (after first tenant).
    FIXED: Uses prefetched data and handles OPEN vacancies correctly.
    """
    from datetime import datetime
    
    # Get property's first tenant date using PREFETCHED data
    first_tenant_date = get_property_first_tenant_date_optimized(prop)
    
    if not first_tenant_date:
        return {
            'occupancy_rate': None,
            'avg_days_to_fill': None,
            'is_currently_vacant': False,
            'current_vacancy_days': 0
        }
    
    # Calculate effective period (never before first tenant)
    effective_start, effective_end = calculate_effective_period(
        first_tenant_date, period_start, period_end
    )
    
    if not effective_start:
        return {
            'occupancy_rate': None,
            'avg_days_to_fill': None,
            'is_currently_vacant': False,
            'current_vacancy_days': 0
        }
    
    # Total days in effective period
    total_days = (effective_end - effective_start).days + 1
    
    # Use PREFETCHED vacancy data
    all_vacancies = list(prop.vacancy_periods.all())
    
    # Filter vacancies that overlap with the period
    vacancies = [
        v for v in all_vacancies
        if (v.start_date >= first_tenant_date and  # Only operational vacancies
            v.start_date <= effective_end and
            (v.status == 'FILLED' or v.status == 'OPEN'))
    ]
    
    # Calculate vacant days in period
    vacant_days_in_period = 0
    total_days_to_fill = 0
    filled_vacancy_count = 0
    is_currently_vacant = False
    current_vacancy_days = 0
    
    today = datetime.now().date()
    
    for vacancy in vacancies:
        # Calculate overlap with effective period
        vacancy_start = max(vacancy.start_date, effective_start)
        
        # FIXED: Handle OPEN vacancies correctly
        if vacancy.status == 'OPEN':
            # For OPEN vacancies, use today or effective_end, whichever is earlier
            vacancy_end = min(today, effective_end)
        else:
            # For FILLED vacancies, use the actual end date
            vacancy_end = min(vacancy.end_date, effective_end)
        
        # Only count if vacancy actually overlaps with period
        if vacancy_start <= vacancy_end:
            days_in_period = (vacancy_end - vacancy_start).days + 1
            vacant_days_in_period += days_in_period
            
            # For avg days to fill calculation
            if vacancy.status == 'FILLED':
                # Use the full vacancy duration
                total_days_to_fill += vacancy.days_vacant
                filled_vacancy_count += 1
            elif vacancy.status == 'OPEN':
                # For current vacancy, count days so far
                days_so_far = (today - vacancy.start_date).days + 1
                total_days_to_fill += days_so_far
                filled_vacancy_count += 1
                is_currently_vacant = True
                current_vacancy_days = days_so_far
    
    # SAFETY CHECK: Vacant days cannot exceed total days
    vacant_days_in_period = min(vacant_days_in_period, total_days)
    
    # Calculate occupancy rate
    occupied_days = total_days - vacant_days_in_period
    occupancy_rate = (occupied_days / total_days * 100) if total_days > 0 else 0
    
    # Calculate average days to fill
    avg_days_to_fill = (total_days_to_fill / filled_vacancy_count) if filled_vacancy_count > 0 else 0
    
    return {
        'occupancy_rate': round(occupancy_rate, 1),
        'avg_days_to_fill': round(avg_days_to_fill, 1),
        'is_currently_vacant': is_currently_vacant,
        'current_vacancy_days': current_vacancy_days
    }

def calculate_portfolio_occupancy_with_period(properties, period_start, period_end):
    """
    Calculate portfolio-wide occupancy metrics for a specific time period.
    """
    total_occupied_days = 0
    total_possible_days = 0
    total_days_to_fill = 0
    property_count = 0  # Count ALL properties for average
    
    for prop in properties:
        metrics = calculate_occupancy_metrics_with_period(prop, period_start, period_end)
        
        if metrics['occupancy_rate'] is not None:
            # Get effective period for this property
            first_tenant_date = get_property_first_tenant_date_optimized(prop)
            if first_tenant_date:
                effective_start, effective_end = calculate_effective_period(
                    first_tenant_date, period_start, period_end
                )
                
                if effective_start:
                    property_total_days = (effective_end - effective_start).days + 1
                    property_occupied_days = int(property_total_days * metrics['occupancy_rate'] / 100)
                    
                    total_possible_days += property_total_days
                    total_occupied_days += property_occupied_days
                    
                    # Add ALL properties' days to fill (including 0)
                    total_days_to_fill += metrics['avg_days_to_fill']
                    property_count += 1
    
    portfolio_occupancy = (total_occupied_days / total_possible_days * 100) if total_possible_days > 0 else 0
    # Calculate average across ALL properties (not just those with vacancies)
    portfolio_avg_days = (total_days_to_fill / property_count) if property_count > 0 else 0
    
    return {
        'occupancy_rate': round(portfolio_occupancy, 1),
        'avg_days_to_fill': round(portfolio_avg_days, 1)
    }

@login_required
def occupancy_trends_view(request):
    """
    Display occupancy, days to fill, and vacancy cost trends over time
    """
    from datetime import date
    from django.db.models import Min
    
    # Get the first tenant date across all properties
    first_tenant_date = tenant.objects.filter(
        tenant_lease_start_date__isnull=False
    ).aggregate(
        first_date=Min('tenant_lease_start_date')
    )['first_date']
    
    if not first_tenant_date:
        # No tenants yet
        return render(request, 'occupancy_trends.html', {
            'error': 'No tenant data available yet'
        })
    
    # Get current year
    today = date.today()
    current_year = today.year
    first_year = first_tenant_date.year
    
    # Calculate all metrics for each year
    yearly_data = []
    
    for year in range(first_year, current_year + 1):
        year_metrics = calculate_year_metrics(year)
        yearly_data.append(year_metrics)
    
    context = {
        'yearly_data': yearly_data,
        'yearly_data_json': json.dumps(yearly_data),  # Add this line
        'current_year': current_year,
        'first_year': first_year,
    }
    
    return render(request, 'occupancy_trends.html', context)


def calculate_year_metrics(year):
    """
    OPTIMIZED: Calculate ALL metrics for a specific calendar year.
    Uses prefetch and processes in Python to minimize database queries.
    """
    from datetime import date
    from django.db.models import Prefetch
    
    # Define year boundaries
    year_start = date(year, 1, 1)
    year_end = date(year, 12, 31)
    today = date.today()
    
    # If current year, use today as end date
    if year == today.year:
        year_end = today
        is_current_year = True
    else:
        is_current_year = False
    
    # ==================== FETCH ALL DATA AT ONCE ====================
    
    # Get all active, non-seasonal properties WITH their tenants prefetched
    active_properties = props.objects.filter(
        prop_status='Active'
    ).prefetch_related(
        Prefetch(
            'tenant_set',
            queryset=tenant.objects.filter(
                tenant_lease_start_date__isnull=False
            ).order_by('tenant_lease_start_date'),
            to_attr='all_tenants'
        )
    )
    
    # Get all vacancies for this year WITH related objects prefetched
    year_vacancies = VacancyPeriod.objects.filter(
        status='FILLED',
        end_date__year=year
    ).select_related(
        'prop',
        'previous_lease',
        'next_lease'
    )
    
    # ==================== PROCESS IN PYTHON ====================
    
    total_available_days = 0
    total_occupied_days = 0
    property_count = 0
    property_details = []
    
    for prop in active_properties:
        # Skip seasonal properties
        if hasattr(prop, 'exclude_from_metrics') and prop.exclude_from_metrics:
            continue
        
        # Get tenants from prefetched data
        tenants_list = prop.all_tenants
        
        if not tenants_list:
            continue
        
        # Find first tenant (already sorted by start date)
        first_tenant = tenants_list[0]
        
        # Determine operational start for this property in this year
        operational_start = max(
            first_tenant.tenant_lease_start_date,
            year_start
        )
        
        # Skip if property wasn't operational yet in this year
        if operational_start > year_end:
            continue
        
        # This property has data for this year
        property_count += 1
        
        # Calculate available days (operational days only)
        available_days = (year_end - operational_start).days + 1
        total_available_days += available_days
        
        # Calculate occupied days for this property in this year
        occupied_days = 0
        
        for t in tenants_list:
            # Calculate overlap between tenant lease and this year period
            lease_start = max(
                t.tenant_lease_start_date,
                operational_start
            )
            
            lease_end = min(
                t.tenant_lease_end_date if t.tenant_lease_end_date else year_end,
                year_end
            )
            
            # Add occupied days if there's an overlap
            if lease_start <= lease_end:
                days = (lease_end - lease_start).days + 1
                occupied_days += days
        
        total_occupied_days += occupied_days
        vacant_days = available_days - occupied_days
        
        # Calculate property occupancy
        prop_occupancy = (occupied_days / available_days * 100) if available_days > 0 else 0
        
        # Store property details for modal
        property_details.append({
            'name': prop.prop_name,
            'available_days': available_days,
            'occupied_days': occupied_days,
            'vacant_days': vacant_days,
            'occupancy_rate': round(prop_occupancy, 1)
        })
    
    # Calculate overall occupancy rate
    if total_available_days > 0:
        occupancy_rate = (total_occupied_days / total_available_days) * 100
    else:
        occupancy_rate = 0
    
    # ==================== PROCESS VACANCIES (Already Prefetched) ====================
    
    vacancy_days_list = []
    vacancy_details = []
    total_vacancy_cost = 0
    
    for v in year_vacancies:
        # Skip seasonal properties (already checked in queryset, but double-check)
        if hasattr(v.prop, 'exclude_from_metrics') and v.prop.exclude_from_metrics:
            continue
        
        vacancy_days_list.append(v.days_vacant)
        
        # Calculate vacancy cost
        monthly_rent = 0
        if v.next_lease:
            monthly_rent = v.next_lease.tenant_rent
        elif v.previous_lease:
            monthly_rent = v.previous_lease.tenant_rent
        
        vacancy_cost = (v.days_vacant / 30) * monthly_rent
        total_vacancy_cost += vacancy_cost
        
        # Store vacancy details for modal
        vacancy_details.append({
            'property': v.prop.prop_name,
            'start_date': v.start_date.strftime('%Y-%m-%d'),
            'end_date': v.end_date.strftime('%Y-%m-%d'),
            'days': v.days_vacant,
            'previous_tenant': v.previous_lease.tenant_name if v.previous_lease else 'None',
            'next_tenant': v.next_lease.tenant_name if v.next_lease else 'None',
            'cost': round(vacancy_cost, 2)
        })
    
    # Calculate average days to fill - SIMPLE AVERAGE across all properties
    # This matches the methodology in vacancy_management_view
    # Properties with no vacancies contribute 0 days to the average
    if property_count > 0:
        avg_days_to_fill = sum(vacancy_days_list) / property_count
    else:
        avg_days_to_fill = 0
    
    # ==================== RETURN ALL METRICS ====================
    
    return {
        'year': year,
        'is_current_year': is_current_year,
        
        # Summary metrics (for chart display)
        'occupancy_rate': round(occupancy_rate, 1),
        'avg_days_to_fill': round(avg_days_to_fill, 1),
        'vacancy_cost': round(total_vacancy_cost, 0),
        
        # Occupancy details (for modal)
        'total_available_days': total_available_days,
        'total_occupied_days': total_occupied_days,
        'total_vacant_days': total_available_days - total_occupied_days,
        'property_count': property_count,
        'property_details': sorted(property_details, key=lambda x: x['occupancy_rate'], reverse=True),
        
        # Vacancy details (for modal)
        'vacancy_count': len(vacancy_days_list),
        'vacancy_details': sorted(vacancy_details, key=lambda x: x['start_date']),
    }

@login_required
def financial_indicators_view(request):
    """
    Display the Financial Indicators Dashboard - ONLY for Active Properties
    Using Portfolio-Wide Calculations with Occupancy Metrics
    OPTIMIZED: Reduced queries with prefetch_related
    """
    if request.method == 'GET' and request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        # AJAX request for property data
        try:
            from datetime import datetime, timedelta
            
            # Financial indicators show annual data (no time period selection)
            today = datetime.now().date()
            current_year = datetime.now().year
            
            # OPTIMIZATION: Prefetch related data in one query
            properties = props.objects.filter(prop_status='Active').prefetch_related(
                Prefetch('tenant_set', queryset=tenant.objects.all()),
                Prefetch('vacancy_periods', queryset=VacancyPeriod.objects.select_related('previous_lease', 'next_lease').all()),
                'prop_values_set'
            )
            
            properties_data = []
            
            # Portfolio-wide totals for all active properties
            portfolio_totals = {
                'total_revenue': Decimal('0.00'),
                'total_budgeted_expenses': Decimal('0.00'),
                'total_purchase_price': Decimal('0.00'),
                'total_current_value': Decimal('0.00'),
                'total_floor_area': 0,
                'property_count': 0
            }
            
            for prop in properties:
                # Get revenue totals using your existing revenue model structure
                revenue_total = calculate_property_revenue(prop)
                
                # Get ONLY budgeted expense totals using your existing expense model
                budgeted_expense_total = calculate_property_budgeted_expenses(prop)
                
                # Get property values - ONLY for active properties
                # Use prefetched data
                property_values_list = list(prop.prop_values_set.all())
                property_values = property_values_list[0] if property_values_list else None
                purchase_price = property_values.prop_values_purchase_price if property_values else 0
                current_value = property_values.prop_values_current_value if property_values else 0
                
                # Add to portfolio totals
                portfolio_totals['total_revenue'] += revenue_total
                portfolio_totals['total_budgeted_expenses'] += budgeted_expense_total
                portfolio_totals['total_purchase_price'] += purchase_price or 0
                portfolio_totals['total_current_value'] += current_value or 0
                portfolio_totals['total_floor_area'] += prop.prop_floor_area or 0
                portfolio_totals['property_count'] += 1
                
                # Calculate individual property indicators for display purposes
                gross_roi = (revenue_total / purchase_price * 100) if purchase_price > 0 else 0
                net_roi = ((revenue_total - budgeted_expense_total) / purchase_price * 100) if purchase_price > 0 else 0
                expense_ratio = (budgeted_expense_total / revenue_total * 100) if revenue_total > 0 else 0
                rent_per_sqm = (revenue_total / 12 / prop.prop_floor_area) if prop.prop_floor_area and prop.prop_floor_area > 0 else 0
                value_increase = ((current_value - purchase_price) / purchase_price * 100) if purchase_price > 0 and current_value > 0 else 0
                
                # Store individual property data
                properties_data.append({
                    'id': prop.prop_id,
                    'name': prop.prop_name or f"Property {prop.prop_id}",
                    'status': prop.prop_status,
                    'grossROI': round(float(gross_roi), 2),
                    'netROI': round(float(net_roi), 2),
                    'expensesToRevenue': round(float(expense_ratio), 2),
                    'rentPerSqm': round(float(rent_per_sqm), 2),
                    'valueIncrease': round(float(value_increase), 2),
                    'revenue': float(revenue_total),
                    'expenses': float(budgeted_expense_total),
                    'profit': float(revenue_total - budgeted_expense_total)
                })
            
            # Calculate TRUE PORTFOLIO-WIDE indicators (FINANCIAL ONLY)
            portfolio_indicators = {
                'grossROI': round(float(
                    (portfolio_totals['total_revenue'] / portfolio_totals['total_purchase_price'] * 100) 
                    if portfolio_totals['total_purchase_price'] > 0 else 0
                ), 2),
                'netROI': round(float(
                    ((portfolio_totals['total_revenue'] - portfolio_totals['total_budgeted_expenses']) / 
                     portfolio_totals['total_purchase_price'] * 100) 
                    if portfolio_totals['total_purchase_price'] > 0 else 0
                ), 2),
                'expensesToRevenue': round(float(
                    (portfolio_totals['total_budgeted_expenses'] / portfolio_totals['total_revenue'] * 100) 
                    if portfolio_totals['total_revenue'] > 0 else 0
                ), 2),
                'rentPerSqm': round(float(
                    (portfolio_totals['total_revenue'] / 12 / portfolio_totals['total_floor_area']) 
                    if portfolio_totals['total_floor_area'] > 0 else 0
                ), 2),
                'valueIncrease': round(float(
                    ((portfolio_totals['total_current_value'] - portfolio_totals['total_purchase_price']) / 
                     portfolio_totals['total_purchase_price'] * 100) 
                    if portfolio_totals['total_purchase_price'] > 0 and portfolio_totals['total_current_value'] > 0 else 0
                ), 2)
            }
            
            return JsonResponse({
                'properties': properties_data,
                'portfolio_indicators': portfolio_indicators,
                'portfolio_totals': {
                    'total_revenue': float(portfolio_totals['total_revenue']),
                    'total_expenses': float(portfolio_totals['total_budgeted_expenses']),
                    'total_purchase_price': float(portfolio_totals['total_purchase_price']),
                    'total_current_value': float(portfolio_totals['total_current_value']),
                    'total_floor_area': portfolio_totals['total_floor_area'],
                    'property_count': portfolio_totals['property_count']
                },
                'total_active_properties': len(properties_data),
                'message': f'Showing {len(properties_data)} active properties - Financial Indicators (annual data)'
            })
            
        except Exception as e:
            import traceback
            print(traceback.format_exc())
            return JsonResponse({'error': str(e)}, status=500)
    
    # Regular page load
    context = {
        'page_title': 'Financial Indicators Dashboard - Portfolio-Wide Analysis (Active Properties)'
    }
    return render(request, 'finance/financial_indicators.html', context)

@login_required
def vacancy_management_view(request):
    """
    Display the Vacancy Management Dashboard - Occupancy Metrics Only
    Shows Occupancy Rate, Avg Days to Fill, and Vacancy Cost
    ONLY for properties included in occupancy tracking
    """
    if request.method == 'GET' and request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        # AJAX request for property data
        try:
            from datetime import datetime, timedelta
            
            # GET TIME PERIOD PARAMETER
            time_period = request.GET.get('period', 'past_year')  # Default: past_year
            today = datetime.now().date()
            current_year = datetime.now().year
            
            # Calculate period boundaries
            if time_period == 'past_year':
                period_start = today - timedelta(days=365)
                period_end = today
            elif time_period == 'last_3_years':
                period_start = today - timedelta(days=365 * 3)
                period_end = today
            else:  # 'all_time'
                # For all_time, we'll use property-specific start dates
                period_start = datetime(2000, 1, 1).date()
                period_end = today
            
            # OPTIMIZATION: Prefetch related data in one query
            properties = props.objects.filter(
                prop_status='Active',
                prop_include_in_occupancy=True  # ONLY properties included in occupancy tracking
            ).prefetch_related(
                Prefetch('tenant_set', queryset=tenant.objects.all()),
                Prefetch('vacancy_periods', queryset=VacancyPeriod.objects.select_related('previous_lease', 'next_lease').all())
            )
            
            properties_data = []
            
            # Calculate vacancy costs
            vacancy_costs = {}
            total_vacancy_cost = Decimal('0.00')
            
            for prop in properties:
                # Calculate occupancy metrics
                occupancy_metrics = calculate_occupancy_metrics_with_period(
                    prop, period_start, period_end
                )
                
                # CALCULATE VACANCY COSTS for this property
                property_vacancy_cost = Decimal('0.00')
                total_days_vacant = 0
                
                # Get property's first tenant date
                first_tenant_date = get_property_first_tenant_date_optimized(prop)
                
                if first_tenant_date:
                    # Calculate effective period for this property
                    effective_start, effective_end = calculate_effective_period(
                        first_tenant_date, period_start, period_end
                    )
                    
                    if effective_start:
                        # Use PREFETCHED vacancy data
                        all_vacancies = list(prop.vacancy_periods.all())
                        vacancies = [
                            v for v in all_vacancies
                            if (v.start_date >= first_tenant_date and
                                v.start_date <= effective_end and
                                (v.status == 'FILLED' or v.status == 'OPEN'))
                        ]

                        for vacancy in vacancies:
                            # Calculate overlap with effective period
                            vacancy_start = max(vacancy.start_date, effective_start)
                            vacancy_end = min(
                                vacancy.end_date if vacancy.end_date else datetime.now().date(),
                                effective_end
                            )
                            
                            if vacancy_start <= vacancy_end:
                                days_in_period = (vacancy_end - vacancy_start).days + 1
                            else:
                                continue
                            
                            # Determine which rent to use
                            rent_to_use = None
                            
                            if vacancy.next_lease and vacancy.next_lease.tenant_rent:
                                rent_to_use = vacancy.next_lease.tenant_rent
                            elif vacancy.previous_lease and vacancy.previous_lease.tenant_rent:
                                rent_to_use = vacancy.previous_lease.tenant_rent
                            
                            if rent_to_use:
                                daily_rent = Decimal(str(rent_to_use)) / Decimal('30')
                                vacancy_cost = Decimal(str(days_in_period)) * daily_rent
                                property_vacancy_cost += vacancy_cost
                                total_days_vacant += days_in_period
                
                total_vacancy_cost += property_vacancy_cost
                
                vacancy_costs[prop.prop_id] = {
                    'total_cost': float(property_vacancy_cost),
                    'days_vacant': total_days_vacant
                }
                
                # Store individual property data
                properties_data.append({
                    'id': prop.prop_id,
                    'name': prop.prop_name or f"Property {prop.prop_id}",
                    'status': prop.prop_status,
                    'occupancyRate': occupancy_metrics['occupancy_rate'],
                    'avgDaysToFill': occupancy_metrics['avg_days_to_fill'],
                    'isCurrentlyVacant': occupancy_metrics['is_currently_vacant'],
                    'currentVacancyDays': occupancy_metrics['current_vacancy_days'],
                    'vacancyCost': vacancy_costs[prop.prop_id]['total_cost']
                })
            
            # Calculate portfolio occupancy
            portfolio_occupancy = calculate_portfolio_occupancy_with_period(
                properties, period_start, period_end
            )
            
            # Calculate vacancy cost average (for modal comparison)
            num_properties = len(properties)
            vacancy_cost_average = (
                float(total_vacancy_cost) / num_properties 
                if num_properties > 0 else 0
            )
            
            # Portfolio indicators - OCCUPANCY ONLY
            portfolio_indicators = {
                'occupancyRate': portfolio_occupancy['occupancy_rate'],
                'avgDaysToFill': portfolio_occupancy['avg_days_to_fill'],
                'vacancyCost': round(float(total_vacancy_cost), 2),
                'vacancyCostAverage': round(vacancy_cost_average, 2),
            }
            
            return JsonResponse({
                'properties': properties_data,
                'portfolio_indicators': portfolio_indicators,
                'vacancy_costs': vacancy_costs,
                'total_vacancy_cost': float(total_vacancy_cost),
                'time_period': time_period,
                'period_start': period_start.strftime('%Y-%m-%d'),
                'period_end': period_end.strftime('%Y-%m-%d'),
                'total_active_properties': len(properties_data),
                'message': f'Showing {len(properties_data)} properties included in occupancy tracking'
            })
            
        except Exception as e:
            import traceback
            print(traceback.format_exc())
            return JsonResponse({'error': str(e)}, status=500)
    
    # Regular page load
    context = {
        'page_title': 'Vacancy Management Dashboard - Occupancy Performance'
    }
    return render(request, 'finance/vacancy_management.html', context)

def calculate_property_revenue(property_obj):
    """
    Calculate total annual revenue for a property using your revenue model
    ONLY processes Active properties
    """
    # Additional safety check - only calculate for active properties
    if property_obj.prop_status != 'Active':
        return Decimal('0.00')
        
    # Get all revenue records for this active property
    revenue_records = revenue.objects.filter(prop=property_obj)
    
    total_revenue = Decimal('0.00')
    
    for record in revenue_records:
        # Sum all monthly revenue amounts
        monthly_total = (
            (record.revenue_jan or Decimal('0.00')) +
            (record.revenue_feb or Decimal('0.00')) +
            (record.revenue_mar or Decimal('0.00')) +
            (record.revenue_apr or Decimal('0.00')) +
            (record.revenue_may or Decimal('0.00')) +
            (record.revenue_jun or Decimal('0.00')) +
            (record.revenue_jul or Decimal('0.00')) +
            (record.revenue_aug or Decimal('0.00')) +
            (record.revenue_sep or Decimal('0.00')) +
            (record.revenue_oct or Decimal('0.00')) +
            (record.revenue_nov or Decimal('0.00')) +
            (record.revenue_dec or Decimal('0.00'))
        )
        total_revenue += monthly_total
    
    return total_revenue

def calculate_property_budgeted_expenses(property_obj):
    """
    Calculate total annual budgeted expenses for a property using your expense model
    ONLY processes Active properties
    """
    # Additional safety check - only calculate for active properties
    if property_obj.prop_status != 'Active':
        return Decimal('0.00')
        
    # Get all budgeted expense records for this active property
    expense_records = expense.objects.filter(prop=property_obj)
    
    total_expenses = Decimal('0.00')
    
    for record in expense_records:
        # Sum all monthly expense amounts
        monthly_total = (
            (record.expense_jan or Decimal('0.00')) +
            (record.expense_feb or Decimal('0.00')) +
            (record.expense_mar or Decimal('0.00')) +
            (record.expense_apr or Decimal('0.00')) +
            (record.expense_may or Decimal('0.00')) +
            (record.expense_jun or Decimal('0.00')) +
            (record.expense_jul or Decimal('0.00')) +
            (record.expense_aug or Decimal('0.00')) +
            (record.expense_sep or Decimal('0.00')) +
            (record.expense_oct or Decimal('0.00')) +
            (record.expense_nov or Decimal('0.00')) +
            (record.expense_dec or Decimal('0.00'))
        )
        total_expenses += monthly_total
    
    return total_expenses

def calculate_property_actual_expenses(property_obj):
    """
    Calculate total actual expenses for a property using your act_expense model
    ONLY processes Active properties
    """
    # Additional safety check - only calculate for active properties
    if property_obj.prop_status != 'Active':
        return Decimal('0.00')
        
    # Get all actual expense records for this active property
    actual_expenses = act_expense.objects.filter(prop=property_obj)
    
    # Sum all actual expense amounts
    total_actual = actual_expenses.aggregate(
        total=Sum('act_expense_amount')
    )['total'] or Decimal('0.00')
    
    return total_actual

# Additional helper function for year-specific calculations if needed
def calculate_property_revenue_for_year(property_obj, year):
    """
    Calculate revenue for a specific year (if you need year filtering later)
    This would require adding year fields to your revenue model or 
    filtering by revenue_types that have year information
    """
    # This is a placeholder - you'd need to modify based on how you handle years
    # in your revenue_types model or add year fields to your models
    return calculate_property_revenue(property_obj)

def calculate_property_expenses_for_year(property_obj, year):
    """
    Calculate expenses for a specific year (if you need year filtering later)
    """
    # For budgeted expenses
    budgeted = calculate_property_budgeted_expenses(property_obj)
    
    # For actual expenses, you can filter by year using the date field
    from django.db.models import Q
    actual_expenses = act_expense.objects.filter(
        prop=property_obj,
        act_expense_date__year=year
    )
    actual_total = actual_expenses.aggregate(
        total=Sum('act_expense_amount')
    )['total'] or Decimal('0.00')
    
    return budgeted + actual_total

### PROJECTS ###
@login_required
def projects_list(request):
    """Display list of projects with filtering and handle modal-based deletion - FULLY OPTIMIZED"""
    
    # Handle POST request for modal-based deletion
    if request.method == 'POST' and 'delete_project_id' in request.POST:
        if not request.user.is_superuser:
            messages.error(request, "You don't have permission to delete projects.")
            return redirect('projects')
        
        project_id = request.POST.get('delete_project_id')
        # OPTIMIZED: Use select_related and prefetch_related for deletion
        project = get_object_or_404(
            Project.objects.select_related('prop').prefetch_related('projecttask_set', 'project_documents'), 
            project_id=project_id
        )
        
        try:
            with transaction.atomic():
                logger.info(f"User {request.user.username} deleting project via modal: {project.project_name}")
                
                # OPTIMIZED: Use prefetched data for counts (no additional queries)
                all_tasks = list(project.projecttask_set.all())
                main_task_count = sum(1 for task in all_tasks if task.parent_task is None)
                subtask_count = sum(1 for task in all_tasks if task.parent_task is not None)
                document_count = len(list(project.project_documents.all())) if hasattr(project, 'project_documents') else 0
                project_name = project.project_name
                
                # Delete the project
                project.delete()
                
                # Success message
                if subtask_count > 0:
                    messages.success(
                        request, 
                        f"Project '{project_name}' has been permanently deleted along with {main_task_count} main tasks, {subtask_count} subtasks, and {document_count} documents."
                    )
                else:
                    messages.success(
                        request, 
                        f"Project '{project_name}' has been permanently deleted along with {main_task_count} tasks and {document_count} documents."
                    )
                
        except Exception as e:
            logger.error(f"Error deleting project {project_id}: {str(e)}")
            messages.error(request, f"An error occurred while deleting the project '{project.project_name}'.")
        
        return redirect('projects')
    
    # Initialize filter variables from GET parameters FIRST
    search_query = request.GET.get('search', '').strip()
    selected_property = request.GET.get('property', '')
    selected_status = request.GET.get('status', '')
    
    # FULLY OPTIMIZED: Build the base queryset WITHOUT prefetching unnecessary data
    # Only prefetch what's absolutely needed for the list view
    projects_queryset = Project.objects.select_related('prop')
    
    # Apply filters BEFORE any prefetching to reduce the dataset
    if search_query:
        projects_queryset = projects_queryset.filter(
            Q(project_name__icontains=search_query) |
            Q(project_description__icontains=search_query)
        )
    
    if selected_property:
        try:
            property_id = int(selected_property)
            projects_queryset = projects_queryset.filter(prop_id=property_id)
        except (ValueError, TypeError):
            selected_property = ""
    
    if selected_status:
        valid_statuses = [choice[0] for choice in Project.PROJECT_STATUS_CHOICES]
        if selected_status in valid_statuses:
            projects_queryset = projects_queryset.filter(project_status=selected_status)
        else:
            selected_status = ""
    
    # Apply ordering AFTER filtering
    projects_queryset = projects_queryset.order_by(F('project_start_date').desc(nulls_last=True))
    
    # REQUIRED: Template calls calculated methods that need task data
    # The template uses: get_calculated_status, get_progress_percentage, 
    # get_calculated_start_date, get_calculated_expected_completion
    # Use comprehensive prefetching to load ALL related task data in minimal queries
    projects_queryset = projects_queryset.prefetch_related(
        Prefetch('projecttask_set', 
            queryset=ProjectTask.objects.select_related().prefetch_related('subtasks')
        )
    )
    
    # If template only shows basic project info (name, description, dates, status), 
    # then DON'T prefetch tasks at all
    
    # OPTIMIZED: Single query for all properties (only if dropdown is used)
    properties = props.objects.only('prop_id', 'prop_name').order_by('prop_name')
    
    # Pagination with filter preservation
    paginator = Paginator(projects_queryset, 25)
    page_number = request.GET.get('page')
    projects_page = paginator.get_page(page_number)
    
    context = {
        'projects': projects_page,
        'properties': properties,
        'search_query': search_query,
        'selected_property': selected_property,
        'selected_status': selected_status,
        'status_choices': Project.PROJECT_STATUS_CHOICES,
    }
    
    return render(request, 'projects/projects.html', context)

@login_required
def projects_add(request):
    """Add new project"""
    if not request.user.is_superuser:
        messages.error(request, "You don't have permission to add projects.")
        return redirect('projects')
    
    if request.method == 'POST':
        project_name = request.POST.get('project_name')
        prop_id = request.POST.get('prop_id')
        project_start_date = request.POST.get('project_start_date')
        project_expected_completion_date = request.POST.get('project_expected_completion_date')
        project_status = request.POST.get('project_status', 'Pending')
        project_actual_completion_date = request.POST.get('project_actual_completion_date')
        project_description = request.POST.get('project_description')
        
        try:
            # Get the property
            property_obj = get_object_or_404(props, prop_id=prop_id)
            
            # Create the project
            project = Project(  # Updated model name
                project_name=project_name,
                prop=property_obj,
                project_start_date=project_start_date if project_start_date else None,
                project_expected_completion_date=project_expected_completion_date if project_expected_completion_date else None,
                project_status=project_status,
                project_actual_completion_date=project_actual_completion_date if project_actual_completion_date else None,
                project_description=project_description
            )
            project.save()
            
            messages.success(request, f"Project '{project_name}' has been created successfully.")
            return redirect('projects')
            
        except Exception as e:
            messages.error(request, f"Error creating project: {str(e)}")
    
    # Get all properties for dropdown
    properties = props.objects.all().order_by('prop_name')
    
    context = {
        'properties': properties,
        'status_choices': Project.PROJECT_STATUS_CHOICES,  # Updated model name
    }
    
    return render(request, 'projects/projects_add.html', context)

@login_required
def projects_edit(request, project_id):
    """Edit existing project - enhanced to handle Gantt chart returns"""
    if not request.user.is_superuser:
        messages.error(request, "You don't have permission to edit projects.")
        return redirect('projects')
    
    project = get_object_or_404(Project, project_id=project_id)
    
    # Check if coming from Gantt chart
    from_gantt = request.GET.get('from_gantt', 'false') == 'true'
    
    if request.method == 'POST':
        project.project_name = request.POST.get('project_name')
        prop_id = request.POST.get('prop_id')
        project.project_start_date = request.POST.get('project_start_date') if request.POST.get('project_start_date') else None
        project.project_expected_completion_date = request.POST.get('project_expected_completion_date') if request.POST.get('project_expected_completion_date') else None
        project.project_status = request.POST.get('project_status', 'Pending')
        project.project_actual_completion_date = request.POST.get('project_actual_completion_date') if request.POST.get('project_actual_completion_date') else None
        project.project_description = request.POST.get('project_description')
        
        # Add Greek translation fields
        project.project_name_greek = request.POST.get('project_name_greek', '').strip()
        project.project_description_greek = request.POST.get('project_description_greek', '').strip()
        
        try:
            # Update the property
            property_obj = get_object_or_404(props, prop_id=prop_id)
            project.prop = property_obj
            
            project.save()
            
            messages.success(request, f"Project '{project.project_name}' has been updated successfully.")
            
            # Redirect based on where user came from
            if from_gantt:
                return redirect('project_gantt', project_id=project_id)
            else:
                return redirect('projects')
                
        except Exception as e:
            messages.error(request, f"Error updating project: {str(e)}")
    
    # Get all properties for dropdown
    properties = props.objects.all().order_by('prop_name')
    
    context = {
        'project': project,
        'properties': properties,
        'status_choices': Project.PROJECT_STATUS_CHOICES,
        'from_gantt': from_gantt,  # Pass this to template for form action
    }
    
    return render(request, 'projects/projects_edit.html', context)

@login_required
def projects_delete(request, project_id):
    """Delete project with enhanced cascade deletion and warnings - OPTIMIZED"""
    if not request.user.is_superuser:
        messages.error(request, "You don't have permission to delete projects.")
        return redirect('projects')
    
    # OPTIMIZED: Single query with prefetching for counts
    project = get_object_or_404(
        Project.objects.select_related('prop').prefetch_related(
            'projecttask_set', 'project_documents'
        ), 
        project_id=project_id
    )
    
    if request.method == 'POST':
        try:
            with transaction.atomic():
                logger.info(f"User {request.user.username} attempting to delete project: {project.project_name} (ID: {project_id})")
                
                # OPTIMIZED: Use prefetched data for counts (no additional queries)
                all_tasks = list(project.projecttask_set.all())
                main_task_count = sum(1 for task in all_tasks if task.parent_task is None)
                subtask_count = sum(1 for task in all_tasks if task.parent_task is not None)
                total_task_count = len(all_tasks)
                document_count = len(list(project.project_documents.all())) if hasattr(project, 'project_documents') else 0
                
                project_name = project.project_name
                
                # Delete the project (this will cascade to delete all related tasks and subtasks)
                project.delete()
                
                logger.info(f"Successfully deleted project: {project_name} (ID: {project_id}) with {main_task_count} main tasks, {subtask_count} subtasks, and {document_count} documents")
                
                # Success message with detailed information
                if subtask_count > 0:
                    messages.success(
                        request, 
                        f"Project '{project_name}' has been permanently deleted along with {main_task_count} main tasks, {subtask_count} subtasks, and {document_count} documents."
                    )
                else:
                    messages.success(
                        request, 
                        f"Project '{project_name}' has been permanently deleted along with {total_task_count} tasks and {document_count} documents."
                    )
                
        except Exception as e:
            logger.error(f"Error deleting project {project_id}: {str(e)}")
            messages.error(
                request, 
                f"An error occurred while deleting the project '{project.project_name}'. Please try again or contact support."
            )
            return render(request, 'projects/projects_delete.html', {'project': project})
        
        return redirect('projects')
    
    # OPTIMIZED: Use prefetched data for confirmation page (no additional queries)
    all_tasks = list(project.projecttask_set.all())
    main_tasks = [task for task in all_tasks if task.parent_task is None]
    subtasks = [task for task in all_tasks if task.parent_task is not None]
    documents = list(project.project_documents.all()) if hasattr(project, 'project_documents') else []
    
    context = {
        'project': project,
        'main_task_count': len(main_tasks),
        'subtask_count': len(subtasks),
        'document_count': len(documents),
        'main_tasks': main_tasks[:5],  # Show first 5 main tasks as examples
        'subtasks': subtasks[:10],     # Show first 10 subtasks as examples
        'documents': documents[:5],    # Show first 5 documents as examples
    }
    
    return render(request, 'projects/projects_delete.html', context)

@login_required
def projects_detail(request, project_id):
    """Display project details with tasks and subtasks - OPTIMIZED"""
    # OPTIMIZED: Single query with comprehensive prefetching
    project = get_object_or_404(
        Project.objects.select_related('prop').prefetch_related(
            Prefetch('projecttask_set', 
                queryset=ProjectTask.objects.filter(parent_task__isnull=True).prefetch_related(
                    Prefetch('subtasks', queryset=ProjectTask.objects.all())
                ).order_by('task_start_date', 'task_id')
            )
        ), 
        project_id=project_id
    )
    
    # OPTIMIZED: Use prefetched data (no additional queries)
    main_tasks = list(project.projecttask_set.all())  # These are already filtered and ordered
    
    context = {
        'project': project,
        'main_tasks': main_tasks,
        'task_status_choices': ProjectTask.TASK_STATUS_CHOICES,
        'task_priority_choices': ProjectTask.TASK_PRIORITY_CHOICES,
    }
    
    return render(request, 'projects/projects_detail.html', context)

@login_required
def project_tasks_add(request, project_id):
    """Add new task to project"""
    if not request.user.is_superuser:
        messages.error(request, "You don't have permission to add tasks.")
        return redirect('projects_detail', project_id=project_id)
    
    project = get_object_or_404(Project, project_id=project_id)  # Updated model name
    
    if request.method == 'POST':
        task_name = request.POST.get('task_name')
        task_description = request.POST.get('task_description')
        task_start_date = request.POST.get('task_start_date')
        task_expected_completion_date = request.POST.get('task_expected_completion_date')
        task_status = request.POST.get('task_status', 'Pending')
        task_priority = request.POST.get('task_priority', 'Medium')
        task_budgeted_cost = request.POST.get('task_budgeted_cost')
        task_actual_cost = request.POST.get('task_actual_cost')
        task_assigned_to = request.POST.get('task_assigned_to')
        task_actual_completion_date = request.POST.get('task_actual_completion_date')
        
        try:
            task = ProjectTask(  # Updated model name
                project=project,
                task_name=task_name,
                task_description=task_description,
                task_start_date=task_start_date if task_start_date else None,
                task_expected_completion_date=task_expected_completion_date if task_expected_completion_date else None,
                task_status=task_status,
                task_priority=task_priority,
                task_budgeted_cost=task_budgeted_cost if task_budgeted_cost else 0.00,
                task_actual_cost=task_actual_cost if task_actual_cost else 0.00,
                task_assigned_to=task_assigned_to,
                task_actual_completion_date=task_actual_completion_date if task_actual_completion_date else None
            )
            task.save()
            
            messages.success(request, f"Task '{task_name}' has been added successfully.")
            return redirect('projects_detail', project_id=project_id)
            
        except Exception as e:
            messages.error(request, f"Error adding task: {str(e)}")
    
    context = {
        'project': project,
        'task_status_choices': ProjectTask.TASK_STATUS_CHOICES,  # Updated model name
        'task_priority_choices': ProjectTask.TASK_PRIORITY_CHOICES,  # Updated model name
    }
    
    return render(request, 'projects/project_tasks_add.html', context)

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.core.exceptions import ValidationError
from django.utils import timezone
from datetime import datetime
from decimal import Decimal
from .models import Project, ProjectTask

def project_tasks_edit(request, project_id, task_id):
    """
    Edit a project task or subtask with support for Greek language fields
    """
    project = get_object_or_404(Project, project_id=project_id)
    task = get_object_or_404(ProjectTask, task_id=task_id, project=project)
    
    # Check if coming from Gantt chart
    from_gantt = request.GET.get('from_gantt', False)
    
    if request.method == 'POST':
        try:
            # Update basic task fields (always editable)
            task.task_name = request.POST.get('task_name', '').strip()
            task.task_description = request.POST.get('task_description', '').strip()
            
            # Update Greek fields
            task.task_name_greek = request.POST.get('task_name_greek', '').strip()
            task.task_description_greek = request.POST.get('task_description_greek', '').strip()
            
            # Update priority (always editable)
            task.task_priority = request.POST.get('task_priority')
            
            # Handle different logic for main tasks vs subtasks
            if task.parent_task:  # This is a subtask - most fields are editable
                # Status is editable for subtasks
                task.task_status = request.POST.get('task_status')
                
                # Dates are editable for subtasks
                start_date = request.POST.get('task_start_date')
                if start_date:
                    task.task_start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
                else:
                    task.task_start_date = None
                
                expected_date = request.POST.get('task_expected_completion_date')
                if expected_date:
                    task.task_expected_completion_date = datetime.strptime(expected_date, '%Y-%m-%d').date()
                else:
                    task.task_expected_completion_date = None
                
                # Handle actual completion date (only for subtasks when status is Completed)
                actual_date = request.POST.get('task_actual_completion_date')
                if actual_date:
                    task.task_actual_completion_date = datetime.strptime(actual_date, '%Y-%m-%d').date()
                else:
                    task.task_actual_completion_date = None
                
                # Costs are editable for subtasks
                budgeted_cost = request.POST.get('task_budgeted_cost')
                if budgeted_cost:
                    task.task_budgeted_cost = Decimal(budgeted_cost)
                else:
                    task.task_budgeted_cost = Decimal('0.00')
                
                actual_cost = request.POST.get('task_actual_cost')
                if actual_cost:
                    task.task_actual_cost = Decimal(actual_cost)
                else:
                    task.task_actual_cost = Decimal('0.00')
                
                # Progress percentage is editable for subtasks
                progress = request.POST.get('task_progress_percentage')
                if progress:
                    task.task_progress_percentage = int(progress)
                else:
                    task.task_progress_percentage = 0
                
                # Assigned to is editable for subtasks
                task.task_assigned_to = request.POST.get('task_assigned_to', '').strip()
                
            else:  # This is a main task - most fields are auto-calculated
                # For main tasks, only basic info and priority are directly editable
                # Status, dates, and costs are calculated from subtasks
                # The model's calculation methods will handle the auto-calculation
                pass
            
            # Validate the task before saving
            task.full_clean()
            
            # Save the task
            task.save()
            
            # Success message
            if task.parent_task:
                messages.success(request, f'Subtask "{task.task_name}" updated successfully!')
            else:
                messages.success(request, f'Main task "{task.task_name}" updated successfully!')
            
            # Redirect based on where we came from
            if from_gantt:
                return redirect('project_gantt', project_id=project.project_id)
            else:
                return redirect('projects_detail', project_id=project.project_id)
                
        except ValidationError as e:
            # Handle Django model validation errors
            error_messages = []
            if hasattr(e, 'error_dict'):
                for field, errors in e.error_dict.items():
                    field_name = field.replace('_', ' ').title()
                    for error in errors:
                        error_messages.append(f"{field_name}: {error}")
            else:
                error_messages = e.messages if hasattr(e, 'messages') else [str(e)]
            
            for error_msg in error_messages:
                messages.error(request, error_msg)
                
        except ValueError as e:
            # Handle value conversion errors (dates, decimals, etc.)
            if 'time data' in str(e):
                messages.error(request, 'Invalid date format. Please use the date picker.')
            elif 'invalid literal' in str(e):
                messages.error(request, 'Invalid number format. Please enter valid numbers for costs and percentages.')
            else:
                messages.error(request, f'Invalid data: {str(e)}')
                
        except Exception as e:
            # Handle any other unexpected errors
            messages.error(request, f'Error updating task: {str(e)}')
    
    # Prepare context for the template
    context = {
        'project': project,
        'task': task,
        'task_status_choices': ProjectTask.TASK_STATUS_CHOICES,
        'task_priority_choices': ProjectTask.TASK_PRIORITY_CHOICES,
        'from_gantt': from_gantt,
    }
    
    return render(request, 'projects/project_tasks_edit.html', context)

@login_required
@require_http_methods(["POST"])
def translate_text(request):
    """
    Translate text using Google Translate API
    """
    try:
        data = json.loads(request.body)
        text = data.get('text', '').strip()
        target_language = data.get('target_language', 'greek')
        source_language = data.get('source_language', 'english')
        
        if not text:
            return JsonResponse({'success': False, 'error': 'No text provided'})
        
        # Use Google Translate service
        if target_language == 'greek':
            translated_text = translate_to_greek_service(text)
        else:
            translated_text = text
        
        return JsonResponse({
            'success': True,
            'translated_text': translated_text,
            'source_language': source_language,
            'target_language': target_language
        })
        
    except Exception as e:
        print(f"Translation view error: {e}")
        return JsonResponse({'success': False, 'error': str(e)})

def translate_to_greek_service(text):
    """
    Use Google Translate API to translate English text to Greek
    """
    try:
        from googletrans import Translator
        
        # Initialize Google Translator
        translator = Translator()
        
        # Translate from English to Greek
        result = translator.translate(text, dest='el', src='en')
        
        return result.text
        
    except Exception as e:
        print(f"Google Translation service error: {e}")
        return text  # Return original text if translation fails

@login_required
def project_tasks_delete(request, project_id, task_id):
    """Delete task"""
    if not request.user.is_superuser:
        messages.error(request, "You don't have permission to delete tasks.")
        return redirect('projects_detail', project_id=project_id)
    
    project = get_object_or_404(Project, project_id=project_id)  # Updated model name
    task = get_object_or_404(ProjectTask, task_id=task_id, project=project)  # Updated model name
    task_name = task.task_name
    
    if request.method == 'POST':
        task.delete()
        messages.success(request, f"Task '{task_name}' has been deleted successfully.")
        return redirect('projects_detail', project_id=project_id)
    
    context = {
        'project': project,
        'task': task,
    }
    
    return render(request, 'projects/project_tasks_delete.html', context)

@login_required
def project_subtasks_add(request, project_id, parent_task_id):
    """Add subtask to a main task"""
    if not request.user.is_superuser:
        messages.error(request, "You don't have permission to add subtasks.")
        return redirect('projects_detail', project_id=project_id)
    
    project = get_object_or_404(Project, project_id=project_id)  # Updated model name
    parent_task = get_object_or_404(ProjectTask, task_id=parent_task_id, project=project)  # Updated model name
    
    if request.method == 'POST':
        task_name = request.POST.get('task_name')
        task_description = request.POST.get('task_description')
        task_start_date = request.POST.get('task_start_date')
        task_expected_completion_date = request.POST.get('task_expected_completion_date')
        task_status = request.POST.get('task_status', 'Pending')
        task_priority = request.POST.get('task_priority', 'Medium')
        task_budgeted_cost = request.POST.get('task_budgeted_cost')
        task_actual_cost = request.POST.get('task_actual_cost')
        task_assigned_to = request.POST.get('task_assigned_to')
        task_actual_completion_date = request.POST.get('task_actual_completion_date')
        
        try:
            subtask = ProjectTask(  # Updated model name
                project=project,
                parent_task=parent_task,
                task_name=task_name,
                task_description=task_description,
                task_start_date=task_start_date if task_start_date else None,
                task_expected_completion_date=task_expected_completion_date if task_expected_completion_date else None,
                task_status=task_status,
                task_priority=task_priority,
                task_budgeted_cost=task_budgeted_cost if task_budgeted_cost else 0.00,
                task_actual_cost=task_actual_cost if task_actual_cost else 0.00,
                task_assigned_to=task_assigned_to,
                task_actual_completion_date=task_actual_completion_date if task_actual_completion_date else None
            )
            subtask.save()
            
            messages.success(request, f"Subtask '{task_name}' has been added successfully.")
            return redirect('projects_detail', project_id=project_id)
            
        except Exception as e:
            messages.error(request, f"Error adding subtask: {str(e)}")
    
    context = {
        'project': project,
        'parent_task': parent_task,
        'task_status_choices': ProjectTask.TASK_STATUS_CHOICES,  # Updated model name
        'task_priority_choices': ProjectTask.TASK_PRIORITY_CHOICES,  # Updated model name
    }
    
    return render(request, 'projects/project_subtasks_add.html', context)

@login_required
def project_gantt(request, project_id):
    """Display Gantt chart for project with tasks and subtasks - OPTIMIZED"""
    # OPTIMIZED: Single query with comprehensive prefetching
    project = get_object_or_404(
        Project.objects.select_related('prop').prefetch_related(
            Prefetch('projecttask_set', 
                queryset=ProjectTask.objects.filter(parent_task__isnull=True).prefetch_related(
                    Prefetch('subtasks', 
                        queryset=ProjectTask.objects.filter(
                            task_start_date__isnull=False,
                            task_expected_completion_date__isnull=False
                        ).order_by('task_start_date', 'task_id')
                    )
                ).order_by('task_start_date', 'task_id')
            )
        ), 
        project_id=project_id
    )
    
    # Check if returning from edit page
    from_edit = request.GET.get('from_edit', False)
    if from_edit:
        messages.success(request, "Changes saved successfully. Gantt chart has been refreshed.")
    
    # OPTIMIZED: Use prefetched data (no additional queries)
    main_tasks = list(project.projecttask_set.all())
    
    # Build Gantt data structure using prefetched data
    gantt_data = []
    
    # Add project as the main item
    project_start = project.get_calculated_start_date()
    project_end = project.get_calculated_expected_completion()
    
    if project_start and project_end:
        project_item = {
            'id': f'project_{project.project_id}',
            'text': project.project_name,
            'start_date': project_start.strftime('%Y-%m-%d'),
            'end_date': project_end.strftime('%Y-%m-%d'),
            'duration': (project_end - project_start).days + 1,
            'progress': project.get_progress_percentage() / 100,
            'type': 'project',
            'status': project.get_calculated_status(),
            'budgeted_cost': float(project.get_calculated_budgeted_cost() or 0),
            'actual_cost': float(project.get_calculated_actual_cost() or 0),
            'open': True
        }
        gantt_data.append(project_item)
    
    # Add main tasks and their subtasks using prefetched data
    for task in main_tasks:
        task_start = task.get_calculated_start_date()
        task_end = task.get_calculated_expected_completion()
        
        if task_start and task_end:
            # Add main task
            task_item = {
                'id': f'task_{task.task_id}',
                'text': task.task_name,
                'start_date': task_start.strftime('%Y-%m-%d'),
                'end_date': task_end.strftime('%Y-%m-%d'),
                'duration': (task_end - task_start).days + 1,
                'progress': task.get_subtask_progress() / 100 if task.subtasks.all() else (1.0 if task.get_calculated_status() == 'Completed' else 0.0),
                'type': 'task',
                'status': task.get_calculated_status(),
                'budgeted_cost': float(task.get_calculated_budgeted_cost() or 0),
                'actual_cost': float(task.get_calculated_actual_cost() or 0),
                'assigned_to': task.task_assigned_to or '',
                'parent': f'project_{project.project_id}' if project_start and project_end else None,
                'open': True,
                'calculated_progress_percentage': round(task.get_subtask_progress(), 1)
            }
            gantt_data.append(task_item)
            
            # Add subtasks for this main task using prefetched data
            subtasks = list(task.subtasks.all())  # Already filtered in prefetch
            
            for subtask in subtasks:
                subtask_start = subtask.task_start_date
                subtask_end = subtask.task_expected_completion_date
                
                if subtask_start and subtask_end:
                    # Calculate progress for subtask
                    subtask_progress = 0
                    if subtask.task_status == 'Completed':
                        subtask_progress = 1.0
                    elif subtask.task_status == 'In Progress':
                        subtask_progress = (subtask.task_progress_percentage or 0) / 100
                    
                    subtask_item = {
                        'id': f'subtask_{subtask.task_id}',
                        'text': subtask.task_name,
                        'start_date': subtask_start.strftime('%Y-%m-%d'),
                        'end_date': subtask_end.strftime('%Y-%m-%d'),
                        'duration': (subtask_end - subtask_start).days + 1,
                        'progress': subtask_progress,
                        'type': 'subtask',
                        'status': subtask.task_status,
                        'budgeted_cost': float(subtask.task_budgeted_cost or 0),
                        'actual_cost': float(subtask.task_actual_cost or 0),
                        'assigned_to': subtask.task_assigned_to or '',
                        'parent': f'task_{task.task_id}',
                        'priority': subtask.task_priority,
                        'progress_percentage': subtask.task_progress_percentage or 0
                    }
                    gantt_data.append(subtask_item)
    
    # If no tasks have dates, create a placeholder message
    if not gantt_data:
        from datetime import datetime
        today = datetime.now().date()
        placeholder_item = {
            'id': 'placeholder_1',
            'text': f'{project.project_name} (No dates set)',
            'start_date': today.strftime('%Y-%m-%d'),
            'duration': 30,
            'progress': 0,
            'type': 'project',
            'status': 'Pending'
        }
        gantt_data.append(placeholder_item)
    
    context = {
        'project': project,
        'main_tasks': main_tasks,
        'gantt_data': json.dumps(gantt_data),
    }
    
    return render(request, 'projects/project_gantt.html', context)

@login_required
def ajax_update_project_status(request):
    """AJAX view to update project status"""
    if request.method == 'POST' and request.user.is_superuser:
        try:
            data = json.loads(request.body)
            project_id = data.get('project_id')
            new_status = data.get('status')
            actual_completion_date = data.get('actual_completion_date')
            
            project = get_object_or_404(Project, project_id=project_id)
            project.project_status = new_status
            
            if new_status == 'Completed' and actual_completion_date:
                project.project_actual_completion_date = actual_completion_date
            elif new_status != 'Completed':
                project.project_actual_completion_date = None
            
            project.save()
            
            return JsonResponse({
                'success': True,
                'message': f"Project status updated to {new_status}"
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': f"Error updating project status: {str(e)}"
            })
    
    return JsonResponse({'success': False, 'message': 'Invalid request'})

@login_required
def ajax_update_task_status(request):
    """AJAX view to update task status"""
    if request.method == 'POST' and request.user.is_superuser:
        try:
            data = json.loads(request.body)
            task_id = data.get('task_id')
            new_status = data.get('status')
            actual_completion_date = data.get('actual_completion_date')
            
            task = get_object_or_404(ProjectTask, task_id=task_id)
            task.task_status = new_status
            
            if new_status == 'Completed' and actual_completion_date:
                task.task_actual_completion_date = actual_completion_date
            elif new_status != 'Completed':
                task.task_actual_completion_date = None
            
            task.save()
            
            return JsonResponse({
                'success': True,
                'message': f"Task status updated to {new_status}"
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': f"Error updating task status: {str(e)}"
            })
    
    return JsonResponse({'success': False, 'message': 'Invalid request'})

@login_required
def ajax_duplicate_project(request):
    """
    OPTIMIZED: AJAX view to duplicate a project with all its tasks and subtasks,
    adjusting all dates based on the new project start date and handling budget copy options
    """
    if request.method != 'POST':
        return JsonResponse({'success': False, 'message': 'Only POST method allowed'})
    
    if not request.user.is_superuser:
        return JsonResponse({
            'success': False,
            'message': 'You do not have permission to duplicate projects'
        })
    
    try:
        # Parse JSON data
        data = json.loads(request.body)
        project_id = data.get('project_id')
        new_project_name = data.get('new_project_name', '').strip()
        new_project_description = data.get('new_project_description', '').strip()
        new_project_start_date_str = data.get('new_project_start_date', '').strip()
        budget_copy_option = data.get('budget_copy_option', 'budgeted')
        clear_greek_translations = data.get('clear_greek_translations', False)
        
        # Validate required fields
        if not project_id or not new_project_name or not new_project_description or not new_project_start_date_str:
            return JsonResponse({
                'success': False,
                'message': 'Project ID, new project name, description, and start date are required'
            })
        
        # Validate budget copy option
        if budget_copy_option not in ['budgeted', 'actual']:
            budget_copy_option = 'budgeted'
        
        # Parse the new start date
        try:
            from datetime import datetime, timedelta
            new_project_start_date = datetime.strptime(new_project_start_date_str, '%Y-%m-%d').date()
        except ValueError:
            return JsonResponse({
                'success': False,
                'message': 'Invalid date format. Please use YYYY-MM-DD format.'
            })
        
        # Ensure project_id is an integer
        try:
            project_id = int(project_id)
        except (ValueError, TypeError):
            return JsonResponse({
                'success': False,
                'message': f'Invalid project ID: {project_id}'
            })
        
        # OPTIMIZED: Get the original project with comprehensive prefetching
        try:
            original_project = Project.objects.select_related('prop').prefetch_related(
                Prefetch('projecttask_set', 
                    queryset=ProjectTask.objects.select_related().prefetch_related('subtasks')
                ),
                'project_documents'
            ).get(project_id=project_id)
        except Project.DoesNotExist:
            return JsonResponse({
                'success': False,
                'message': f'Project with ID {project_id} was not found'
            })
        
        # Check if project name already exists
        if Project.objects.filter(project_name=new_project_name).exists():
            return JsonResponse({
                'success': False,
                'message': f'A project with the name "{new_project_name}" already exists'
            })
        
        # Calculate date offset if original project has a start date
        date_offset = None
        original_start_date = original_project.get_calculated_start_date()
        
        if original_start_date:
            date_offset = (new_project_start_date - original_start_date).days
        
        def adjust_date(original_date, offset_days):
            """Helper function to adjust a date by the offset"""
            if original_date and offset_days is not None:
                return original_date + timedelta(days=offset_days)
            return original_date
        
        def get_cost_for_budget(original_task, budget_option):
            """Helper function to determine which cost to use as the new budget"""
            if budget_option == 'actual':
                return original_task.task_actual_cost
            else:
                return original_task.task_budgeted_cost
        
        # Use transaction to ensure all-or-nothing duplication
        with transaction.atomic():
            # Calculate new project dates
            new_project_expected_completion = None
            if original_project.get_calculated_expected_completion() and date_offset is not None:
                new_project_expected_completion = adjust_date(
                    original_project.get_calculated_expected_completion(), 
                    date_offset
                )
            
            # Calculate new project budget based on option
            new_project_total_budgeted_cost = original_project.project_total_budgeted_cost
            if budget_copy_option == 'actual':
                new_project_total_budgeted_cost = original_project.get_calculated_actual_cost()
            
            # Create new project
            new_project = Project.objects.create(
                project_name=new_project_name,
                project_description=new_project_description,
                prop=original_project.prop,
                project_start_date=new_project_start_date,
                project_expected_completion_date=new_project_expected_completion,
                project_status='Pending',
                project_actual_completion_date=None,
                project_total_budgeted_cost=new_project_total_budgeted_cost,
                project_total_actual_cost=Decimal('0.00'),
                project_name_greek=None if clear_greek_translations else getattr(original_project, 'project_name_greek', None),
                project_description_greek=None if clear_greek_translations else getattr(original_project, 'project_description_greek', None),
            )
            
            # OPTIMIZED: Get all tasks using prefetched data
            all_original_tasks = list(original_project.projecttask_set.all())
            main_tasks = [task for task in all_original_tasks if task.parent_task_id is None]
            
            # OPTIMIZED: Prepare bulk data for main tasks
            main_tasks_to_create = []
            task_id_mapping = {}  # Map old task ID to new task index
            
            # First pass: Prepare main tasks for bulk creation
            for original_task in main_tasks:
                new_task_start_date = adjust_date(original_task.task_start_date, date_offset)
                new_task_expected_completion = adjust_date(original_task.task_expected_completion_date, date_offset)
                new_task_budgeted_cost = get_cost_for_budget(original_task, budget_copy_option)
                
                new_task = ProjectTask(
                    project=new_project,
                    task_name=original_task.task_name,
                    task_description=original_task.task_description,
                    task_start_date=new_task_start_date,
                    task_expected_completion_date=new_task_expected_completion,
                    task_budgeted_cost=new_task_budgeted_cost,
                    task_actual_cost=Decimal('0.00'),
                    task_priority=original_task.task_priority,
                    task_status='Pending',
                    task_actual_completion_date=None,
                    task_assigned_to=original_task.task_assigned_to,
                    parent_task=None,
                    task_progress_percentage=0,
                    task_name_greek=getattr(original_task, 'task_name_greek', None),
                    task_description_greek=getattr(original_task, 'task_description_greek', None),
                )
                main_tasks_to_create.append(new_task)
                # Store mapping for later subtask creation
                task_id_mapping[original_task.task_id] = len(main_tasks_to_create) - 1
            
            # OPTIMIZED: Bulk create main tasks
            created_main_tasks = ProjectTask.objects.bulk_create(main_tasks_to_create)
            
            # IMPORTANT: After bulk_create, we need to fetch the tasks with their IDs
            # because bulk_create doesn't populate the ID field on the returned objects
            created_main_tasks_with_ids = list(
                ProjectTask.objects.filter(
                    project=new_project, 
                    parent_task__isnull=True
                ).order_by('task_id')
            )
            
            # Create mapping from original task ID to new task object (with ID)
            task_object_mapping = {}
            for i, original_task in enumerate(main_tasks):
                if i < len(created_main_tasks_with_ids):
                    task_object_mapping[original_task.task_id] = created_main_tasks_with_ids[i]
            
            # OPTIMIZED: Prepare bulk data for subtasks
            subtasks_to_create = []
            
            # Get all subtasks using prefetched data and group by parent
            for original_main_task in main_tasks:
                # Use prefetched subtasks
                subtasks = list(original_main_task.subtasks.all())
                
                if original_main_task.task_id in task_object_mapping:
                    new_main_task = task_object_mapping[original_main_task.task_id]
                    
                    for original_subtask in subtasks:
                        new_subtask_start_date = adjust_date(original_subtask.task_start_date, date_offset)
                        new_subtask_expected_completion = adjust_date(original_subtask.task_expected_completion_date, date_offset)
                        new_subtask_budgeted_cost = get_cost_for_budget(original_subtask, budget_copy_option)
                        
                        new_subtask = ProjectTask(
                            project=new_project,
                            task_name=original_subtask.task_name,
                            task_description=original_subtask.task_description,
                            task_start_date=new_subtask_start_date,
                            task_expected_completion_date=new_subtask_expected_completion,
                            task_budgeted_cost=new_subtask_budgeted_cost,
                            task_actual_cost=Decimal('0.00'),
                            task_priority=original_subtask.task_priority,
                            task_status='Pending',
                            task_actual_completion_date=None,
                            task_assigned_to=original_subtask.task_assigned_to,
                            parent_task=new_main_task,
                            task_progress_percentage=0,
                            task_name_greek=getattr(original_subtask, 'task_name_greek', None),
                            task_description_greek=getattr(original_subtask, 'task_description_greek', None),
                        )
                        subtasks_to_create.append(new_subtask)
            
            # OPTIMIZED: Bulk create subtasks
            if subtasks_to_create:
                ProjectTask.objects.bulk_create(subtasks_to_create)
            
            # OPTIMIZED: Copy project documents using prefetched data
            documents_to_create = []
            try:
                original_documents = list(original_project.project_documents.all())
                for original_doc in original_documents:
                    new_document = ProjectDocument(
                        project=new_project,
                        task=None,
                        document_name=f"Copy of {original_doc.document_name}" if original_doc.document_name else None,
                        document_description=original_doc.document_description,
                        document_file=original_doc.document_file,
                        document_uploaded_by=request.user.username,
                    )
                    documents_to_create.append(new_document)
                
                # Bulk create documents
                if documents_to_create:
                    ProjectDocument.objects.bulk_create(documents_to_create)
                    
            except Exception as doc_error:
                # Silent fail for document copying
                pass
        
        # Build success message
        budget_message = ""
        if budget_copy_option == 'actual':
            budget_message = " with actual costs copied as budgeted costs"
        else:
            budget_message = " with budgeted costs copied"
        
        translation_message = ""
        if clear_greek_translations:
            translation_message = " and Greek translations cleared for project name and description"
            
        success_message = f'Project "{new_project_name}" created successfully{budget_message}{translation_message}'
        if date_offset is not None:
            success_message += f' and all dates adjusted by {date_offset} days'
        
        return JsonResponse({
            'success': True,
            'message': success_message,
            'new_project_id': new_project.project_id,
            'date_offset': date_offset,
            'budget_copy_option': budget_copy_option,
            'greek_translations_cleared': clear_greek_translations
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'message': 'Invalid JSON data'
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'message': f'An error occurred while duplicating the project: {str(e)}'
        })

@login_required
def project_task_list(request, project_id):
    """Display task list for a specific project and assignee - OPTIMIZED"""
    # OPTIMIZED: Single query with comprehensive prefetching
    project = get_object_or_404(
        Project.objects.select_related('prop').prefetch_related(
            Prefetch('projecttask_set', 
                queryset=ProjectTask.objects.filter(parent_task__isnull=True).prefetch_related(
                    'subtasks'
                ).order_by('task_start_date', 'task_id')
            )
        ), 
        project_id=project_id
    )
    
    # Get parameters
    assigned_to = request.GET.get('assigned_to', '')
    language = request.GET.get('language', 'english')
    
    # Ensure Greek translations if language is Greek
    if language == 'greek':
        ensure_project_translations(project)
    
    # OPTIMIZED: Use prefetched data
    main_tasks = list(project.projecttask_set.all())
    
    # OPTIMIZED: Filter in Python using prefetched data instead of additional queries
    if assigned_to:
        filtered_main_tasks = []
        for task in main_tasks:
            task_matches = task.task_assigned_to == assigned_to
            subtask_matches = any(subtask.task_assigned_to == assigned_to for subtask in task.subtasks.all())
            if task_matches or subtask_matches:
                filtered_main_tasks.append(task)
        main_tasks = filtered_main_tasks
    
    # Build task list with hierarchy using prefetched data
    task_list = []
    total_tasks = 0
    completed_tasks = 0
    pending_tasks = 0
    
    # Add project as root item
    project_start_date = project.get_calculated_start_date()
    project_end_date = project.get_calculated_expected_completion()
    project_is_overdue = (
        project_end_date and 
        project_end_date < timezone.now().date() and 
        project.get_calculated_status() != 'Completed'
    )
    
    project_item = {
        'name': project.project_name,
        'name_greek': get_translated_text(
            project.project_name, 
            getattr(project, 'project_name_greek', None), 
            language
        ) if language == 'greek' else project.project_name,
        'description': project.project_description,
        'description_greek': get_translated_text(
            project.project_description, 
            getattr(project, 'project_description_greek', None), 
            language
        ) if language == 'greek' else project.project_description,
        'type': 'project',
        'status': project.get_calculated_status(),
        'start_date': project_start_date,
        'end_date': project_end_date,
        'priority': None,
        'indent_level': 0,
        'is_overdue': project_is_overdue,
        'project_obj': project
    }
    task_list.append(project_item)
    
    # Process main tasks and subtasks using prefetched data
    for main_task in main_tasks:
        # For main tasks, use calculated dates
        task_start_date = main_task.get_calculated_start_date()
        task_end_date = main_task.get_calculated_expected_completion()
        task_is_overdue = (
            task_end_date and 
            task_end_date < timezone.now().date() and 
            main_task.get_calculated_status() != 'Completed'
        )
        
        main_task_item = {
            'name': main_task.task_name,
            'name_greek': get_translated_text(
                main_task.task_name, 
                getattr(main_task, 'task_name_greek', None), 
                language
            ) if language == 'greek' else main_task.task_name,
            'description': main_task.task_description,
            'description_greek': get_translated_text(
                main_task.task_description, 
                getattr(main_task, 'task_description_greek', None), 
                language
            ) if language == 'greek' else main_task.task_description,
            'type': 'task',
            'status': main_task.get_calculated_status(),
            'start_date': task_start_date,
            'end_date': task_end_date,
            'priority': main_task.task_priority,
            'indent_level': 1,
            'is_overdue': task_is_overdue,
            'task_obj': main_task
        }
        task_list.append(main_task_item)
        
        # Add subtasks using prefetched data
        subtasks = list(main_task.subtasks.all())
        if assigned_to:
            subtasks = [subtask for subtask in subtasks if subtask.task_assigned_to == assigned_to]
        
        # Sort subtasks by start date and task_id
        subtasks.sort(key=lambda x: (x.task_start_date or timezone.now().date(), x.task_id))
        
        for subtask in subtasks:
            is_overdue = (
                subtask.task_expected_completion_date and 
                subtask.task_expected_completion_date < timezone.now().date() and 
                subtask.task_status != 'Completed'
            )
            
            subtask_item = {
                'name': subtask.task_name,
                'name_greek': get_translated_text(
                    subtask.task_name, 
                    getattr(subtask, 'task_name_greek', None), 
                    language
                ) if language == 'greek' else subtask.task_name,
                'description': subtask.task_description,
                'description_greek': get_translated_text(
                    subtask.task_description, 
                    getattr(subtask, 'task_description_greek', None), 
                    language
                ) if language == 'greek' else subtask.task_description,
                'type': 'subtask',
                'status': subtask.task_status,
                'start_date': subtask.task_start_date,
                'end_date': subtask.task_expected_completion_date,
                'priority': subtask.task_priority,
                'indent_level': 2,
                'is_overdue': is_overdue,
                'task_obj': subtask
            }
            task_list.append(subtask_item)
            
            # ONLY count subtasks in totals
            total_tasks += 1
            
            if subtask.task_status == 'Completed':
                completed_tasks += 1
            else:
                pending_tasks += 1
    
    # Calculate completion percentage
    completion_percentage = (completed_tasks / total_tasks * 100) if total_tasks > 0 else 0
    
    context = {
        'project': project,
        'task_list': task_list,
        'assigned_to': assigned_to,
        'language': language,
        'total_tasks': total_tasks,
        'completed_tasks': completed_tasks,
        'pending_tasks': pending_tasks,
        'completion_percentage': completion_percentage,
        'current_date': timezone.now(),
    }
    
    return render(request, 'projects/project_task_list.html', context)

@login_required
def ajax_delete_task(request):
    """
    AJAX view to delete a task or subtask
    """
    if request.method != 'POST':
        return JsonResponse({'success': False, 'message': 'Only POST method allowed'})
    
    if not request.user.is_superuser:
        return JsonResponse({
            'success': False,
            'message': 'You do not have permission to delete tasks'
        })
    
    try:
        data = json.loads(request.body)
        task_id = data.get('task_id')
        task_type = data.get('task_type')  # 'task' or 'subtask'
        
        if not task_id:
            return JsonResponse({
                'success': False,
                'message': 'Task ID is required'
            })
        
        try:
            task_id = int(task_id)
        except (ValueError, TypeError):
            return JsonResponse({
                'success': False,
                'message': f'Invalid task ID: {task_id}'
            })
        
        # Get the task to delete
        try:
            task_to_delete = ProjectTask.objects.get(task_id=task_id)
        except ProjectTask.DoesNotExist:
            return JsonResponse({
                'success': False,
                'message': f'Task with ID {task_id} was not found'
            })
        
        # Use transaction to ensure all-or-nothing deletion
        with transaction.atomic():
            if task_type == 'task':
                # Delete main task and all its subtasks
                subtasks = ProjectTask.objects.filter(parent_task=task_to_delete)
                subtask_count = subtasks.count()
                
                # Delete subtasks first
                subtasks.delete()
                
                # Delete the main task
                task_name = task_to_delete.task_name
                task_to_delete.delete()
                
                message = f'Task "{task_name}" and {subtask_count} subtask(s) deleted successfully'
                
            elif task_type == 'subtask':
                # Delete only the subtask
                task_name = task_to_delete.task_name
                task_to_delete.delete()
                
                message = f'Subtask "{task_name}" deleted successfully'
            
            else:
                return JsonResponse({
                    'success': False,
                    'message': 'Invalid task type. Must be "task" or "subtask"'
                })
        
        return JsonResponse({
            'success': True,
            'message': message
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'message': 'Invalid JSON data'
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'message': f'An error occurred while deleting: {str(e)}'
        })

@login_required
def get_project_assignees(request, project_id):
    """AJAX endpoint to get all assignees for a project - OPTIMIZED"""
    # OPTIMIZED: Single query with prefetching
    project = get_object_or_404(
        Project.objects.prefetch_related('projecttask_set'), 
        project_id=project_id
    )
    
    # OPTIMIZED: Use prefetched data to get assignees
    assignees = set()
    
    for task in project.projecttask_set.all():
        if task.task_assigned_to and task.task_assigned_to.strip():
            assignees.add(task.task_assigned_to.strip())
    
    # Convert to sorted list
    assignees_list = sorted(list(assignees))
    
    return JsonResponse({
        'success': True,
        'assignees': assignees_list,
        'project_name': project.project_name
    })


def render_to_pdf(template_src, context_dict):
    template = get_template(template_src)
    html = template.render(context_dict)
    response = HttpResponse(content_type='application/pdf')
    pisa_status = pisa.CreatePDF(html, dest=response)
    if pisa_status.err:
        return HttpResponse('PDF generation failed', status=500)
    return response

### HOME ###
def home(request):
	results = props.objects.all().order_by('prop_country','prop_name')
	tresults = tenant.objects.filter(tenant_current="Yes")
	sresults = supplier.objects.all().order_by('supplier_country','supplier_contact_person')
	return render (request, "home.html", {"props":results, "tenant":tresults, "supplier":sresults})

### ADMIN ###
@login_required
def admin_apms(request):
    results = props.objects.all().order_by('prop_country', 'prop_name')
    tresults = tenant.objects.select_related('prop').all().order_by('tenant_name')
    return render(request, "admin_apms.html", {
        "props": results, 
        "tenant": tresults
    })

@login_required
def personal_page(request):
    """Personal management page"""
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to access this page.')
        return redirect('home')
    
    return render(request, "personal.html")

@login_required
def lease_agreement_report(request, tenant_id):
    try:
        # Get tenant and property info
        tenant_obj = get_object_or_404(tenant, pk=tenant_id)
        if not hasattr(tenant_obj, 'prop') or not tenant_obj.prop:
            raise Http404("Tenant has no property assigned")
        
        property_name = tenant_obj.prop.prop_name
        filename = f"{property_name} - Lease Agreement.pdf"
        file_path = os.path.join(settings.MEDIA_ROOT, 'lease_agreements', filename)
        
        context = {
            'tenant': tenant_obj,
            'property': tenant_obj.prop,
            'filename': filename,
            'file_exists': os.path.exists(file_path),
            'file_url': os.path.join(settings.MEDIA_URL, 'lease_agreements', filename)
        }
        return render(request, 'lease_agreement_report.html', context)
        
    except Exception as e:
        return render(request, 'error.html', {'error': str(e)})

@login_required
def serve_lease(request, filename):
    try:
        # Security validation
        if not filename.endswith(' - Lease Agreement.pdf'):
            raise Http404("Invalid filename format")
            
        file_path = os.path.join(settings.MEDIA_ROOT, 'lease_agreements', filename)
        
        if not os.path.exists(file_path):
            raise Http404("File not found")
            
        response = FileResponse(open(file_path, 'rb'), content_type='application/pdf')
        response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        return response
        
    except Exception as e:
        return Http404(str(e))

@login_required
def upload_lease_agreement(request):
    if request.method == 'POST':
        tenant_id = request.POST.get('tenant')
        uploaded_file = request.FILES.get('lease_agreement')
        
        if not uploaded_file:
            messages.error(request, "No file was uploaded.")
            return redirect('admin_apms')
        
        try:
            # Validate file
            if not uploaded_file.name.lower().endswith('.pdf'):
                raise ValueError("Only PDF files are allowed")
                
            tenant_obj = tenant.objects.get(pk=tenant_id)
            if not hasattr(tenant_obj, 'prop') or not tenant_obj.prop:
                raise ValueError("No property assigned to tenant")
                
            property_name = tenant_obj.prop.prop_name
            lease_dir = os.path.join(settings.STATIC_ROOT, 'lease_agreements')
            os.makedirs(lease_dir, exist_ok=True)
            
            filename = f"{property_name} - Lease Agreement.pdf"
            file_path = os.path.join(lease_dir, filename)
            
            # Save file
            with open(file_path, 'wb+') as destination:
                for chunk in uploaded_file.chunks():
                    destination.write(chunk)
            
            messages.success(request, f"Lease agreement uploaded successfully!")
            return redirect('admin_apms')
            
        except Exception as e:
            messages.error(request, f"Error: {str(e)}")
        
        return redirect('admin_apms')
    return redirect('admin_apms')

@login_required
def serve_lease(request, filename):
    """Secure file serving for exact filename format"""
    try:
        # Verify filename format
        if not filename.endswith(' - Lease Agreement.pdf'):
            raise Http404("Invalid filename format")
            
        file_path = os.path.join(settings.STATIC_ROOT, 'lease_agreements', filename)
        
        if not os.path.exists(file_path):
            raise Http404("Lease agreement not found")
            
        # Serve with cache-control headers
        response = FileResponse(open(file_path, 'rb'), content_type='application/pdf')
        response['Cache-Control'] = 'no-store, must-revalidate'
        response['Pragma'] = 'no-cache'
        return response
        
    except Exception as e:
        messages.error(request, f"Error serving file: {str(e)}")
        return redirect('admin_apms')

@login_required
def upload_title_deed(request):
    if request.method == 'POST':
        # Get the selected property name
        property_name = request.POST.get('property')
        
        # Get the uploaded file
        uploaded_file = request.FILES.get('title_deed')
        
        if not uploaded_file:
            messages.error(request, "No file was uploaded.")
            return redirect('admin_apms')
        
        # Validate file extension
        if not uploaded_file.name.lower().endswith('.pdf'):
            messages.error(request, "Only PDF files are allowed.")
            return redirect('admin_apms')
        
        # Create the title_deeds directory if it doesn't exist
        title_deeds_dir = os.path.join(settings.STATIC_ROOT, 'title_deeds')
        os.makedirs(title_deeds_dir, exist_ok=True)
        
        # Create the filename
        filename = f'{property_name} - Title Deed.pdf'
        file_path = os.path.join(title_deeds_dir, filename)
        
        # Save the file
        try:
            # Delete existing file if it exists
            if os.path.exists(file_path):
                os.remove(file_path)
                
            with open(file_path, 'wb+') as destination:
                for chunk in uploaded_file.chunks():
                    destination.write(chunk)
            messages.success(request, f"Title deed for {property_name} uploaded successfully!")
        except Exception as e:
            messages.error(request, f"Error saving file: {str(e)}")
        
        return redirect('admin_apms')
    
    return redirect('admin_apms')

@login_required
def admin_clear(request):
	import os
	import glob
	file_path = "C:/Users/DemetrisManias/Desktop/code/djangoproject/static/reports/*.pdf"
	files = glob.glob(file_path)
	for f in files:
		os.remove(f)
	return redirect("admin_apms")

@login_required
def admin_unpaid(request):
	import open_invoices
	rep_output = "Email"
	check = "Yes"
	email = "demetrimanias@gmail.com"
	fname = "Demetri"
	open_invoices.open_invoices(rep_output, check, email, fname)
#	email = "stella.simitopoulos@alivente.com"
#	fname = "Stella"
#	open_invoices.open_invoices(rep_output, check, email, fname)
	return redirect("admin_apms")

@login_required
def admin_renewals(request):
	import lease_renewal
	rep_output = "Email"
	check = "Yes"
	email = "demetrimanias@gmail.com"
	fname = "Demetri"
	lease_renewal.lease_renewal(rep_output,check, email, fname)
#	email = "stella.simitopoulos@alivente.com"
#	fname = "Stella"
#	lease_renewal.lease_renewal(rep_output,check, email, fname)
	return redirect("admin_apms")

@login_required
def admin_invoices(request):
	import open_invoices
	today = date.today()
	months = ('Month','January','February','March','April','May','June','July','August','September','October','November','December')
	open_invoices.create_invoices(months[today.month],today.year,request)
	return redirect("admin_apms")

### DASHBOARD ###
@login_required
def property_management_dashboard(request):
    """
    Main property dashboard view with spoke-and-wheel interface
    """
    try:
        # Get all properties for the dropdown
        properties = props.objects.filter(prop_status='Active').order_by('prop_name')
        
        # Check if a specific property was selected
        selected_property_id = request.GET.get('property')
        selected_property = None
        
        if selected_property_id:
            try:
                selected_property = props.objects.get(prop_id=selected_property_id)
            except props.DoesNotExist:
                messages.error(request, f"Property with ID {selected_property_id} not found.")
        
        context = {
            'properties': properties,
            'selected_property': selected_property,
        }
        
        return render(request, 'property_management_dashboard.html', context)
        
    except Exception as e:
        messages.error(request, f"Error loading dashboard: {str(e)}")
        return redirect('properties')

@login_required
def property_detail(request, property_id, box_type):
    property_obj = get_object_or_404(props, prop_id=property_id)
    
    # Get the active tenant for this property (there should only be one)
    active_tenant = tenant.objects.filter(
        prop=property_obj, 
        tenant_current='Yes'
    ).first()
    
    # Get open invoices data for this property
    open_invoices_data = None
    total_invoices_amount = 0
    
    # Lease renewal data
    lease_renewal_data = None
    
    # Issues data for this specific property
    property_issues = None
    resolved_count = 0
    unresolved_count = 0
    total_issues_count = 0
    
    # Valuation data for this specific property
    property_valuation = None
    
    # Revenue data for this specific property
    property_revenues = None
    total_revenue_amount = 0
    
    # Budgeted expenses data for this specific property
    property_budgeted_expenses = None
    total_budgeted_expense_amount = 0
    
    # Actual expenses data for this specific property
    property_actual_expenses = None
    actual_expense_years = []
    selected_actual_year = None
    
    if active_tenant:
        # Get all unpaid invoices for this tenant
        unpaid_invoices = invoices.objects.filter(
            tenant=active_tenant
        ).exclude(
            invoice_paid='Yes'  # Exclude paid invoices
        ).order_by('invoice_date')
        
        # Calculate days overdue and prepare data
        open_invoices_data = []
        today = timezone.now().date()
        
        for invoice in unpaid_invoices:
            # Calculate due date (invoice_date + payment_terms)
            payment_terms = active_tenant.tenant_payment_terms or 0
            due_date = invoice.invoice_date + timedelta(days=payment_terms)
            
            # Calculate days overdue
            days_overdue = (today - due_date).days if today > due_date else 0
            
            # Use the actual invoice amount if available, otherwise fall back to tenant rent
            invoice_amount = getattr(invoice, 'invoice_amount', None) or active_tenant.tenant_rent
            
            open_invoices_data.append({
                'invoice_date': invoice.invoice_date,
                'due_date': due_date,
                'days_overdue': days_overdue,
                'overdue': days_overdue > 0,
                'amount': invoice_amount
            })
            
            # Add to total amount
            total_invoices_amount += invoice_amount
        
        # Lease renewal logic for this specific property
        if box_type == 'lease-renewals':
            lease_renewal_data = {
                'tenant': active_tenant,
                'property': property_obj,
                'needs_renewal': False,
                'renewal_date': None,
                'status': 'current',
                'message': None
            }
            
            if active_tenant.tenant_lease_end_date:
                # Calculate renewal contact date
                renewal_period = active_tenant.tenant_renewal_period or 30
                renewal_contact_date = active_tenant.tenant_lease_end_date - timedelta(days=renewal_period)
                
                # Check if renewal is needed
                if today >= renewal_contact_date:
                    lease_renewal_data['needs_renewal'] = True
                    lease_renewal_data['renewal_date'] = renewal_contact_date
                    
                    # Check renewal status
                    if active_tenant.tenant_renewal_status == 'declined':
                        lease_renewal_data['status'] = 'declined'
                        lease_renewal_data['message'] = f"TENANT DECLINED RENEWAL - LEASE EXPIRES {active_tenant.tenant_lease_end_date}"
                    elif active_tenant.tenant_renewal_status == 'new_lease_signed':
                        lease_renewal_data['status'] = 'renewed'
                        lease_renewal_data['message'] = "NEW LEASE SIGNED"
                    else:
                        lease_renewal_data['status'] = 'pending'
                        lease_renewal_data['message'] = f"RENEWAL CONTACT REQUIRED BY {renewal_contact_date}"
    
    elif box_type == 'lease-renewals':
        # No active tenant - property is vacant
        lease_renewal_data = {
            'tenant': None,
            'property': property_obj,
            'status': 'vacant',
            'message': "NO CURRENT TENANT - NEED NEW TENANT"
        }
    
    # Issues logic - process for any box_type but only use data when box_type is 'issues'
    if box_type == 'issues':
        # Get all issues for this property, ordered by date (most recent first)
        property_issues = issues.objects.filter(
            prop=property_obj
        ).order_by('-issues_date_logged')
        
        # Calculate issue counts
        total_issues_count = property_issues.count()
        resolved_count = property_issues.filter(issues_status='Resolved').count()
        
        # Updated logic: Include both "Unresolved" AND "Issue" status as unresolved
        unresolved_count = property_issues.filter(
            issues_status__in=['Unresolved', 'Issue']
        ).count()
    
    # Valuation logic - process when box_type is 'valuation'
    if box_type == 'valuation':
        # Get valuation data for this property
        try:
            property_valuation = prop_values.objects.get(prop=property_obj)
            
            # Calculate value change in the view
            if property_valuation.prop_values_current_value and property_valuation.prop_values_purchase_price:
                difference = property_valuation.prop_values_current_value - property_valuation.prop_values_purchase_price
                if property_valuation.prop_values_purchase_price > 0:
                    percentage = (difference / property_valuation.prop_values_purchase_price) * 100
                else:
                    percentage = 0
                
                # Add calculated values to the valuation object
                property_valuation.value_difference = difference
                property_valuation.value_percentage = percentage
            else:
                property_valuation.value_difference = 0
                property_valuation.value_percentage = 0
                
        except prop_values.DoesNotExist:
            property_valuation = None
    
    # Revenue logic - process when box_type is 'revenues'
    if box_type == 'revenues':
        # Get all revenue data for this property
        property_revenues = property_obj.revenue_set.all().order_by('revenue_line_types__revenue_line_types_name', 'revenue_types__revenue_types_name')
        
        # Calculate total revenue amount
        total_revenue_amount = sum(rev.revenue_amount for rev in property_revenues)
    
    # Budgeted Expenses logic - process when box_type is 'budgeted-expenses'
    if box_type == 'budgeted-expenses':
        # Get all budgeted expense data for this property, sorted by expense line type
        property_budgeted_expenses = property_obj.expense_set.all().order_by('expense_line_types__expense_line_types_name', 'expense_types__expense_types_name')
        
        # Calculate total budgeted expense amount
        total_budgeted_expense_amount = sum(exp.expense_amount for exp in property_budgeted_expenses)
    
    # Actual Expenses logic - process when box_type is 'actual-expenses'
    if box_type == 'actual-expenses':
        from django.db.models import Q
        
        # Get selected year from request or default to current year
        selected_actual_year = request.GET.get('year')
        current_year = timezone.now().year
        
        # Get all years that have actual expenses for this property (approved and paid only)
        actual_expense_years = list(
            property_obj.act_expense_set.filter(
                act_expense_approved='Yes',
                act_expense_paid='Yes'
            ).dates('act_expense_date', 'year', order='DESC').distinct()
        )
        actual_expense_years = [date.year for date in actual_expense_years]
        
        # Default to the latest year if no year selected
        if not selected_actual_year and actual_expense_years:
            selected_actual_year = actual_expense_years[0]
        elif not selected_actual_year:
            selected_actual_year = current_year
        else:
            selected_actual_year = int(selected_actual_year)
        
        # Get actual expenses for the selected year (only approved and paid)
        property_actual_expenses = property_obj.act_expense_set.filter(
            act_expense_date__year=selected_actual_year,
            act_expense_approved='Yes',
            act_expense_paid='Yes'
        ).order_by('-act_expense_date')
        
        # Calculate total actual expenses amount
        total_actual_expense_amount = sum(exp.act_expense_amount for exp in property_actual_expenses)
    
    # Map box types to display names
    box_type_display_map = {
        'title-deed': 'Title Deed',
        'property-report': 'Property Report',
        'tenant': 'Tenant Information',
        'actual-expenses': 'Actual Expenses',
        'issues': 'Property Issues',
        'valuation': 'Property Valuation',
        'profit-loss': 'Profit & Loss',
        'revenues': 'Revenues',
        'expenses': 'Budgeted Expenses',
        'open-invoices': 'Open Invoices',
        'lease-renewals': 'Lease Renewals',
        'lease': 'Lease Details',
    }
    
    context = {
        'property': property_obj,
        'active_tenant': active_tenant,
        'open_invoices_data': open_invoices_data,
        'total_invoices_amount': total_invoices_amount,
        'lease_renewal_data': lease_renewal_data,
        'property_issues': property_issues,
        'resolved_count': resolved_count,
        'unresolved_count': unresolved_count,
        'total_issues_count': total_issues_count,
        'property_valuation': property_valuation,
        'property_revenues': property_revenues,
        'total_revenue_amount': total_revenue_amount,
        'property_budgeted_expenses': property_budgeted_expenses,
        'total_budgeted_expense_amount': total_budgeted_expense_amount,
        'property_actual_expenses': property_actual_expenses,
        'actual_expense_years': actual_expense_years,
        'selected_actual_year': selected_actual_year,
        'total_actual_expense_amount': locals().get('total_actual_expense_amount', 0),
        'box_type': box_type,
        'box_type_display': box_type_display_map.get(box_type, box_type.title()),
        'today': timezone.now().date(),
    }
    
    return render(request, 'property_detail.html', context)

@login_required
def dashboard_pl(request, property_id):
    """
    Dedicated view for Profit & Loss dashboard
    """
    property_obj = get_object_or_404(props, prop_id=property_id)

    from django.db.models import Sum, Q
    from collections import defaultdict
    
    # Get selected year from request
    selected_year = request.GET.get('year', 'budget')
    
    # Get available years for this property (from actual expenses only since revenues/expenses are budget data)
    actual_expense_years_obj = set(property_obj.act_expense_set.filter(
        act_expense_approved='Yes',
        act_expense_paid='Yes'
    ).dates('act_expense_date', 'year', order='DESC').distinct())
    
    # Convert to integers and sort
    available_years = sorted([date.year for date in actual_expense_years_obj], reverse=True)

    # Set display name for selected year
    if selected_year == 'budget':
        selected_year_display = 'Budget'
    else:
        try:
            selected_year = int(selected_year)
            selected_year_display = str(selected_year)
        except (ValueError, TypeError):
            selected_year = 'budget'
            selected_year_display = 'Budget'
    
    # Get revenue and expense line types - using correct model names
    revenue_line_types_queryset = revenue_line_types.objects.all().order_by('revenue_line_types_name')
    expense_line_types_queryset = expense_line_types.objects.all().order_by('expense_line_types_name')
    
    # Initialize totals dictionaries
    property_revenue_totals = {}
    property_expense_totals = {}
    
    # Initialize monthly totals
    months = ['jan', 'feb', 'mar', 'apr', 'may', 'jun', 'jul', 'aug', 'sep', 'oct', 'nov', 'dec']
    property_revenue_total = {month: 0 for month in months}
    property_revenue_total['year'] = 0
    property_expense_total = {month: 0 for month in months}
    property_expense_total['year'] = 0
    property_actual_expense_total = {month: 0 for month in months}
    property_actual_expense_total['year'] = 0
    
    # Process revenues for this property (using monthly fields)
    for line_type in revenue_line_types_queryset:
        line_totals = {month: 0 for month in months}
        line_totals['total'] = 0
        
        # Get revenues for this line type and property
        revenues = property_obj.revenue_set.filter(
            revenue_line_types=line_type
        )
        
        # Sum revenues by month using the monthly fields
        for rev in revenues:
            line_totals['jan'] += rev.revenue_jan or 0
            line_totals['feb'] += rev.revenue_feb or 0
            line_totals['mar'] += rev.revenue_mar or 0
            line_totals['apr'] += rev.revenue_apr or 0
            line_totals['may'] += rev.revenue_may or 0
            line_totals['jun'] += rev.revenue_jun or 0
            line_totals['jul'] += rev.revenue_jul or 0
            line_totals['aug'] += rev.revenue_aug or 0
            line_totals['sep'] += rev.revenue_sep or 0
            line_totals['oct'] += rev.revenue_oct or 0
            line_totals['nov'] += rev.revenue_nov or 0
            line_totals['dec'] += rev.revenue_dec or 0
        
        # Calculate total for this line type
        line_totals['total'] = sum(line_totals[month] for month in months)
        
        property_revenue_totals[line_type.revenue_line_types_id] = line_totals
        
        # Add to property totals
        for month in months:
            property_revenue_total[month] += line_totals[month]
        property_revenue_total['year'] += line_totals['total']
    
    # Process budgeted expenses for this property (using monthly fields)
    for line_type in expense_line_types_queryset:
        line_totals = {month: 0 for month in months}
        line_totals['total'] = 0
        
        # Get expenses for this line type and property
        expenses = property_obj.expense_set.filter(
            expense_line_types=line_type
        )
        
        # Sum expenses by month using the monthly fields
        for exp in expenses:
            # Your expense model has monthly fields
            line_totals['jan'] += exp.expense_jan or 0
            line_totals['feb'] += exp.expense_feb or 0
            line_totals['mar'] += exp.expense_mar or 0
            line_totals['apr'] += exp.expense_apr or 0
            line_totals['may'] += exp.expense_may or 0
            line_totals['jun'] += exp.expense_jun or 0
            line_totals['jul'] += exp.expense_jul or 0
            line_totals['aug'] += exp.expense_aug or 0
            line_totals['sep'] += exp.expense_sep or 0
            line_totals['oct'] += exp.expense_oct or 0
            line_totals['nov'] += exp.expense_nov or 0
            line_totals['dec'] += exp.expense_dec or 0
        
        # Calculate total for this line type
        line_totals['total'] = sum(line_totals[month] for month in months)
        
        property_expense_totals[line_type.expense_line_types_id] = line_totals
        
        # Add to property totals
        for month in months:
            property_expense_total[month] += line_totals[month]
        property_expense_total['year'] += line_totals['total']
    
    # Process actual expenses for this property (only if not budget view)
    if selected_year != 'budget':
        actual_expenses = property_obj.act_expense_set.filter(
            act_expense_date__year=selected_year,
            act_expense_approved='Yes',
            act_expense_paid='Yes'
        ).values('act_expense_date', 'act_expense_amount')
        
        # Sum actual expenses by month
        for exp in actual_expenses:
            month_name = months[exp['act_expense_date'].month - 1]
            property_actual_expense_total[month_name] += exp['act_expense_amount']
            property_actual_expense_total['year'] += exp['act_expense_amount']
    
    context = {
        'property': property_obj,
        'available_years': available_years,
        'selected_year': selected_year,
        'selected_year_display': selected_year_display,
        'revenue_line_types': revenue_line_types_queryset,  # Updated variable name
        'expense_line_types': expense_line_types_queryset,  # Updated variable name
        'property_revenue_totals': property_revenue_totals,
        'property_expense_totals': property_expense_totals,
        'property_revenue_total': property_revenue_total,
        'property_expense_total': property_expense_total,
        'property_actual_expense_total': property_actual_expense_total,
        'today': timezone.now().date(),
    }
    
    return render(request, 'dashboard_pl.html', context)

### FINANCE ###
@login_required
def finance(request):
#	return redirect("finance")
	return render (request, "finance.html", {})

@login_required
def finance_revenue(request):
	prop_output = request.POST.get('propname')
	if prop_output is None or prop_output == "All":
			props_data = props.objects.prefetch_related(
				Prefetch(
					'revenue_set',
					queryset=revenue.objects.select_related('revenue_line_types', 'revenue_types')
				)
			).all().order_by('prop_country', 'prop_name')
	else:
		if prop_output is not None:
				props_data = props.objects.prefetch_related(
					Prefetch(
						'revenue_set',
						queryset=revenue.objects.select_related('revenue_line_types', 'revenue_types')
					)
				).all().order_by('prop_country', 'prop_name').filter(prop_name=prop_output)
	return render(request, "finance_revenue.html", {
		"props_data": props_data,
	})

@login_required
def finance_revenue_add(request):
    props_data = props.objects.all().order_by('prop_country', 'prop_name')
    revenue_types_list = revenue_types.objects.all()  # Fetch all revenue types
    revenue_line_types_list = revenue_line_types.objects.all()  # Fetch all revenue line types

    return render(request, "finance_revenue_add.html", {
        "props_data": props_data,
        "revenue_types": revenue_types_list,  # Pass to template
        "revenue_line_types": revenue_line_types_list,  # Pass to template
    })

@login_required
def finance_revenue_commit(request):
    if request.method == "POST":
        # Extract form data
        prop_id = request.POST.get('prop')
        rlt_id = request.POST.get('revenue_line_types')  # revenue_line_types_id
        rt_id = request.POST.get('revenue_types')    # revenue_types_id
        revenue_amount = request.POST.get('revenue_amount')

        # Fetch the revenue_type to check monthly flags
        try:
            revenue_type = revenue_types.objects.get(revenue_types_id=rt_id)
        except revenue_types.DoesNotExist:
            messages.error(request, "Invalid Revenue Type")
            return redirect('finance_revenue_add')
        # Initialize monthly revenue data
        monthly_data = {
            'prop_id': prop_id,
            'revenue_line_types_id': rlt_id,
            'revenue_types_id': rt_id,
            'revenue_amount': revenue_amount,
        }
        # Check each month and set revenue_jan, revenue_feb, etc. if "YES"
        months = ['jan', 'feb', 'mar', 'apr', 'may', 'jun', 
                 'jul', 'aug', 'sep', 'oct', 'nov', 'dec']
        
        for month in months:
            if getattr(revenue_type, f'revenue_types_{month}') == "Yes":
                monthly_data[f'revenue_{month}'] = revenue_amount
                print(monthly_data[f'revenue_{month}'])
        # Create or update the revenue record
        revenue.objects.update_or_create(
            prop_id=prop_id,
            revenue_line_types_id=rlt_id,
            revenue_types_id=rt_id,
            defaults=monthly_data
        )
        messages.success(request, "Revenue Updated Successfully")
        return redirect('finance_revenue')
    # If not a POST request, redirect back
    return redirect('finance_revenue_add')

@login_required
def finance_revenue_edit(request, revenue_id):
    rev = get_object_or_404(revenue, pk=revenue_id)
    props_data = props.objects.all().order_by('prop_country', 'prop_name')
    revenue_types_list = revenue_types.objects.all()
    revenue_line_types_list = revenue_line_types.objects.all()
    form = RevenueForm(instance=rev)  # Use your actual form class
    
    return render(request, "finance_revenue_edit.html", {
        "rev": rev,  # Changed from rresults to rev for clarity
        "props_data": props_data,
        "revenue_types": revenue_types_list,
        "revenue_line_types": revenue_line_types_list,
        "form": form,  # Pass the form to template
    })

@login_required
def finance_revenue_edit_commit(request, revenue_id):
    rev = get_object_or_404(revenue, pk=revenue_id)
    
    if request.method == "POST":
        # Extract form data
        prop_id = request.POST.get('prop')
        rlt_id = request.POST.get('revenue_line_types')
        rt_id = request.POST.get('revenue_types')
        revenue_amount = request.POST.get('revenue_amount')

        # Fetch the revenue_type to check monthly flags
        try:
            revenue_type = revenue_types.objects.get(revenue_types_id=rt_id)
        except revenue_types.DoesNotExist:
            messages.error(request, "Invalid Revenue Type")
            return redirect('finance_revenue_edit', revenue_id=revenue_id)

        # Initialize monthly revenue data
        monthly_data = {
            'prop_id': prop_id,
            'revenue_line_types_id': rlt_id,
            'revenue_types_id': rt_id,
            'revenue_amount': revenue_amount,
        }

        # Check each month and set revenue_jan, revenue_feb, etc. if "YES"
        months = ['jan', 'feb', 'mar', 'apr', 'may', 'jun', 
                 'jul', 'aug', 'sep', 'oct', 'nov', 'dec']
        
        for month in months:
            if getattr(revenue_type, f'revenue_types_{month}') == "Yes":
                monthly_data[f'revenue_{month}'] = revenue_amount
            else:
                monthly_data[f'revenue_{month}'] = None  # Clear if not applicable

        # Update the revenue record
        for key, value in monthly_data.items():
            setattr(rev, key, value)
        rev.save()

        messages.success(request, "Revenue Updated Successfully")
        return redirect('finance_revenue')

    # If GET request, show the form
    props_data = props.objects.all().order_by('prop_country', 'prop_name')
    revenue_types_list = revenue_types.objects.all()
    revenue_line_types_list = revenue_line_types.objects.all()
    form = RevenueForm(instance=rev)
    
    return render(request, "finance_revenue_edit.html", {
        "rev": rev,
        "props_data": props_data,
        "revenue_types": revenue_types_list,
        "revenue_line_types": revenue_line_types_list,
        "form": form,
    })

@login_required
def finance_revenue_types(request):
    rev_types = revenue_types.objects.all()
    return render(request, "finance_revenue_types.html", {
        "rtresults": rev_types,
    })

@login_required
def finance_revenue_types_add(request):
    rev_types = revenue_types.objects.all().order_by('revenue_types_name')
    return render(request, "finance_revenue_types_add.html", {"rtresults":rev_types})

@login_required
def finance_revenue_types_commit(request):
    if request.method == "POST":
        form = RevenueTypesForm(request.POST or None)
        if form.is_valid():
            form.save()
            messages.success(request, "Revenue Types Added Successfully")
    rev_types = revenue_types.objects.all()
    return render(request, "finance_revenue_types.html", {"rtresults":rev_types})

@login_required
def finance_revenue_types_edit(request, revenue_types_id):
    rev_types = revenue_types.objects.filter(pk=revenue_types_id)
    return render(request, "finance_revenue_types_edit.html", {"rtresults":rev_types})

@login_required
def finance_revenue_types_edit_commit(request, revenue_types_id):
    rev = get_object_or_404(revenue_types, pk=revenue_types_id)
    all_types = revenue_types.objects.all().order_by('revenue_types_name')
    if request.method == "POST":
        name = request.POST.get('revenue_types_name')
        # Check for duplicates (case-insensitive, excluding current record)
        if revenue_types.objects.filter(
            revenue_types_name__iexact=name
        ).exclude(
            pk=revenue_types_id
        ).exists():
            messages.error(request, "No duplicate Revenue Types Allowed")
            return render(request, "finance_revenue_types.html", {
                "rtresults": all_types,
                "rev": rev,
                "name_error": True
            })
        form = RevenueTypesForm(request.POST, instance=rev)
        if form.is_valid():
            form.save()
            messages.success(request, "Revenue Type Edited Successfully")
            return redirect('finance_revenue_types')
    # If GET request or form invalid
    return render(request, "finance_revenue_types.html", {
        "rtresults": all_types,
        "rev": rev
    })

@login_required
def finance_revenue_line_types(request):
    rev_line_types = revenue_line_types.objects.all()
    return render(request, "finance_revenue_line_types.html", {
        "rltresults": rev_line_types,
    })

@login_required
def finance_revenue_line_types_add(request):
    rev_line_types = revenue_line_types.objects.all().order_by('revenue_line_types_name')
    return render(request, "finance_revenue_line_types_add.html", {"rltresults":rev_line_types})

@login_required
def finance_revenue_line_types_commit(request):
    if request.method == "POST":
        form = RevenueLineForm(request.POST or None)
        if form.is_valid():
            form.save()
            messages.success(request, "Revenue Line Types Added Successfully")
    rev_line_types = revenue_line_types.objects.all()
    return render(request, "finance_revenue_line_types.html", {"rltresults":rev_line_types})

@login_required
def finance_revenue_line_types_edit(request, revenue_line_types_id):
    rev_line_types = revenue_line_types.objects.filter(pk=revenue_line_types_id)
    return render(request, "finance_revenue_line_types_edit.html", {"rltresults":rev_line_types})

@login_required
def finance_revenue_line_types_edit_commit(request, revenue_line_types_id):
    rev = get_object_or_404(revenue_line_types, pk=revenue_line_types_id)
    all_types = revenue_line_types.objects.all().order_by('revenue_line_types_name')
    if request.method == "POST":
        name = request.POST.get('revenue_line_types_name')
        # Check for duplicates (case-insensitive, excluding current record)
        if revenue_line_types.objects.filter(
            revenue_line_types_name__iexact=name
        ).exclude(
            pk=revenue_line_types_id
        ).exists():
            messages.error(request, "No duplicate Revenue Line Types Allowed")
            return render(request, "finance_revenue_line_types.html", {
                "rltresults": all_types,
                "rev": rev,
                "name_error": True
            })
        form = RevenueLineForm(request.POST, instance=rev)
        if form.is_valid():
            form.save()
            messages.success(request, "Revenue Line Type Edited Successfully")
            return redirect('finance_revenue_line_types')
    
    # If GET request or form invalid
    return render(request, "finance_revenue_line_types.html", {
        "rltresults": all_types,
        "rev": rev
    })

@login_required
def finance_expense(request):
    prop_output = request.POST.get('propname')
    if prop_output is None or prop_output == "All":
            props_data = props.objects.prefetch_related(
                Prefetch(
                    'expense_set',
                    queryset=expense.objects.select_related('expense_line_types', 'expense_types')
                )
            ).all().order_by('prop_country', 'prop_name')
    else:
        if prop_output is not None:
                props_data = props.objects.prefetch_related(
                    Prefetch(
                        'expense_set',
                        queryset=expense.objects.select_related('expense_line_types', 'expense_types')
                    )
                ).all().order_by('prop_country', 'prop_name').filter(prop_name=prop_output)
    return render(request, "finance_expense.html", {
        "props_data": props_data,
    })

@login_required
def finance_expense_add(request):
    # Get properties with their values (using select_related if it's a ForeignKey)
    props_data = props.objects.all().order_by('prop_country', 'prop_name')
    # Annotate each property with its current value (0 if none exists)
    props_data = props_data.annotate(
        current_value=Coalesce(
            Subquery(
                prop_values.objects.filter(prop_id=OuterRef('prop_id'))
                .values('prop_values_current_value')[:1]
            ),
            0
        )
    )
    expense_types_list = expense_types.objects.all()
    expense_line_types_list = expense_line_types.objects.all().order_by('expense_line_types_name')
    return render(request, "finance_expense_add.html", {
        "props_data": props_data,
        "expense_types": expense_types_list,
        "expense_line_types": expense_line_types_list,
    })

@login_required
def finance_expense_commit(request):
    if request.method == "POST":
        # Extract form data
        prop_id = request.POST.get('prop')
        elt_id = request.POST.get('expense_line_types')
        et_id = request.POST.get('expense_types')  # Fix: Changed from 'expense_types' to match your form
        expense_amount = request.POST.get('expense_amount')
        prorata_data = request.POST.get('prorata_calculation_data')

        # Define months list outside the prorata block
        months = ['jan', 'feb', 'mar', 'apr', 'may', 'jun', 
                 'jul', 'aug', 'sep', 'oct', 'nov', 'dec']

        try:
            expense_type = expense_types.objects.get(expense_types_id=et_id)
        except expense_types.DoesNotExist:
            messages.error(request, "Invalid Expense Type")
            return redirect('finance_expense_add')

        # Check if this is a pro-rata expense with multiple properties
        if prorata_data:
            try:
                prorata_data = json.loads(prorata_data)
                selected_properties = prorata_data.get('selected_properties', [])
                
                # Create an expense for each selected property
                for property_data in selected_properties:
                    monthly_data = {
                        'prop_id': property_data['prop_id'],
                        'expense_line_types_id': elt_id,
                        'expense_types_id': et_id,
                        'expense_amount': property_data['calculated_amount'],
                    }
                    
                    for month in months:
                        if getattr(expense_type, f'expense_types_{month}') == "Yes":
                            monthly_data[f'expense_{month}'] = property_data['calculated_amount']
                    
                    expense.objects.update_or_create(
                        prop_id=property_data['prop_id'],
                        expense_line_types_id=elt_id,
                        expense_types_id=et_id,
                        defaults=monthly_data
                    )
                
                messages.success(request, f"{len(selected_properties)} pro-rata expenses created successfully")
                return redirect('finance_expense')
                
            except json.JSONDecodeError:
                messages.error(request, "Invalid pro-rata data")
                return redirect('finance_expense_add')

        # Handle non-pro-rata or single property expense
        monthly_data = {
            'prop_id': prop_id,
            'expense_line_types_id': elt_id,
            'expense_types_id': et_id,
            'expense_amount': expense_amount,
        }
        
        for month in months:
            if getattr(expense_type, f'expense_types_{month}') == "Yes":
                monthly_data[f'expense_{month}'] = expense_amount
        
        expense.objects.update_or_create(
            prop_id=prop_id,
            expense_line_types_id=elt_id,
            expense_types_id=et_id,
            defaults=monthly_data
        )
        
        messages.success(request, "Expense Updated Successfully")
        return redirect('finance_expense')
    
    return redirect('finance_expense_add')

@login_required
def finance_expense_edit(request, expense_id):
    # Get the existing expense
    try:
        existing_expense = expense.objects.get(expense_id=expense_id)
    except expense.DoesNotExist:
        messages.error(request, "Expense not found")
        return redirect('finance_expense')

    # Get properties with their values (using select_related if it's a ForeignKey)
    props_data = props.objects.all().order_by('prop_country', 'prop_name')
    # Annotate each property with its current value (0 if none exists)
    props_data = props_data.annotate(
        current_value=Coalesce(
            Subquery(
                prop_values.objects.filter(prop_id=OuterRef('prop_id'))
                .values('prop_values_current_value')[:1]
            ),
            0
        )
    )
    
    expense_types_list = expense_types.objects.all()
    expense_line_types_list = expense_line_types.objects.all().order_by('expense_line_types_name')
    
    return render(request, "finance_expense_edit.html", {
        "props_data": props_data,
        "expense_types": expense_types_list,
        "expense_line_types": expense_line_types_list,
        "existing_expense": existing_expense,
    })

@login_required
def finance_expense_edit_commit(request, expense_id):
    # Get the existing expense first
    try:
        existing_expense = expense.objects.get(expense_id=expense_id)
    except expense.DoesNotExist:
        messages.error(request, "Expense not found")
        return redirect('finance_expense')

    if request.method == "POST":
        # Extract form data
        prop_id = request.POST.get('prop')
        elt_id = request.POST.get('expense_line_types')
        et_id = request.POST.get('expense_types')
        expense_amount = request.POST.get('expense_amount')
        prorata_data = request.POST.get('prorata_calculation_data')

        # Define months list outside the prorata block
        months = ['jan', 'feb', 'mar', 'apr', 'may', 'jun', 
                 'jul', 'aug', 'sep', 'oct', 'nov', 'dec']

        try:
            expense_type = expense_types.objects.get(expense_types_id=et_id)
        except expense_types.DoesNotExist:
            messages.error(request, "Invalid Expense Type")
            return redirect('finance_expense_edit', expense_id=expense_id)

        # Check if this is a pro-rata expense with multiple properties
        if prorata_data and prorata_data != 'undefined':
            try:
                prorata_data = json.loads(prorata_data)
                selected_properties = prorata_data.get('selected_properties', [])
                
                if not selected_properties:
                    messages.error(request, "No properties selected for pro-rata distribution")
                    return redirect('finance_expense_edit', expense_id=expense_id)
                
                # For pro-rata expenses, we need to handle the original expense differently
                # First, get all existing expenses with the ORIGINAL line type and expense type
                original_expenses = expense.objects.filter(
                    expense_line_types_id=existing_expense.expense_line_types_id,
                    expense_types_id=existing_expense.expense_types_id
                )
                
                # Delete all original pro-rata expenses (they will be recreated)
                original_expenses.delete()
                
                # Create new expenses for each selected property
                for property_data in selected_properties:
                    monthly_data = {
                        'prop_id': property_data['prop_id'],
                        'expense_line_types_id': elt_id,
                        'expense_types_id': et_id,
                        'expense_amount': property_data['calculated_amount'],
                    }
                    
                    for month in months:
                        if getattr(expense_type, f'expense_types_{month}') == "Yes":
                            monthly_data[f'expense_{month}'] = property_data['calculated_amount']
                    
                    # Create new expense
                    expense.objects.create(**monthly_data)
                
                messages.success(request, f"{len(selected_properties)} pro-rata expenses updated successfully")
                return redirect('finance_expense')
                
            except json.JSONDecodeError:
                messages.error(request, "Invalid pro-rata data")
                return redirect('finance_expense_edit', expense_id=expense_id)
            except Exception as e:
                messages.error(request, f"Error processing pro-rata expense: {str(e)}")
                return redirect('finance_expense_edit', expense_id=expense_id)

        # Handle non-pro-rata or single property expense
        # IMPORTANT: Clear all monthly amounts first, then set only the active ones
        monthly_data = {
            'prop_id': prop_id,
            'expense_line_types_id': elt_id,
            'expense_types_id': et_id,
            'expense_amount': expense_amount,
            # Clear all monthly amounts first
            'expense_jan': None,
            'expense_feb': None,
            'expense_mar': None,
            'expense_apr': None,
            'expense_may': None,
            'expense_jun': None,
            'expense_jul': None,
            'expense_aug': None,
            'expense_sep': None,
            'expense_oct': None,
            'expense_nov': None,
            'expense_dec': None,
        }
        
        # Set only the active months based on the NEW expense type
        for month in months:
            if getattr(expense_type, f'expense_types_{month}') == "Yes":
                monthly_data[f'expense_{month}'] = expense_amount
        
        # Update the existing expense directly (don't use update_or_create)
        for field, value in monthly_data.items():
            setattr(existing_expense, field, value)
        existing_expense.save()
        
        messages.success(request, "Expense Updated Successfully")
        return redirect('finance_expense')
    
    return redirect('finance_expense_edit', expense_id=expense_id)

@login_required
def finance_expense_types(request):
    exp_types = expense_types.objects.all()
    return render(request, "finance_expense_types.html", {
        "etresults": exp_types,
    })

@login_required
def finance_expense_types_add(request):
    exp_types = expense_types.objects.all().order_by('expense_types_name')
    return render(request, "finance_expense_types_add.html", {"etresults":exp_types})

@login_required
def finance_expense_types_commit(request):
    if request.method == "POST":
        form = ExpenseTypesForm(request.POST or None)
        if form.is_valid():
            form.save()
            messages.success(request, "Expense Types Added Successfully")
    exp_types = expense_types.objects.all()
    return render(request, "finance_expense_types.html", {"etresults":exp_types})

@login_required
def finance_expense_types_edit(request, expense_types_id):
    exp_types = expense_types.objects.filter(pk=expense_types_id)
    return render(request, "finance_expense_types_edit.html", {"etresults":exp_types})

@login_required
def finance_expense_types_edit_commit(request, expense_types_id):
    exp = get_object_or_404(expense_types, pk=expense_types_id)
    all_types = expense_types.objects.all().order_by('expense_types_name')
    if request.method == "POST":
        name = request.POST.get('expense_types_name')
        # Check for duplicates (case-insensitive, excluding current record)
        if expense_types.objects.filter(
            expense_types_name__iexact=name
        ).exclude(
            pk=expense_types_id
        ).exists():
            messages.error(request, "No duplicate Expense Types Allowed")
            return render(request, "finance_expense_types.html", {
                "etresults": all_types,
                "exp": exp,
                "name_error": True
            })
        form = ExpenseTypesForm(request.POST, instance=exp)
        if form.is_valid():
            form.save()
            messages.success(request, "Expense Type Edited Successfully")
            return redirect('finance_expense_types')
    
    # If GET request or form invalid
    return render(request, "finance_expense_types.html", {
        "etresults": all_types,
        "exp": exp
    })

@login_required
def finance_expense_line_types(request):
    exp_line_types = expense_line_types.objects.all().order_by('expense_line_types_name')
    return render(request, "finance_expense_line_types.html", {
        "eltresults": exp_line_types,
    })

@login_required
def finance_expense_line_types_add(request):
    exp_line_types = expense_line_types.objects.all()
    return render(request, "finance_expense_line_types_add.html", {"eltresults":exp_line_types})

@login_required
def finance_expense_line_types_commit(request):
    if request.method == "POST":
        form = ExpenseLineForm(request.POST or None)
        if form.is_valid():
            form.save()
            messages.success(request, "Expense Line Type Added Successfully")
    exp_line_types = expense_line_types.objects.all()
    return render(request, "finance_expense_line_types.html", {"eltresults":exp_line_types})

@login_required
def finance_expense_line_types_edit(request, expense_line_types_id):
    exp_line_types = expense_line_types.objects.filter(pk=expense_line_types_id)
    return render(request, "finance_expense_line_types_edit.html", {"eltresults":exp_line_types})

@login_required
def finance_expense_line_types_edit_commit(request, expense_line_types_id):
    exp = get_object_or_404(expense_line_types, pk=expense_line_types_id)
    all_types = expense_line_types.objects.all().order_by('expense_line_types_name')
    if request.method == "POST":
        name = request.POST.get('expense_line_types_name')
        # Check for duplicates (case-insensitive, excluding current record)
        if expense_line_types.objects.filter(
            expense_line_types_name__iexact=name
        ).exclude(
            pk=expense_line_types_id
        ).exists():
            messages.error(request, "No duplicate Expense Line Types Allowed")
            return render(request, "finance_expense_line_types.html", {
                "eltresults": all_types,
                "exp": exp,
                "name_error": True
            })
        form = ExpenseLineForm(request.POST, instance=exp)
        if form.is_valid():
            form.save()
            messages.success(request, "Expense Line Type Edited Successfully")
            return redirect('finance_expense_line_types')
    
    # If GET request or form invalid
    return render(request, "finance_expense_line_types.html", {
        "eltresults": all_types,
        "exp": exp
    })

@login_required
def check_expenses_for_line_type(request, expense_line_type_id):
    """
    Check if there are expenses linked to this expense line type
    Returns JSON with expense details if any exist
    """
    if not request.user.is_superuser:
        return JsonResponse({'error': 'Unauthorized'}, status=403)
    
    try:
        # Get the expense line type
        expense_line_type = get_object_or_404(expense_line_types, expense_line_types_id=expense_line_type_id)
        
        # Check for linked expenses using the correct foreign key field name
        linked_expenses = expense.objects.filter(expense_line_types=expense_line_type)
        
        if linked_expenses.exists():
            # Prepare expense data for the frontend
            expenses_data = []
            for exp in linked_expenses:
                # Calculate total amount from all months
                total_amount = 0
                monthly_amounts = []
                
                months = ['jan', 'feb', 'mar', 'apr', 'may', 'jun', 
                         'jul', 'aug', 'sep', 'oct', 'nov', 'dec']
                
                for month in months:
                    month_value = getattr(exp, f'expense_{month}', None)
                    if month_value:
                        total_amount += month_value
                        monthly_amounts.append(f'{month.capitalize()}: {month_value}')
                
                # Use base expense_amount if available, otherwise use calculated total
                display_amount = exp.expense_amount if exp.expense_amount else total_amount
                
                expenses_data.append({
                    'id': exp.expense_id,
                    'expense_type': str(exp.expense_types) if exp.expense_types else 'N/A',
                    'property': str(exp.prop) if exp.prop else 'N/A',
                    'base_amount': str(exp.expense_amount) if exp.expense_amount else '0.00',
                    'total_monthly': str(total_amount),
                    'display_amount': str(display_amount),
                    'monthly_breakdown': monthly_amounts
                })
            
            return JsonResponse({
                'has_expenses': True,
                'expense_count': linked_expenses.count(),
                'expenses': expenses_data
            })
        else:
            return JsonResponse({
                'has_expenses': False,
                'expense_count': 0,
                'expenses': []
            })
            
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@login_required
def delete_expense_line_type(request, expense_line_type_id):
    """
    Delete an expense line type and all its linked expenses
    """
    if not request.user.is_superuser:
        return JsonResponse({'error': 'Unauthorized'}, status=403)
    
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    
    try:
        with transaction.atomic():
            # Get the expense line type
            expense_line_type = get_object_or_404(expense_line_types, expense_line_types_id=expense_line_type_id)
            
            # Get linked expenses before deletion
            linked_expenses = expense.objects.filter(expense_line_types=expense_line_type)
            expense_count = linked_expenses.count()
            
            # Delete all linked expenses first
            linked_expenses.delete()
            
            # Delete the expense line type
            expense_line_type_name = expense_line_type.expense_line_types_name
            expense_line_type.delete()
            
            # Create success message
            if expense_count > 0:
                message = f'Expense line type "{expense_line_type_name}" and {expense_count} linked expense(s) have been deleted successfully.'
            else:
                message = f'Expense line type "{expense_line_type_name}" has been deleted successfully.'
            
            messages.success(request, message)
            
            return JsonResponse({
                'success': True,
                'message': message,
                'deleted_expenses': expense_count
            })
            
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

@login_required
def finance_valuations(request):
    props_list = props.objects.all().order_by('prop_country', 'prop_name')
    valuations = prop_values.objects.all()
    
    # Create a dictionary for easy lookup
    valuations_dict = {v.prop_id: v for v in valuations}
    
    # Calculate totals
    pur_balance = sum(
        v.prop_values_purchase_price 
        for v in valuations 
        if v.prop_values_purchase_price is not None
    )
    cur_balance = sum(
        v.prop_values_current_value 
        for v in valuations 
        if v.prop_values_current_value is not None
    )
    
    return render(request, "finance_valuations.html", {
        "props": props_list,
        "prop_values": valuations_dict,
        "pur_balance": pur_balance,
        "cur_balance": cur_balance
    })

@login_required
def finance_valuations_add(request):
	results = props.objects.all().order_by('prop_country', 'prop_name')
	vresults = prop_values.objects.all()
	context = {
		'props': results,
		'prop_values': vresults,
	}
	return render(request, "finance_valuations_add.html", context)

@login_required
def finance_valuations_commit(request):
    if request.method == "POST":
        prop_id = request.POST.get('prop_id')  # Get property ID from form
        
        # Check if valuation already exists for this property
        if prop_values.objects.filter(prop_id=prop_id).exists():
            messages.error(request, "A valuation already exists for this property. Please edit the existing valuation.")
            return redirect('finance_valuations')
            
        form = ValuesForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Valuation Added Successfully")
            return redirect('finance_valuations')
        else:
            messages.error(request, "Please correct the errors below")
    
    # Rest of your view remains the same...
    results = props.objects.all().order_by('prop_country','prop_name')
    vresults = prop_values.objects.all().order_by('prop_values_purchase_price')    
    
    pur_balance = sum(x.prop_values_purchase_price for x in vresults if x.prop_values_purchase_price is not None)
    cur_balance = sum(x.prop_values_current_value for x in vresults if x.prop_values_current_value is not None)

    context = {
        'pur_balance': pur_balance,
        'cur_balance': cur_balance,        
        'props': results,
        'prop_values': vresults,
    }
    return render(request, "finance_valuations.html", context)

@login_required
def finance_valuations_edit(request, prop_values_id):
	try:
		vresults = prop_values.objects.get(pk=prop_values_id)
	except prop_values.DoesNotExist:
		print(f"ERROR: No prop_values record found for ID {prop_values_id}")
		raise Http404("Valuation not found")
	results = props.objects.all().order_by('prop_country','prop_name')
	return render(request, "finance_valuations_edit.html", {
		"props": results,
		"vresults": vresults
	})

@login_required
def finance_valuations_edit_commit(request, prop_values_id):
    print("Form data received:", request.POST)
    vresult = prop_values.objects.get(pk=prop_values_id)
    
    if request.method == "POST":
        form = ValuesForm(request.POST, instance=vresult)
        if form.is_valid():
            print("Form is valid, saving...")
            form.save()
            print("Saved values:", vresult.prop_values_purchase_price, vresult.prop_values_current_value)
            messages.success(request, "Valuations Edited Successfully")
            return redirect('finance_valuations')  # Redirect instead of render
        else:
        	print("Form errors:", form.errors)
    
    # If GET or invalid form, show the valuations page
    return redirect('finance_valuations')

### TENANTS ###
@login_required
def tenant_page(request):
    # Get filter values from the new form
    selected_property = request.POST.get('propname', '').strip()
    selected_tenant = request.POST.get('tenantname', '').strip()
    selected_status = request.POST.get('act', '').strip()
    
    # Start with all properties and tenants
    all_properties = props.objects.all().order_by('prop_country', 'prop_name')
    all_tenants = tenant.objects.all().order_by('tenant_name')
    
    # Filter tenants based on the selected criteria
    filtered_tenants = all_tenants
    
    # Apply tenant name filter
    if selected_tenant:
        filtered_tenants = filtered_tenants.filter(tenant_name=selected_tenant)
    
    # Apply status filter
    if selected_status:
        filtered_tenants = filtered_tenants.filter(tenant_current=selected_status)
    
    # Filter properties based on the selected property
    filtered_properties = all_properties
    if selected_property:
        filtered_properties = filtered_properties.filter(prop_name=selected_property)
    
    # Pass filter values back to template for form persistence
    context = {
        'tenant': filtered_tenants,
        'props': filtered_properties,
        'selected_property': selected_property,
        'selected_tenant': selected_tenant,
        'selected_status': selected_status,
    }
    
    return render(request, "tenant.html", context)

@login_required
def tenant_add(request):
	results = props.objects.all().order_by('prop_country','prop_name')
	tresults = tenant.objects.all().order_by('tenant_name')
	return render(request, "tenant_add.html", {"props":results, "tenant":tresults})

@login_required
def tenant_edit(request, tenant_id):
	tresults = tenant.objects.filter(pk=tenant_id)
	results = props.objects.all().order_by('prop_country','prop_name')
	return render (request, "tenant_edit.html", {"props":results, "tenant":tresults})

@login_required
def tenant_commit(request):
    props_list = props.objects.all().order_by('prop_country','prop_name')
    
    if request.method == "POST":
        form = TenantForm(request.POST)
        
        if form.is_valid():
            try:
                new_tenant = form.save()
                messages.success(request, f"Tenant {new_tenant.tenant_name} added successfully")
                return redirect('tenant')
            except ValidationError as e:
                # Clean up the error message
                clean_error = str(e).replace('__all__: ', '')
                messages.error(request, clean_error)
                return render(request, "tenant_add.html", {
                    'form': form,
                    'props': props_list,
                    'form_data': request.POST
                })
            except Exception as e:
                messages.error(request, f"Error saving tenant: {str(e)}")
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    clean_error = str(error).replace('__all__: ', '')
                    messages.error(request, clean_error)
    
    return render(request, "tenant_add.html", {
        'form': TenantForm(request.POST if request.method == "POST" else None),
        'props': props_list,
        'form_data': request.POST if request.method == "POST" else None
    })

@login_required
def tenant_edit_commit(request, tenant_id):
	ten = tenant.objects.get(pk=tenant_id)
	if request.method == "POST":
		form = TenantForm(request.POST or None, instance=ten)
		if form.is_valid():
			form.save()
			messages.success(request, "Tenant Edited Successfully")
	results = props.objects.all().order_by('prop_country','prop_name')
	tresults = tenant.objects.all().order_by('tenant_name')
	return render (request, "tenant.html", {"tenant":tresults, "props":results})

@login_required
def tenant_lease_agreement(request):
    tenants = tenant.objects.all().order_by('prop__prop_country', 'prop__prop_name', 'tenant_name')
    
    if request.method == 'POST':
        action = request.POST.get('action')
        tenant_id = request.POST.get('tenant_id')
        
        if not tenant_id:
            messages.error(request, 'No tenant selected')
            return redirect('tenant_lease_agreement')
        
        try:
            tenant_obj = get_object_or_404(tenant, pk=tenant_id)
            
            if action == 'delete':
                if tenant_obj.tenant_lease_agreement:
                    # Delete the file from storage
                    tenant_obj.tenant_lease_agreement.delete()
                    tenant_obj.tenant_lease_agreement_status = "No Lease Agreement"
                    tenant_obj.save()
                    messages.success(request, f'Lease agreement deleted for {tenant_obj.tenant_name}!')
                else:
                    messages.warning(request, 'No lease agreement found to delete.')
                    

            elif action == 'upload':
                if 'lease_agreement' in request.FILES:
                    uploaded_file = request.FILES['lease_agreement']
                    
                    # Validate file size (10MB limit)
                    if uploaded_file.size > 10 * 1024 * 1024:
                        messages.error(request, 'File size exceeds 10MB limit')
                        return redirect('tenant_lease_agreement')
                    
                    # Validate file type
                    allowed_extensions = ['.pdf', '.jpg', '.jpeg', '.png']
                    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
                    
                    if file_extension not in allowed_extensions:
                        messages.error(request, 'Invalid file type. Please upload PDF or image files only.')
                        return redirect('tenant_lease_agreement')
                    
                    try:
                        # Ensure directory exists
                        upload_path = os.path.join(settings.MEDIA_ROOT, 'tenants', 'lease_agreements')
                        os.makedirs(upload_path, exist_ok=True)
                        
                        # Delete old file if exists
                        if tenant_obj.tenant_lease_agreement:
                            tenant_obj.tenant_lease_agreement.delete(save=False)
                        
                        # Save new file
                        tenant_obj.tenant_lease_agreement = uploaded_file
                        tenant_obj.tenant_lease_agreement_status = "Lease Agreement Uploaded"
                        tenant_obj.save()
                        
                        messages.success(request, f'Lease agreement uploaded successfully for {tenant_obj.tenant_name}!')
                    except Exception as e:
                        messages.error(request, f'Error saving file: {str(e)}')
                else:
                    messages.error(request, 'Please select a file to upload')
                    
        except Exception as e:
            messages.error(request, f'Error processing request: {str(e)}')
    
    context = {
        'tenants': tenants,
    }
    return render(request, 'tenant_lease_agreement.html', context)

@login_required
def lease_timeline_view(request):
    # Get all properties with their tenants
    properties = props.objects.filter(prop_status='Active').prefetch_related(
        'tenant_set'
    ).order_by('prop_country', 'prop_name')
    
    # Build timeline data
    timeline_data = []
    
    for prop in properties:
        tenants = prop.tenant_set.all().order_by('tenant_lease_start_date')
        
        for t in tenants:
            if t.tenant_lease_start_date:
                timeline_data.append({
                    'property_id': prop.prop_id,
                    'property_name': prop.prop_name,
                    'tenant_id': t.tenant_id,
                    'tenant_name': t.tenant_name,
                    'start_date': t.tenant_lease_start_date.isoformat(),
                    'end_date': t.tenant_lease_end_date.isoformat() if t.tenant_lease_end_date else None,
                    'is_active': t.tenant_current == 'Yes',
                    'rent': float(t.tenant_rent) if t.tenant_rent else 0,
                })
    
    context = {
        'properties': properties,
        'timeline_data': timeline_data,
    }
    
    return render(request, 'lease_timeline.html', context)

@login_required
def duplicate_tenant_view(request, tenant_id):
    """
    Duplicate an existing tenant to create a renewal or new lease.
    Copies all tenant details except ID and lease dates.
    """
    try:
        # Get the original tenant
        original_tenant = tenant.objects.get(pk=tenant_id)
        
        # Create a new tenant instance with copied data
        new_tenant = tenant()
        
        # Copy all fields EXCEPT primary key and dates
        new_tenant.prop = original_tenant.prop
        new_tenant.tenant_type = original_tenant.tenant_type
        new_tenant.tenant_name = original_tenant.tenant_name
        new_tenant.tenant_contact_person = original_tenant.tenant_contact_person
        new_tenant.tenant_contact_number = original_tenant.tenant_contact_number
        new_tenant.tenant_email = original_tenant.tenant_email
        new_tenant.tenant_deposit = original_tenant.tenant_deposit
        new_tenant.tenant_rental_type = original_tenant.tenant_rental_type
        new_tenant.tenant_renewal = original_tenant.tenant_renewal
        new_tenant.tenant_renewal_period = original_tenant.tenant_renewal_period
        new_tenant.tenant_rent = original_tenant.tenant_rent
        new_tenant.tenant_levies = original_tenant.tenant_levies
        new_tenant.tenant_payment_terms = original_tenant.tenant_payment_terms
        
        # Set lease dates to None - user will fill these in
        new_tenant.tenant_lease_start_date = None
        new_tenant.tenant_lease_end_date = None
        
        # Set as inactive initially - user will activate after setting dates
        new_tenant.tenant_current = 'No'
        
        # Reset renewal status
        new_tenant.tenant_renewal_status = 'pending'
        
        # Don't copy lease agreement - user may need to upload new one
        new_tenant.tenant_lease_agreement = None
        new_tenant.tenant_lease_agreement_status = None
        
        # Save the new tenant
        new_tenant.save()
        
        # Redirect to edit page for the new tenant
        messages.success(
            request, 
            f'Tenant duplicated successfully! Please update the lease dates and set to Active when ready.'
        )
        return redirect('tenant_edit', tenant_id=new_tenant.tenant_id)
        
    except tenant.DoesNotExist:
        messages.error(request, 'Tenant not found.')
        return redirect('tenant')  # ← FIXED
    except Exception as e:
        messages.error(request, f'Error duplicating tenant: {str(e)}')
        return redirect('tenant')  # ← FIXED

@login_required
def delete_tenant_view(request, tenant_id):
    """
    Delete a tenant and automatically recalculate vacancy periods.
    Only superusers can delete tenants.
    """
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to delete tenants.')
        return redirect('tenant')
    
    try:
        tenant_to_delete = tenant.objects.get(pk=tenant_id)
        tenant_name = tenant_to_delete.tenant_name
        property_name = tenant_to_delete.prop.prop_name
        
        # Delete the tenant (signal will handle vacancy cleanup)
        tenant_to_delete.delete()
        
        messages.success(
            request,
            f'Tenant "{tenant_name}" from {property_name} has been deleted. Vacancy periods have been automatically recalculated.'
        )
        return redirect('tenant')
        
    except tenant.DoesNotExist:
        messages.error(request, 'Tenant not found.')
        return redirect('tenant')
    except Exception as e:
        messages.error(request, f'Error deleting tenant: {str(e)}')
        return redirect('tenant')

### SUPPLIERS ###
@login_required
def suppliers(request):
    sup_output = request.POST.get('supname')
    sup_count = request.POST.get('supcount')
    
    # Start with all suppliers
    sresults = supplier.objects.all().order_by('supplier_country','supplier_contact_person')
    
    # Apply search filter if provided and not "All"
    if sup_output and sup_output != "All":
        sresults = sresults.filter(supplier_contact_person__icontains=sup_output)
    
    # Apply country filter if provided and not "All"  
    if sup_count and sup_count != "All":
        sresults = sresults.filter(supplier_country=sup_count)
    
    # Pass the search values back to template for form preservation
    context = {
        "supplier": sresults,
        "selected_supplier": sup_output if sup_output and sup_output != "All" else "",
        "selected_country": sup_count if sup_count and sup_count != "All" else "All"
    }
    
    return render(request, "suppliers.html", context)

@login_required
def suppliers_add(request):
	sresults = supplier.objects.all().order_by('supplier_country','supplier_contact_person')
	return render(request, "suppliers_add.html", {"supplier":sresults})

@login_required
def suppliers_edit(request, supplier_id):
	sresults = supplier.objects.filter(pk=supplier_id)
	return render (request, "suppliers_edit.html", {"supplier":sresults})

@login_required
def suppliers_commit(request):
	if request.method == "POST":
		form = SupplierForm(request.POST or None)
		if form.is_valid():
			form.save()
			messages.success(request, "Supplier Added Successfully")
		else:
			print(form.errors.as_data())
	sresults = supplier.objects.all().order_by('supplier_country','supplier_contact_person')
	return render (request, "suppliers.html", {"supplier":sresults})

@login_required
def suppliers_edit_commit(request, supplier_id):
	sup = supplier.objects.get(pk=supplier_id)
	if request.method == "POST":
		form = SupplierForm(request.POST or None, instance=sup)
		if form.is_valid():
			form.save()
			messages.success(request, "Supplier Edited Successfully")
	sresults = supplier.objects.all().order_by('supplier_country','supplier_contact_person')
	return render (request, "suppliers.html", {"supplier":sresults})

@login_required
def suppliers_delete(request, supplier_id):
    if request.method == 'POST' and request.user.is_superuser:
        try:
            supplier_instance = supplier.objects.get(supplier_id=supplier_id)
            contact_person = supplier_instance.supplier_contact_person
            supplier_instance.delete()
            messages.success(request, f'Supplier "{contact_person}" has been deleted successfully.')
        except supplier.DoesNotExist:
            messages.error(request, 'Supplier not found.')
        except Exception as e:
            messages.error(request, f'Error deleting supplier: {str(e)}')
    else:
        messages.error(request, 'Unauthorized action.')
    
    return redirect('suppliers')

### INVOICES ###
@login_required
def invoices_page(request):
    # Get filter values from POST request
    prop_output = request.POST.get('propname', '')
    tenant_output = request.POST.get('tenantname', '')
    
    # Always get all props for the dropdown
    all_props = props.objects.all().order_by('prop_country', 'prop_name')
    
    # Always get all tenants for the dropdown  
    all_tenants = tenant.objects.all().order_by('tenant_name')
    
    # Get unpaid invoices
    iresults = invoices.objects.filter(invoice_paid="No").order_by('invoice_date')
    
    # Filter props based on selection
    if prop_output and prop_output != "All":
        filtered_props = props.objects.filter(prop_name=prop_output)
    else:
        filtered_props = all_props
    
    # Filter tenants based on selection
    if tenant_output and tenant_output != "All":
        filtered_tenants = tenant.objects.filter(tenant_name=tenant_output)
    else:
        filtered_tenants = all_tenants
    
    context = {
        "invoices": iresults,
        "tenant": filtered_tenants,  # Filtered tenants for display
        "props": filtered_props,     # Filtered props for display
        "all_props": all_props,      # All props for dropdown
        "all_tenants": all_tenants,  # All tenants for dropdown
        "selected_property": prop_output if prop_output != "All" else "",
        "selected_tenant": tenant_output if tenant_output != "All" else "",
    }
    
    return render(request, "invoices.html", context)

@login_required
def invoices_commit(request, invoice_id):
    inv_tbp = invoices.objects.filter(pk=invoice_id).update(invoice_paid="Yes")
    iresults = invoices.objects.get(pk=invoice_id)
    tresults = tenant.objects.get(pk=iresults.tenant_id)
    
    # Get property information - FIXED: Changed 'prop' to 'props'
    presults = props.objects.get(pk=tresults.prop_id)
    
    # Attempt to send the notification email
    if send_invoices_paid_email(tresults, presults, iresults.invoice_date):
        messages.info(request, "Invoice marked as Paid notification email sent.")
    else:
        messages.warning(request, "Invoice marked as Paid, but email could not be sent.")
    return redirect('invoices')

def send_invoices_paid_email(tenant_obj, property_obj, invoice_date):
    """
    Send email notification of an invoice payment for a specific tenant
    """
    from django.db import connection
    from pages.management.commands.email_utils import get_email_recipients, format_email_recipients_for_header
    import logging
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    import os
    
    logger = logging.getLogger(__name__)
    smtp_object = None
    
    try:
        # Get email recipients for invoice paid notifications (returns dict with to/cc/all)
        recipients = get_email_recipients('invoice_paid')
        
        # Get email credentials and settings from environment variables
        email_password = os.environ.get('EMAIL_PASSWORD')
        email_host = os.environ.get('EMAIL_HOST', 'smtp.gmail.com')
        email_port = int(os.environ.get('EMAIL_PORT', 465))
        email_use_ssl = os.environ.get('EMAIL_USE_SSL', 'True').lower() == 'true'
        email_use_tls = os.environ.get('EMAIL_USE_TLS', 'False').lower() == 'true'
        email_user = os.environ.get('EMAIL_USER', 'demetrimanias@gmail.com')
        
        if not email_password:
            logger.error('EMAIL_PASSWORD environment variable not set')
            return False
        
        # Create message
        msg = MIMEMultipart('alternative')
        msg['From'] = email_user
        msg['To'] = format_email_recipients_for_header(recipients['to'])
        if recipients['cc']:
            msg['Cc'] = format_email_recipients_for_header(recipients['cc'])
        msg['Subject'] = f"Invoice Paid - {property_obj.prop_name} - {tenant_obj.tenant_name}"
        
        # Build HTML email body
        html_body = f"""
        <html>
        <head>
        <style>
        p {{ margin: 0; padding: 0; }}
        .info-section {{ margin: 15px 0; }}
        .label {{ font-weight: bold; color: #2c3e50; }}
        .value {{ color: #495057; }}
        .success {{ color: #28a745; font-weight: bold; }}
        </style>
        </head>
        <body>
            <p>Dear User,</p>
            <br>
            <p class="success">✅ INVOICE MARKED AS PAID</p>
            <br>
            <div class="info-section">
                <p><span class="label">Property:</span> <span class="value">{property_obj.prop_name} ({property_obj.prop_country})</span></p>
                <p><span class="label">Tenant:</span> <span class="value">{tenant_obj.tenant_name}</span></p>
                <p><span class="label">Rental Amount:</span> <span class="value">€{tenant_obj.tenant_rent:,.2f}</span></p>
                <p><span class="label">Invoice Date:</span> <span class="value">{invoice_date.strftime('%Y-%m-%d')}</span></p>
            </div>
            <br>
            <p>This invoice has been successfully marked as paid in the Alivente Online System.</p>
            <br>
            <p>You can view all invoice records at <a href="https://alivente.online">alivente.online</a> in the Financial Management section.</p>
            <br>
            <p>Best regards,<br>
            Alivente Property Management System<br>
            Automated Invoice Tracking</p>
        </body>
        </html>
        """
        
        # Create plain text version
        text_body = f"""Dear User,

✅ INVOICE MARKED AS PAID

Property: {property_obj.prop_name} ({property_obj.prop_country})
Tenant: {tenant_obj.tenant_name}
Rental Amount: €{tenant_obj.tenant_rent:,.2f}
Invoice Date: {invoice_date.strftime('%Y-%m-%d')}

This invoice has been successfully marked as paid in the Alivente Online System.

You can view all invoice records at alivente.online in the Financial Management section.

Best regards,
Alivente Property Management System
Automated Invoice Tracking"""
        
        # Attach both HTML and plain text versions
        part1 = MIMEText(text_body, 'plain')
        part2 = MIMEText(html_body, 'html')
        
        msg.attach(part1)
        msg.attach(part2)
        
        # SMTP setup with environment variable configuration
        if email_use_ssl:
            smtp_object = smtplib.SMTP_SSL(email_host, email_port, timeout=10)
        else:
            smtp_object = smtplib.SMTP(email_host, email_port, timeout=10)
            smtp_object.ehlo()
            if email_use_tls:
                smtp_object.starttls()
        
        smtp_object.login(email_user, email_password)
        
        # Send email to all recipients (TO + CC)
        text = msg.as_string()
        smtp_object.sendmail(email_user, recipients['all'], text)
        
        logger.info(f'Invoice paid notification sent for {property_obj.prop_name} - {tenant_obj.tenant_name}')
        return True
        
    except smtplib.SMTPAuthenticationError as e:
        logger.error(f"SMTP Authentication Error: {e}")
        return False
    except smtplib.SMTPException as e:
        logger.error(f"SMTP Error: {e}")
        return False
    except Exception as e:
        logger.error(f"Error sending invoice paid email: {e}")
        return False
    finally:
        if smtp_object:
            try:
                smtp_object.quit()
            except:
                pass
        # Close database connection
        connection.close()

### PROPERTIES ###
@login_required
def properties_page(request):
    # Get filter values from the new form
    search_query = request.POST.get('search', '').strip()
    selected_country = request.POST.get('country', '')
    selected_status = request.POST.get('status', '')
    
    # Start with all properties
    results = props.objects.all()
    
    # Apply cumulative filters (all work together)
    if search_query:
        results = results.filter(prop_name__icontains=search_query)
    
    if selected_country:
        results = results.filter(prop_country=selected_country)
    
    if selected_status:
        results = results.filter(prop_status=selected_status)
    
    # Always order the results
    results = results.order_by('prop_country', 'prop_name')
    
    # Pass filter values back to template for form persistence
    context = {
        'props': results,
        'search_query': search_query,
        'selected_country': selected_country,
        'selected_status': selected_status,
    }

    return render(request, "properties.html", context)

@login_required
def properties_map_view(request):
    """Display all properties on an interactive map"""
    
    # Get all properties from the database
    properties = props.objects.all()
    
    # Convert properties to JSON format for JavaScript
    properties_data = []
    for prop in properties:
        # Handle Decimal fields properly
        latitude = None
        longitude = None
        
        if prop.prop_latitude is not None:
            latitude = float(prop.prop_latitude)
        if prop.prop_longitude is not None:
            longitude = float(prop.prop_longitude)
        
        property_dict = {
            'id': prop.prop_id,
            'name': prop.prop_name,
            'address1': prop.prop_address1,
            'address2': prop.prop_address2,
            'suburb': prop.prop_suburb,
            'city': prop.prop_city,
            'province': prop.prop_province,
            'country': prop.prop_country,
            'pcode': prop.prop_pcode,
            'latitude': latitude,
            'longitude': longitude,
            'floor_area': prop.prop_floor_area,
            'year_built': prop.prop_year_built,
            'status': prop.prop_status,
            'available_for_rent': prop.prop_available_for_rent,
        }
        properties_data.append(property_dict)
    
    context = {
        'properties_json': json.dumps(properties_data),
        'properties_count': len(properties_data)
    }
    
    return render(request, 'map_view.html', context)

@login_required
def properties_title_deed(request):
    properties = props.objects.all().order_by('prop_country','prop_name')
    
    if request.method == 'POST':
        action = request.POST.get('action')
        prop_id = request.POST.get('property_id')  # Changed from 'prop_id' to match form
        
        if not prop_id:
            messages.error(request, 'No property selected')
            return redirect('properties_title_deed')
        
        try:
            property_obj = get_object_or_404(props, pk=prop_id)
            
            if action == 'delete':
                if property_obj.prop_title_deed:
                    # Delete the file from storage
                    property_obj.prop_title_deed.delete()
                    property_obj.prop_title_deed_status = "No Title Deed"
                    property_obj.save()
                    messages.success(request, f'Title deed deleted for {property_obj.prop_name}!')
                else:
                    messages.warning(request, 'No title deed found to delete.')
                    

            elif action == 'upload':
                if 'title_deed' in request.FILES:
                    uploaded_file = request.FILES['title_deed']
                    
                    # Validate file size (10MB limit)
                    if uploaded_file.size > 10 * 1024 * 1024:
                        messages.error(request, 'File size exceeds 10MB limit')
                        return redirect('properties_title_deed')
                    
                    # Validate file type
                    allowed_extensions = ['.pdf', '.jpg', '.jpeg', '.png']
                    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
                    
                    if file_extension not in allowed_extensions:
                        messages.error(request, 'Invalid file type. Please upload PDF or image files only.')
                        return redirect('properties_title_deed')
                    
                    try:
                        # Ensure directory exists
                        upload_path = os.path.join(settings.MEDIA_ROOT, 'properties', 'title_deeds')
                        os.makedirs(upload_path, exist_ok=True)
                        
                        # Delete old file if exists
                        if property_obj.prop_title_deed:
                            property_obj.prop_title_deed.delete(save=False)
                        
                        # Save new file
                        property_obj.prop_title_deed = uploaded_file
                        property_obj.prop_title_deed_status = "Title Deed Uploaded"
                        property_obj.save()
                        
                        messages.success(request, f'Title deed uploaded successfully for {property_obj.prop_name}!')
                    except Exception as e:
                        messages.error(request, f'Error saving file: {str(e)}')
                else:
                    messages.error(request, 'Please select a file to upload')
                    
        except Exception as e:
            messages.error(request, f'Error processing request: {str(e)}')
    
    context = {
        'properties': properties,
    }
    return render(request, 'properties_title_deed.html', context)

@login_required
def properties_add(request):
	results = props.objects.all().order_by('prop_country','prop_name')
	existing_names = list(props.objects.values_list('prop_name', flat=True))
	return render(request, "properties_add.html", {"props":results, "existing_names": existing_names})

@login_required
def properties_commit(request):
	if request.method == "POST":
		form = PropForm(request.POST or None)
		if form.is_valid():
			form.save()
	results = props.objects.all().order_by('prop_country','prop_name')
	messages.success(request, "Property Added Successfully")
	return render (request, "properties.html", {"props":results})

@login_required
def properties_edit(request, prop_id):
    # Get the current property being edited
    current_property = get_object_or_404(props, pk=prop_id)
    
    # Get all other property names (excluding the current one)
    existing_names = props.objects.exclude(prop_id=prop_id).values_list('prop_name', flat=True)
    
    return render(request, "properties_edit.html", {
        "props": [current_property],  # Maintain your existing structure
        "existing_names": list(existing_names)  # Add this for client-side validation
    })

@login_required
def properties_edit_commit(request, prop_id):
    prop = get_object_or_404(props, pk=prop_id)
    existing_names = props.objects.exclude(prop_id=prop_id).values_list('prop_name', flat=True)
    
    if request.method == "POST":
        form = PropForm(request.POST, instance=prop)
        
        if form.is_valid():
            new_name = form.cleaned_data.get('prop_name')
            current_name = prop.prop_name
            
            if new_name.lower() != current_name.lower():
                if props.objects.exclude(prop_id=prop_id).filter(prop_name__iexact=new_name).exists():
                    messages.error(request, "A property with this name already exists.")
                    return render(request, "properties_edit.html", {
                        'props': [prop],
                        'existing_names': list(existing_names)
                    })
            
            form.save()
            messages.success(request, "Property Edited Successfully")
            results = props.objects.all().order_by('prop_country','prop_name')
            return redirect('properties')  # Better to redirect after POST
        
        # Form is invalid
        messages.error(request, "Please correct the errors below.")
        return render(request, "properties_edit.html", {
            'props': [prop],
            'existing_names': list(existing_names)
        })
    
    # If not POST, redirect to properties page
    return redirect('properties')


### ACTUAL EXPENSES ###
@login_required
def act_expense_manage_document(request):
    """
    Handle document upload, replacement, and deletion within the main expense page
    """
    if request.method == 'POST':
        action = request.POST.get('action')
        document_action = request.POST.get('document_action')  # Get the document action type
        expense_id = request.POST.get('expense_id')
        
        if not expense_id:
            messages.error(request, 'No expense selected')
            return redirect('act_expense_all')
        
        try:
            expense = get_object_or_404(act_expense, pk=expense_id)
            
            if action == 'delete_document':
                # Handle document deletion only (not the entire expense)
                if expense.act_expense_document:
                    # Delete the physical file
                    if expense.act_expense_document.storage.exists(expense.act_expense_document.name):
                        expense.act_expense_document.delete(save=False)
                    
                    # Clear the database field
                    expense.act_expense_document = None
                    expense.save()
                    
                    messages.success(request, f'Invoice document deleted successfully for expense on {expense.act_expense_date}!')
                else:
                    messages.warning(request, 'No document found to delete.')
                    
            elif action == 'upload':
                # Handle file upload/replacement
                if 'act_expense_document' in request.FILES:
                    uploaded_file = request.FILES['act_expense_document']
                    
                    # Validate file size (5MB limit)
                    if uploaded_file.size > 5 * 1024 * 1024:
                        messages.error(request, 'File size exceeds 5MB limit')
                        return redirect('act_expense_all')
                    
                    # Validate file type
                    allowed_extensions = ['.pdf', '.jpg', '.jpeg', '.png', '.xlsx', '.xls', '.doc', '.docx']
                    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
                    
                    if file_extension not in allowed_extensions:
                        messages.error(request, 'Invalid file type. Please upload PDF, JPG, PNG, Excel, or Word files only.')
                        return redirect('act_expense_all')
                    
                    # Check if we're adding to existing or replacing
                    if document_action == 'add_to_existing' and expense.act_expense_document:
                        # For merge, existing file must be PDF
                        if not is_pdf(expense.act_expense_document):
                            messages.error(request, 'Cannot merge: Existing document is not a PDF. Please use Replace instead.')
                            return redirect('act_expense_all')
                        
                        # Convert uploaded file to PDF first if necessary
                        try:
                            pdf_content, pdf_filename = convert_to_pdf(uploaded_file)
                            
                            # Merge the PDFs (pdf_content is already a ContentFile)
                            merged_pdf = merge_pdfs(expense.act_expense_document, pdf_content)
                            
                            # Generate a new filename
                            original_name = os.path.splitext(os.path.basename(expense.act_expense_document.name))[0]
                            new_filename = f"{original_name}_merged.pdf"
                            
                            # Delete the old file
                            if expense.act_expense_document.storage.exists(expense.act_expense_document.name):
                                expense.act_expense_document.delete(save=False)
                            
                            # Save the merged PDF
                            expense.act_expense_document.save(new_filename, merged_pdf, save=True)
                            
                            messages.success(request, f'Documents merged successfully for expense on {expense.act_expense_date}!')
                        except ValueError as e:
                            messages.error(request, f'Error: {str(e)}')
                            return redirect('act_expense_all')
                        except Exception as e:
                            messages.error(request, f'Error merging documents: {str(e)}')
                            return redirect('act_expense_all')
                    else:
                        # Regular upload/replace with automatic PDF conversion
                        # Delete existing file if present
                        if expense.act_expense_document:
                            if expense.act_expense_document.storage.exists(expense.act_expense_document.name):
                                expense.act_expense_document.delete(save=False)
                        
                        # Convert to PDF if necessary
                        try:
                            pdf_content, pdf_filename = convert_to_pdf(uploaded_file)
                            expense.act_expense_document.save(pdf_filename, pdf_content, save=True)
                            
                            # Show different message if conversion happened
                            if file_extension != '.pdf':
                                messages.success(request, f'Document uploaded and converted to PDF successfully for expense on {expense.act_expense_date}!')
                            else:
                                messages.success(request, f'Document uploaded successfully for expense on {expense.act_expense_date}!')
                        except Exception as e:
                            messages.error(request, f'Error processing document: {str(e)}')
                            return redirect('act_expense_all')
                else:
                    messages.error(request, 'Please select a file to upload')
                    
        except Exception as e:
            messages.error(request, f'Error processing request: {str(e)}')
    
    return redirect('act_expense_all')

from datetime import datetime

@login_required
def act_expense_all(request):
    # Get filter parameters from request
    search_query = request.GET.get('search', '').strip()
    property_filter = request.GET.get('property', '').strip()
    status_filter = request.GET.get('status', '').strip()
    from_date = request.GET.get('from_date', '').strip()
    to_date = request.GET.get('to_date', '').strip()

    # Base queryset - all expenses, ordered by date (most recent first)
    expenses = act_expense.objects.select_related('prop').order_by('-act_expense_date')
    
    # Apply filters one by one
    
    # 1. Search filter - search in description
    if search_query:
        expenses = expenses.filter(
            act_expense_description__icontains=search_query
        )
    
    # 2. Property filter
    if property_filter:
        try:
            property_id = int(property_filter)
            expenses = expenses.filter(prop_id=property_id)
        except (ValueError, TypeError):
            pass
    
    # 3. Status filter
    if status_filter:
        if status_filter == 'require_approval':
            expenses = expenses.filter(act_expense_approved='No', act_expense_paid='No')
        elif status_filter == 'approved_not_paid':
            expenses = expenses.filter(act_expense_approved='Yes', act_expense_paid='No')
        elif status_filter == 'approved_and_paid':
            expenses = expenses.filter(act_expense_approved='Yes', act_expense_paid='Yes')
    
    # 4. Date range filtering
    if from_date:
        try:
            # Ensure proper date format
            parsed_from_date = datetime.strptime(from_date, '%Y-%m-%d').date()
            expenses = expenses.filter(act_expense_date__gte=parsed_from_date)
        except ValueError:
            pass
    
    if to_date:
        try:
            # Ensure proper date format
            parsed_to_date = datetime.strptime(to_date, '%Y-%m-%d').date()
            expenses = expenses.filter(act_expense_date__lte=parsed_to_date)
        except ValueError:
            pass
    
    # Get properties for filter dropdown
    properties = props.objects.filter(prop_status="Active").order_by('prop_country', 'prop_name')
    
    # Determine navigation context
    came_from = request.GET.get('from', None)
    from_finance_pl_act = request.GET.get('from_finance_pl_act', False)
    
    # Convert string 'True'/'False' to boolean if needed
    if isinstance(from_finance_pl_act, str):
        from_finance_pl_act = from_finance_pl_act.lower() == 'true'

    return render(request, 'act_expense.html', {
        'expenses': expenses,
        'props': properties,
        'current_year': datetime.now().year,
        'from_finance_pl_act': from_finance_pl_act,
        'came_from': came_from,
        # Pass filter values back to template to maintain state
        'search_query': search_query,
        'selected_property': property_filter,
        'selected_status': status_filter,
        'selected_from_date': from_date,
        'selected_to_date': to_date,
    })

@login_required
def act_expense_upload_inv(request):
    # Get all expenses to display in the table
    expenses = act_expense.objects.all().order_by('-act_expense_date')
    
    if request.method == 'POST':
        action = request.POST.get('action')
        document_action = request.POST.get('document_action')  # Get the document action type
        expense_id = request.POST.get('expense_id')
        
        if not expense_id:
            messages.error(request, 'No expense selected')
            return redirect('act_expense_upload_inv')
        
        try:
            expense = get_object_or_404(act_expense, pk=expense_id)
            
            if action == 'delete':
                # Handle file deletion
                if expense.act_expense_document:
                    # Delete the physical file
                    if expense.act_expense_document.storage.exists(expense.act_expense_document.name):
                        expense.act_expense_document.delete(save=False)
                    
                    # Clear the database field
                    expense.act_expense_document = None
                    expense.save()
                    
                    messages.success(request, f'Document deleted successfully for expense on {expense.act_expense_date}!')
                else:
                    messages.warning(request, 'No document found to delete.')
                    
            elif action == 'upload':
                # Handle file upload
                if 'act_expense_document' in request.FILES:
                    uploaded_file = request.FILES['act_expense_document']
                    
                    # Validate file size (5MB limit)
                    if uploaded_file.size > 5 * 1024 * 1024:
                        messages.error(request, 'File size exceeds 5MB limit')
                        return redirect('act_expense_upload_inv')
                    
                    # Validate file type
                    allowed_extensions = ['.pdf', '.jpg', '.jpeg', '.png', '.xlsx', '.xls', '.doc', '.docx']
                    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
                    
                    if file_extension not in allowed_extensions:
                        messages.error(request, 'Invalid file type. Please upload PDF, JPG, PNG, Excel, or Word files only.')
                        return redirect('act_expense_upload_inv')
                    
                    # Check if we're adding to existing or replacing
                    if document_action == 'add_to_existing' and expense.act_expense_document:
                        # For merge, existing file must be PDF
                        if not is_pdf(expense.act_expense_document):
                            messages.error(request, 'Cannot merge: Existing document is not a PDF. Please use Replace instead.')
                            return redirect('act_expense_upload_inv')
                        
                        # Convert uploaded file to PDF first if necessary
                        try:
                            pdf_content, pdf_filename = convert_to_pdf(uploaded_file)
                            
                            # Merge the PDFs (pdf_content is already a ContentFile)
                            merged_pdf = merge_pdfs(expense.act_expense_document, pdf_content)
                            
                            # Generate a new filename
                            original_name = os.path.splitext(os.path.basename(expense.act_expense_document.name))[0]
                            new_filename = f"{original_name}_merged.pdf"
                            
                            # Delete the old file
                            if expense.act_expense_document.storage.exists(expense.act_expense_document.name):
                                expense.act_expense_document.delete(save=False)
                            
                            # Save the merged PDF
                            expense.act_expense_document.save(new_filename, merged_pdf, save=True)
                            
                            messages.success(request, f'Documents merged successfully for expense on {expense.act_expense_date}!')
                        except ValueError as e:
                            messages.error(request, f'Error: {str(e)}')
                            return redirect('act_expense_upload_inv')
                        except Exception as e:
                            messages.error(request, f'Error merging documents: {str(e)}')
                            return redirect('act_expense_upload_inv')
                    else:
                        # Regular upload/replace with automatic PDF conversion
                        # Delete existing file if present
                        if expense.act_expense_document:
                            if expense.act_expense_document.storage.exists(expense.act_expense_document.name):
                                expense.act_expense_document.delete(save=False)
                        
                        # Convert to PDF if necessary
                        try:
                            pdf_content, pdf_filename = convert_to_pdf(uploaded_file)
                            expense.act_expense_document.save(pdf_filename, pdf_content, save=True)
                            
                            # Show different message if conversion happened
                            if file_extension != '.pdf':
                                messages.success(request, f'Document uploaded and converted to PDF successfully for expense on {expense.act_expense_date}!')
                            else:
                                messages.success(request, f'Document uploaded successfully for expense on {expense.act_expense_date}!')
                        except Exception as e:
                            messages.error(request, f'Error processing document: {str(e)}')
                            return redirect('act_expense_upload_inv')
                else:
                    messages.error(request, 'Please select a file to upload')
                    
        except Exception as e:
            messages.error(request, f'Error processing request: {str(e)}')
    
    context = {
        'expenses': expenses,
    }
    return render(request, 'act_expense_upload_inv.html', context)

@login_required
def act_expense_view(request):
    # Get year/month from request or use current year as default
    selected_year = request.GET.get('year', datetime.now().year)
    selected_month = request.GET.get('month')
    from_finance_pl_act = request.GET.get('from_finance_pl_act', False)
    property_id = request.GET.get('property_id')
    properties = request.GET.get('properties', '')  # NEW: Handle comma-separated properties
    
    # Base queryset - only approved and paid expenses, ordered by date
    expenses = act_expense.objects.select_related('prop').filter(
        act_expense_approved="Yes",
        act_expense_paid="Yes"
    ).order_by('-act_expense_date')
    
    # Filter by property - UPDATED LOGIC
    if properties:  # NEW: Handle comma-separated properties
        try:
            property_ids = [int(prop_id.strip()) for prop_id in properties.split(',') if prop_id.strip()]
            expenses = expenses.filter(prop_id__in=property_ids)
        except ValueError:
            pass  # Invalid property IDs, skip filtering
    elif property_id:  # Keep existing single property logic for backward compatibility
        try:
            expenses = expenses.filter(prop_id=int(property_id))
        except (ValueError, TypeError):
            pass  # Skip if property_id is invalid
    
    # Handle YEAR/MONTH filtering (convert to int safely)
    try:
        year = int(request.GET.get('year', 0)) if request.GET.get('year') else None
        month = int(request.GET.get('month', 0)) if request.GET.get('month') else None
    except (ValueError, TypeError):
        year, month = None, None  # Fallback if invalid input
    
    if year:
        expenses = expenses.filter(act_expense_date__year=year)
        if month:
            expenses = expenses.filter(act_expense_date__month=month)
    
    # Handle DATE RANGE filtering
    from_date = request.GET.get('from_date')
    to_date = request.GET.get('to_date')
    
    if from_date and to_date:
        expenses = expenses.filter(
            act_expense_date__gte=from_date,
            act_expense_date__lte=to_date
        )
    
    # Get available years for filter dropdown
    available_years = act_expense.objects.filter(
        act_expense_approved="Yes",
        act_expense_paid="Yes"
    ).dates('act_expense_date', 'year').order_by('-act_expense_date')
    
    return render(request, 'act_expense.html', {
        'expenses': expenses,
        'selected_year': year if year else int(selected_year),
        'selected_month': month,
        'current_year': datetime.now().year,
        'available_years': [y.year for y in available_years],
        'from_finance_pl_act': from_finance_pl_act,
        'selected_property_id': property_id
    })

@login_required
def act_expense_edit(request, expense_id):
    # Get the current expense being edited
    current_expense = get_object_or_404(act_expense, pk=expense_id)
    
    # Get property details from props table
    results = props.objects.filter(prop_status="Active").order_by('prop_country','prop_name')

    return render(request, "act_expense_edit.html", {
        "props": results,
        "current_expense": current_expense,
    })

@login_required
def act_expense_edit_commit(request, expense_id):
    if request.method == 'POST':
        try:
            expense = act_expense.objects.get(act_expense_id=expense_id)
            
            # Update expense fields
            expense.act_expense_date = request.POST.get('act_expense_date')
            expense.prop_id = request.POST.get('prop')
            expense.act_expense_description = request.POST.get('act_expense_description')
            expense.act_expense_amount = request.POST.get('act_expense_amount')
            
            if request.user.is_superuser:
                expense.act_expense_approved = request.POST.get('act_expense_approved')
                
                # Handle the paid field - check for hidden field if main field is missing
                paid_value = request.POST.get('act_expense_paid')
                if not paid_value:  # If main field is empty (disabled)
                    paid_value = request.POST.get('act_expense_paid_hidden')
                
                expense.act_expense_paid = paid_value
            
            expense.save()
            
            messages.success(request, 'Expense updated successfully!')
            
        except act_expense.DoesNotExist:
            messages.error(request, 'Expense not found.')
        except Exception as e:
            messages.error(request, f'Error updating expense: {str(e)}')
    
    return redirect('act_expense_all')

@login_required
def get_expense_invoice(request, expense_id):
    try:
        # Adjust this query based on your Expense model
        # expense_id might be the date or actual expense ID
        expense = YourExpenseModel.objects.filter(
            # Add your filter logic here - could be by date, ID, etc.
            date=expense_id  # or id=expense_id
        ).first()
        
        if expense and expense.invoice_file:  # Adjust field name
            response = HttpResponse(
                expense.invoice_file.read(), 
                content_type='application/pdf'  # or detect content type
            )
            response['Content-Disposition'] = f'inline; filename="invoice_{expense_id}.pdf"'
            return response
        else:
            raise Http404("Invoice not found")
            
    except Exception as e:
        raise Http404("Invoice not found")

@login_required
def mark_approved(request, expense_id):
    expense = get_object_or_404(act_expense, pk=expense_id)
    if expense.act_expense_approved != 'Yes':  # Only update if not already approved
        expense.act_expense_approved = 'Yes'
        expense.save()
        # Attempt to send the notification email with enhanced details
        from datetime import date
        if send_expense_approved_email(
            expense.act_expense_date, 
            expense.prop.prop_name,  # Access through the foreign key relationship
            expense.act_expense_description, 
            expense.act_expense_amount,
            date.today()
        ):
            messages.info(request, "Expense approved and notification email sent.")
        else:
            messages.warning(request, "Expense approved, but email could not be sent.")
    return redirect('act_expense_all')

@login_required
def mark_paid(request, expense_id):
    expense = get_object_or_404(act_expense, pk=expense_id)
    if expense.act_expense_paid != 'Yes':  # Only update if not already paid
        expense.act_expense_paid = 'Yes'
        expense.save()
        # Attempt to send the notification email with enhanced details
        from datetime import date
        if send_expense_paid_email(
            expense.act_expense_date,
            expense.prop.prop_name,  # Access through the foreign key relationship
            expense.act_expense_description,
            expense.act_expense_amount,
            date.today()
        ):
            messages.info(request, "Expense marked as paid and notification email sent.")
        else:
            messages.warning(request, "Expense marked as paid, but email could not be sent.")
    return redirect('act_expense_all')

@login_required
def mark_deleted(request, expense_id):
    try:
        expense = get_object_or_404(act_expense, pk=expense_id)
        expense.delete()  # Permanently deletes the record
        messages.success(request, "Expense deleted successfully")
    except Exception as e:
        messages.error(request, f"Error deleting expense: {str(e)}")
    return redirect('act_expense_all')

@login_required
def act_expense_add(request):
    results = props.objects.filter(prop_status="Active").order_by('prop_country','prop_name')
    return render(request, "act_expense_add.html", {'props': results})

def send_expense_approved_email(expense_date, property_name, description, amount, approved_date):
    """
    Send email notification of an expense approval for a specific expense
    """
    from django.db import connection
    from pages.management.commands.email_utils import get_email_recipients, format_email_recipients_for_header
    import logging
    
    logger = logging.getLogger(__name__)
    smtp_object = None
    
    try:
        # Get recipients with TO/CC split
        recipients = get_email_recipients('expense_approved')
        
        # Create message
        msg = MIMEMultipart()
        msg['From'] = "demetrimanias@gmail.com"
        msg['To'] = format_email_recipients_for_header(recipients['to'])
        if recipients['cc']:
            msg['Cc'] = format_email_recipients_for_header(recipients['cc'])
        msg['Subject'] = f"Expense Approved - €{amount} for {property_name}"
        
        # Email body with proper formatting
        body = f"""Dear User,

An expense has been APPROVED. The details are as follows:

- Expense Date: {expense_date.strftime('%d/%m/%Y')}
- Property: {property_name}
- Description: {description}
- Amount: €{amount}
- Approved Date: {approved_date.strftime('%d/%m/%Y')}
- Status: Approved (Pending Payment)

You can view this expense in the Alivente Property Management System.

Thanks,

Alivente Property Management System"""
        
        msg.attach(MIMEText(body, 'plain'))
        
        # Get email credentials and settings from environment variables
        email_password = os.environ.get('EMAIL_PASSWORD')
        email_host = os.environ.get('EMAIL_HOST', 'smtp.gmail.com')
        email_port = int(os.environ.get('EMAIL_PORT', 465))
        email_use_ssl = os.environ.get('EMAIL_USE_SSL', 'True').lower() == 'true'
        email_use_tls = os.environ.get('EMAIL_USE_TLS', 'False').lower() == 'true'
        
        if not email_password:
            logger.error('EMAIL_PASSWORD environment variable not set')
            return False
        
        # SMTP setup with environment variable configuration
        if email_use_ssl:
            smtp_object = smtplib.SMTP_SSL(email_host, email_port, timeout=10)
        else:
            smtp_object = smtplib.SMTP(email_host, email_port, timeout=10)
            smtp_object.ehlo()
            if email_use_tls:
                smtp_object.starttls()
        
        email = "demetrimanias@gmail.com"
        smtp_object.login(email, email_password)
        
        # Send email
        text = msg.as_string()
        smtp_object.sendmail(email, recipients['all'], text)
        return True
        
    except smtplib.SMTPAuthenticationError as e:
        logger.error(f"SMTP Authentication Error: {e}")
        return False
    except smtplib.SMTPException as e:
        logger.error(f"SMTP Error: {e}")
        return False
    except Exception as e:
        logger.error(f"Error sending email: {e}")
        return False
    finally:
        if smtp_object:
            try:
                smtp_object.quit()
            except:
                pass
        # Close database connection
        connection.close()

def send_expense_paid_email(expense_date, property_name, description, amount, paid_date):
    """
    Send email notification of an expense payment for a specific expense
    """
    from django.db import connection
    from pages.management.commands.email_utils import get_email_recipients, format_email_recipients_for_header
    import logging
    
    logger = logging.getLogger(__name__)
    smtp_object = None
    
    try:
        # Get recipients with TO/CC split
        recipients = get_email_recipients('expense_paid')
        
        # Create message
        msg = MIMEMultipart()
        msg['From'] = "demetrimanias@gmail.com"
        msg['To'] = format_email_recipients_for_header(recipients['to'])
        if recipients['cc']:
            msg['Cc'] = format_email_recipients_for_header(recipients['cc'])
        msg['Subject'] = f"Expense Paid - €{amount} for {property_name}"
        
        # Email body with proper formatting
        body = f"""Dear User,

An expense has been PAID. The details are as follows:

- Expense Date: {expense_date.strftime('%d/%m/%Y')}
- Property: {property_name}
- Description: {description}
- Amount: €{amount}
- Paid Date: {paid_date.strftime('%d/%m/%Y')}
- Status: Fully Processed

This expense has been completed and processed.

Thanks,

Alivente Property Management System"""
        
        msg.attach(MIMEText(body, 'plain'))
        
        # Get email credentials and settings from environment variables
        email_password = os.environ.get('EMAIL_PASSWORD')
        email_host = os.environ.get('EMAIL_HOST', 'smtp.gmail.com')
        email_port = int(os.environ.get('EMAIL_PORT', 465))
        email_use_ssl = os.environ.get('EMAIL_USE_SSL', 'True').lower() == 'true'
        email_use_tls = os.environ.get('EMAIL_USE_TLS', 'False').lower() == 'true'
        
        if not email_password:
            logger.error('EMAIL_PASSWORD environment variable not set')
            return False
        
        # SMTP setup with environment variable configuration
        if email_use_ssl:
            smtp_object = smtplib.SMTP_SSL(email_host, email_port, timeout=10)
        else:
            smtp_object = smtplib.SMTP(email_host, email_port, timeout=10)
            smtp_object.ehlo()
            if email_use_tls:
                smtp_object.starttls()
        
        email = "demetrimanias@gmail.com"
        smtp_object.login(email, email_password)
        
        # Send email
        text = msg.as_string()
        smtp_object.sendmail(email, recipients['all'], text)
        return True
        
    except smtplib.SMTPAuthenticationError as e:
        logger.error(f"SMTP Authentication Error: {e}")
        return False
    except smtplib.SMTPException as e:
        logger.error(f"SMTP Error: {e}")
        return False
    except Exception as e:
        logger.error(f"Error sending email: {e}")
        return False
    finally:
        if smtp_object:
            try:
                smtp_object.quit()
            except:
                pass
        # Close database connection
        connection.close()

def send_expense_approval_email_with_link(expense_date, property_name, description, amount, created_date):
    """
    Send email notification for expense approval with enhanced details
    """
    from django.db import connection
    from pages.management.commands.email_utils import get_email_recipients, format_email_recipients_for_header
    import logging
    
    logger = logging.getLogger(__name__)
    smtp_object = None
    
    try:
        # Get recipients with TO/CC split
        recipients = get_email_recipients('expense_needs_approval')
        
        # Create message
        msg = MIMEMultipart()
        msg['From'] = "demetrimanias@gmail.com"
        msg['To'] = format_email_recipients_for_header(recipients['to'])
        if recipients['cc']:
            msg['Cc'] = format_email_recipients_for_header(recipients['cc'])
        msg['Subject'] = f"New Expense Requires Approval - €{amount} for {property_name}"
        
        # Email body with proper formatting
        body = f"""Dear User,

A new Actual Expense has been created that requires your approval. The details are as follows:

- Expense Date: {expense_date.strftime('%d/%m/%Y')}
- Property: {property_name}
- Description: {description}
- Amount: €{amount}
- Created Date: {created_date.strftime('%d/%m/%Y')}
- Status: Pending Approval

You can view this expense in the Alivente Property Management System.

Thanks,

Alivente Property Management System"""
        
        msg.attach(MIMEText(body, 'plain'))
        
        # Get email credentials and settings from environment variables
        email_password = os.environ.get('EMAIL_PASSWORD')
        email_host = os.environ.get('EMAIL_HOST', 'smtp.gmail.com')
        email_port = int(os.environ.get('EMAIL_PORT', 465))
        email_use_ssl = os.environ.get('EMAIL_USE_SSL', 'True').lower() == 'true'
        email_use_tls = os.environ.get('EMAIL_USE_TLS', 'False').lower() == 'true'
        
        if not email_password:
            logger.error('EMAIL_PASSWORD environment variable not set')
            return False
        
        # SMTP setup with environment variable configuration
        if email_use_ssl:
            smtp_object = smtplib.SMTP_SSL(email_host, email_port, timeout=10)
        else:
            smtp_object = smtplib.SMTP(email_host, email_port, timeout=10)
            smtp_object.ehlo()
            if email_use_tls:
                smtp_object.starttls()
        
        email = "demetrimanias@gmail.com"
        smtp_object.login(email, email_password)
        
        # Send email
        text = msg.as_string()
        smtp_object.sendmail(email, recipients['all'], text)
        return True
        
    except smtplib.SMTPAuthenticationError as e:
        logger.error(f"SMTP Authentication Error: {e}")
        return False
    except smtplib.SMTPException as e:
        logger.error(f"SMTP Error: {e}")
        return False
    except Exception as e:
        logger.error(f"Error sending email: {e}")
        return False
    finally:
        if smtp_object:
            try:
                smtp_object.quit()
            except:
                pass
        # Close database connection
        connection.close()

@login_required
def act_expense_commit(request):
    if request.method == 'POST':
        try:
            # Get data from the form
            expense_date = request.POST.get('act_expense_date')
            expense_prop = request.POST.get('prop')
            expense_description = request.POST.get('act_expense_description')
            expense_amount = request.POST.get('act_expense_amount')
            expense_approved = request.POST.get('act_expense_approved', 'No')
            expense_paid = request.POST.get('act_expense_paid', 'No')
            
            # Validate required fields
            if not expense_date or not expense_description or not expense_amount or not expense_prop:
                messages.error(request, 'All fields are required.')
                return redirect('act_expense_add')
            
            # Create and save the expense record
            expense = act_expense(
                act_expense_date=expense_date,
                act_expense_description=expense_description,
                act_expense_amount=float(expense_amount),
                act_expense_approved=expense_approved,
                act_expense_paid=expense_paid,
                prop_id=expense_prop
            )
            expense.save()
            
            # Check if user is not a superuser and send email
            if not request.user.is_superuser:
                from datetime import date
                from django.utils.dateparse import parse_date
                
                # Parse the expense date for email
                parsed_expense_date = parse_date(expense_date)
                
                email_sent = send_expense_approval_email_with_link(
                    parsed_expense_date,
                    expense.prop.prop_name,  # Get property name through foreign key
                    expense_description,
                    expense_amount,
                    date.today()  # Created date
                )
                if email_sent:
                    messages.success(request, 'Expense added successfully and approval email sent!')
                else:
                    messages.warning(request, 'Expense added successfully but failed to send approval email.')
            else:
                messages.success(request, 'Expense added successfully!')
            
            return redirect('act_expense_all')
            
        except ValueError as e:
            messages.error(request, 'Please enter a valid amount.')
            return redirect('act_expense_add')
        except Exception as e:
            messages.error(request, f'An error occurred: {str(e)}')
            return redirect('act_expense_add')
    
    return redirect('act_expense_add')

### PETTY CASH ###
@login_required
def petty_cash(request):
	presults = petty.objects.all().order_by('petty_cash_date')
	pvalues = petty.objects.values()
	balance = 0
	for x in pvalues:
		if x['petty_cash_dr_cr'] == "DR":
			balance = balance + x['petty_cash_amount']
		elif x['petty_cash_dr_cr'] == "CR":
			balance = balance - x['petty_cash_amount']
	return render (request, "petty_cash.html", {"petty":presults, "balance":balance})

@login_required
def petty_cash_commit(request):
	if request.method == "POST":
		form = PettyForm(request.POST or None)
		print(form)
		if form.is_valid():
			form.save()
			messages.success(request, "Transaction Added Successfully")
	presults = petty.objects.all().order_by('petty_cash_date')
	pvalues = petty.objects.values()
	balance = 0
	for x in pvalues:
		if x['petty_cash_dr_cr'] == "DR":
			balance = balance + x['petty_cash_amount']
		elif x['petty_cash_dr_cr'] == "CR":
			balance = balance - x['petty_cash_amount']
	return render (request, "petty_cash.html", {"petty":presults, "balance":balance})

@login_required
def petty_cash_add(request):
	presults = petty.objects.all().order_by('petty_cash_date')
	return render(request, "petty_cash_add.html", {"petty":presults})


### ISSUES - FRIDAY STATUS REPORT ###
@login_required
def fsr(request):
    # Get filter parameters
    prop_output = request.POST.get('propname', '').strip()
    country_output = request.POST.get('propcountry', '').strip()
    status_output = request.POST.get('issuestatus', '').strip()
    search_query = request.POST.get('search', '').strip()
    
    # Start with all objects
    results = props.objects.all().order_by('prop_country', 'prop_name')
    isresults = issues.objects.all().order_by('issues_date_logged', 'issues_status')
    idresults = issues_details.objects.all().order_by('issues_details_date', 'issues_details_id')
    
    # Apply filters to properties based on country
    if country_output and country_output != 'All':
        results = results.filter(prop_country=country_output)
    
    # Apply filters to properties based on property name
    if prop_output and prop_output != 'All':
        results = results.filter(prop_name=prop_output)
    
    # Apply filters to issues based on status
    if status_output and status_output != 'All':
        isresults = isresults.filter(issues_status=status_output)
    
    # Apply search filter to issues (search in heading and description)
    if search_query:
        isresults = isresults.filter(
            Q(issues_heading__icontains=search_query) | 
            Q(issues_description__icontains=search_query)
        )
    
    # Get the property IDs from filtered results to ensure issues match filtered properties
    if country_output and country_output != 'All':
        property_ids = results.values_list('prop_id', flat=True)
        isresults = isresults.filter(prop_id__in=property_ids)
    
    if prop_output and prop_output != 'All':
        property_ids = results.values_list('prop_id', flat=True)
        isresults = isresults.filter(prop_id__in=property_ids)
    
    # Pass search query to template for displaying in search input
    context = {
        "props": results, 
        "issues": isresults, 
        "issues_details": idresults,
        "search_query": search_query,
        "selected_country": country_output,
        "selected_property": prop_output,
        "selected_status": status_output,
    }
    
    return render(request, "fsr.html", context)

@login_required
@require_POST
def delete_issue(request, issue_id):
    if not request.user.is_superuser:
        return JsonResponse({'success': False, 'error': 'Permission denied'})
    
    try:
        with transaction.atomic():
            # Get the issue (using your actual model name)
            issue_obj = get_object_or_404(issues, issues_id=issue_id)
            
            # Delete all related details first
            issues_details.objects.filter(issues=issue_obj).delete()
            
            # Delete the issue
            issue_obj.delete()
            
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})

@login_required
def fsr_add(request):
	results = props.objects.all().order_by('prop_country','prop_name')
	isresults = issues.objects.all().order_by('issues_date_logged','issues_status')
	idresults = issues_details.objects.all().order_by('issues_details_date','issues_details_id')
	log_date = date.today()
	return render(request, "fsr_add.html", {"props":results, "issues":isresults, "issues_details":idresults, "log_date":log_date})

@login_required
def fsr_commit(request):
    if request.method == "POST":
        form = IssuesForm(request.POST or None)
        if form.is_valid():
            form.save()
            messages.success(request, "Issue Added Successfully")
    temp_results = issues.objects.all().order_by('-issues_id')
    is_id = temp_results[0].issues_id
    return redirect(reverse("fsr_details", args=[is_id]) + "?from=fsr_add&origin=fsr")

@login_required
def fsr_details(request, issues_id):
    isresults = issues.objects.filter(pk=issues_id)
    results = props.objects.all().order_by('prop_country','prop_name')
    idresults = issues_details.objects.all().order_by('issues_details_date','issues_details_id').reverse()
    
    # Get the HTTP_REFERER if it exists
    referrer = request.META.get('HTTP_REFERER', '')
    
    # Determine the clean redirect URL
    if 'fsr_details' in referrer:
        # If coming from another details page, go back to main FSR
        redirect_url = reverse('fsr')
    elif 'status_report' in referrer:
        # If coming from status report, go back there
        redirect_url = reverse('friday_status_report')
    else:
        # Default to the main FSR page
        redirect_url = reverse('fsr')
    
    context = {
        "props": results,
        "issues": isresults,
        "issues_details": idresults,
        "redirect_url": redirect_url,
    }
    
    return render(request, "fsr_details.html", context)

@login_required
def fsr_commit_status_change(request):
    if request.method == "POST":
        # Get form data
        issues_id = request.POST.get('issues_id')
        new_status = request.POST.get('issues_status')
        next_url = request.POST.get('next', '')
        
        # Get return parameters from hidden fields
        from_param = request.POST.get('from', 'fsr')
        property_id = request.POST.get('property_id')
        box_type = request.POST.get('box_type')
        
        # Update the issue
        issue = issues.objects.get(pk=issues_id)
        issue.issues_status = new_status
        if new_status == "Resolved":
            issue.issues_resolution_date = date.today()
        issue.save()
        
        # Handle property_detail navigation
        if from_param == 'property_detail' and property_id and box_type:
            # Redirect back to the same fsr_details page with property_detail parameters
            redirect_url = reverse('fsr_details', args=[issues_id])
            redirect_url += f"?from=property_detail&property_id={property_id}&box_type={box_type}"
            return redirect(redirect_url)
        
        # Handle other cases
        elif from_param == 'fsr':
            return redirect(reverse('fsr') + "?refresh=true")
        elif from_param == 'status_report':
            return redirect(reverse('friday_status_report') + "?refresh=true")
        else:
            # Fallback - try to use the next_url if available
            if next_url:
                return redirect(next_url)
            else:
                return redirect(reverse('fsr') + "?refresh=true")

@login_required
def fsr_comment_add(request, issues_id):
    if request.method == 'POST':
        # Get comment text from form
        comment_text = request.POST.get('issues_details_comment', '').strip()
        
        # Get navigation context efficiently
        next_url = request.POST.get('next', '')
        from_param = request.GET.get('from', '')
        
        # Build redirect URL early to avoid complex logic later
        if next_url:
            redirect_url = next_url
        elif from_param:
            redirect_url = reverse('fsr_details', args=[issues_id]) + f"?from={from_param}"
        else:
            redirect_url = reverse('fsr_details', args=[issues_id])
        
        # Validate comment exists
        if not comment_text:
            messages.error(request, "Comment cannot be empty")
            return redirect(redirect_url)
        
        # Get user info if authenticated
        user_initials = ''
        if request.user.is_authenticated:
            user_initials = f"{request.user.first_name[:1]}{request.user.last_name[:1]}"
        
        try:
            # Create the comment - single database operation
            issues_details.objects.create(
                issues_details_comment=comment_text,
                issues_details_user=user_initials,
                issues_details_date=date.today(),
                issues_id=issues_id
            )
            
            messages.success(request, "Comment added successfully")
            
        except Exception as e:
            messages.error(request, "Failed to add comment. Please try again.")
            
        # Use the pre-built redirect URL
        return redirect(redirect_url)
    
    # If not POST, redirect to details page
    return redirect('fsr_details', issues_id=issues_id)

def fsr_pdf(request):
    """
    Generate PDF version of FSR report
    """
    from django.db import connection
    
    try:
        context = get_fsr_context_data(request)
        return render_to_pdf('fsr_email.html', context)
    finally:
        # Close database connection
        connection.close()

def get_fsr_context_data(request):
    """
    Generate context data for Friday Status Report (used by both web view and email)
    Rewritten to use Django ORM instead of raw SQL
    """
    from django.db import connection
    import logging
    
    logger = logging.getLogger(__name__)
    
    try:
        today = date.today()
        
        # Get max_comments parameter for summarized reports
        max_comments = request.GET.get('max_comments', None) if hasattr(request, 'GET') else None
        is_summarized_report = max_comments is not None
        
        if is_summarized_report:
            try:
                max_comments = int(max_comments)
            except (ValueError, TypeError):
                max_comments = None
                is_summarized_report = False
        
        # Get all properties ordered by country and name
        properties = props.objects.all().order_by('prop_country', 'prop_name').values('prop_name')
        
        # Get all issues with their details, using select_related and prefetch_related for optimization
        issues_queryset = issues.objects.select_related('prop').prefetch_related(
            Prefetch(
                'issues_details_set',
                queryset=issues_details.objects.all().order_by('-issues_details_id'),
                to_attr='details_list'
            )
        ).order_by('issues_id')
        
        # Process issues data
        issues_data = []
        for issue_obj in issues_queryset:
            # Build the issue dictionary
            issue_dict = {
                'prop_name': issue_obj.prop.prop_name,
                'issues_id': issue_obj.issues_id,
                'issues_heading': issue_obj.issues_heading,
                'issues_description': issue_obj.issues_description,
                'issues_status': issue_obj.issues_status,
                'issues_date_logged': issue_obj.issues_date_logged,
                'issues_resolution_date': issue_obj.issues_resolution_date,
                'days_to_resolve': None,
                'days_open': None,
                'details': []
            }
            
            # Calculate days metrics based on status
            if issue_dict['issues_date_logged']:
                if issue_dict['issues_status'] == 'Resolved':
                    if (issue_dict['issues_resolution_date'] and 
                        issue_dict['issues_resolution_date'] != date(1900, 1, 1)):
                        issue_dict['days_to_resolve'] = (issue_dict['issues_resolution_date'] - issue_dict['issues_date_logged']).days
                else:
                    issue_dict['days_open'] = (today - issue_dict['issues_date_logged']).days
            
            # Process details
            details_data = []
            for detail in issue_obj.details_list:
                details_data.append({
                    'issues_details_id': detail.issues_details_id,
                    'issues_details_comment': detail.issues_details_comment,
                    'issues_details_user': detail.issues_details_user,
                    'issues_details_date': detail.issues_details_date
                })
            
            # Apply comment limiting for summarized reports
            if is_summarized_report and max_comments and len(details_data) > max_comments:
                total_comments_before_limit = len(details_data)
                issue_dict['details'] = details_data[:max_comments]
                issue_dict['has_more_comments'] = True
                issue_dict['total_comments'] = total_comments_before_limit
            else:
                issue_dict['details'] = details_data
                issue_dict['has_more_comments'] = False
                issue_dict['total_comments'] = len(details_data)
            
            issues_data.append(issue_dict)
        
        # Process data by status and property
        processed_data = {}
        for status in ['Resolved', 'Unresolved', 'Issue']:
            processed_data[status] = {}
            for prop in properties:
                prop_name = prop['prop_name']
                processed_data[status][prop_name] = []

                unique_issues = set()

                for issue in issues_data:
                    if (issue['prop_name'] == prop_name and 
                        issue['issues_status'] == status and 
                        (issue['issues_heading'], issue['issues_description']) not in unique_issues):

                        if status == 'Resolved':
                            if (issue['issues_resolution_date'] != date(1900, 1, 1) and 
                                issue['issues_resolution_date'] >= (date.today() - timedelta(days=7))):
                                processed_data[status][prop_name].append(issue)
                                unique_issues.add((issue['issues_heading'], issue['issues_description']))
                        else:
                            processed_data[status][prop_name].append(issue)
                            unique_issues.add((issue['issues_heading'], issue['issues_description']))
        
        context = {
            'today': today,
            'statuses': ['Resolved', 'Unresolved', 'Issue'],
            'properties': properties,
            'is_summarized_report': is_summarized_report,
            'max_comments': max_comments,
            'status_groups': [
                {
                    'status': status,
                    'property_issues': [
                        {
                            'prop_name': prop['prop_name'],
                            'issues': processed_data[status][prop['prop_name']]
                        }
                        for prop in properties
                        if processed_data[status][prop['prop_name']]
                    ]
                }
                for status in ['Resolved', 'Unresolved', 'Issue']
            ]
        }
        
        return context
        
    except Exception as e:
        logger.error(f"Error in get_fsr_context_data: {e}")
        # Return minimal context on error
        return {
            'today': date.today(),
            'statuses': ['Resolved', 'Unresolved', 'Issue'],
            'properties': [],
            'is_summarized_report': False,
            'max_comments': None,
            'status_groups': []
        }
        
    finally:
        # Close database connection
        connection.close()

@login_required
def fsr_notification(request):
    from django.db import connection
    from django.db.utils import OperationalError, InterfaceError
    from pages.management.commands.email_utils import get_email_recipients, format_email_recipients_for_header
    import time
    import logging
    
    logger = logging.getLogger(__name__)
    smtp_object = None
    
    # Close any stale connections before starting
    connection.close()
    
    max_retries = 3
    for attempt in range(max_retries):
        try:
            # Check if there's a max_comments parameter in the session or request
            # This indicates the user wants a summarized report
            max_comments = None
            is_summarized_report = False
            
            # Check for max_comments in various places:
            # 1. Direct GET parameter (if coming from Friday report page)
            # 2. Session storage (if user navigated from a summarized report)
            # 3. HTTP_REFERER analysis (check if previous page had max_comments)
            
            if 'max_comments' in request.GET:
                max_comments = request.GET.get('max_comments')
                is_summarized_report = True
            elif 'last_report_type' in request.session:
                # If we stored the last report type in session
                if request.session['last_report_type'] == 'summarized':
                    max_comments = request.session.get('max_comments', '2')
                    is_summarized_report = True
            else:
                # Check the HTTP referer to see if it came from a summarized report
                referer = request.META.get('HTTP_REFERER', '')
                if 'max_comments=' in referer:
                    # Extract max_comments from referer URL
                    import re
                    match = re.search(r'max_comments=(\d+)', referer)
                    if match:
                        max_comments = match.group(1)
                        is_summarized_report = True
            
            # Validate max_comments
            if is_summarized_report and max_comments:
                try:
                    max_comments = int(max_comments)
                    if max_comments < 1:
                        max_comments = 2
                        is_summarized_report = False
                except (ValueError, TypeError):
                    max_comments = 2
                    is_summarized_report = False
            
            # Create a mock request object with the appropriate parameters for context generation
            mock_request = type('MockRequest', (), {})()
            mock_request.user = request.user
            mock_request.session = request.session
            mock_request.META = request.META
            
            if is_summarized_report:
                # Create GET parameters for summarized report
                mock_request.GET = {'max_comments': str(max_comments)}
                report_type_text = f"Summarized Report (Max {max_comments} comments per issue)"
            else:
                # No parameters for detailed report
                mock_request.GET = {}
                report_type_text = "Detailed Report (All comments)"
            
            # Fetch context data for the report with appropriate parameters
            # This is the critical database operation that needs protection
            context = get_fsr_context_data(mock_request)
            
            # Add report type information to context for email template
            context['is_summarized_report'] = is_summarized_report
            context['max_comments'] = max_comments if is_summarized_report else None
            context['report_type_text'] = report_type_text
            
            # Render HTML content
            html_content = render_to_string("fsr_email.html", context, request=request)
            text_content = strip_tags(html_content)
            
            # Get recipients from database based on who is submitting
            if request.user.is_superuser:
                # Supervisor submitting - send to Stella
                recipients = get_email_recipients('friday_status_report_supervisor')
            else:
                # Staff submitting - send to Demetri
                recipients = get_email_recipients('friday_status_report_staff')
            
            # Prepare email with report type in subject
            msg = MIMEMultipart("alternative")
            msg['From'] = "demetrimanias@gmail.com"
            msg['To'] = format_email_recipients_for_header(recipients['to'])
            if recipients['cc']:
                msg['Cc'] = format_email_recipients_for_header(recipients['cc'])
            
            # Include report type in subject
            if is_summarized_report:
                msg['Subject'] = f"Friday Status Report - Summarized ({max_comments} comments/issue)"
            else:
                msg['Subject'] = "Friday Status Report - Detailed"
            
            # Attach both plain text and HTML
            msg.attach(MIMEText(text_content, 'plain'))
            msg.attach(MIMEText(html_content, 'html'))
            
            # Get email credentials and settings from environment variables
            email_password = os.environ.get('EMAIL_PASSWORD')
            email_host = os.environ.get('EMAIL_HOST', 'smtp.gmail.com')
            email_port = int(os.environ.get('EMAIL_PORT', 465))
            email_use_ssl = os.environ.get('EMAIL_USE_SSL', 'True').lower() == 'true'
            email_use_tls = os.environ.get('EMAIL_USE_TLS', 'False').lower() == 'true'
            
            if not email_password:
                logger.error('❌ EMAIL_PASSWORD environment variable not set')
                messages.error(request, "Failed to send email - No password configured.")
                return redirect('fsr')
            
            # SMTP setup with environment variable configuration
            if email_use_ssl:
                # Use SSL connection (typically port 465)
                smtp_object = smtplib.SMTP_SSL(email_host, email_port, timeout=10)
            else:
                # Use regular SMTP connection (typically port 587)
                smtp_object = smtplib.SMTP(email_host, email_port, timeout=10)
                smtp_object.ehlo()
                if email_use_tls:
                    smtp_object.starttls()
            
            email = "demetrimanias@gmail.com"
            smtp_object.login(email, email_password)
            
            # Send email to all recipients
            smtp_object.sendmail(email, recipients['all'], msg.as_string())
            
            success_message = f"Friday Status Report ({report_type_text}) sent successfully!"
            messages.success(request, success_message)
            
            # If we get here, everything worked - break out of retry loop
            break
            
        except (OperationalError, InterfaceError) as e:
            if attempt < max_retries - 1:
                # Close connection and wait before retry
                connection.close()
                time.sleep(2)  # Wait 2 seconds before retry
                logger.warning(f"Database connection error on attempt {attempt + 1}, retrying: {e}")
                continue
            else:
                # Final attempt failed
                logger.error(f"Database connection failed after {max_retries} attempts: {e}")
                messages.error(request, "Database connection error. Please try again in a moment.")
                return redirect('fsr')
                
        except smtplib.SMTPAuthenticationError as e:
            logger.error(f"SMTP Authentication Error: {e}")
            messages.error(request, "Failed to send email - Authentication error.")
            break
            
        except smtplib.SMTPException as e:
            logger.error(f"SMTP Error: {e}")
            messages.error(request, "Failed to send email - SMTP error.")
            break
            
        except Exception as e:
            logger.error(f"Error sending email: {e}")
            messages.error(request, f"Failed to send email notification: {str(e)}")
            break
            
        finally:
            if smtp_object:
                try:
                    smtp_object.quit()
                except:
                    pass
            # Close database connection
            connection.close()
    
    return redirect('fsr')

@login_required
def comments_report(request):
    """Generate a report of all comments with filtering by time period"""
    
    period = request.GET.get('period', '30')
    
    # Get all comments with related data (optimize with select_related)
    comments = issues_details.objects.select_related('issues', 'issues__prop').all()
    
    # Apply time filter
    if period != 'all':
        days = int(period)
        cutoff_date = timezone.now().date() - timedelta(days=days)
        comments = comments.filter(issues_details_date__gte=cutoff_date)
    
    # Order by date descending (most recent first)
    comments = comments.order_by('-issues_details_date', '-issues_details_id')
    
    # Build the report data with property information
    report_data = []
    for comment in comments:
        issue = comment.issues
        if issue and issue.prop:
            property_name = issue.prop.prop_name
            issue_heading = issue.issues_heading
            issue_id = issue.issues_id
            issue_status = issue.issues_status
            issue_description = issue.issues_description
        elif issue:
            property_name = 'Unknown'
            issue_heading = issue.issues_heading
            issue_id = issue.issues_id
            issue_status = issue.issues_status
            issue_description = issue.issues_description
        else:
            property_name = 'Unknown'
            issue_heading = 'Unknown'
            issue_id = None
            issue_status = None
            issue_description = None
        
        # Define admin users (add initials of admin users here)
        admin_users = ['DM']  # Add other admin initials as needed
        user_initials = comment.issues_details_user or ''
        is_admin = user_initials.upper() in [u.upper() for u in admin_users]
        
        report_data.append({
            'comment_id': comment.issues_details_id,        # ADD THIS - needed for delete
            'comment': comment.issues_details_comment,
            'date': comment.issues_details_date,
            'property': property_name,
            'user': comment.issues_details_user,
            'issue_heading': issue_heading,
            'issue_id': issue_id,
            'issue_status': issue_status,
            'issue_description': issue_description,
            'is_admin': is_admin,
        })
    
    # Period display text
    period_labels = {
        '7': 'Last 7 Days',
        '30': 'Last 30 Days',
        '90': 'Last 90 Days',
        'all': 'All Comments'
    }
    period_label = period_labels.get(period, 'Last 30 Days')
    
    context = {
        'report_data': report_data,
        'period': period,
        'period_label': period_label,
        'comment_count': len(report_data),
    }
    
    return render(request, 'comments_report.html', context)

@login_required
def delete_comment(request, comment_id):
    """Delete a comment from the issues_details table (admin only)"""
    from django.urls import reverse
    from urllib.parse import urlencode
    
    # Check if user is superuser
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to delete comments.')
        return HttpResponseForbidden("You do not have permission to delete comments.")
    
    # Only allow POST requests for deletion
    if request.method == 'POST':
        try:
            # Get the comment
            comment = issues_details.objects.get(issues_details_id=comment_id)
            
            # Store info for success message
            comment_text = comment.issues_details_comment[:50]  # First 50 chars
            
            # Delete the comment
            comment.delete()
            
            # Success message
            messages.success(request, f'Comment "{comment_text}..." has been successfully deleted.')
            
        except issues_details.DoesNotExist:
            messages.error(request, 'Comment not found.')
        except Exception as e:
            messages.error(request, f'Error deleting comment: {str(e)}')
    
    # FIXED: Properly redirect back with period parameter
    # Get period from the POST data or default to '30'
    period = request.POST.get('period', request.GET.get('period', '30'))
    
    # Build the redirect URL with proper query string
    url = reverse('comments_report') + '?' + urlencode({'period': period})
    return redirect(url)

@login_required
def get_issue_details(request, issue_id):
    """
    Fetch issue details and all comments for modal display
    Returns JSON data
    """
    try:
        # Get the issue
        issue = issues.objects.select_related('prop').get(issues_id=issue_id)
        
        # Get all comments for this issue (most recent first)
        comments = issues_details.objects.filter(
            issues=issue
        ).order_by('-issues_details_date', '-issues_details_id')
        
        # Define admin users (same as in comments_report view)
        admin_users = ['DM']  # Add other admin initials as needed
        
        # Build comments list
        comments_list = []
        for comment in comments:
            user_initials = comment.issues_details_user or ''
            is_admin = user_initials.upper() in [u.upper() for u in admin_users]
            
            comments_list.append({
                'comment': comment.issues_details_comment,
                'date': comment.issues_details_date.strftime('%d/%m/%Y'),
                'user': comment.issues_details_user,
                'is_admin': is_admin,
            })
        
        # Build response data
        data = {
            'issue_id': issue.issues_id,
            'issue_heading': issue.issues_heading,
            'description': issue.issues_description,
            'property': issue.prop.prop_name if issue.prop else 'Unknown',
            'status': issue.issues_status,
            'date_logged': issue.issues_date_logged.strftime('%d/%m/%Y') if issue.issues_date_logged else '—',
            'resolution_date': issue.issues_resolution_date.strftime('%d/%m/%Y') if issue.issues_resolution_date else None,
            'comments': comments_list,
        }
        
        return JsonResponse(data)
        
    except issues.DoesNotExist:
        return JsonResponse({'error': 'Issue not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

### REPORTS - DASHBOARD (FROM HOME PAGE) ###
# Add these views to your views.py file
@login_required
def revenue_details_view(request):
    """
    View to show revenue details breakdown for budgeted/fixed revenues
    """
    year = request.GET.get('year', datetime.now().year)  # Just for display
    month = request.GET.get('month')
    line_type = request.GET.get('line_type')
    property_id = request.GET.get('property_id')
    prop = request.GET.get('prop', 'all')
    properties = request.GET.get('properties', '')  # NEW: Handle comma-separated properties
    
    # Get all revenue records (no year filtering needed for budgeted revenues)
    revenues = revenue.objects.all().select_related('prop', 'revenue_line_types', 'revenue_types')
    
    # Filter by line type if specified
    if line_type:
        revenues = revenues.filter(revenue_line_types_id=line_type)
    
    # Filter by property - UPDATED LOGIC
    if properties:  # NEW: Handle comma-separated properties
        try:
            property_ids = [int(prop_id.strip()) for prop_id in properties.split(',') if prop_id.strip()]
            revenues = revenues.filter(prop_id__in=property_ids)
        except ValueError:
            pass  # Invalid property IDs, show no results
    elif property_id and property_id != 'all':
        revenues = revenues.filter(prop_id=property_id)
    elif prop and prop != 'all':
        revenues = revenues.filter(prop_id=prop)
    
    # Get line type name for header
    line_type_name = "Revenue"
    if line_type:
        try:
            line_type_obj = revenue_line_types.objects.get(revenue_line_types_id=line_type)
            line_type_name = line_type_obj.revenue_line_types_name
        except revenue_line_types.DoesNotExist:
            line_type_name = "Unknown"
    
    # Get month name for subtitle
    month_names = {
        '1': 'January', '2': 'February', '3': 'March', '4': 'April',
        '5': 'May', '6': 'June', '7': 'July', '8': 'August',
        '9': 'September', '10': 'October', '11': 'November', '12': 'December'
    }
    month_name = month_names.get(str(month), "All Months") if month else "All Months"
    
    # Create a list of revenue items with monthly breakdown
    revenue_items = []
    total_amount = 0
    
    for rev in revenues:
        months = ['jan', 'feb', 'mar', 'apr', 'may', 'jun', 
                 'jul', 'aug', 'sep', 'oct', 'nov', 'dec']
        
        for i, month_name_field in enumerate(months, 1):
            month_value = getattr(rev, f'revenue_{month_name_field}', 0)
            
            if month_value and month_value > 0:
                # If specific month is requested, only show that month
                if month and int(month) != i:
                    continue
                
                revenue_items.append({
                    'revenue_id': rev.revenue_id,
                    'property': rev.prop,
                    'amount': float(month_value),
                })
                total_amount += float(month_value)
    
    context = {
        'revenue_items': revenue_items,
        'total_amount': total_amount,
        'selected_year': year,
        'selected_month': month,
        'selected_line_type': line_type,
        'line_type_name': line_type_name,
        'month_name': month_name,
    }
    
    return render(request, 'revenue_details.html', context)

@login_required
def budget_expense_details_view(request):
    """
    View to show budgeted expense details breakdown
    """
    year = request.GET.get('year', datetime.now().year)  # Just for display
    month = request.GET.get('month')
    line_type = request.GET.get('line_type')
    property_id = request.GET.get('property_id')
    prop = request.GET.get('prop', 'all')
    properties = request.GET.get('properties', '')  # NEW: Handle comma-separated properties
    
    # Get all budgeted expense records (no year filtering needed for budgeted expenses)
    expenses = expense.objects.all().select_related('prop', 'expense_line_types', 'expense_types')
    
    # Filter by line type if specified
    if line_type:
        expenses = expenses.filter(expense_line_types_id=line_type)
    
    # Filter by property - UPDATED LOGIC
    if properties:  # NEW: Handle comma-separated properties
        try:
            property_ids = [int(prop_id.strip()) for prop_id in properties.split(',') if prop_id.strip()]
            expenses = expenses.filter(prop_id__in=property_ids)
        except ValueError:
            pass  # Invalid property IDs, show no results
    elif property_id and property_id != 'all':
        expenses = expenses.filter(prop_id=property_id)
    elif prop and prop != 'all':
        expenses = expenses.filter(prop_id=prop)
    
    # Get line type name for header
    line_type_name = "Budget Expenses"
    if line_type:
        try:
            line_type_obj = expense_line_types.objects.get(expense_line_types_id=line_type)
            line_type_name = line_type_obj.expense_line_types_name
        except expense_line_types.DoesNotExist:
            line_type_name = "Unknown"
    
    # Get month name for subtitle
    month_names = {
        '1': 'January', '2': 'February', '3': 'March', '4': 'April',
        '5': 'May', '6': 'June', '7': 'July', '8': 'August',
        '9': 'September', '10': 'October', '11': 'November', '12': 'December'
    }
    month_name = month_names.get(str(month), "All Months") if month else "All Months"
    
    # Create a list of expense items with monthly breakdown
    expense_items = []
    total_amount = 0
    
    for exp in expenses:
        months = ['jan', 'feb', 'mar', 'apr', 'may', 'jun', 
                 'jul', 'aug', 'sep', 'oct', 'nov', 'dec']
        
        for i, month_field in enumerate(months, 1):
            month_value = getattr(exp, f'expense_{month_field}', 0)
            
            if month_value and month_value > 0:
                # If specific month is requested, only show that month
                if month and int(month) != i:
                    continue
                
                expense_items.append({
                    'expense_id': exp.expense_id,
                    'property': exp.prop,
                    'amount': float(month_value),
                })
                total_amount += float(month_value)
    
    context = {
        'expense_items': expense_items,
        'total_amount': total_amount,
        'selected_year': year,
        'selected_month': month,
        'selected_line_type': line_type,
        'line_type_name': line_type_name,
        'month_name': month_name,
    }
    
    return render(request, 'budget_expense_details.html', context)

@login_required
def total_expense_details_view(request):
    """
    View to show combined actual + budgeted expense details
    """
    year = request.GET.get('year', datetime.now().year)
    month = request.GET.get('month')
    property_id = request.GET.get('property_id')
    prop = request.GET.get('prop', 'all')
    
    # Get actual expenses
    actual_expenses = act_expense.objects.filter(
        act_expense_date__year=year
    ).select_related('prop')
    
    # Get budget expenses
    budget_expenses = expense.objects.filter(
        expense_types__expense_types_name__icontains=str(year)
    ).select_related('prop', 'expense_line_types', 'expense_types')
    
    # Filter by month if specified
    if month:
        actual_expenses = actual_expenses.filter(act_expense_date__month=month)
    
    # Filter by property
    if property_id:
        actual_expenses = actual_expenses.filter(prop_id=property_id)
        budget_expenses = budget_expenses.filter(prop_id=property_id)
    elif prop != 'all':
        actual_expenses = actual_expenses.filter(prop_id=prop)
        budget_expenses = budget_expenses.filter(prop_id=prop)
    
    # Order by date
    actual_expenses = actual_expenses.order_by('-act_expense_date')
    
    # Create budget expense items with monthly breakdown (similar to budget_expense_details_view)
    budget_expense_items = []
    for exp in budget_expenses:
        months = ['jan', 'feb', 'mar', 'apr', 'may', 'jun', 
                 'jul', 'aug', 'sep', 'oct', 'nov', 'dec']
        
        for i, month_name in enumerate(months, 1):
            month_value = getattr(exp, f'expense_{month_name}', 0)
            if month_value and month_value > 0:
                # If specific month is requested, only show that month
                if month and int(month) != i:
                    continue
                    
                budget_expense_items.append({
                    'expense_id': exp.expense_id,
                    'property': exp.prop,
                    'expense_line_type': exp.expense_line_types,
                    'expense_type': exp.expense_types,
                    'month': i,
                    'month_name': month_name.capitalize(),
                    'amount': month_value,
                    'description': f"{exp.expense_line_types.expense_line_types_name} - {month_name.capitalize()} {year}",
                    'type': 'budget'
                })
    
    # Get line types and properties for context
    expense_line_types_list = expense_line_types.objects.all()
    properties = props.objects.all()
    
    context = {
        'actual_expenses': actual_expenses,
        'budget_expense_items': budget_expense_items,
        'expense_line_types': expense_line_types_list,
        'properties': properties,
        'selected_year': year,
        'selected_month': month,
        'selected_property': property_id,
        'prop': prop,
    }
    
    return render(request, 'total_expense_details.html', context)

@login_required
def finance_pl(request):
    # Get selected properties from request
    selected_properties = request.GET.getlist('properties')
    
    # Get all active properties with prefetched prop_values to optimize queries
    all_properties = props.objects.filter(prop_status="Active").prefetch_related('prop_values_set')
    
    # If no properties selected, return empty data
    if not selected_properties:
        # Return minimal context with empty data
        context = {
            'properties': [],
            'all_properties': all_properties,
            'revenue_line_types': revenue_line_types.objects.all(),
            'expense_line_types': expense_line_types.objects.all(),
            'revenue_totals': {'jan': 0, 'feb': 0, 'mar': 0, 'apr': 0, 'may': 0, 'jun': 0,
                             'jul': 0, 'aug': 0, 'sep': 0, 'oct': 0, 'nov': 0, 'dec': 0, 'year': 0},
            'expense_totals': {'jan': 0, 'feb': 0, 'mar': 0, 'apr': 0, 'may': 0, 'jun': 0,
                             'jul': 0, 'aug': 0, 'sep': 0, 'oct': 0, 'nov': 0, 'dec': 0, 'year': 0},
            'profit_totals': {'jan': 0, 'feb': 0, 'mar': 0, 'apr': 0, 'may': 0, 'jun': 0,
                            'jul': 0, 'aug': 0, 'sep': 0, 'oct': 0, 'nov': 0, 'dec': 0, 'year': 0},
            'revenue_totals_by_line': {'all': {}},
            'expense_totals_by_line': {'all': {}},
            'revenue_prop_totals': {},
            'expense_prop_totals': {},
            'total_current_value': 0,
            'selected_properties': selected_properties,
        }
        return render(request, 'finance_pl.html', context)
    
    # Convert to integers and filter
    selected_prop_ids = [int(pid) for pid in selected_properties if pid.isdigit()]
    properties = all_properties.filter(prop_id__in=selected_prop_ids)
    
    # Revenue Section
    revenue_line_types_list = revenue_line_types.objects.all()
    
    # Filter revenues by selected properties only
    revenues = revenue.objects.filter(prop_id__in=selected_prop_ids)
    
    # Calculate revenue totals for selected properties
    revenue_totals = {
        'jan': sum(r.revenue_jan or 0 for r in revenues),
        'feb': sum(r.revenue_feb or 0 for r in revenues),
        'mar': sum(r.revenue_mar or 0 for r in revenues),
        'apr': sum(r.revenue_apr or 0 for r in revenues),
        'may': sum(r.revenue_may or 0 for r in revenues),
        'jun': sum(r.revenue_jun or 0 for r in revenues),
        'jul': sum(r.revenue_jul or 0 for r in revenues),
        'aug': sum(r.revenue_aug or 0 for r in revenues),
        'sep': sum(r.revenue_sep or 0 for r in revenues),
        'oct': sum(r.revenue_oct or 0 for r in revenues),
        'nov': sum(r.revenue_nov or 0 for r in revenues),
        'dec': sum(r.revenue_dec or 0 for r in revenues),
    }
    revenue_totals['year'] = sum(revenue_totals.values())
    
    # Calculate revenue totals by line type for selected properties
    revenue_totals_by_line = {'all': {}}
    for lt in revenue_line_types_list:
        line_revenues = revenues.filter(revenue_line_types=lt)
        monthly_totals = {
            'jan': sum(r.revenue_jan or 0 for r in line_revenues),
            'feb': sum(r.revenue_feb or 0 for r in line_revenues),
            'mar': sum(r.revenue_mar or 0 for r in line_revenues),
            'apr': sum(r.revenue_apr or 0 for r in line_revenues),
            'may': sum(r.revenue_may or 0 for r in line_revenues),
            'jun': sum(r.revenue_jun or 0 for r in line_revenues),
            'jul': sum(r.revenue_jul or 0 for r in line_revenues),
            'aug': sum(r.revenue_aug or 0 for r in line_revenues),
            'sep': sum(r.revenue_sep or 0 for r in line_revenues),
            'oct': sum(r.revenue_oct or 0 for r in line_revenues),
            'nov': sum(r.revenue_nov or 0 for r in line_revenues),
            'dec': sum(r.revenue_dec or 0 for r in line_revenues),
        }
        monthly_totals['total'] = sum(monthly_totals.values())
        revenue_totals_by_line['all'][lt.revenue_line_types_id] = monthly_totals
    
    # Calculate property-specific revenue totals for selected properties
    revenue_prop_totals = {}
    for prop in properties:
        prop_revenues = revenues.filter(prop=prop)
        monthly_totals = {
            'jan': sum(r.revenue_jan or 0 for r in prop_revenues),
            'feb': sum(r.revenue_feb or 0 for r in prop_revenues),
            'mar': sum(r.revenue_mar or 0 for r in prop_revenues),
            'apr': sum(r.revenue_apr or 0 for r in prop_revenues),
            'may': sum(r.revenue_may or 0 for r in prop_revenues),
            'jun': sum(r.revenue_jun or 0 for r in prop_revenues),
            'jul': sum(r.revenue_jul or 0 for r in prop_revenues),
            'aug': sum(r.revenue_aug or 0 for r in prop_revenues),
            'sep': sum(r.revenue_sep or 0 for r in prop_revenues),
            'oct': sum(r.revenue_oct or 0 for r in prop_revenues),
            'nov': sum(r.revenue_nov or 0 for r in prop_revenues),
            'dec': sum(r.revenue_dec or 0 for r in prop_revenues),
        }
        monthly_totals['year'] = sum(monthly_totals.values())
        revenue_prop_totals[prop.prop_id] = monthly_totals
        
        # Add property-specific revenue line type totals
        revenue_totals_by_line[prop.prop_id] = {}
        for lt in revenue_line_types_list:
            prop_line_revenues = prop_revenues.filter(revenue_line_types=lt)
            line_monthly_totals = {
                'jan': sum(r.revenue_jan or 0 for r in prop_line_revenues),
                'feb': sum(r.revenue_feb or 0 for r in prop_line_revenues),
                'mar': sum(r.revenue_mar or 0 for r in prop_line_revenues),
                'apr': sum(r.revenue_apr or 0 for r in prop_line_revenues),
                'may': sum(r.revenue_may or 0 for r in prop_line_revenues),
                'jun': sum(r.revenue_jun or 0 for r in prop_line_revenues),
                'jul': sum(r.revenue_jul or 0 for r in prop_line_revenues),
                'aug': sum(r.revenue_aug or 0 for r in prop_line_revenues),
                'sep': sum(r.revenue_sep or 0 for r in prop_line_revenues),
                'oct': sum(r.revenue_oct or 0 for r in prop_line_revenues),
                'nov': sum(r.revenue_nov or 0 for r in prop_line_revenues),
                'dec': sum(r.revenue_dec or 0 for r in prop_line_revenues),
            }
            line_monthly_totals['total'] = sum(line_monthly_totals.values())
            revenue_totals_by_line[prop.prop_id][lt.revenue_line_types_id] = line_monthly_totals

    # Expense Section
    expense_line_types_list = expense_line_types.objects.all()
    
    # Filter expenses by selected properties only
    expenses = expense.objects.filter(prop_id__in=selected_prop_ids)
    
    # Calculate expense totals for selected properties
    expense_totals = {
        'jan': sum(e.expense_jan or 0 for e in expenses),
        'feb': sum(e.expense_feb or 0 for e in expenses),
        'mar': sum(e.expense_mar or 0 for e in expenses),
        'apr': sum(e.expense_apr or 0 for e in expenses),
        'may': sum(e.expense_may or 0 for e in expenses),
        'jun': sum(e.expense_jun or 0 for e in expenses),
        'jul': sum(e.expense_jul or 0 for e in expenses),
        'aug': sum(e.expense_aug or 0 for e in expenses),
        'sep': sum(e.expense_sep or 0 for e in expenses),
        'oct': sum(e.expense_oct or 0 for e in expenses),
        'nov': sum(e.expense_nov or 0 for e in expenses),
        'dec': sum(e.expense_dec or 0 for e in expenses),
    }
    expense_totals['year'] = sum(expense_totals.values())
    
    # Calculate expense totals by line type for selected properties
    expense_totals_by_line = {'all': {}}
    for elt in expense_line_types_list:
        line_expenses = expenses.filter(expense_line_types=elt)
        monthly_totals = {
            'jan': sum(e.expense_jan or 0 for e in line_expenses),
            'feb': sum(e.expense_feb or 0 for e in line_expenses),
            'mar': sum(e.expense_mar or 0 for e in line_expenses),
            'apr': sum(e.expense_apr or 0 for e in line_expenses),
            'may': sum(e.expense_may or 0 for e in line_expenses),
            'jun': sum(e.expense_jun or 0 for e in line_expenses),
            'jul': sum(e.expense_jul or 0 for e in line_expenses),
            'aug': sum(e.expense_aug or 0 for e in line_expenses),
            'sep': sum(e.expense_sep or 0 for e in line_expenses),
            'oct': sum(e.expense_oct or 0 for e in line_expenses),
            'nov': sum(e.expense_nov or 0 for e in line_expenses),
            'dec': sum(e.expense_dec or 0 for e in line_expenses),
        }
        monthly_totals['total'] = sum(monthly_totals.values())
        expense_totals_by_line['all'][elt.expense_line_types_id] = monthly_totals
    
    # Calculate property-specific expense totals for selected properties
    expense_prop_totals = {}
    for prop in properties:
        prop_expenses = expenses.filter(prop=prop)
        monthly_totals = {
            'jan': sum(e.expense_jan or 0 for e in prop_expenses),
            'feb': sum(e.expense_feb or 0 for e in prop_expenses),
            'mar': sum(e.expense_mar or 0 for e in prop_expenses),
            'apr': sum(e.expense_apr or 0 for e in prop_expenses),
            'may': sum(e.expense_may or 0 for e in prop_expenses),
            'jun': sum(e.expense_jun or 0 for e in prop_expenses),
            'jul': sum(e.expense_jul or 0 for e in prop_expenses),
            'aug': sum(e.expense_aug or 0 for e in prop_expenses),
            'sep': sum(e.expense_sep or 0 for e in prop_expenses),
            'oct': sum(e.expense_oct or 0 for e in prop_expenses),
            'nov': sum(e.expense_nov or 0 for e in prop_expenses),
            'dec': sum(e.expense_dec or 0 for e in prop_expenses),
        }
        monthly_totals['year'] = sum(monthly_totals.values())
        expense_prop_totals[prop.prop_id] = monthly_totals
        
        # Add property-specific expense line type totals
        expense_totals_by_line[prop.prop_id] = {}
        for elt in expense_line_types_list:
            prop_line_expenses = prop_expenses.filter(expense_line_types=elt)
            line_monthly_totals = {
                'jan': sum(e.expense_jan or 0 for e in prop_line_expenses),
                'feb': sum(e.expense_feb or 0 for e in prop_line_expenses),
                'mar': sum(e.expense_mar or 0 for e in prop_line_expenses),
                'apr': sum(e.expense_apr or 0 for e in prop_line_expenses),
                'may': sum(e.expense_may or 0 for e in prop_line_expenses),
                'jun': sum(e.expense_jun or 0 for e in prop_line_expenses),
                'jul': sum(e.expense_jul or 0 for e in prop_line_expenses),
                'aug': sum(e.expense_aug or 0 for e in prop_line_expenses),
                'sep': sum(e.expense_sep or 0 for e in prop_line_expenses),
                'oct': sum(e.expense_oct or 0 for e in prop_line_expenses),
                'nov': sum(e.expense_nov or 0 for e in prop_line_expenses),
                'dec': sum(e.expense_dec or 0 for e in prop_line_expenses),
            }
            line_monthly_totals['total'] = sum(line_monthly_totals.values())
            expense_totals_by_line[prop.prop_id][elt.expense_line_types_id] = line_monthly_totals

    # Calculate Profit (Revenue - Expenses) for selected properties
    profit_totals = {
        'jan': revenue_totals['jan'] - expense_totals['jan'],
        'feb': revenue_totals['feb'] - expense_totals['feb'],
        'mar': revenue_totals['mar'] - expense_totals['mar'],
        'apr': revenue_totals['apr'] - expense_totals['apr'],
        'may': revenue_totals['may'] - expense_totals['may'],
        'jun': revenue_totals['jun'] - expense_totals['jun'],
        'jul': revenue_totals['jul'] - expense_totals['jul'],
        'aug': revenue_totals['aug'] - expense_totals['aug'],
        'sep': revenue_totals['sep'] - expense_totals['sep'],
        'oct': revenue_totals['oct'] - expense_totals['oct'],
        'nov': revenue_totals['nov'] - expense_totals['nov'],
        'dec': revenue_totals['dec'] - expense_totals['dec'],
        'year': revenue_totals['year'] - expense_totals['year']
    }

    # Prepare property values mapping for easy access in template
    prop_values_map = {prop.prop_id: prop.prop_values_set.first() for prop in properties}
    total_current_value = 0
    for prop in properties:
        prop_values = prop.prop_values_set.first()
        if prop_values and prop_values.prop_values_current_value is not None:
            total_current_value += prop_values.prop_values_current_value

    # Handle AJAX requests
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return JsonResponse({
            'success': True,
            'properties': [{'id': p.prop_id, 'name': p.prop_name} for p in properties],
            'revenue_totals': revenue_totals,
            'expense_totals': expense_totals,
            'profit_totals': profit_totals,
            'selected_properties': selected_properties,
        })

    context = {
        'properties': properties,
        'all_properties': all_properties,  # For the property selector
        'revenue_line_types': revenue_line_types_list,
        'revenue_totals': revenue_totals,
        'revenue_totals_by_line': revenue_totals_by_line,
        'revenue_prop_totals': revenue_prop_totals,
        'expense_line_types': expense_line_types_list,
        'expense_totals': expense_totals,
        'expense_totals_by_line': expense_totals_by_line,
        'expense_prop_totals': expense_prop_totals,
        'profit_totals': profit_totals,
        'prop_values_map': prop_values_map,
        'total_current_value': total_current_value,
        'selected_properties': selected_properties,
    }
    
    return render(request, 'finance_pl.html', context)

@login_required
def finance_pl_act(request):
    # Get selected year from request (default to 'budget')
    selected_year = request.GET.get('year', 'budget')
    
    # Get selected properties from request
    selected_properties = request.GET.getlist('properties')

    # Handle year parameter - can be 'budget' or a year number
    if selected_year != 'budget':
        try:
            selected_year = int(selected_year)
            # Ensure only 2024 or 2025 is selectable
            if selected_year not in [2024, 2025]:
                selected_year = 'budget'
        except (ValueError, TypeError):
            selected_year = 'budget'
    
    # FULLY OPTIMIZED: Single query with comprehensive prefetching
    all_properties = props.objects.filter(prop_status="Active").select_related().prefetch_related(
        'prop_values_set',  # Simplified - no explicit queryset
        Prefetch('revenue_set', queryset=revenue.objects.select_related('revenue_line_types', 'revenue_types')),
        Prefetch('expense_set', queryset=expense.objects.select_related('expense_line_types', 'expense_types'))
    )
    
    # If no properties selected, default to ALL properties
    if not selected_properties:
        selected_properties = [str(prop.prop_id) for prop in all_properties]
    
    # Convert to integers and filter
    selected_prop_ids = [int(pid) for pid in selected_properties if pid.isdigit()]
    properties = all_properties.filter(prop_id__in=selected_prop_ids)
    
    # OPTIMIZED: Single queries for line types
    revenue_line_types_list = list(revenue_line_types.objects.all())
    expense_line_types_list = list(expense_line_types.objects.all())
    
    # OPTIMIZED: Use prefetched data instead of separate queries
    revenues = []
    expenses = []
    
    for prop in properties:
        revenues.extend(prop.revenue_set.all())
        expenses.extend(prop.expense_set.all())
    
    # ========= REVENUE SECTION ========= (optimized calculations)
    revenue_totals = {
        'jan': sum(r.revenue_jan or 0 for r in revenues),
        'feb': sum(r.revenue_feb or 0 for r in revenues),
        'mar': sum(r.revenue_mar or 0 for r in revenues),
        'apr': sum(r.revenue_apr or 0 for r in revenues),
        'may': sum(r.revenue_may or 0 for r in revenues),
        'jun': sum(r.revenue_jun or 0 for r in revenues),
        'jul': sum(r.revenue_jul or 0 for r in revenues),
        'aug': sum(r.revenue_aug or 0 for r in revenues),
        'sep': sum(r.revenue_sep or 0 for r in revenues),
        'oct': sum(r.revenue_oct or 0 for r in revenues),
        'nov': sum(r.revenue_nov or 0 for r in revenues),
        'dec': sum(r.revenue_dec or 0 for r in revenues),
    }
    revenue_totals['year'] = sum(revenue_totals.values())
    
    # OPTIMIZED: Pre-group revenues by line type
    revenues_by_line_type = {}
    for rev in revenues:
        line_type_id = rev.revenue_line_types.revenue_line_types_id
        if line_type_id not in revenues_by_line_type:
            revenues_by_line_type[line_type_id] = []
        revenues_by_line_type[line_type_id].append(rev)
    
    # Calculate revenue totals by line type
    revenue_totals_by_line = {'all': {}}
    for lt in revenue_line_types_list:
        line_revenues = revenues_by_line_type.get(lt.revenue_line_types_id, [])
        monthly_totals = {
            'jan': sum(r.revenue_jan or 0 for r in line_revenues),
            'feb': sum(r.revenue_feb or 0 for r in line_revenues),
            'mar': sum(r.revenue_mar or 0 for r in line_revenues),
            'apr': sum(r.revenue_apr or 0 for r in line_revenues),
            'may': sum(r.revenue_may or 0 for r in line_revenues),
            'jun': sum(r.revenue_jun or 0 for r in line_revenues),
            'jul': sum(r.revenue_jul or 0 for r in line_revenues),
            'aug': sum(r.revenue_aug or 0 for r in line_revenues),
            'sep': sum(r.revenue_sep or 0 for r in line_revenues),
            'oct': sum(r.revenue_oct or 0 for r in line_revenues),
            'nov': sum(r.revenue_nov or 0 for r in line_revenues),
            'dec': sum(r.revenue_dec or 0 for r in line_revenues),
        }
        monthly_totals['total'] = sum(monthly_totals.values())
        revenue_totals_by_line['all'][lt.revenue_line_types_id] = monthly_totals
    
    # OPTIMIZED: Pre-group revenues by property
    revenues_by_property = {}
    for rev in revenues:
        prop_id = rev.prop.prop_id
        if prop_id not in revenues_by_property:
            revenues_by_property[prop_id] = []
        revenues_by_property[prop_id].append(rev)
    
    # Calculate property-specific revenue totals
    revenue_prop_totals = {}
    for prop in properties:
        prop_revenues = revenues_by_property.get(prop.prop_id, [])
        monthly_totals = {
            'jan': sum(r.revenue_jan or 0 for r in prop_revenues),
            'feb': sum(r.revenue_feb or 0 for r in prop_revenues),
            'mar': sum(r.revenue_mar or 0 for r in prop_revenues),
            'apr': sum(r.revenue_apr or 0 for r in prop_revenues),
            'may': sum(r.revenue_may or 0 for r in prop_revenues),
            'jun': sum(r.revenue_jun or 0 for r in prop_revenues),
            'jul': sum(r.revenue_jul or 0 for r in prop_revenues),
            'aug': sum(r.revenue_aug or 0 for r in prop_revenues),
            'sep': sum(r.revenue_sep or 0 for r in prop_revenues),
            'oct': sum(r.revenue_oct or 0 for r in prop_revenues),
            'nov': sum(r.revenue_nov or 0 for r in prop_revenues),
            'dec': sum(r.revenue_dec or 0 for r in prop_revenues),
        }
        monthly_totals['year'] = sum(monthly_totals.values())
        revenue_prop_totals[prop.prop_id] = monthly_totals
        
        # Add property-specific revenue line type totals
        revenue_totals_by_line[prop.prop_id] = {}
        for lt in revenue_line_types_list:
            prop_line_revenues = [r for r in prop_revenues if r.revenue_line_types.revenue_line_types_id == lt.revenue_line_types_id]
            line_monthly_totals = {
                'jan': sum(r.revenue_jan or 0 for r in prop_line_revenues),
                'feb': sum(r.revenue_feb or 0 for r in prop_line_revenues),
                'mar': sum(r.revenue_mar or 0 for r in prop_line_revenues),
                'apr': sum(r.revenue_apr or 0 for r in prop_line_revenues),
                'may': sum(r.revenue_may or 0 for r in prop_line_revenues),
                'jun': sum(r.revenue_jun or 0 for r in prop_line_revenues),
                'jul': sum(r.revenue_jul or 0 for r in prop_line_revenues),
                'aug': sum(r.revenue_aug or 0 for r in prop_line_revenues),
                'sep': sum(r.revenue_sep or 0 for r in prop_line_revenues),
                'oct': sum(r.revenue_oct or 0 for r in prop_line_revenues),
                'nov': sum(r.revenue_nov or 0 for r in prop_line_revenues),
                'dec': sum(r.revenue_dec or 0 for r in prop_line_revenues),
            }
            line_monthly_totals['total'] = sum(line_monthly_totals.values())
            revenue_totals_by_line[prop.prop_id][lt.revenue_line_types_id] = line_monthly_totals

    # ========= EXPENSE SECTION ========= (optimized)
    # Initialize actual expense totals
    actual_expense_totals = {
        'jan': 0, 'feb': 0, 'mar': 0, 'apr': 0, 'may': 0, 'jun': 0,
        'jul': 0, 'aug': 0, 'sep': 0, 'oct': 0, 'nov': 0, 'dec': 0, 'year': 0
    }
    
    # FULLY OPTIMIZED: Actual expenses with single aggregate query
    actual_expense_prop_totals = {}
    if selected_year != 'budget':
        from django.db.models import Sum
        
        # Single query to get all actual expenses with month grouping
        actual_expenses_aggregated = act_expense.objects.filter(
            act_expense_date__year=selected_year,
            act_expense_approved="Yes",
            act_expense_paid="Yes",
            prop_id__in=selected_prop_ids
        ).values('prop_id', 'act_expense_date__month').annotate(
            total_amount=Sum('act_expense_amount')
        ).order_by('prop_id', 'act_expense_date__month')
        
        # Calculate monthly totals for all properties in one pass
        month_mapping = {
            1: 'jan', 2: 'feb', 3: 'mar', 4: 'apr', 5: 'may', 6: 'jun',
            7: 'jul', 8: 'aug', 9: 'sep', 10: 'oct', 11: 'nov', 12: 'dec'
        }
        
        # Initialize all property totals
        for prop in properties:
            actual_expense_prop_totals[prop.prop_id] = {month: 0 for month in month_mapping.values()}
            actual_expense_prop_totals[prop.prop_id]['year'] = 0
        
        # Process aggregated results
        for result in actual_expenses_aggregated:
            prop_id = result['prop_id']
            month_num = result['act_expense_date__month']
            amount = result['total_amount'] or 0
            
            if prop_id in actual_expense_prop_totals:
                month_name = month_mapping[month_num]
                actual_expense_prop_totals[prop_id][month_name] = amount
                actual_expense_prop_totals[prop_id]['year'] += amount
        
        # Calculate overall monthly totals
        for month_name in month_mapping.values():
            actual_expense_totals[month_name] = sum(
                actual_expense_prop_totals[prop.prop_id][month_name] 
                for prop in properties
            )
        
        actual_expense_totals['year'] = sum(actual_expense_totals.values())
    
    # Calculate budgeted expense totals
    expense_totals = {
        'jan': sum(e.expense_jan or 0 for e in expenses),
        'feb': sum(e.expense_feb or 0 for e in expenses),
        'mar': sum(e.expense_mar or 0 for e in expenses),
        'apr': sum(e.expense_apr or 0 for e in expenses),
        'may': sum(e.expense_may or 0 for e in expenses),
        'jun': sum(e.expense_jun or 0 for e in expenses),
        'jul': sum(e.expense_jul or 0 for e in expenses),
        'aug': sum(e.expense_aug or 0 for e in expenses),
        'sep': sum(e.expense_sep or 0 for e in expenses),
        'oct': sum(e.expense_oct or 0 for e in expenses),
        'nov': sum(e.expense_nov or 0 for e in expenses),
        'dec': sum(e.expense_dec or 0 for e in expenses),
    }
    expense_totals['year'] = sum(expense_totals.values())
    
    # OPTIMIZED: Pre-group expenses by line type
    expenses_by_line_type = {}
    for exp in expenses:
        line_type_id = exp.expense_line_types.expense_line_types_id
        if line_type_id not in expenses_by_line_type:
            expenses_by_line_type[line_type_id] = []
        expenses_by_line_type[line_type_id].append(exp)
    
    # Calculate expense totals by line type
    expense_totals_by_line = {'all': {}}
    for elt in expense_line_types_list:
        line_expenses = expenses_by_line_type.get(elt.expense_line_types_id, [])
        monthly_totals = {
            'jan': sum(e.expense_jan or 0 for e in line_expenses),
            'feb': sum(e.expense_feb or 0 for e in line_expenses),
            'mar': sum(e.expense_mar or 0 for e in line_expenses),
            'apr': sum(e.expense_apr or 0 for e in line_expenses),
            'may': sum(e.expense_may or 0 for e in line_expenses),
            'jun': sum(e.expense_jun or 0 for e in line_expenses),
            'jul': sum(e.expense_jul or 0 for e in line_expenses),
            'aug': sum(e.expense_aug or 0 for e in line_expenses),
            'sep': sum(e.expense_sep or 0 for e in line_expenses),
            'oct': sum(e.expense_oct or 0 for e in line_expenses),
            'nov': sum(e.expense_nov or 0 for e in line_expenses),
            'dec': sum(e.expense_dec or 0 for e in line_expenses),
        }
        monthly_totals['total'] = sum(monthly_totals.values())
        expense_totals_by_line['all'][elt.expense_line_types_id] = monthly_totals
    
    # OPTIMIZED: Pre-group expenses by property
    expenses_by_property = {}
    for exp in expenses:
        prop_id = exp.prop.prop_id
        if prop_id not in expenses_by_property:
            expenses_by_property[prop_id] = []
        expenses_by_property[prop_id].append(exp)
    
    # Calculate property-specific expense totals
    expense_prop_totals = {}
    for prop in properties:
        prop_expenses = expenses_by_property.get(prop.prop_id, [])
        monthly_totals = {
            'jan': sum(e.expense_jan or 0 for e in prop_expenses),
            'feb': sum(e.expense_feb or 0 for e in prop_expenses),
            'mar': sum(e.expense_mar or 0 for e in prop_expenses),
            'apr': sum(e.expense_apr or 0 for e in prop_expenses),
            'may': sum(e.expense_may or 0 for e in prop_expenses),
            'jun': sum(e.expense_jun or 0 for e in prop_expenses),
            'jul': sum(e.expense_jul or 0 for e in prop_expenses),
            'aug': sum(e.expense_aug or 0 for e in prop_expenses),
            'sep': sum(e.expense_sep or 0 for e in prop_expenses),
            'oct': sum(e.expense_oct or 0 for e in prop_expenses),
            'nov': sum(e.expense_nov or 0 for e in prop_expenses),
            'dec': sum(e.expense_dec or 0 for e in prop_expenses),
        }
        monthly_totals['year'] = sum(monthly_totals.values())
        expense_prop_totals[prop.prop_id] = monthly_totals
        
        # Add property-specific expense line type totals
        expense_totals_by_line[prop.prop_id] = {}
        for elt in expense_line_types_list:
            prop_line_expenses = [e for e in prop_expenses if e.expense_line_types.expense_line_types_id == elt.expense_line_types_id]
            line_monthly_totals = {
                'jan': sum(e.expense_jan or 0 for e in prop_line_expenses),
                'feb': sum(e.expense_feb or 0 for e in prop_line_expenses),
                'mar': sum(e.expense_mar or 0 for e in prop_line_expenses),
                'apr': sum(e.expense_apr or 0 for e in prop_line_expenses),
                'may': sum(e.expense_may or 0 for e in prop_line_expenses),
                'jun': sum(e.expense_jun or 0 for e in prop_line_expenses),
                'jul': sum(e.expense_jul or 0 for e in prop_line_expenses),
                'aug': sum(e.expense_aug or 0 for e in prop_line_expenses),
                'sep': sum(e.expense_sep or 0 for e in prop_line_expenses),
                'oct': sum(e.expense_oct or 0 for e in prop_line_expenses),
                'nov': sum(e.expense_nov or 0 for e in prop_line_expenses),
                'dec': sum(e.expense_dec or 0 for e in prop_line_expenses),
            }
            line_monthly_totals['total'] = sum(line_monthly_totals.values())
            expense_totals_by_line[prop.prop_id][elt.expense_line_types_id] = line_monthly_totals

    # ========= PROFIT CALCULATION =========
    if selected_year == 'budget':
        profit_totals = {
            'jan': revenue_totals['jan'] - expense_totals['jan'],
            'feb': revenue_totals['feb'] - expense_totals['feb'],
            'mar': revenue_totals['mar'] - expense_totals['mar'],
            'apr': revenue_totals['apr'] - expense_totals['apr'],
            'may': revenue_totals['may'] - expense_totals['may'],
            'jun': revenue_totals['jun'] - expense_totals['jun'],
            'jul': revenue_totals['jul'] - expense_totals['jul'],
            'aug': revenue_totals['aug'] - expense_totals['aug'],
            'sep': revenue_totals['sep'] - expense_totals['sep'],
            'oct': revenue_totals['oct'] - expense_totals['oct'],
            'nov': revenue_totals['nov'] - expense_totals['nov'],
            'dec': revenue_totals['dec'] - expense_totals['dec'],
            'year': revenue_totals['year'] - expense_totals['year']
        }
    else:
        profit_totals = {
            'jan': revenue_totals['jan'] - expense_totals['jan'] - actual_expense_totals['jan'],
            'feb': revenue_totals['feb'] - expense_totals['feb'] - actual_expense_totals['feb'],
            'mar': revenue_totals['mar'] - expense_totals['mar'] - actual_expense_totals['mar'],
            'apr': revenue_totals['apr'] - expense_totals['apr'] - actual_expense_totals['apr'],
            'may': revenue_totals['may'] - expense_totals['may'] - actual_expense_totals['may'],
            'jun': revenue_totals['jun'] - expense_totals['jun'] - actual_expense_totals['jun'],
            'jul': revenue_totals['jul'] - expense_totals['jul'] - actual_expense_totals['jul'],
            'aug': revenue_totals['aug'] - expense_totals['aug'] - actual_expense_totals['aug'],
            'sep': revenue_totals['sep'] - expense_totals['sep'] - actual_expense_totals['sep'],
            'oct': revenue_totals['oct'] - expense_totals['oct'] - actual_expense_totals['oct'],
            'nov': revenue_totals['nov'] - expense_totals['nov'] - actual_expense_totals['nov'],
            'dec': revenue_totals['dec'] - expense_totals['dec'] - actual_expense_totals['dec'],
            'year': revenue_totals['year'] - expense_totals['year'] - actual_expense_totals['year']
        }

    # FULLY OPTIMIZED: Property values using prefetched data - NO additional queries
    prop_values_map = {}
    total_current_value = 0
    
    for prop in properties:
        # Use the prefetched prop_values_set data
        prop_values_list = list(prop.prop_values_set.all())
        if prop_values_list:
            prop_values = prop_values_list[0]
            prop_values_map[prop.prop_id] = prop_values
            if prop_values.prop_values_current_value is not None:
                total_current_value += prop_values.prop_values_current_value
        else:
            prop_values_map[prop.prop_id] = None

    # Handle AJAX requests
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return JsonResponse({
            'success': True,
            'properties': [{'id': p.prop_id, 'name': p.prop_name} for p in properties],
            'revenue_totals': revenue_totals,
            'expense_totals': expense_totals,
            'actual_expense_totals': actual_expense_totals,
            'profit_totals': profit_totals,
            'selected_properties': selected_properties,
        })

    return render(request, 'finance_pl_act.html', {
        'properties': properties,
        'all_properties': all_properties,
        'revenue_line_types': revenue_line_types_list,
        'revenue_totals': revenue_totals,
        'revenue_totals_by_line': revenue_totals_by_line,
        'revenue_prop_totals': revenue_prop_totals,
        'expense_line_types': expense_line_types_list,
        'expense_totals': expense_totals,
        'expense_totals_by_line': expense_totals_by_line,
        'expense_prop_totals': expense_prop_totals,
        'profit_totals': profit_totals,
        'prop_values_map': prop_values_map,
        'total_current_value': total_current_value,
        'actual_expense_totals': actual_expense_totals,
        'actual_expense_prop_totals': actual_expense_prop_totals,
        'selected_year': selected_year,
        'selected_properties': selected_properties,
        'available_years': [2025, 2024],
    })

@login_required
def petty_cash_rep(request):
	import petty_cash
	rep_output = request.POST.get('d_e')
	if request.user.is_authenticated:
		email = request.user.email
		fname = request.user.first_name
	petty_cash.petty_cash(rep_output, email, fname)
	messages.success(request, "Report Created Successfully")
	return redirect('home')

@login_required
def lease_agreements(request):
	import print_lease
	prop = request.POST.get('propname')
	rep_output = request.POST.get('d_e')
	if request.user.is_authenticated:
		email = request.user.email
		fname = request.user.first_name
	print_lease.lease_report(prop, rep_output, email, fname)
	messages.success(request, "Report Created Successfully")
	return redirect('home')

@login_required
def title_deeds(request):
	import print_title
	prop = request.POST.get('propname')
	rep_output = request.POST.get('d_e')
	if request.user.is_authenticated:
		email = request.user.email
		fname = request.user.first_name
	print_title.title_report(prop, rep_output, email, fname)
	messages.success(request, "Report Created Successfully")
	return redirect('home')

@login_required
def prop_rep(request):
	import print_prop
	prop = request.POST.get('propname')
	rep_output = request.POST.get('d_e')
	if request.user.is_authenticated:
		email = request.user.email
		fname = request.user.first_name
	print_prop.prop_report(prop, rep_output, email, fname)
	messages.success(request, "Report Created Successfully")
	return redirect('home')

@login_required
def property_report(request, prop_id):
	today = date.today()
	property = get_object_or_404(props.objects.only(
    	'prop_id', 'prop_name', 'prop_address1', 'prop_address2', 'prop_suburb', 
		'prop_city', 'prop_province', 'prop_country', 'prop_pcode',
		'prop_floor_area', 'prop_year_built', 'prop_status',
		'prop_available_for_rent', 'prop_title_deed',
		'prop_title_deed_status', 'prop_electricity', 'prop_water',
		'prop_refuse', 'prop_property_tax', 'prop_sewerage', 'prop_insurance'
	), pk=prop_id)
	context = {
		'today': today,
		'property': property,
	}
	return render(request, 'property_report.html', context)

@login_required
def title_deed_report(request, prop_id):
    property = get_object_or_404(props, pk=prop_id)
    
    if not property.prop_title_deed:
        return JsonResponse({'error': 'No title deed available for this property'}, status=404)
    
    # Return JSON with the file URL and type
    return JsonResponse({
        'file_url': property.prop_title_deed.url,
        'file_name': property.prop_title_deed.name.split('/')[-1],
        'file_type': property.prop_title_deed.name.split('.')[-1].lower()
    })

@login_required
def lease_agreement_report(request, tenant_id):
	today = date.today()
	tenant_obj = get_object_or_404(tenant.objects.only(
		'tenant_id', 'prop_id', 'tenant_type', 'tenant_name', 'tenant_contact_person', 'tenant_contact_number', 
		'tenant_email', 'tenant_deposit', 'tenant_lease_start_date', 'tenant_lease_end_date',
		'tenant_rental_type', 'tenant_renewal', 'tenant_renewal_period',
		'tenant_rent', 'tenant_levies',
		'tenant_payment_terms', 'tenant_current', 'tenant_lease_agreement'
	), pk=tenant_id)
	property = get_object_or_404(props.objects.only(
		'prop_id', 'prop_name', 'prop_address1', 'prop_address2', 'prop_suburb', 
		'prop_city', 'prop_province', 'prop_country', 'prop_pcode',
		'prop_floor_area', 'prop_year_built', 'prop_status',
		'prop_available_for_rent', 'prop_title_deed',
		'prop_title_deed_status', 'prop_electricity', 'prop_water',
		'prop_refuse', 'prop_property_tax', 'prop_sewerage', 'prop_insurance'
	), pk=tenant_obj.prop_id)
	context = {
		'today': today,
		'tenant': tenant_obj,
		'property': property,
	}
	return render(request, 'lease_agreement_report.html', context)

@login_required
def tenant_report(request, tenant_id):
	today = date.today()
	tenant_obj = get_object_or_404(tenant.objects.only(
		'tenant_id', 'prop_id', 'tenant_type', 'tenant_name', 'tenant_contact_person', 'tenant_contact_number', 
		'tenant_email', 'tenant_deposit', 'tenant_lease_start_date', 'tenant_lease_end_date',
		'tenant_rental_type', 'tenant_renewal', 'tenant_renewal_period',
		'tenant_rent', 'tenant_levies',
		'tenant_payment_terms', 'tenant_current', 'tenant_lease_agreement'
	), pk=tenant_id)
	context = {
		'today': today,
		'tenant': tenant_obj,
	}
	return render(request, 'tenant_report.html', context)

@login_required
def supplier_report(request, supplier_id):
	today = date.today()
	supplier_obj = get_object_or_404(supplier.objects.only(
		'supplier_id', 'supplier_contact_person', 'supplier_contact_number', 
		'supplier_email', 'supplier_company_name', 'supplier_role',
		'supplier_country'
	), pk=supplier_id)
	context = {
		'today': today,
		'supplier': supplier_obj,
	}
	return render(request, 'supplier_report.html', context)

@login_required
def tenant_rep(request):
	import print_tenant
	prop = request.POST.get('propname')
	rep_output = request.POST.get('d_e')
	if request.user.is_authenticated:
		email = request.user.email
		fname = request.user.first_name
	print_tenant.tenant_report(prop, rep_output, email, fname)
	messages.success(request, "Report Created Successfully")
	return redirect('home')

@login_required
def suppliers_rep(request):
	import print_supplier
	sup = request.POST.get('supname')
	rep_output = request.POST.get('d_e')
	if request.user.is_authenticated:
		email = request.user.email
		fname = request.user.first_name
	print_supplier.supplier_report(sup, rep_output, email, fname)
	messages.success(request, "Report Created Successfully")
	return redirect('home')

@login_required
def fsr_rep(request):
	import fsr
	rep_type = request.POST.get('d_s')
	rep_output = request.POST.get('d_e')
	rep_date = date.today()
	if request.user.is_authenticated:
		email = request.user.email
		fname = request.user.first_name
	fsr.fsr_report(rep_type, rep_date, rep_output, email, fname)
	messages.success(request, "Report Created Successfully")
	return redirect('home')

@login_required
def friday_status_report(request):
    from django.db import connection
    from django.db.utils import OperationalError, InterfaceError
    import time
    
    # Close any stale connections before starting
    connection.close()
    
    max_retries = 3
    for attempt in range(max_retries):
        try:
            today = date.today()
            
            # Get max_comments parameter for summarized reports
            max_comments = request.GET.get('max_comments', None)
            is_summarized_report = max_comments is not None
            
            # Store report type in session for use by fsr_notification
            if is_summarized_report:
                try:
                    max_comments = int(max_comments)
                    request.session['last_report_type'] = 'summarized'
                    request.session['max_comments'] = max_comments
                except (ValueError, TypeError):
                    max_comments = None
                    is_summarized_report = False
                    request.session['last_report_type'] = 'detailed'
                    request.session.pop('max_comments', None)
            else:
                request.session['last_report_type'] = 'detailed'
                request.session.pop('max_comments', None)
            
            # OPTIMIZED APPROACH: Process data in smaller chunks
            
            # 1. Get properties first (lightweight query)
            properties = list(props.objects.all().order_by('prop_country', 'prop_name').values('prop_id', 'prop_name'))
            
            # 2. Get issues in smaller batches
            all_issues = []
            batch_size = 50  # Process 50 issues at a time
            
            # Get total count first
            total_issues = issues.objects.count()
            
            for offset in range(0, total_issues, batch_size):
                # Get a batch of issues with minimal related data
                issues_batch = issues.objects.select_related('prop').filter(
                ).order_by('issues_id')[offset:offset + batch_size]
                
                for issue_obj in issues_batch:
                    # Build basic issue data
                    issue_dict = {
                        'prop_name': issue_obj.prop.prop_name,
                        'issues_id': issue_obj.issues_id,
                        'issues_heading': issue_obj.issues_heading,
                        'issues_description': issue_obj.issues_description,
                        'issues_status': issue_obj.issues_status,
                        'issues_date_logged': issue_obj.issues_date_logged,
                        'issues_resolution_date': issue_obj.issues_resolution_date,
                        'days_to_resolve': None,
                        'days_open': None,
                        'details': [],
                        'has_more_comments': False,
                        'total_comments': 0
                    }
                    
                    # Calculate days metrics
                    if issue_dict['issues_date_logged']:
                        if issue_dict['issues_status'] == 'Resolved':
                            if (issue_dict['issues_resolution_date'] and 
                                issue_dict['issues_resolution_date'] != date(1900, 1, 1)):
                                issue_dict['days_to_resolve'] = (issue_dict['issues_resolution_date'] - issue_dict['issues_date_logged']).days
                        else:
                            issue_dict['days_open'] = (today - issue_dict['issues_date_logged']).days
                    
                    # Get details separately for this issue
                    if is_summarized_report and max_comments:
                        # For summarized reports, get limited details
                        details_queryset = issues_details.objects.filter(
                            issues_id=issue_obj.issues_id
                        ).order_by('-issues_details_id')[:max_comments + 1]  # Get one extra to check if there are more
                        
                        details_list = list(details_queryset)
                        
                        if len(details_list) > max_comments:
                            # There are more comments than the limit
                            issue_dict['details'] = [{
                                'issues_details_id': detail.issues_details_id,
                                'issues_details_comment': detail.issues_details_comment,
                                'issues_details_user': detail.issues_details_user,
                                'issues_details_date': detail.issues_details_date
                            } for detail in details_list[:max_comments]]
                            issue_dict['has_more_comments'] = True
                            issue_dict['total_comments'] = issues_details.objects.filter(issues_id=issue_obj.issues_id).count()
                        else:
                            issue_dict['details'] = [{
                                'issues_details_id': detail.issues_details_id,
                                'issues_details_comment': detail.issues_details_comment,
                                'issues_details_user': detail.issues_details_user,
                                'issues_details_date': detail.issues_details_date
                            } for detail in details_list]
                            issue_dict['has_more_comments'] = False
                            issue_dict['total_comments'] = len(details_list)
                    else:
                        # For detailed reports, get all details
                        details_queryset = issues_details.objects.filter(
                            issues_id=issue_obj.issues_id
                        ).order_by('-issues_details_id')
                        
                        issue_dict['details'] = [{
                            'issues_details_id': detail.issues_details_id,
                            'issues_details_comment': detail.issues_details_comment,
                            'issues_details_user': detail.issues_details_user,
                            'issues_details_date': detail.issues_details_date
                        } for detail in details_queryset]
                        issue_dict['has_more_comments'] = False
                        issue_dict['total_comments'] = len(issue_dict['details'])
                    
                    all_issues.append(issue_dict)
                
                # Small pause between batches to prevent overwhelming the DB
                time.sleep(0.1)
            
            # 3. Process data by status and property
            processed_data = {}
            cutoff_date = today - timedelta(days=7)
            
            for status in ['Resolved', 'Unresolved', 'Issue']:
                processed_data[status] = {}
                for prop in properties:
                    prop_name = prop['prop_name']
                    processed_data[status][prop_name] = []
                    
                    unique_issues = set()
                    
                    for issue in all_issues:
                        if (issue['prop_name'] == prop_name and 
                            issue['issues_status'] == status and 
                            (issue['issues_heading'], issue['issues_description']) not in unique_issues):
                            
                            # For Resolved issues, check if:
                            # 1. Resolved within last 7 days, OR
                            # 2. Has a comment added within last 7 days
                            if status == 'Resolved':
                                show_issue = False
                                
                                # Check if resolved within last 7 days
                                if (issue['issues_resolution_date'] and
                                    issue['issues_resolution_date'] != date(1900, 1, 1) and 
                                    issue['issues_resolution_date'] >= cutoff_date):
                                    show_issue = True
                                
                                # Check if any comment was added within last 7 days
                                if not show_issue and issue['details']:
                                    for detail in issue['details']:
                                        if (detail['issues_details_date'] and 
                                            detail['issues_details_date'] >= cutoff_date):
                                            show_issue = True
                                            break
                                
                                if show_issue:
                                    processed_data[status][prop_name].append(issue)
                                    unique_issues.add((issue['issues_heading'], issue['issues_description']))
                            else:
                                processed_data[status][prop_name].append(issue)
                                unique_issues.add((issue['issues_heading'], issue['issues_description']))
            
            # 4. Build context
            context = {
                'today': today,
                'statuses': ['Resolved', 'Unresolved', 'Issue'],
                'properties': [{'prop_name': prop['prop_name']} for prop in properties],
                'is_summarized_report': is_summarized_report,
                'max_comments': max_comments,
                'status_groups': [
                    {
                        'status': status,
                        'property_issues': [
                            {
                                'prop_name': prop['prop_name'],
                                'issues': processed_data[status][prop['prop_name']]
                            }
                            for prop in properties
                            if processed_data[status][prop['prop_name']]
                        ]
                    }
                    for status in ['Resolved', 'Unresolved', 'Issue']
                ]
            }
            
            return render(request, 'friday_status_report.html', context)
            
        except (OperationalError, InterfaceError) as e:
            if attempt < max_retries - 1:
                connection.close()
                time.sleep(3)  # Increased wait time
                continue
            else:
                messages.error(request, "Database connection error. Please try again in a moment.")
                return redirect('fsr')
        except Exception as e:
            messages.error(request, f"An error occurred while generating the report: {str(e)}")
            return redirect('fsr')

@login_required
def resolved_issues_report(request):
    # Get dates from GET parameters
    f_date_str = request.GET.get('f_date')
    t_date_str = request.GET.get('t_date')

    # Validate dates
    if not f_date_str or not t_date_str:
        messages.error(request, "Both date ranges are required")
        return redirect('fsr')

    try:
        f_date = parse_date(f_date_str)
        t_date = parse_date(t_date_str)
        
        if not f_date or not t_date:
            raise ValueError("Invalid date format")
            
        if t_date < f_date:
            messages.error(request, "End date cannot be before start date")
            return redirect('fsr')

    except (ValueError, TypeError) as e:
        messages.error(request, f"Invalid date format: {str(e)}")
        return redirect('fsr')

    try:
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT 
                    prop.prop_name, 
                    issues.issues_heading, 
                    issues.issues_description, 
                    issues.issues_status,
                    issues_details.issues_details_comment,
                    issues_details.issues_details_user,
                    issues_details.issues_details_date,
                    issues.issues_resolution_date,
                    issues.issues_date_logged
                FROM 
                    prop
                    JOIN issues ON prop.prop_id = issues.prop_id
                    JOIN issues_details ON issues.issues_id = issues_details.issues_id
                WHERE 
                    issues.issues_status = 'Resolved'
                    AND issues.issues_resolution_date BETWEEN %s AND %s
                ORDER BY 
                    prop.prop_name ASC,
                    issues.issues_heading ASC,
                    issues_details.issues_details_date DESC
            """, [f_date_str, t_date_str])

            rows = cursor.fetchall()

        # Helper function to parse dates
        def parse_db_date(date_value):
            if isinstance(date_value, date):
                return date_value
            elif isinstance(date_value, str):
                return datetime.strptime(date_value, '%Y-%m-%d').date()
            elif isinstance(date_value, datetime):
                return date_value.date()
            else:
                raise ValueError(f"Unsupported date format: {type(date_value)}")

        # Structure the data
        properties = defaultdict(lambda: {
            'prop_name': '',
            'issues': defaultdict(list)
        })

        for row in rows:
            prop_name = row[0]
            issue_heading = row[1]
            
            try:
                resolution_date = parse_db_date(row[7])
                date_logged = parse_db_date(row[8])
                days_to_resolve = (resolution_date - date_logged).days
            except Exception as e:
                days_to_resolve = 0  # Default value if date parsing fails

            properties[prop_name]['prop_name'] = prop_name
            properties[prop_name]['issues'][issue_heading].append({
                'issues_description': row[2],
                'comment': row[4],
                'user': row[5],
                'comment_date': row[6],
                'resolution_date': row[7],
                'date_logged': row[8],
                'days_to_resolve': days_to_resolve
            })

        # Convert to list format for template
        properties_list = []
        for prop_name, prop_data in properties.items():
            issues_list = []
            for issue_heading, comments in prop_data['issues'].items():
                issues_list.append({
                    'heading': issue_heading,
                    'description': comments[0]['issues_description'],
                    'issues_date_logged': comments[0]['date_logged'],
                    'issues_resolution_date': comments[0]['resolution_date'],
                    'days_to_resolve': comments[0]['days_to_resolve'],
                    'comments': sorted(comments, key=lambda x: x['comment_date'], reverse=True)[:20]
                })

            properties_list.append({
                'prop_name': prop_name,
                'issues': sorted(issues_list, key=lambda x: x['heading'])
            })

        context = {
            'f_date': f_date_str,
            't_date': t_date_str,
            'properties': sorted(properties_list, key=lambda x: x['prop_name'])
        }

        return render(request, 'resolved_issues_report.html', context)

    except Exception as e:
        messages.error(request, f"Error generating report: {str(e)}")
        return redirect('fsr')

@login_required
def issues_rep(request):
	import issues
	f_d = request.POST.get('from_date')
	f_date = datetime.strptime(f_d, "%Y-%m-%d")
	from_date = f_date.date()
	t_d = request.POST.get('to_date')
	t_date = datetime.strptime(t_d, "%Y-%m-%d")
	to_date = t_date.date()
	rep_output = request.POST.get('d_e')
	rep_date = date.today()
	if request.user.is_authenticated:
		email = request.user.email
		fname = request.user.first_name
	issues.issues_report(from_date, to_date, rep_output, email, fname)
	messages.success(request, "Report Created Successfully")
	return redirect('home')

@login_required
def open_invoices(request):
	import open_invoices
	rep_output = request.POST.get('d_e')
	check = 'No'
	if request.user.is_authenticated:
		email = request.user.email
		fname = request.user.first_name
	open_invoices.open_invoices(rep_output, check, email, fname)
	messages.success(request, "Report Created Successfully")
	return redirect('home')

@login_required
def open_invoices_report(request):
    
    today = date.today()
    properties_with_invoices = []
    
    # Get all current tenants with their property details
    current_tenants = tenant.objects.filter(
        tenant_current='Yes'
    ).select_related('prop').order_by('prop__prop_country', 'prop__prop_name')
    
    # Get all unpaid invoices with tenant details in one query
    unpaid_invoices = invoices.objects.filter(
        invoice_paid='No'
    ).select_related('tenant', 'tenant__prop').order_by('invoice_date')
    
    # Process detailed invoice breakdown
    for tenant_obj in current_tenants:
        tenant_invoices = []
        
        # Get unpaid invoices for this tenant
        tenant_unpaid_invoices = [inv for inv in unpaid_invoices if inv.tenant.tenant_id == tenant_obj.tenant_id]
        
        for invoice_obj in tenant_unpaid_invoices:
            payment_terms = tenant_obj.tenant_payment_terms or 0
            due_date = invoice_obj.invoice_date + timedelta(days=payment_terms)
            days_overdue = (today - due_date).days if today > due_date else 0
            
            tenant_invoices.append({
                'invoice_id': invoice_obj.invoice_id,
                'invoice_date': invoice_obj.invoice_date.strftime('%Y-%m-%d'),
                'due_date': due_date.strftime('%Y-%m-%d'),
                'days_overdue': days_overdue,
                'overdue': days_overdue > 0
            })
        
        # Only include tenants with unpaid invoices
        if tenant_invoices:
            properties_with_invoices.append({
                'prop_name': tenant_obj.prop.prop_name,
                'prop_country': tenant_obj.prop.prop_country,
                'tenant_id': tenant_obj.tenant_id,
                'tenant_name': tenant_obj.tenant_name,
                'tenant_contact_person': tenant_obj.tenant_contact_person,
                'tenant_contact_number': tenant_obj.tenant_contact_number,
                'tenant_email': tenant_obj.tenant_email,
                'tenant_rent': tenant_obj.tenant_rent,
                'tenant_payment_terms': tenant_obj.tenant_payment_terms,
                'invoices': tenant_invoices
            })
    
    # Calculate Debtors Age Analysis
    debtors_age_analysis = []
    totals = {
        'total_outstanding': 0,
        'current_0_30': 0,
        'past_due_31_60': 0,
        'past_due_61_90': 0,
        'past_due_91_plus': 0
    }
    
    for tenant_obj in current_tenants:
        tenant_analysis = {
            'tenant_name': tenant_obj.tenant_name,
            'tenant_id': tenant_obj.tenant_id,  # Add tenant_id here too
            'total_outstanding': 0,
            'current_0_30': 0,
            'past_due_31_60': 0,
            'past_due_61_90': 0,
            'past_due_91_plus': 0
        }
        
        # Get unpaid invoices for this tenant
        tenant_unpaid_invoices = [inv for inv in unpaid_invoices if inv.tenant.tenant_id == tenant_obj.tenant_id]
        
        # Calculate aging for this tenant's invoices
        for invoice_obj in tenant_unpaid_invoices:
            payment_terms = tenant_obj.tenant_payment_terms or 0
            due_date = invoice_obj.invoice_date + timedelta(days=payment_terms)
            days_overdue = (today - due_date).days if today > due_date else 0
            amount = float(tenant_obj.tenant_rent or 0)
            
            tenant_analysis['total_outstanding'] += amount
            
            if days_overdue <= 30:
                # Current (0-30 days - includes not yet due and up to 30 days overdue)
                tenant_analysis['current_0_30'] += amount
            elif 31 <= days_overdue <= 60:
                # Past due 31-60 days
                tenant_analysis['past_due_31_60'] += amount
            elif 61 <= days_overdue <= 90:
                # Past due 61-90 days
                tenant_analysis['past_due_61_90'] += amount
            else:
                # Past due 91+ days
                tenant_analysis['past_due_91_plus'] += amount
        
        # Only include tenants with outstanding invoices
        if tenant_analysis['total_outstanding'] > 0:
            debtors_age_analysis.append(tenant_analysis)
            
            # Add to totals
            totals['total_outstanding'] += tenant_analysis['total_outstanding']
            totals['current_0_30'] += tenant_analysis['current_0_30']
            totals['past_due_31_60'] += tenant_analysis['past_due_31_60']
            totals['past_due_61_90'] += tenant_analysis['past_due_61_90']
            totals['past_due_91_plus'] += tenant_analysis['past_due_91_plus']
    
    # Sort debtors by total outstanding (highest first)
    debtors_age_analysis.sort(key=lambda x: x['total_outstanding'], reverse=True)

    context = {
        'today': today.strftime('%Y-%m-%d'),
        'properties_with_invoices': properties_with_invoices,
        'debtors_age_analysis': debtors_age_analysis,
        'totals': totals
    }
    
    return render(request, 'open_invoices_report.html', context)

@login_required
def lease_renewal_report(request):
    
    today = date.today()
    tenants_for_renewal = []
    vacant_properties = []
    declined_renewals = []
    
    # Get all active tenants with their property details using select_related for efficiency
    active_tenants = tenant.objects.filter(
        tenant_current='Yes'
    ).select_related('prop').order_by('prop__prop_country', 'prop__prop_name')
    
    # Get list of property names that have active tenants
    prop_active_tenant = list(active_tenants.values_list('prop__prop_name', flat=True))
    
    # Get all active properties available for rent
    active_properties = props.objects.filter(
        prop_status='Active',
        prop_available_for_rent='Yes'
    ).order_by('prop_country', 'prop_name')
    
    # Process each active tenant for renewal logic
    for tenant_obj in active_tenants:
        lease_end_date = tenant_obj.tenant_lease_end_date
        renewal_period = tenant_obj.tenant_renewal_period or 30  # Default to 30 days if None
        
        if lease_end_date:  # Make sure lease_end_date exists
            renewal_date = lease_end_date - timedelta(days=renewal_period)
            warning_date = renewal_date
#           This was for the old notification which was 30 days before the renewal date
#           warning_date = renewal_date - timedelta(days=30)
            renewal_status = tenant_obj.tenant_renewal_status or 'pending'  # Default to pending
            
            if today >= warning_date:
                if renewal_status == 'pending':
                    # Normal renewal case - add to tenants list
                    tenants_for_renewal.append({
                        'prop_name': tenant_obj.prop.prop_name,
                        'prop_country': tenant_obj.prop.prop_country,
                        'tenant_type': tenant_obj.tenant_type,
                        'tenant_name': tenant_obj.tenant_name,
                        'tenant_contact_person': tenant_obj.tenant_contact_person,
                        'tenant_contact_number': tenant_obj.tenant_contact_number,
                        'tenant_email': tenant_obj.tenant_email,
                        'tenant_deposit': tenant_obj.tenant_deposit,
                        'tenant_lease_start_date': tenant_obj.tenant_lease_start_date.strftime('%Y-%m-%d') if tenant_obj.tenant_lease_start_date else '',
                        'tenant_lease_end_date': tenant_obj.tenant_lease_end_date.strftime('%Y-%m-%d') if tenant_obj.tenant_lease_end_date else '',
                        'tenant_rental_type': tenant_obj.tenant_rental_type,
                        'tenant_renewal': tenant_obj.tenant_renewal,
                        'tenant_renewal_period': tenant_obj.tenant_renewal_period,
                        'tenant_rent': tenant_obj.tenant_rent,
                        'tenant_levies': tenant_obj.tenant_levies,
                        'tenant_payment_terms': tenant_obj.tenant_payment_terms,
                        'renewal_date': renewal_date.strftime('%Y-%m-%d'),
                        'needs_renewal': True
                    })
                elif renewal_status == 'declined':
                    # Tenant declined renewal - add to declined_renewals list
                    declined_renewals.append({
                        'prop_name': tenant_obj.prop.prop_name,
                        'tenant_name': tenant_obj.tenant_name,
                        'lease_end_date': tenant_obj.tenant_lease_end_date.strftime('%Y-%m-%d') if tenant_obj.tenant_lease_end_date else '',
                        'message': 'CURRENT TENANT NOT RENEWING LEASE - NEED NEW TENANT'
                    })
                # If renewal_status == 'new_lease_signed', do nothing (exclude from report)
    
    # Find vacant properties (properties without active tenants)
    vacant_properties = []
    for prop in active_properties:
        if prop.prop_name not in prop_active_tenant:
            vacant_properties.append({
                'prop_name': prop.prop_name,
                'prop_country': prop.prop_country
            })
    
    context = {
        'tenants': tenants_for_renewal,
        'vacant_properties': vacant_properties,
        'declined_renewals': declined_renewals,
        'today': today.strftime('%Y-%m-%d')
    }
    return render(request, 'lease_renewal_report.html', context)

@login_required
def lease_renewal(request):
	import lease_renewal
	rep_output = request.POST.get('d_e')
	check = 'No'
	if request.user.is_authenticated:
		email = request.user.email
		fname = request.user.first_name
	lease_renewal.lease_renewal(rep_output, check, email, fname)
	messages.success(request, "Report Created Successfully")
	return redirect('home')


### USER ADMIN AND LOGIN AND LOGOUT ###
def login_user(request):
	if request.method =="POST":
	    username = request.POST["username"]
	    password = request.POST["password"]
	    user = authenticate(request, username=username, password=password)
	    if user is not None:
	        login(request, user)
	        messages.success(request, ('You Have Successfully Logged In.'))
	        return redirect('home')
	    else:
	        messages.success(request, ('Error Logging In - Please Try Again !!'))
	        return redirect('login')
	else:
		return render(request, 'login.html', {})

def logout_user(request):
    logout(request)
    messages.success(request, ('You Have Succefully Logged Out.'))
    return redirect('home')

### RECIPE MANAGEMENT ###
# HELPER FUNCTION

def convert_to_decimal(quantity_str):
    """Convert fractions like '1/4' to Decimal(0.25)"""
    try:
        quantity_str = str(quantity_str).strip()
        
        if '/' in quantity_str:
            if ' ' in quantity_str:  # Mixed number like "1 1/2"
                parts = quantity_str.split(' ')
                whole = Decimal(parts[0])
                frac = Fraction(parts[1])
                return whole + Decimal(frac.numerator) / Decimal(frac.denominator)
            else:  # Simple fraction like "1/4"
                frac = Fraction(quantity_str)
                return Decimal(frac.numerator) / Decimal(frac.denominator)
        else:
            return Decimal(quantity_str)
    except Exception as e:
        print(f"Error converting '{quantity_str}': {e}")
        return Decimal("1")

def format_quantity(quantity_str):
    """Format for display - keep fractions as-is"""
    try:
        quantity_str = str(quantity_str).strip()
        
        # Keep fractions as fractions
        if '/' in quantity_str:
            return quantity_str
        
        # Remove trailing zeros from decimals
        num = float(quantity_str)
        return '{:g}'.format(num)
        
    except:
        return str(quantity_str)

def get_or_create_ingredient(name):
    """Get or create an ingredient by name (case-insensitive)"""
    name = name.strip()
    # Capitalize each word
    name = ' '.join(word.capitalize() for word in name.split())
    
    ingredient, created = Ingredient.objects.get_or_create(
        name__iexact=name,
        defaults={'name': name}
    )
    return ingredient


def get_or_create_unit(name):
    """Get or create a measurement unit by name (case-insensitive)"""
    name = name.strip().lower()
    
    # Try to find existing
    unit = MeasurementUnit.objects.filter(name__iexact=name).first()
    if unit:
        return unit
    
    # Create new with abbreviation
    abbr = name[:5] if len(name) <= 5 else name[:4] + '.'
    unit = MeasurementUnit.objects.create(
        name=name,
        abbreviation=abbr,
        unit_type='other'
    )
    return unit


def get_or_create_preparation(name):
    """Get or create a preparation method by name (case-insensitive)"""
    if not name or not name.strip():
        return None
    
    name = name.strip().lower()
    
    prep, created = PreparationMethod.objects.get_or_create(
        name__iexact=name,
        defaults={'name': name}
    )
    return prep

# RECIPE MANAGEMENT

@login_required
@require_POST
def send_shopping_list(request):
    """Generate shopping list with unit conversion - DEBUG VERSION"""
    if not request.user.is_superuser:
        return JsonResponse({'success': False, 'error': 'Permission denied'}, status=403)
    
    try:
        data = json.loads(request.body)
        
        recipe_id = data.get('recipe_id')
        servings = data.get('servings')
        original_servings = data.get('original_servings')
        
        # Get recipe
        recipe = Recipe.objects.prefetch_related(
            'recipe_ingredients__ingredient__category',
            'recipe_ingredients__ingredient__default_unit',
            'recipe_ingredients__unit'
        ).get(recipe_id=recipe_id)
        
        servings_multiplier = Decimal(servings) / Decimal(original_servings)
        
        # Build shopping list
        shopping_list_categorized = {}
        
        print("\n" + "="*80)
        print(f"DEBUGGING SHOPPING LIST FOR: {recipe.recipe_name}")
        print("="*80)
        
        for recipe_ingredient in recipe.recipe_ingredients.all():
            ingredient = recipe_ingredient.ingredient
            quantity = Decimal(recipe_ingredient.amount or 0) * servings_multiplier
            from_unit = recipe_ingredient.unit
            shopping_unit = ingredient.default_unit
            
            print(f"\nIngredient: {ingredient.name}")
            print(f"  Recipe amount: {quantity} {from_unit.name if from_unit else 'None'}")
            print(f"  Shopping unit: {shopping_unit.name if shopping_unit else 'NOT SET ❌'}")
            
            if not from_unit:
                continue
            
            # Determine final unit and amount
            if not shopping_unit:
                print(f"  ❌ No shopping unit - using recipe unit")
                final_unit = from_unit
                final_amount = quantity
            elif from_unit.measurement_unit_id == shopping_unit.measurement_unit_id:
                print(f"  ✅ Same unit - no conversion needed")
                final_unit = shopping_unit
                final_amount = quantity
            else:
                print(f"  🔄 Trying conversion: {from_unit.name} → {shopping_unit.name}")
                converted_qty, _ = convert_quantity(quantity, from_unit, shopping_unit)
                if converted_qty is not None:
                    print(f"  ✅ Conversion SUCCESS: {quantity} {from_unit.name} = {converted_qty} {shopping_unit.name}")
                    final_unit = shopping_unit
                    final_amount = converted_qty
                else:
                    print(f"  ❌ Conversion FAILED - using recipe unit")
                    final_unit = from_unit
                    final_amount = quantity
            
            print(f"  Final: {final_amount} {final_unit.name}")
            
            # Format amount
            qty = float(final_amount)
            if qty % 1 == 0:
                qty_str = f"{int(qty)}"
            else:
                qty_str = f"{qty:.2f}".rstrip('0').rstrip('.')
            
            # Get category
            category = ingredient.category.name if ingredient.category else 'Other'
            
            if category not in shopping_list_categorized:
                shopping_list_categorized[category] = []
            
            shopping_list_categorized[category].append(
                f"{qty_str} {final_unit.name} {ingredient.name}"
            )
        
        print("\n" + "="*80 + "\n")
        
        return JsonResponse({
            'success': True,
            'shopping_list_categorized': shopping_list_categorized,
            'recipe_name': recipe.recipe_name,
            'servings': servings,
            'original_servings': original_servings
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

@login_required
def recipe_management(request):
    """Recipe management page with multi-select filtering, A-Z filter, and pagination"""
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to access this page.')
        return redirect('home')
    
    # Handle delete action
    if request.method == 'POST' and request.POST.get('action') == 'delete':
        recipe_id = request.POST.get('recipe_id')
        try:
            recipe = Recipe.objects.get(recipe_id=recipe_id)
            recipe_name = recipe.recipe_name
            recipe.delete()
            messages.success(request, f'Recipe "{recipe_name}" has been deleted successfully.')
        except Recipe.DoesNotExist:
            messages.error(request, 'Recipe not found.')
        return redirect('recipe_management')
    
    # Get all recipes with prefetch
    recipes = Recipe.objects.all().prefetch_related('courses', 'categories', 'proteins')
    
    # Get filter parameters
    search_query = request.GET.get('search', '')
    selected_courses = request.GET.getlist('course')
    selected_categories = request.GET.getlist('category')
    selected_proteins = request.GET.getlist('protein')
    selected_authors = request.GET.getlist('author')
    selected_letter = request.GET.get('letter', '')
    
    # Apply filters
    if search_query:
        recipes = recipes.filter(recipe_name__icontains=search_query)
    
    if selected_courses:
        recipes = recipes.filter(courses__recipe_course_id__in=selected_courses)
    
    if selected_categories:
        recipes = recipes.filter(categories__recipe_category_id__in=selected_categories)
    
    if selected_proteins:
        if 'vegetarian' in selected_proteins:
            protein_ids = [p for p in selected_proteins if p != 'vegetarian']
            if protein_ids:
                recipes = recipes.filter(
                    Q(is_vegetarian=True) | 
                    Q(proteins__custom_protein_id__in=protein_ids)
                )
            else:
                recipes = recipes.filter(is_vegetarian=True)
        else:
            recipes = recipes.filter(proteins__custom_protein_id__in=selected_proteins)
    
    if selected_authors:
        recipes = recipes.filter(author__in=selected_authors)
    
    recipes = recipes.distinct().order_by('recipe_name')
    
    # Calculate available letters BEFORE applying letter filter
    # This allows users to switch between letters without deselecting first
    all_letters = list(string.ascii_uppercase)
    available_letters = set()
    
    # Get first letter of each recipe name (without letter filter applied)
    for recipe in recipes:
        if recipe.recipe_name:
            first_letter = recipe.recipe_name[0].upper()
            if first_letter.isalpha():
                available_letters.add(first_letter)
    
    # Create letter data for template
    letter_data = []
    for letter in all_letters:
        letter_data.append({
            'letter': letter,
            'available': letter in available_letters
        })
    
    # NOW apply letter filter (after calculating available letters)
    # This allows users to see all available letters and click them to switch
    if selected_letter:
        recipes = recipes.filter(recipe_name__istartswith=selected_letter)
    
    # Handle AJAX request for Load More
    page = request.GET.get('page', 1)
    is_ajax = request.headers.get('X-Requested-With') == 'XMLHttpRequest'
    
    # Pagination: 48 recipes per page
    paginator = Paginator(recipes, 48)
    page_obj = paginator.get_page(page)
    
    if is_ajax:
        # Return JSON for AJAX requests
        recipes_data = []
        for recipe in page_obj:
            recipes_data.append({
                'recipe_id': recipe.recipe_id,
                'recipe_name': recipe.recipe_name,
                'recipe_image': recipe.recipe_image.url if recipe.recipe_image else None,
                'prep_time': recipe.prep_time,
                'cook_time': recipe.cook_time,
                'servings': recipe.servings,
                'difficulty_level': recipe.difficulty_level,
                'is_vegetarian': recipe.is_vegetarian,
                'author': recipe.author,
                'courses': [{'name': c.name} for c in recipe.courses.all()],
                'categories': [{'name': c.name} for c in recipe.categories.all()],
                'proteins': [{'name': p.name} for p in recipe.proteins.all()],
            })
        
        return JsonResponse({
            'recipes': recipes_data,
            'has_next': page_obj.has_next(),
            'next_page': page_obj.next_page_number() if page_obj.has_next() else None,
            'total_count': paginator.count,
        })
    
    # Get all for filter dropdowns
    courses = RecipeCourse.objects.all().order_by('name')
    categories = RecipeCategory.objects.all().order_by('name')
    proteins = CustomProtein.objects.all().order_by('name')
    
    # Get distinct authors from Recipe.AUTHOR_CHOICES
    authors = [
        {'value': 'General', 'name': 'General'},
        {'value': 'Demetri & Angy', 'name': 'Demetri & Angy'},
        {'value': 'Erene', 'name': 'Erene'},
        {'value': 'Alexandra', 'name': 'Alexandra'},
    ]
    
    context = {
        'recipes': page_obj,
        'total_recipe_count': paginator.count,
        'courses': courses,
        'categories': categories,
        'proteins': proteins,
        'authors': authors,
        'search_query': search_query,
        'selected_courses': selected_courses,
        'selected_categories': selected_categories,
        'selected_proteins': selected_proteins,
        'selected_authors': selected_authors,
        'selected_letter': selected_letter,
        'letter_data': letter_data,
        'has_next': page_obj.has_next(),
        'next_page': page_obj.next_page_number() if page_obj.has_next() else None,
        'ingredient_categories': IngredientCategory.objects.prefetch_related('ingredient_set').all().order_by('name'),  # ← ADD THIS LINE
    }
    
    return render(request, 'recipe_management.html', context)

@login_required
def duplicate_recipe(request, recipe_id):
    """Duplicate a recipe with all its ingredients and related data"""
    if not request.user.is_superuser:
        return JsonResponse({'success': False, 'error': 'Permission denied'}, status=403)
    
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'POST required'}, status=405)
    
    try:
        from django.db import transaction
        
        # Get the original recipe
        original_recipe = Recipe.objects.get(recipe_id=recipe_id)
        
        with transaction.atomic():
            # Store the original recipe ID
            original_id = original_recipe.recipe_id
            
            # Store many-to-many relationships before modifying
            original_courses = list(original_recipe.courses.all())
            original_categories = list(original_recipe.categories.all())
            original_proteins = list(original_recipe.proteins.all())
            
            # Get related data
            original_recipe_ingredients = list(RecipeIngredient.objects.filter(recipe_id=original_id))
            original_text_ingredients = list(RecipeIngredientText.objects.filter(recipe_id=original_id))
            original_instructions = list(RecipeInstruction.objects.filter(recipe_id=original_id))
            
            # Create new recipe by setting pk to None
            original_recipe.pk = None
            original_recipe.recipe_id = None  # Let it auto-generate
            original_recipe.recipe_name = f"{original_recipe.recipe_name} (Copy)"
            original_recipe.save()
            
            new_recipe_id = original_recipe.recipe_id
            new_recipe_name = original_recipe.recipe_name
            
            # Copy many-to-many relationships
            original_recipe.courses.set(original_courses)
            original_recipe.categories.set(original_categories)
            original_recipe.proteins.set(original_proteins)
            
            # Bulk create structured recipe ingredients
            new_recipe_ingredients = [
                RecipeIngredient(
                    recipe=original_recipe,
                    ingredient=ri.ingredient,
                    amount=ri.amount,
                    unit=ri.unit,
                    preparation=ri.preparation,
                    preparation_note=ri.preparation_note,
                    ingredient_order=ri.ingredient_order,
                    ingredient_group=ri.ingredient_group,
                )
                for ri in original_recipe_ingredients
            ]
            if new_recipe_ingredients:
                RecipeIngredient.objects.bulk_create(new_recipe_ingredients)
            
            # Bulk create text-based ingredients
            new_text_ingredients = [
                RecipeIngredientText(
                    recipe=original_recipe,
                    ingredient_text=ti.ingredient_text,
                    ingredient_group=ti.ingredient_group,
                    order=ti.order,
                )
                for ti in original_text_ingredients
            ]
            if new_text_ingredients:
                RecipeIngredientText.objects.bulk_create(new_text_ingredients)
            
            # Bulk create instructions
            new_instructions = [
                RecipeInstruction(
                    recipe=original_recipe,
                    step_number=inst.step_number,
                    instruction_text=inst.instruction_text,
                    instruction_group=inst.instruction_group,
                    time_estimate=inst.time_estimate,
                    step_image=inst.step_image,
                )
                for inst in original_instructions
            ]
            if new_instructions:
                RecipeInstruction.objects.bulk_create(new_instructions)
        
        return JsonResponse({
            'success': True,
            'new_recipe_id': new_recipe_id,
            'new_recipe_name': new_recipe_name
        })
        
    except Recipe.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'Recipe not found'
        }, status=404)
    
    except Exception as e:
        import traceback
        print(traceback.format_exc())
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

# ============================================
# VIEW: View Recipe
# ============================================

@login_required
def view_recipe(request, recipe_id):
    """View recipe detail page"""
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to access this page.')
        return redirect('home')
    
    recipe = get_object_or_404(Recipe, recipe_id=recipe_id)
    
    # Get ingredients - ONLY from normalized table
    ingredients = RecipeIngredient.objects.filter(recipe=recipe).select_related(
        'ingredient', 'ingredient__category', 'unit', 'preparation'
    ).order_by('ingredient_group', 'ingredient_order')

    # Format amounts
    for ingredient in ingredients:
        ingredient.formatted_amount = format_quantity(ingredient.amount)

    # Get instructions
    instructions = RecipeInstruction.objects.filter(recipe=recipe).order_by('step_number')
    
    context = {
        'recipe': recipe,
        'ingredients': ingredients,
        'instructions': instructions,
    }
    
    return render(request, 'view_recipe.html', context)


# ============================================
# VIEW: Create Recipe
# ============================================

@login_required
def create_recipe(request):
    """Create new recipe - same for manual and AI import"""
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to access this page.')
        return redirect('home')
    
    if request.method == 'POST':
        try:
            # Basic info
            recipe = Recipe()
            recipe.recipe_name = request.POST.get('recipe_name')
            recipe.recipe_description = request.POST.get('recipe_description', '')
            recipe.author = request.POST.get('author', 'General')
            recipe.prep_time = request.POST.get('prep_time') or None
            recipe.cook_time = request.POST.get('cook_time') or None
            recipe.total_time = request.POST.get('total_time') or None
            recipe.servings = int(request.POST.get('servings', 4))
            recipe.difficulty_level = request.POST.get('difficulty_level', '')
            recipe.is_vegetarian = request.POST.get('is_vegetarian') == '1'
            
            if 'recipe_image' in request.FILES:
                recipe.recipe_image = request.FILES['recipe_image']
            
            recipe.created_by = request.user.username
            recipe.save()
            
            # Many-to-many
            course_ids = request.POST.getlist('course[]')
            if course_ids:
                recipe.courses.set(course_ids)
            
            category_ids = request.POST.getlist('category[]')
            if category_ids:
                recipe.categories.set(category_ids)
            
            if not recipe.is_vegetarian:
                protein_ids = request.POST.getlist('protein[]')
                if protein_ids:
                    recipe.proteins.set(protein_ids)
            
            # ========== SAVE INGREDIENTS (NORMALIZED) ==========
            ingredient_quantities = request.POST.getlist('ingredient_quantity[]')
            ingredient_measurements = request.POST.getlist('ingredient_measurement[]')
            ingredient_names = request.POST.getlist('ingredient_name[]')
            ingredient_preparations = request.POST.getlist('ingredient_preparation[]')
            ingredient_groups = request.POST.getlist('ingredient_group[]')
            
            for i in range(len(ingredient_names)):
                if ingredient_names[i].strip():
                    # Get or create related objects
                    ingredient = get_or_create_ingredient(ingredient_names[i])
                    unit = get_or_create_unit(ingredient_measurements[i])
                    
                    prep = None
                    if i < len(ingredient_preparations) and ingredient_preparations[i].strip():
                        prep = get_or_create_preparation(ingredient_preparations[i])
                    
                    # Convert quantity (handles fractions!)
                    quantity_str = ingredient_quantities[i]

                    # Create normalized record
                    RecipeIngredient.objects.create(
                        recipe=recipe,
                        ingredient=ingredient,
                        unit=unit,
                        amount=convert_to_decimal(quantity_str),  # ← CHANGED
                        preparation=prep,
                        ingredient_group=ingredient_groups[i] if i < len(ingredient_groups) else '',
                        ingredient_order=i
                    )
            
            # ========== SAVE INSTRUCTIONS ==========
            instructions = request.POST.getlist('instruction[]')
            instruction_groups = request.POST.getlist('instruction_group[]')
            
            for idx, instruction_text in enumerate(instructions):
                if instruction_text.strip():
                    group = instruction_groups[idx] if idx < len(instruction_groups) else ''
                    RecipeInstruction.objects.create(
                        recipe=recipe,
                        step_number=idx + 1,
                        instruction_text=instruction_text,
                        instruction_group=group
                    )
            
            messages.success(request, f'Recipe "{recipe.recipe_name}" has been created successfully!')
            return redirect('recipe_management')
            
        except Exception as e:
            messages.error(request, f'Error creating recipe: {str(e)}')
            return redirect('create_recipe')
    
    # GET request - show empty form
    extracted_data = {
        'recipe_name': '',
        'description': '',
        'author': 'General',
        'prep_time': 0,
        'cook_time': 0,
        'total_time': 0,
        'servings': 4,
        'difficulty_level': '',
        'is_vegetarian': False,
        'ingredients': [],
        'instructions': [],
    }
    
    # Get lookups
    existing_measurements = list(MeasurementUnit.objects.values_list('name', flat=True))
    existing_ingredients_list = list(Ingredient.objects.values_list('name', flat=True))
    existing_preparations = list(PreparationMethod.objects.values_list('name', flat=True))
    courses = RecipeCourse.objects.all().order_by('name')
    categories = RecipeCategory.objects.all().order_by('name')
    proteins = CustomProtein.objects.all().order_by('name')
    ingredient_categories = IngredientCategory.objects.all().order_by('name')  # ← ADD THIS
    
    # Load all units for the modal
    all_units = MeasurementUnit.objects.all().order_by('name')

    context = {
        'mode': 'create',
        'temp_recipe_id': None,
        'extracted_data': extracted_data,
        'existing_measurements': json.dumps(existing_measurements),
        'existing_ingredients': json.dumps(existing_ingredients_list),
        'existing_preparations': json.dumps(existing_preparations),
        'courses': courses,
        'categories': categories,
        'proteins': proteins,
        'ingredient_categories': ingredient_categories,
        'all_units': all_units,  # ← ADD THIS
    }
    
    return render(request, 'preview_imported_recipe.html', context)


# ============================================
# VIEW: Edit Recipe
# ============================================

@login_required
def edit_recipe(request, recipe_id):
    """Edit an existing recipe"""
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to access this page.')
        return redirect('home')
    
    # OPTIMIZED: Prefetch all related data upfront
    recipe = get_object_or_404(
        Recipe.objects.prefetch_related(
            'courses', 
            'categories', 
            'proteins',
            'recipe_ingredients__ingredient',
            'recipe_ingredients__unit',
            'recipe_ingredients__preparation',
            'instructions'
        ),
        recipe_id=recipe_id
    )
    
    if request.method == 'POST':
        try:
            # Basic information
            recipe.recipe_name = request.POST.get('recipe_name')
            recipe.recipe_description = request.POST.get('recipe_description', '')
            recipe.author = request.POST.get('author', 'General')
            recipe.prep_time = request.POST.get('prep_time') or None
            recipe.cook_time = request.POST.get('cook_time') or None
            recipe.total_time = request.POST.get('total_time') or None
            recipe.servings = request.POST.get('servings')
            
            # Image handling
            if request.POST.get('remove_image') == '1':
                if recipe.recipe_image:
                    recipe.recipe_image.delete(save=False)
                    recipe.recipe_image = None
            
            if 'recipe_image' in request.FILES:
                if recipe.recipe_image:
                    recipe.recipe_image.delete(save=False)
                recipe.recipe_image = request.FILES['recipe_image']
            
            # Classification
            course_ids = request.POST.getlist('course[]')
            category_ids = request.POST.getlist('category[]')
            recipe.difficulty_level = request.POST.get('difficulty_level')
            
            # Dietary
            recipe.is_vegetarian = request.POST.get('is_vegetarian') == '1'
            protein_ids = request.POST.getlist('protein[]')
            
            # Save recipe
            recipe.save()
            
            # Update many-to-many
            if course_ids:
                recipe.courses.set(course_ids)
            if category_ids:
                recipe.categories.set(category_ids)
            if recipe.is_vegetarian:
                recipe.proteins.clear()
            elif protein_ids:
                recipe.proteins.set(protein_ids)
            
            # ========== OPTIMIZED INGREDIENTS SECTION ==========
            # Delete old ingredients
            recipe.recipe_ingredients.all().delete()
            
            # Get form data
            ingredient_quantities = request.POST.getlist('ingredient_quantity[]')
            ingredient_measurements = request.POST.getlist('ingredient_measurement[]')
            ingredient_names = request.POST.getlist('ingredient_name[]')
            ingredient_preparations = request.POST.getlist('ingredient_preparation[]')
            ingredient_groups = request.POST.getlist('ingredient_group[]')
            
            # OPTIMIZATION: Bulk fetch all existing ingredients, units, and preparations
            ingredient_name_set = {name.strip() for name in ingredient_names if name.strip()}
            measurement_name_set = {meas.strip() for meas in ingredient_measurements if meas.strip()}
            preparation_name_set = {prep.strip() for prep in ingredient_preparations if prep.strip()}
            
            # Fetch existing records in bulk
            existing_ingredients = {
                ing.name: ing 
                for ing in Ingredient.objects.filter(name__in=ingredient_name_set)
            }
            existing_units = {
                unit.name: unit 
                for unit in MeasurementUnit.objects.filter(name__in=measurement_name_set)
            }
            existing_preparations = {
                prep.name: prep 
                for prep in PreparationMethod.objects.filter(name__in=preparation_name_set)
            }
            
            # Create new ingredients/units/preparations in bulk
            new_ingredients = []
            for name in ingredient_name_set:
                if name not in existing_ingredients:
                    new_ingredients.append(Ingredient(name=name))
            if new_ingredients:
                Ingredient.objects.bulk_create(new_ingredients)
                # Re-fetch to get IDs
                existing_ingredients = {
                    ing.name: ing 
                    for ing in Ingredient.objects.filter(name__in=ingredient_name_set)
                }
            
            new_units = []
            for name in measurement_name_set:
                if name not in existing_units:
                    new_units.append(MeasurementUnit(name=name))
            if new_units:
                MeasurementUnit.objects.bulk_create(new_units)
                # Re-fetch to get IDs
                existing_units = {
                    unit.name: unit 
                    for unit in MeasurementUnit.objects.filter(name__in=measurement_name_set)
                }
            
            new_preparations = []
            for name in preparation_name_set:
                if name not in existing_preparations:
                    new_preparations.append(PreparationMethod(name=name))
            if new_preparations:
                PreparationMethod.objects.bulk_create(new_preparations)
                # Re-fetch to get IDs
                existing_preparations = {
                    prep.name: prep 
                    for prep in PreparationMethod.objects.filter(name__in=preparation_name_set)
                }
            
            # Build RecipeIngredient objects
            recipe_ingredients = []
            for i in range(len(ingredient_names)):
                if ingredient_names[i].strip():
                    ingredient = existing_ingredients.get(ingredient_names[i].strip())
                    unit = existing_units.get(ingredient_measurements[i].strip())
                    
                    prep = None
                    if i < len(ingredient_preparations) and ingredient_preparations[i].strip():
                        prep = existing_preparations.get(ingredient_preparations[i].strip())
                    
                    quantity_str = ingredient_quantities[i]
                    
                    recipe_ingredients.append(RecipeIngredient(
                        recipe=recipe,
                        ingredient=ingredient,
                        unit=unit,
                        amount=convert_to_decimal(quantity_str),
                        preparation=prep,
                        ingredient_group=ingredient_groups[i] if i < len(ingredient_groups) else '',
                        ingredient_order=i
                    ))
            
            # Bulk create all recipe ingredients
            if recipe_ingredients:
                RecipeIngredient.objects.bulk_create(recipe_ingredients)
            
            # ========== OPTIMIZED INSTRUCTIONS SECTION ==========
            recipe.instructions.all().delete()
            
            instructions = request.POST.getlist('instruction[]')
            instruction_groups = request.POST.getlist('instruction_group[]')
            
            # Build instruction objects
            instruction_objects = []
            for idx, instruction_text in enumerate(instructions):
                if instruction_text.strip():
                    group = instruction_groups[idx] if idx < len(instruction_groups) else ''
                    instruction_objects.append(RecipeInstruction(
                        recipe=recipe,
                        step_number=idx + 1,
                        instruction_text=instruction_text,
                        instruction_group=group
                    ))
            
            # Bulk create all instructions
            if instruction_objects:
                RecipeInstruction.objects.bulk_create(instruction_objects)
            
            messages.success(request, f'Recipe "{recipe.recipe_name}" has been updated successfully!')
            return redirect('recipe_management')
            
        except Exception as e:
            messages.error(request, f'Error updating recipe: {str(e)}')
            return redirect('recipe_management')
    
    # ========== GET request - prepare data for editing ==========
    # OPTIMIZED: Use select_related to fetch related objects in one query
    existing_ingredients = []
    
    for ing in recipe.recipe_ingredients.select_related(
        'ingredient', 'unit', 'preparation'
    ).order_by('ingredient_order'):
        existing_ingredients.append({
            'quantity': format_quantity(ing.amount),
            'measurement': ing.unit.name if ing.unit else '',
            'ingredient': ing.ingredient.name,
            'preparation': ing.preparation.name if ing.preparation else '',
            'group': ing.ingredient_group or ''
        })
    
    existing_instructions = [
        {
            'instruction': inst.instruction_text,
            'group': inst.instruction_group or ''
        }
        for inst in recipe.instructions.all().order_by('step_number')
    ]
    
    extracted_data = {
        'recipe_name': recipe.recipe_name,
        'description': recipe.recipe_description or '',
        'author': recipe.author,
        'prep_time': recipe.prep_time or 0,
        'cook_time': recipe.cook_time or 0,
        'total_time': recipe.total_time or 0,
        'servings': recipe.servings or 1,
        'difficulty_level': recipe.difficulty_level or '',
        'is_vegetarian': recipe.is_vegetarian,
        'recipe_image': recipe.recipe_image.url if recipe.recipe_image else None,
        'ingredients': existing_ingredients,
        'instructions': existing_instructions,
    }
    
    # Get lookups
    existing_measurements = list(MeasurementUnit.objects.values_list('name', flat=True))
    existing_ingredients_list = list(Ingredient.objects.values_list('name', flat=True))
    existing_preparations = list(PreparationMethod.objects.values_list('name', flat=True))
    courses = RecipeCourse.objects.all().order_by('name')
    categories = RecipeCategory.objects.all().order_by('name')
    proteins = CustomProtein.objects.all().order_by('name')
    ingredient_categories = IngredientCategory.objects.all().order_by('name')  # ← ADD THIS
    
    # OPTIMIZED: Use values_list to get IDs without additional queries
    selected_course_ids = list(recipe.courses.values_list('recipe_course_id', flat=True))
    selected_category_ids = list(recipe.categories.values_list('recipe_category_id', flat=True))
    selected_protein_ids = list(recipe.proteins.values_list('custom_protein_id', flat=True))

    # Load all units for the modal
    all_units = MeasurementUnit.objects.all().order_by('name')

    context = {
        'mode': 'edit',
        'temp_recipe_id': recipe_id,
        'recipe': recipe,
        'extracted_data': extracted_data,
        'existing_measurements': json.dumps(existing_measurements),
        'existing_ingredients': json.dumps(existing_ingredients_list),
        'existing_preparations': json.dumps(existing_preparations),
        'courses': courses,
        'categories': categories,
        'proteins': proteins,
        'ingredient_categories': ingredient_categories,
        'all_units': all_units,  # ← ADD THIS
        'selected_courses': json.dumps(selected_course_ids),
        'selected_categories': json.dumps(selected_category_ids),
        'selected_proteins': json.dumps(selected_protein_ids),
    }

    return render(request, 'preview_imported_recipe.html', context)

@login_required
@require_POST
def add_recipe_course(request):
    """AJAX view to add a new recipe course"""
    if not request.user.is_superuser:
        return JsonResponse({'success': False, 'error': 'Permission denied'})
    
    try:
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        display_order = data.get('display_order', 0)
        
        if not name:
            return JsonResponse({'success': False, 'error': 'Course name is required'})
        
        if RecipeCourse.objects.filter(name__iexact=name).exists():
            return JsonResponse({'success': False, 'error': 'A course with this name already exists'})
        
        course = RecipeCourse.objects.create(
            name=name,
            display_order=int(display_order)
        )
        
        return JsonResponse({
            'success': True,
            'course_id': course.recipe_course_id,
            'name': course.name,
            'display_order': course.display_order
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@login_required
@require_POST
def add_recipe_category(request):
    """AJAX view to add a new recipe category"""
    if not request.user.is_superuser:
        return JsonResponse({'success': False, 'error': 'Permission denied'})
    
    try:
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        
        if not name:
            return JsonResponse({'success': False, 'error': 'Category name is required'})
        
        if RecipeCategory.objects.filter(name__iexact=name).exists():
            return JsonResponse({'success': False, 'error': 'A category with this name already exists'})
        
        category = RecipeCategory.objects.create(name=name)
        
        return JsonResponse({
            'success': True,
            'category_id': category.recipe_category_id,
            'name': category.name
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@login_required
@require_POST
def add_recipe_ingredient(request):
    """AJAX view to add a new ingredient"""
    if not request.user.is_superuser:
        return JsonResponse({'success': False, 'error': 'Permission denied'})
    
    try:
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        category_id = data.get('category_id', None)
        
        if not name:
            return JsonResponse({'success': False, 'error': 'Ingredient name is required'})
        
        if Ingredient.objects.filter(name__iexact=name).exists():
            return JsonResponse({'success': False, 'error': 'An ingredient with this name already exists'})
        
        ingredient = Ingredient.objects.create(name=name)
        
        if category_id:
            ingredient.category_id = category_id
            ingredient.save()
        
        return JsonResponse({
            'success': True,
            'ingredient_id': ingredient.ingredient_id,
            'name': ingredient.name
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@login_required
@require_POST
def add_recipe_protein(request):
    """AJAX view to add a new custom protein - UPDATED"""
    if not request.user.is_superuser:
        return JsonResponse({'success': False, 'error': 'Permission denied'})
    
    try:
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        
        if not name:
            return JsonResponse({'success': False, 'error': 'Protein name is required'})
        
        # CHANGED: Only check if it exists in CustomProtein table
        if CustomProtein.objects.filter(name__iexact=name).exists():
            return JsonResponse({'success': False, 'error': 'This protein already exists'})
        
        # REMOVED: Check against MAIN_PROTEIN_CHOICES (no longer exists)
        
        protein = CustomProtein.objects.create(name=name)
        
        return JsonResponse({
            'success': True,
            'protein_id': protein.custom_protein_id,  # CHANGED: Return ID
            'name': protein.name
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})

# Temporary storage for extracted recipe data (use session or database)
class TempRecipeData:
    """Temporary storage for AI-extracted recipe data"""
    def __init__(self, recipe_id, data):
        self.recipe_id = recipe_id
        self.data = data

@login_required
def import_recipe(request):
    """Upload recipe file for AI extraction"""
    
    if request.method == 'POST':
        uploaded_file = request.FILES.get('recipe_file')
        
        if not uploaded_file:
            messages.error(request, 'Please select a file to upload.')
            return redirect('import_recipe')
        
        try:
            file_name = uploaded_file.name
            file_ext = file_name.split('.')[-1].lower()
            
            # Extract text based on file type
            if file_ext == 'pdf':
                text_content = extract_text_from_pdf(uploaded_file)
            elif file_ext in ['doc', 'docx']:
                text_content = extract_text_from_docx(uploaded_file)
            elif file_ext in ['jpg', 'jpeg', 'png']:
                text_content = extract_text_from_image(uploaded_file)
            else:
                messages.error(request, 'Unsupported file format.')
                return redirect('import_recipe')
            
            # Use Claude AI to extract
            extracted_data = extract_recipe_with_ai(text_content, file_ext)
            
            if not extracted_data:
                messages.error(request, 'Could not extract recipe data. Please try a different file.')
                return redirect('import_recipe')
            
            # Store in session
            import uuid
            temp_id = str(uuid.uuid4())
            request.session[f'temp_recipe_{temp_id}'] = extracted_data
            request.session[f'temp_recipe_{temp_id}_file'] = file_name
            
            messages.success(request, 'Recipe extracted successfully! Please review and edit as needed.')
            return redirect('preview_imported_recipe', temp_id=temp_id)
            
        except Exception as e:
            messages.error(request, f'Error processing file: {str(e)}')
            return redirect('import_recipe')
    
    return render(request, 'import_recipe.html')


# ============================================
# VIEW: Preview Imported Recipe
# ============================================

@login_required
def preview_imported_recipe(request, temp_id):
    """Preview and edit AI-extracted recipe data - SAME save logic as create_recipe"""
    
    extracted_data = request.session.get(f'temp_recipe_{temp_id}')
    
    if not extracted_data:
        messages.error(request, 'Recipe data not found. Please import again.')
        return redirect('import_recipe')
    
    if request.method == 'POST':
        # ========== SAVE LOGIC - IDENTICAL TO create_recipe ==========
        try:
            recipe = Recipe()
            recipe.recipe_name = request.POST.get('recipe_name')
            recipe.recipe_description = request.POST.get('recipe_description', '')
            recipe.author = request.POST.get('author', 'General')
            recipe.prep_time = request.POST.get('prep_time') or None
            recipe.cook_time = request.POST.get('cook_time') or None
            recipe.total_time = request.POST.get('total_time') or None
            recipe.servings = request.POST.get('servings')
            recipe.difficulty_level = request.POST.get('difficulty_level')
            recipe.is_vegetarian = request.POST.get('is_vegetarian') == '1'
            recipe.created_by = request.user.username if request.user.is_authenticated else 'Anonymous'
            recipe.is_ai_imported = True
            
            if 'recipe_image' in request.FILES:
                recipe.recipe_image = request.FILES['recipe_image']
            
            recipe.save()
            
            # Many-to-many
            course_ids = request.POST.getlist('course[]')
            if course_ids:
                recipe.courses.set(course_ids)
            
            category_ids = request.POST.getlist('category[]')
            if category_ids:
                recipe.categories.set(category_ids)
            
            protein_ids = request.POST.getlist('protein[]')
            if not recipe.is_vegetarian and protein_ids:
                recipe.proteins.set(protein_ids)
            
            # ========== SAVE INGREDIENTS (NORMALIZED) - SAME AS create_recipe ==========
            ingredient_quantities = request.POST.getlist('ingredient_quantity[]')
            ingredient_measurements = request.POST.getlist('ingredient_measurement[]')
            ingredient_names = request.POST.getlist('ingredient_name[]')
            ingredient_preparations = request.POST.getlist('ingredient_preparation[]')
            ingredient_groups = request.POST.getlist('ingredient_group[]')
            
            for i in range(len(ingredient_names)):
                if ingredient_names[i].strip():
                    ingredient = get_or_create_ingredient(ingredient_names[i])
                    unit = get_or_create_unit(ingredient_measurements[i])
                    
                    prep = None
                    if i < len(ingredient_preparations) and ingredient_preparations[i].strip():
                        prep = get_or_create_preparation(ingredient_preparations[i])
                    
                    quantity_str = ingredient_quantities[i]

                    RecipeIngredient.objects.create(
                        recipe=recipe,
                        ingredient=ingredient,
                        unit=unit,
                        amount=convert_to_decimal(quantity_str),  # ← CHANGED
                        preparation=prep,
                        ingredient_group=ingredient_groups[i] if i < len(ingredient_groups) else '',
                        ingredient_order=i
                    )
            
            # ========== SAVE INSTRUCTIONS ==========
            instructions = request.POST.getlist('instruction[]')
            instruction_groups = request.POST.getlist('instruction_group[]')
            
            for idx, instruction_text in enumerate(instructions):
                if instruction_text.strip():
                    group = instruction_groups[idx] if idx < len(instruction_groups) else ''
                    RecipeInstruction.objects.create(
                        recipe=recipe,
                        step_number=idx + 1,
                        instruction_text=instruction_text,
                        instruction_group=group
                    )
            
            # Clear session
            del request.session[f'temp_recipe_{temp_id}']
            if f'temp_recipe_{temp_id}_file' in request.session:
                del request.session[f'temp_recipe_{temp_id}_file']
            
            messages.success(request, f'Recipe "{recipe.recipe_name}" has been imported successfully!')
            return redirect('recipe_management')
            
        except Exception as e:
            messages.error(request, f'Error saving recipe: {str(e)}')
            return redirect('recipe_management')
    
    # GET request - show preview
    existing_measurements = list(MeasurementUnit.objects.values_list('name', flat=True))
    existing_ingredients_list = list(Ingredient.objects.values_list('name', flat=True))
    existing_preparations = list(PreparationMethod.objects.values_list('name', flat=True))
    courses = RecipeCourse.objects.all().order_by('name')
    categories = RecipeCategory.objects.all().order_by('name')
    proteins = CustomProtein.objects.all().order_by('name')
    ingredient_categories = IngredientCategory.objects.all().order_by('name')  # ← ADD THIS LINE

    # Load all units for the modal
    all_units = MeasurementUnit.objects.all().order_by('name')

    context = {
        'mode': 'import',
        'temp_recipe_id': temp_id,
        'extracted_data': extracted_data,
        'existing_measurements': json.dumps(existing_measurements),
        'existing_ingredients': json.dumps(existing_ingredients_list),
        'existing_preparations': json.dumps(existing_preparations),
        'courses': courses,
        'categories': categories,
        'proteins': proteins,
        'ingredient_categories': ingredient_categories,
        'all_units': all_units,  # ← ADD THIS
    }

    return render(request, 'preview_imported_recipe.html', context)

def aggregate_meal_plan_ingredients(meal_plan):
    """
    Aggregate ingredients and convert to shopping units.
    Always uses the ingredient's shopping unit. Prompts for missing conversions.
    Returns: (aggregated_dict, missing_conversions_list, missing_shopping_units_list)
    """
    from collections import defaultdict
    from decimal import Decimal
    import math
    
    def smart_categorize(ingredient_name):
        """Intelligently categorize ingredients based on name"""
        ingredient_lower = ingredient_name.lower()
        
        canned_terms = ['stock', 'broth', 'cube', 'bouillon', 'canned', 'tinned', 'tin', 'paste', 'sauce', 'puree', 'concentrate']
        for term in canned_terms:
            if term in ingredient_lower:
                return 'Canned & Packaged'
        
        beverages = ['wine', 'beer', 'sherry', 'brandy', 'rum', 'vodka', 'whiskey', 'liqueur']
        for term in beverages:
            if term in ingredient_lower:
                return 'Beverages'
        
        herbs_spices = ['oregano', 'origanum', 'basil', 'thyme', 'rosemary', 'sage', 'parsley', 'cilantro', 'pepper', 'salt', 'paprika', 'cumin', 'turmeric', 'cinnamon', 'nutmeg', 'cloves', 'curry', 'chili', 'cayenne']
        for term in herbs_spices:
            if term in ingredient_lower:
                return 'Herbs & Spices'
        
        meats = ['beef', 'pork', 'chicken', 'lamb', 'mince', 'meat', 'bacon', 'sausage', 'fish', 'salmon', 'tuna', 'shrimp', 'prawns']
        for term in meats:
            if term in ingredient_lower:
                return 'Meat & Seafood'
        
        dairy = ['milk', 'cream', 'yogurt', 'cheese', 'butter']
        is_beverage = any(bev in ingredient_lower for bev in beverages)
        if not is_beverage:
            for term in dairy:
                if term in ingredient_lower:
                    return 'Dairy'
        
        oils_fats = ['oil', 'olive oil', 'vegetable oil', 'butter', 'margarine', 'lard', 'ghee']
        for term in oils_fats:
            if term in ingredient_lower:
                return 'Oils & Fats'
        
        grains = ['flour', 'rice', 'pasta', 'spaghetti', 'noodles', 'bread', 'quinoa', 'couscous', 'oats', 'macaroni']
        for term in grains:
            if term in ingredient_lower:
                return 'Grains & Pasta'
        
        vegetables = ['onion', 'garlic', 'tomato', 'potato', 'carrot', 'celery', 'pepper', 'lettuce', 'spinach', 'broccoli', 'mushroom', 'peas', 'corn']
        for term in vegetables:
            if term in ingredient_lower:
                return 'Vegetables'
        
        return 'Other'
    
    def round_shopping_qty(qty, unit):
        """
        Round quantities intelligently based on unit type for shopping lists.
        NEVER rounds down - always rounds UP to ensure you have enough.
        """
        if qty <= 0:
            return qty
        
        unit_type = getattr(unit, 'unit_type', 'other')
        
        if unit_type == 'count':
            # Always round UP to whole number, minimum 1
            return max(1, math.ceil(qty))
        
        elif unit_type == 'weight':
            # Round UP to sensible numbers based on magnitude
            if qty < 10:
                return math.ceil(qty * 2) / 2
            elif qty < 100:
                return math.ceil(qty / 5) * 5
            else:
                return math.ceil(qty / 10) * 10
        
        elif unit_type == 'volume':
            # Round UP to sensible numbers based on magnitude
            if qty < 10:
                return math.ceil(qty * 4) / 4
            elif qty < 100:
                return math.ceil(qty / 5) * 5
            else:
                return math.ceil(qty / 10) * 10
        
        else:
            # OTHER (dash, pinch, to taste): round UP
            return max(1, math.ceil(qty))
    
    def get_unit_display(unit, qty):
        """Get the proper unit display with pluralization"""
        if qty == 1:
            return unit.abbreviation or unit.name
        else:
            return unit.abbreviation_plural or unit.abbreviation or unit.name_plural or unit.name
    
    # Pre-load all unit conversions for fast lookups (prevents N+1 queries)
    conversion_cache = get_conversion_cache()
    
    # Track issues
    missing_conversions = []
    missing_shopping_units = []
    seen_missing_conversions = set()
    seen_missing_units = set()
    
    # Dictionary: ingredient_id -> {amount, unit, ingredient_obj, unconverted_items}
    aggregated = defaultdict(lambda: {
        'unit': None,
        'amount': Decimal('0'),
        'ingredient_obj': None,
        'unconverted_items': []  # Items that couldn't be converted
    })
    
    # Process all recipes in meal plan
    for day in meal_plan.days.all().order_by('date'):
        for meal_recipe in day.recipes.all():
            recipe = meal_recipe.recipe
            servings_multiplier = Decimal(meal_recipe.servings) / Decimal(recipe.servings or 1)
            
            for recipe_ingredient in recipe.recipe_ingredients.all():
                ingredient = recipe_ingredient.ingredient
                ingredient_id = ingredient.ingredient_id
                quantity = Decimal(recipe_ingredient.amount or 0) * servings_multiplier
                from_unit = recipe_ingredient.unit
                
                if not from_unit:
                    continue
                
                # CRITICAL: Always use the ingredient's shopping unit
                shopping_unit = ingredient.default_unit
                
                # Track ingredients without shopping units
                if not shopping_unit:
                    if ingredient_id not in seen_missing_units:
                        seen_missing_units.add(ingredient_id)
                        missing_shopping_units.append({
                            'ingredient_id': ingredient_id,
                            'ingredient_name': ingredient.name,
                            'recipe': recipe.recipe_name
                        })
                    # Add to unconverted items
                    if ingredient_id not in aggregated:
                        aggregated[ingredient_id]['ingredient_obj'] = ingredient
                        aggregated[ingredient_id]['unit'] = from_unit  # Fallback
                    
                    aggregated[ingredient_id]['unconverted_items'].append({
                        'quantity': float(quantity),
                        'unit': from_unit.name,
                        'recipe': recipe.recipe_name,
                        'reason': 'No shopping unit defined'
                    })
                    continue
                
                # Initialize if first time seeing this ingredient
                if aggregated[ingredient_id]['unit'] is None:
                    aggregated[ingredient_id]['unit'] = shopping_unit
                    aggregated[ingredient_id]['ingredient_obj'] = ingredient
                
                shopping_unit = aggregated[ingredient_id]['unit']
                
                # Convert to shopping unit
                if from_unit.measurement_unit_id == shopping_unit.measurement_unit_id:
                    # Same unit - just add
                    aggregated[ingredient_id]['amount'] += quantity
                else:
                    # Need conversion - USE CACHE for fast lookup
                    converted_qty, multiplier = convert_quantity(quantity, from_unit, shopping_unit, ingredient, conversion_cache)
                    
                    if converted_qty is not None:
                        # Conversion successful
                        aggregated[ingredient_id]['amount'] += converted_qty
                    else:
                        # NO CONVERSION EXISTS - Track it
                        conversion_key = f"{from_unit.measurement_unit_id}-{shopping_unit.measurement_unit_id}"
                        if conversion_key not in seen_missing_conversions:
                            seen_missing_conversions.add(conversion_key)
                            missing_conversions.append({
                                'ingredient': ingredient.name,
                                'from_unit': from_unit.name,
                                'from_unit_id': from_unit.measurement_unit_id,
                                'to_unit': shopping_unit.name,
                                'to_unit_id': shopping_unit.measurement_unit_id,
                                'quantity': float(quantity),
                                'recipe': recipe.recipe_name
                            })
                        
                        # Add to unconverted items
                        aggregated[ingredient_id]['unconverted_items'].append({
                            'quantity': float(quantity),
                            'unit': from_unit.name,
                            'recipe': recipe.recipe_name,
                            'reason': f'Missing conversion from {from_unit.name} to {shopping_unit.name}'
                        })
    
    # Build categorized shopping list
    categorized_ingredients = defaultdict(list)
    
    for ingredient_id in sorted(aggregated.keys(), key=lambda x: aggregated[x]['ingredient_obj'].name):
        data = aggregated[ingredient_id]
        ingredient_obj = data['ingredient_obj']
        unit = data['unit']
        
        # Determine category
        if ingredient_obj and ingredient_obj.category:
            category = ingredient_obj.category.name
        else:
            category = smart_categorize(ingredient_obj.name)
        
        # Only add to shopping list if we successfully converted some amount
        if data['amount'] > 0:
            # Apply smart rounding for shopping
            raw_qty = float(data['amount'])
            qty = round_shopping_qty(raw_qty, unit)
            
            # Get proper unit display with pluralization
            unit_display = get_unit_display(unit, qty)
            
            entry = {
                'ingredient': ingredient_obj.name,
                'quantity': qty,
                'unit': unit_display
            }
            
            # Add note about unconverted items
            if data['unconverted_items']:
                entry['has_unconverted'] = True
                entry['unconverted_count'] = len(data['unconverted_items'])
                entry['unconverted_items'] = data['unconverted_items']
            
            categorized_ingredients[category].append(entry)
    
    return (dict(categorized_ingredients), missing_conversions, missing_shopping_units)

@login_required
def meal_plans(request):
    """List all meal plans"""
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to access this page.')
        return redirect('home')
    
    # Get all meal plans with recipe counts, sorted by most recent first
    meal_plans_list = MealPlan.objects.annotate(
        recipe_count=Count('days__recipes')
    ).order_by('-start_date')  # ← This orders newest first
    
    context = {
        'meal_plans': meal_plans_list,
    }
    
    return render(request, 'meal_plans.html', context)

@login_required
def create_meal_plan(request):
    """Create a new meal plan"""
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to access this page.')
        return redirect('home')
    
    if request.method == 'POST':
        # ... existing POST code stays the same ...
        try:
            # Get basic info
            plan_name = request.POST.get('plan_name')
            start_date_str = request.POST.get('start_date')
            end_date_str = request.POST.get('end_date')
            
            # Parse dates
            start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
            end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
            
            # Validate date range
            days_diff = (end_date - start_date).days + 1
            if days_diff < 1 or days_diff > 7:
                messages.error(request, 'Meal plan must be between 1 and 7 days.')
                return redirect('create_meal_plan')
            
            # Create meal plan
            meal_plan = MealPlan.objects.create(
                plan_name=plan_name,
                start_date=start_date,
                end_date=end_date,
                created_by=request.user
            )
            
            # Create days and assign recipes
            current_date = start_date
            while current_date <= end_date:
                # Create day
                meal_day = MealPlanDay.objects.create(
                    meal_plan=meal_plan,
                    date=current_date
                )
                
                # Get recipes for this day using date string (YYYY-MM-DD format)
                date_key = current_date.strftime('%Y-%m-%d')
                recipe_ids = request.POST.getlist(f'recipes_{date_key}[]')
                servings_list = request.POST.getlist(f'servings_{date_key}[]')
                
                # Add recipes to this day
                for idx, recipe_id in enumerate(recipe_ids):
                    if recipe_id:  # Skip empty values
                        recipe = Recipe.objects.get(recipe_id=recipe_id)
                        servings = int(servings_list[idx]) if idx < len(servings_list) else recipe.servings
                        
                        MealPlanRecipe.objects.create(
                            meal_plan_day=meal_day,
                            recipe=recipe,
                            servings=servings,
                            sort_order=idx
                        )
                
                current_date += timedelta(days=1)
            
            messages.success(request, f'Meal plan "{plan_name}" created successfully!')
            return redirect('view_meal_plan', meal_plan_id=meal_plan.meal_plan_id)
            
        except Exception as e:
            print(f"\n!!! ERROR: {str(e)} !!!")
            import traceback
            traceback.print_exc()
            messages.error(request, f'Error creating meal plan: {str(e)}')
            return redirect('create_meal_plan')
    
    # GET request - show form
    recipes_qs = Recipe.objects.prefetch_related(
        'courses', 
        'categories', 
        'proteins'
    ).all().order_by('recipe_name')
    
    recipes = []
    for recipe in recipes_qs:
        recipe_data = {
            'recipe_id': recipe.recipe_id,
            'recipe_name': recipe.recipe_name,
            'servings': recipe.servings,
            'prep_time': recipe.prep_time,
            'cook_time': recipe.cook_time,
            'difficulty_level': recipe.difficulty_level or '',
            'is_vegetarian': recipe.is_vegetarian,
            'author': recipe.author,
            'recipe_image': recipe.recipe_image.url if recipe.recipe_image else None,
            'courses': [course.name for course in recipe.courses.all()],
            'categories': [cat.name for cat in recipe.categories.all()],
            'proteins': [protein.name for protein in recipe.proteins.all()]
        }
        recipes.append(recipe_data)
    
    # Get all filter options
    all_courses = list(RecipeCourse.objects.all().order_by('name').values('recipe_course_id', 'name'))
    all_categories = list(RecipeCategory.objects.all().order_by('name').values('recipe_category_id', 'name'))
    all_proteins = list(CustomProtein.objects.all().order_by('name').values('custom_protein_id', 'name'))
    
    # Add authors
    all_authors = [
        {'value': 'General', 'name': 'General'},
        {'value': 'Demetri & Angy', 'name': 'Demetri & Angy'},
        {'value': 'Erene', 'name': 'Erene'},
        {'value': 'Alexandra', 'name': 'Alexandra'},
    ]

    # Suggest a default date range (today + 6 days)
    today = datetime.now().date()
    default_end = today + timedelta(days=6)
    
    context = {
        'recipes_json': json.dumps(recipes),  # ← Serialize to JSON
        'all_courses_json': json.dumps(all_courses),
        'all_categories_json': json.dumps(all_categories),
        'all_proteins_json': json.dumps(all_proteins),
        'all_authors_json': json.dumps(all_authors),
        'today': today,
        'default_start_date': today.strftime('%Y-%m-%d'),
        'default_end_date': default_end.strftime('%Y-%m-%d'),
    }
    
    return render(request, 'create_meal_plan.html', context)

@login_required
def view_meal_plan(request, meal_plan_id):
    """View a meal plan with all days and recipes"""
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to access this page.')
        return redirect('home')
    
    # Get meal plan with optimized prefetch
    meal_plan = get_object_or_404(
        MealPlan.objects.prefetch_related(
            Prefetch(
                'days',
                queryset=MealPlanDay.objects.order_by('date').prefetch_related(
                    Prefetch(
                        'recipes',
                        queryset=MealPlanRecipe.objects.select_related('recipe')
                    )
                )
            )
        ),
        meal_plan_id=meal_plan_id
    )
    
    # Days are already prefetched and ordered
    days = meal_plan.days.all()
    
    context = {
        'meal_plan': meal_plan,
        'days': days,
    }
    
    return render(request, 'view_meal_plan.html', context)

@login_required
def delete_meal_plan(request, meal_plan_id):
    """Delete a meal plan"""
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to perform this action.')
        return redirect('meal_plans')
    
    # Only allow POST requests for deletion
    if request.method != 'POST':
        messages.error(request, 'Invalid request method.')
        return redirect('meal_plans')
    
    try:
        meal_plan = MealPlan.objects.get(meal_plan_id=meal_plan_id)
        plan_name = meal_plan.plan_name
        
        # Delete the meal plan (cascade will delete days and recipes)
        meal_plan.delete()
        
        messages.success(request, f'Meal plan "{plan_name}" has been deleted successfully.')
        
    except MealPlan.DoesNotExist:
        messages.error(request, 'Meal plan not found.')
    
    # Check where to redirect
    redirect_to = request.POST.get('redirect_to', 'list')
    if redirect_to == 'calendar':
        return redirect('meal_plan_calendar')
    return redirect('meal_plans')

@login_required
def edit_meal_plan(request, meal_plan_id):
    """Edit an existing meal plan"""
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to access this page.')
        return redirect('home')
    
    # Get meal plan with all related data
    meal_plan = get_object_or_404(
        MealPlan.objects.prefetch_related(
            Prefetch(
                'days',
                queryset=MealPlanDay.objects.order_by('date').prefetch_related(
                    Prefetch(
                        'recipes',
                        queryset=MealPlanRecipe.objects.select_related('recipe').order_by('sort_order')
                    )
                )
            )
        ),
        meal_plan_id=meal_plan_id
    )
    
    if request.method == 'POST':
        try:
            from django.db import transaction
            
            # Get updated basic info
            new_plan_name = request.POST.get('plan_name', '').strip()
            new_start_date_str = request.POST.get('start_date')
            new_end_date_str = request.POST.get('end_date')
            
            if not new_plan_name:
                messages.error(request, 'Plan name is required.')
                return redirect('edit_meal_plan', meal_plan_id=meal_plan_id)
            
            # Parse new dates
            new_start_date = datetime.strptime(new_start_date_str, '%Y-%m-%d').date()
            new_end_date = datetime.strptime(new_end_date_str, '%Y-%m-%d').date()
            
            # Validate date range
            days_diff = (new_end_date - new_start_date).days + 1
            if days_diff < 1 or days_diff > 7:
                messages.error(request, 'Meal plan must be between 1 and 7 days.')
                return redirect('edit_meal_plan', meal_plan_id=meal_plan_id)
            
            # Use transaction to ensure data consistency
            with transaction.atomic():
                # Update meal plan basic info
                meal_plan.plan_name = new_plan_name
                meal_plan.start_date = new_start_date
                meal_plan.end_date = new_end_date
                meal_plan.save()
                
                # Get current days
                existing_days = {day.date: day for day in meal_plan.days.all()}
                
                # Generate list of dates in new range
                new_dates = []
                current_date = new_start_date
                while current_date <= new_end_date:
                    new_dates.append(current_date)
                    current_date += timedelta(days=1)
                
                # Delete days that are no longer in range
                dates_to_keep = set(new_dates)
                for date, day in existing_days.items():
                    if date not in dates_to_keep:
                        day.delete()
                
                # Process each day in the new range
                for date in new_dates:
                    # Get or create day
                    if date in existing_days:
                        meal_day = existing_days[date]
                    else:
                        meal_day = MealPlanDay.objects.create(
                            meal_plan=meal_plan,
                            date=date
                        )
                    
                    # Get recipes for this day from form
                    date_key = date.strftime('%Y-%m-%d')
                    recipe_ids = request.POST.getlist(f'recipes_{date_key}[]')
                    servings_list = request.POST.getlist(f'servings_{date_key}[]')
                    
                    # Get existing recipes for this day
                    existing_recipes = {mr.recipe.recipe_id: mr for mr in meal_day.recipes.all()}
                    
                    # Track which recipes we're keeping
                    recipes_to_keep = set()
                    
                    # Process submitted recipes
                    for idx, recipe_id in enumerate(recipe_ids):
                        if recipe_id:  # Skip empty values
                            recipe_id = int(recipe_id)
                            recipes_to_keep.add(recipe_id)
                            
                            recipe = Recipe.objects.get(recipe_id=recipe_id)
                            servings = int(servings_list[idx]) if idx < len(servings_list) else recipe.servings
                            
                            if recipe_id in existing_recipes:
                                # Update existing recipe
                                meal_recipe = existing_recipes[recipe_id]
                                meal_recipe.servings = servings
                                meal_recipe.sort_order = idx
                                meal_recipe.save()
                            else:
                                # Add new recipe
                                MealPlanRecipe.objects.create(
                                    meal_plan_day=meal_day,
                                    recipe=recipe,
                                    servings=servings,
                                    sort_order=idx
                                )
                    
                    # Delete recipes that were removed
                    for recipe_id, meal_recipe in existing_recipes.items():
                        if recipe_id not in recipes_to_keep:
                            meal_recipe.delete()
            
            messages.success(request, f'Meal plan "{new_plan_name}" updated successfully!')
            return redirect('view_meal_plan', meal_plan_id=meal_plan_id)
            
        except Exception as e:
            print(f"\n!!! ERROR: {str(e)} !!!")
            import traceback
            traceback.print_exc()
            messages.error(request, f'Error updating meal plan: {str(e)}')
            return redirect('edit_meal_plan', meal_plan_id=meal_plan_id)
    
    # GET request - prepare data for editing
    
    # Serialize meal plan data for JavaScript
    meal_plan_data = {
        'meal_plan_id': meal_plan.meal_plan_id,
        'plan_name': meal_plan.plan_name,
        'start_date': meal_plan.start_date.strftime('%Y-%m-%d'),
        'end_date': meal_plan.end_date.strftime('%Y-%m-%d'),
        'days': []
    }
    
    # Add each day with its recipes
    for day in meal_plan.days.all():
        day_data = {
            'date': day.date.strftime('%Y-%m-%d'),
            'recipes': []
        }
        
        for meal_recipe in day.recipes.all():
            recipe = meal_recipe.recipe
            day_data['recipes'].append({
                'meal_plan_recipe_id': meal_recipe.meal_plan_recipe_id,
                'recipe_id': recipe.recipe_id,
                'recipe_name': recipe.recipe_name,
                'servings': meal_recipe.servings,
                'sort_order': meal_recipe.sort_order,
                'recipe_image': recipe.recipe_image.url if recipe.recipe_image else None,
                'prep_time': recipe.prep_time,
                'cook_time': recipe.cook_time,
                'difficulty_level': recipe.difficulty_level or '',
                'is_vegetarian': recipe.is_vegetarian,
                'courses': [course.name for course in recipe.courses.all()],
                'categories': [cat.name for cat in recipe.categories.all()],
                'proteins': [protein.name for protein in recipe.proteins.all()]
            })
        
        meal_plan_data['days'].append(day_data)
    
    # Get all available recipes for the selector
    recipes_qs = Recipe.objects.prefetch_related(
        'courses', 
        'categories', 
        'proteins'
    ).all().order_by('recipe_name')
    
    recipes = []
    for recipe in recipes_qs:
        recipe_data = {
            'recipe_id': recipe.recipe_id,
            'recipe_name': recipe.recipe_name,
            'servings': recipe.servings,
            'prep_time': recipe.prep_time,
            'cook_time': recipe.cook_time,
            'difficulty_level': recipe.difficulty_level or '',
            'is_vegetarian': recipe.is_vegetarian,
            'author': recipe.author,
            'recipe_image': recipe.recipe_image.url if recipe.recipe_image else None,
            'courses': [course.name for course in recipe.courses.all()],
            'categories': [cat.name for cat in recipe.categories.all()],
            'proteins': [protein.name for protein in recipe.proteins.all()]
        }
        recipes.append(recipe_data)
    
    # Get all filter options
    all_courses = list(RecipeCourse.objects.all().order_by('name').values('recipe_course_id', 'name'))
    all_categories = list(RecipeCategory.objects.all().order_by('name').values('recipe_category_id', 'name'))
    all_proteins = list(CustomProtein.objects.all().order_by('name').values('custom_protein_id', 'name'))
    
    # Add authors
    all_authors = [
        {'value': 'General', 'name': 'General'},
        {'value': 'Demetri & Angy', 'name': 'Demetri & Angy'},
        {'value': 'Erene', 'name': 'Erene'},
        {'value': 'Alexandra', 'name': 'Alexandra'},
    ]
    
    context = {
        'edit_mode': True,
        'meal_plan': meal_plan,
        'meal_plan_json': json.dumps(meal_plan_data),
        'recipes_json': json.dumps(recipes),
        'all_courses_json': json.dumps(all_courses),
        'all_categories_json': json.dumps(all_categories),
        'all_proteins_json': json.dumps(all_proteins),
        'all_authors_json': json.dumps(all_authors),
    }
    
    return render(request, 'create_meal_plan.html', context)

@login_required
def duplicate_meal_plan(request, meal_plan_id):
    """Duplicate a meal plan to new dates"""
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to perform this action.')
        return redirect('meal_plans')
    
    try:
        # Get the original meal plan with all related data
        original_plan = MealPlan.objects.prefetch_related(
            'days__recipes__recipe'
        ).get(meal_plan_id=meal_plan_id)
        
        if request.method == 'POST':
            new_plan_name = request.POST.get('new_plan_name', '').strip()
            new_start_date = request.POST.get('new_start_date')
            
            if not new_plan_name:
                messages.error(request, 'Please enter a name for the duplicated meal plan.')
                return redirect('view_meal_plan', meal_plan_id=meal_plan_id)
            
            if not new_start_date:
                messages.error(request, 'Please select a start date.')
                return redirect('view_meal_plan', meal_plan_id=meal_plan_id)
            
            try:
                # Parse the new start date
                new_start = datetime.strptime(new_start_date, '%Y-%m-%d').date()
                
                # Calculate duration of original plan
                duration = (original_plan.end_date - original_plan.start_date).days
                new_end = new_start + timedelta(days=duration)
                
                # Create new meal plan with user-provided name
                new_plan = MealPlan.objects.create(
                    plan_name=new_plan_name,
                    start_date=new_start,
                    end_date=new_end,
                    created_by=request.user
                )
                
                # Duplicate all days and recipes
                for day in original_plan.days.all().order_by('date'):
                    # Calculate the offset from original start date
                    day_offset = (day.date - original_plan.start_date).days
                    new_day_date = new_start + timedelta(days=day_offset)
                    
                    # Create new day
                    new_day = MealPlanDay.objects.create(
                        meal_plan=new_plan,
                        date=new_day_date
                    )
                    
                    # Copy all recipes for this day
                    for meal_recipe in day.recipes.all():
                        MealPlanRecipe.objects.create(
                            meal_plan_day=new_day,
                            recipe=meal_recipe.recipe,
                            servings=meal_recipe.servings
                        )
                
                # Check if user wants to delete original (shift functionality)
                delete_original = request.POST.get('delete_original') == 'yes'

                if delete_original:
                    # Delete the original meal plan (this is a "shift" operation)
                    original_plan_name = original_plan.plan_name
                    original_plan.delete()
                    messages.success(request, f'Meal plan "{original_plan_name}" shifted to new dates: {new_start.strftime("%B %d, %Y")} - {new_end.strftime("%B %d, %Y")}')
                else:
                    # Regular duplicate (keep original)
                    messages.success(request, f'Meal plan "{new_plan_name}" created successfully starting on {new_start.strftime("%B %d, %Y")}.')

                return redirect('view_meal_plan', meal_plan_id=new_plan.meal_plan_id)
                
            except ValueError as e:
                messages.error(request, f'Invalid date format: {str(e)}')
                return redirect('view_meal_plan', meal_plan_id=meal_plan_id)
            except Exception as e:
                messages.error(request, f'Error duplicating meal plan: {str(e)}')
                return redirect('view_meal_plan', meal_plan_id=meal_plan_id)
        
        # GET request - should not happen with modal, but redirect just in case
        return redirect('view_meal_plan', meal_plan_id=meal_plan_id)
        
    except MealPlan.DoesNotExist:
        messages.error(request, 'Meal plan not found.')
        return redirect('meal_plans')

@login_required
def meal_plan_calendar(request):
    """Calendar view for meal plans"""
    from datetime import timedelta
    from calendar import monthcalendar
    import json
    
    # Get requested month/year or default to current
    today = timezone.now().date()
    year = int(request.GET.get('year', today.year))
    month = int(request.GET.get('month', today.month))
    
    # Get selected week start date (Monday)
    selected_week_start = request.GET.get('week')
    if selected_week_start:
        selected_week_start = datetime.strptime(selected_week_start, '%Y-%m-%d').date()
    else:
        # Default to current week (find Monday)
        selected_week_start = today - timedelta(days=today.weekday())
    
    selected_week_end = selected_week_start + timedelta(days=6)
    
    # Build calendar data for the month
    # Get first day of month and adjust to start from Monday
    first_of_month = date(year, month, 1)
    
    # Get all days we need to show (including overflow from prev/next months)
    cal = monthcalendar(year, month)  # Returns weeks starting from Monday
    
    # Get all meal plans for this user
    all_meal_plans = MealPlan.objects.filter(created_by=request.user).order_by('-start_date')
    
    # Get meal plans that overlap with this month (for dot indicators)
    month_start = date(year, month, 1)
    if month == 12:
        month_end = date(year + 1, 1, 1) - timedelta(days=1)
    else:
        month_end = date(year, month + 1, 1) - timedelta(days=1)
    
    month_meal_plans = MealPlan.objects.filter(
        created_by=request.user,
        start_date__lte=month_end,
        end_date__gte=month_start
    )
    
    # Build a set of dates that have meals planned
    dates_with_meals = set()
    for plan in month_meal_plans:
        days_with_recipes = MealPlanDay.objects.filter(
            meal_plan=plan,
            recipes__isnull=False
        ).values_list('date', flat=True).distinct()
        dates_with_meals.update(days_with_recipes)
    
    # Build calendar weeks with metadata
    calendar_weeks = []
    for week in cal:
        week_data = []
        week_dates = []
        for day_num in week:
            if day_num == 0:
                week_data.append(None)
            else:
                day_date = date(year, month, day_num)
                week_dates.append(day_date)
                week_data.append({
                    'day': day_num,
                    'date': day_date,
                    'has_meal': day_date in dates_with_meals,
                    'is_today': day_date == today,
                })
        
        # Calculate week start (Monday) for this row
        if week_dates:
            # Find the Monday of this week
            first_valid_date = week_dates[0]
            week_start = first_valid_date - timedelta(days=first_valid_date.weekday())
        else:
            week_start = None
            
        calendar_weeks.append({
            'days': week_data,
            'week_start': week_start,
        })
    
    # Get the meal plan for the selected week (if any)
    selected_meal_plan = MealPlan.objects.filter(
        created_by=request.user,
        start_date__lte=selected_week_end,
        end_date__gte=selected_week_start
    ).first()
    
    # Build week detail data
    week_days = []
    for i in range(7):
        day_date = selected_week_start + timedelta(days=i)
        day_data = {
            'date': day_date,
            'day_name': day_date.strftime('%A'),
            'day_short': day_date.strftime('%a'),
            'day_num': day_date.day,
            'is_today': day_date == today,
            'recipes': [],
            'meal_plan_day_id': None,
        }
        
        # If there's a meal plan for this week, get recipes for this day
        if selected_meal_plan:
            meal_plan_day = MealPlanDay.objects.filter(
                meal_plan=selected_meal_plan,
                date=day_date
            ).first()
            
            if meal_plan_day:
                day_data['meal_plan_day_id'] = meal_plan_day.meal_plan_day_id
                recipes = MealPlanRecipe.objects.filter(
                    meal_plan_day=meal_plan_day
                ).select_related('recipe').order_by('sort_order')
                
                for mpr in recipes:
                    recipe = mpr.recipe
                    day_data['recipes'].append({
                        'meal_plan_recipe_id': mpr.meal_plan_recipe_id,
                        'recipe_id': recipe.recipe_id,
                        'name': recipe.recipe_name,
                        'image': recipe.recipe_image.url if recipe.recipe_image else None,
                        'prep_time': recipe.prep_time,
                        'cook_time': recipe.cook_time,
                        'total_time': (recipe.prep_time or 0) + (recipe.cook_time or 0),
                        'servings': mpr.servings,
                        'difficulty': recipe.difficulty_level,
                    })
        
        week_days.append(day_data)
    
    # Previous and next month for navigation
    if month == 1:
        prev_month = {'year': year - 1, 'month': 12}
    else:
        prev_month = {'year': year, 'month': month - 1}
    
    if month == 12:
        next_month = {'year': year + 1, 'month': 1}
    else:
        next_month = {'year': year, 'month': month + 1}
    
    # Month name for display
    month_name = date(year, month, 1).strftime('%B %Y')
    
    # Get all recipes for the "Add Recipe" modal
    all_recipes = Recipe.objects.all().order_by('recipe_name')
    
    context = {
        'calendar_weeks': calendar_weeks,
        'month_name': month_name,
        'year': year,
        'month': month,
        'prev_month': prev_month,
        'next_month': next_month,
        'today': today,
        'selected_week_start': selected_week_start,
        'selected_week_end': selected_week_end,
        'selected_meal_plan': selected_meal_plan,
        'week_days': week_days,
        'all_meal_plans': all_meal_plans,
        'all_recipes': all_recipes,
    }
    
    return render(request, 'meal_plan_calendar.html', context)

@login_required
def meal_plan_shopping_list(request, meal_plan_id):
    """Display shopping list with unit conversion and prompt for missing conversions"""
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to access this page.')
        return redirect('meal_plans')
    
    try:
        # Get meal plan with optimized prefetching
        meal_plan = get_object_or_404(
            MealPlan.objects.prefetch_related(
                Prefetch(
                    'days',
                    queryset=MealPlanDay.objects.order_by('date').prefetch_related(
                        Prefetch(
                            'recipes',
                            queryset=MealPlanRecipe.objects.select_related('recipe').prefetch_related(
                                Prefetch(
                                    'recipe__recipe_ingredients',
                                    queryset=RecipeIngredient.objects.select_related(
                                        'ingredient',
                                        'ingredient__category',
                                        'ingredient__default_unit',
                                        'unit'
                                    )
                                )
                            )
                        )
                    )
                )
            ),
            meal_plan_id=meal_plan_id
        )
        
        # Aggregate ingredients with conversion
        ingredients, missing_conversions, missing_shopping_units = aggregate_meal_plan_ingredients(meal_plan)
        
        # Check for missing conversions or shopping units
        has_issues = len(missing_conversions) > 0 or len(missing_shopping_units) > 0
        
        context = {
            'meal_plan': meal_plan,
            'ingredients': ingredients,
            'total_ingredients': sum(len(items) for items in ingredients.values()),
            'missing_conversions': json.dumps(missing_conversions) if missing_conversions else '[]',
            'missing_shopping_units': json.dumps(missing_shopping_units) if missing_shopping_units else '[]',
            'has_missing_conversions': len(missing_conversions) > 0,
            'has_missing_shopping_units': len(missing_shopping_units) > 0,
            'has_issues': has_issues,
        }
        
        return render(request, 'meal_plan_shopping_list.html', context)
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        messages.error(request, f'Error generating shopping list: {str(e)}')
        return redirect('view_meal_plan', meal_plan_id=meal_plan_id)

@login_required
@require_POST
def add_recipe_to_meal_plan_day(request):
    """Add a recipe to a meal plan day"""
    meal_plan_day_id = request.POST.get('meal_plan_day_id')
    recipe_id = request.POST.get('recipe_id')
    servings = request.POST.get('servings', 4)
    
    try:
        meal_plan_day = MealPlanDay.objects.get(meal_plan_day_id=meal_plan_day_id)
        recipe = Recipe.objects.get(recipe_id=recipe_id)
        
        # Verify user owns this meal plan
        if meal_plan_day.meal_plan.created_by != request.user:
            messages.error(request, 'You do not have permission to modify this meal plan.')
            return redirect('meal_plan_calendar')
        
        # Get max sort order for this day
        max_order = MealPlanRecipe.objects.filter(
            meal_plan_day=meal_plan_day
        ).aggregate(Max('sort_order'))['sort_order__max'] or 0
        
        # Create the meal plan recipe
        MealPlanRecipe.objects.create(
            meal_plan_day=meal_plan_day,
            recipe=recipe,
            servings=int(servings),
            sort_order=max_order + 1
        )
        
        messages.success(request, f'Added "{recipe.recipe_name}" to {meal_plan_day.date.strftime("%A, %B %d")}')
        
    except (MealPlanDay.DoesNotExist, Recipe.DoesNotExist) as e:
        messages.error(request, 'Error adding recipe. Please try again.')
    
    # Redirect back to calendar with current view
    return redirect(f"{reverse('meal_plan_calendar')}?week={meal_plan_day.date.strftime('%Y-%m-%d')}")


@login_required
@require_POST
def remove_recipe_from_meal_plan(request):
    """Remove a recipe from a meal plan day"""
    meal_plan_recipe_id = request.POST.get('meal_plan_recipe_id')
    
    try:
        meal_plan_recipe = MealPlanRecipe.objects.get(meal_plan_recipe_id=meal_plan_recipe_id)
        meal_plan_day = meal_plan_recipe.meal_plan_day
        
        # Verify user owns this meal plan
        if meal_plan_day.meal_plan.created_by != request.user:
            messages.error(request, 'You do not have permission to modify this meal plan.')
            return redirect('meal_plan_calendar')
        
        recipe_name = meal_plan_recipe.recipe.recipe_name
        week_date = meal_plan_day.date.strftime('%Y-%m-%d')
        
        meal_plan_recipe.delete()
        
        messages.success(request, f'Removed "{recipe_name}" from the meal plan')
        
        return redirect(f"{reverse('meal_plan_calendar')}?week={week_date}")
        
    except MealPlanRecipe.DoesNotExist:
        messages.error(request, 'Recipe not found.')
        return redirect('meal_plan_calendar')

def round_shopping_quantity(qty, unit):
    """
    Round quantities intelligently based on unit type for shopping lists.
    NEVER rounds down - always rounds UP to ensure you have enough.
    - COUNT units: Round UP to whole numbers (min 1)
    - WEIGHT/VOLUME: Round UP to nearest sensible number
    - OTHER: Round UP to whole numbers
    """
    import math
    
    if qty <= 0:
        return qty
    
    unit_type = getattr(unit, 'unit_type', 'other')
    
    if unit_type == 'count':
        # Always round UP to whole number, minimum 1
        return max(1, math.ceil(qty))
    
    elif unit_type == 'weight':
        # Round UP to sensible numbers based on magnitude
        if qty < 10:
            # Small amounts: round UP to nearest 0.5
            return math.ceil(qty * 2) / 2
        elif qty < 100:
            # Medium amounts: round UP to nearest 5
            return math.ceil(qty / 5) * 5
        else:
            # Large amounts: round UP to nearest 10
            return math.ceil(qty / 10) * 10
    
    elif unit_type == 'volume':
        # Round UP to sensible numbers based on magnitude
        if qty < 10:
            # Small amounts (tsp, tbsp): round UP to nearest 0.25
            return math.ceil(qty * 4) / 4
        elif qty < 100:
            # Medium amounts: round UP to nearest 5
            return math.ceil(qty / 5) * 5
        else:
            # Large amounts: round UP to nearest 10
            return math.ceil(qty / 10) * 10
    
    else:
        # OTHER (dash, pinch, to taste): round UP
        return max(1, math.ceil(qty))

@login_required
@require_POST
def generate_recipe_shopping_list(request):
    """Generate shopping list for a single recipe with unit conversion to shopping units"""
    if not request.user.is_superuser:
        return JsonResponse({'success': False, 'error': 'Permission denied'}, status=403)
    
    try:
        data = json.loads(request.body)
        recipe_id = data.get('recipe_id')
        servings = data.get('servings')
        
        recipe = Recipe.objects.select_related().prefetch_related(
            'recipe_ingredients__ingredient__category',
            'recipe_ingredients__ingredient__default_unit',
            'recipe_ingredients__unit'
        ).get(recipe_id=recipe_id)
        
        original_servings = recipe.servings or 4
        servings_multiplier = Decimal(servings) / Decimal(original_servings)
        
        # Track issues
        missing_conversions = []
        missing_shopping_units = []
        seen_missing_conversions = set()
        seen_missing_units = set()

        # Pre-load all unit conversions for fast lookups
        conversion_cache = get_conversion_cache()
        
        # Aggregate by ingredient ID
        aggregated = defaultdict(lambda: {
            'amount': Decimal('0'),
            'unit': None,
            'ingredient_obj': None,
            'unconverted_items': []
        })
        
        for recipe_ingredient in recipe.recipe_ingredients.all():
            ingredient = recipe_ingredient.ingredient
            ingredient_id = ingredient.ingredient_id
            quantity = Decimal(recipe_ingredient.amount or 0) * servings_multiplier
            from_unit = recipe_ingredient.unit
            
            if not from_unit:
                continue
            
            # CRITICAL: Always use shopping unit
            shopping_unit = ingredient.default_unit
            
            # Track ingredients without shopping units
            if not shopping_unit:
                if ingredient_id not in seen_missing_units:
                    seen_missing_units.add(ingredient_id)
                    missing_shopping_units.append({
                        'ingredient_id': ingredient_id,
                        'ingredient_name': ingredient.name
                    })
                
                # Use fallback unit
                if aggregated[ingredient_id]['unit'] is None:
                    aggregated[ingredient_id]['unit'] = from_unit
                    aggregated[ingredient_id]['ingredient_obj'] = ingredient
                
                aggregated[ingredient_id]['unconverted_items'].append({
                    'quantity': float(quantity),
                    'unit': from_unit.name,
                    'reason': 'No shopping unit defined'
                })
                continue
            
            # Initialize
            if aggregated[ingredient_id]['unit'] is None:
                aggregated[ingredient_id]['unit'] = shopping_unit
                aggregated[ingredient_id]['ingredient_obj'] = ingredient
            
            shopping_unit = aggregated[ingredient_id]['unit']
            
            # Convert to shopping unit
            if from_unit.measurement_unit_id == shopping_unit.measurement_unit_id:
                # Same unit - just add
                aggregated[ingredient_id]['amount'] += quantity
            else:
                # Need conversion
                converted_qty, multiplier = convert_quantity(quantity, from_unit, shopping_unit, ingredient, conversion_cache)
                
                if converted_qty is not None:
                    # Conversion successful
                    aggregated[ingredient_id]['amount'] += converted_qty
                else:
                    # Missing conversion
                    conversion_key = f"{from_unit.measurement_unit_id}-{shopping_unit.measurement_unit_id}"
                    if conversion_key not in seen_missing_conversions:
                        seen_missing_conversions.add(conversion_key)
                        missing_conversions.append({
                            'ingredient': ingredient.name,
                            'from_unit': from_unit.name,
                            'from_unit_id': from_unit.measurement_unit_id,
                            'to_unit': shopping_unit.name,
                            'to_unit_id': shopping_unit.measurement_unit_id,
                            'quantity': float(quantity)
                        })
                    
                    # Track unconverted
                    aggregated[ingredient_id]['unconverted_items'].append({
                        'quantity': float(quantity),
                        'unit': from_unit.name,
                        'reason': f'Missing conversion from {from_unit.name} to {shopping_unit.name}'
                    })
        
        # Build categorized shopping list
        shopping_list_categorized = defaultdict(list)
        
        for ingredient_id, agg_data in aggregated.items():
            ingredient_obj = agg_data['ingredient_obj']
            category = ingredient_obj.category.name if ingredient_obj.category else 'Other'
            
            # Only add if we have a converted amount
            if agg_data['amount'] > 0:
                unit = agg_data['unit']
                raw_qty = float(agg_data['amount'])
                
                # Apply smart rounding for shopping
                qty = round_shopping_quantity(raw_qty, unit)
                
                # Format quantity string
                if qty % 1 == 0:
                    qty_str = f"{int(qty)}"
                else:
                    qty_str = f"{qty:.2f}".rstrip('0').rstrip('.')
                
                # Get proper unit display with pluralization
                if qty == 1:
                    unit_display = unit.abbreviation or unit.name
                else:
                    unit_display = unit.abbreviation_plural or unit.abbreviation or unit.name_plural or unit.name
                
                item_str = f"{qty_str} {unit_display} {ingredient_obj.name}"
                
                # Add note if there are unconverted items
                if agg_data['unconverted_items']:
                    item_str += f" (+ unconverted items)"
                
                shopping_list_categorized[category].append(item_str)
        
        # Check if we have issues to prompt for
        has_issues = len(missing_conversions) > 0 or len(missing_shopping_units) > 0
        
        if has_issues:
            return JsonResponse({
                'success': False,
                'needs_conversions': True,
                'missing_conversions': missing_conversions,
                'missing_shopping_units': missing_shopping_units
            })
        
        return JsonResponse({
            'success': True,
            'shopping_list_categorized': dict(shopping_list_categorized),
            'recipe_name': recipe.recipe_name,
            'servings': servings,
            'original_servings': original_servings
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

def get_base_unit_for_ingredient(ingredient):
    """
    Determine the best base unit for an ingredient.
    Priority: 1) ingredient.default_unit, 2) category defaults, 3) gram
    Returns a MeasurementUnit object.
    """
    # PRIORITY 1: Use ingredient's default_unit if set
    if ingredient.default_unit:
        return ingredient.default_unit
    
    # PRIORITY 2: Use category-based defaults
    base_unit_map = {
        'Vegetables': 'gram',
        'Fruits': 'gram',
        'Meat & Seafood': 'gram',
        'Poultry': 'gram',
        'Dairy': 'milliliter',
        'Grains & Pasta': 'gram',
        'Oils & Fats': 'milliliter',
        'Baking': 'gram',
        'Herbs & Spices': 'gram',
        'Canned & Packaged': 'gram',
        'Beverages': 'milliliter',
    }
    
    if ingredient.category:
        preferred_unit_name = base_unit_map.get(ingredient.category.name, 'gram')
    else:
        preferred_unit_name = 'gram'
    
    # PRIORITY 3: Try to get the preferred unit
    try:
        return MeasurementUnit.objects.get(name__iexact=preferred_unit_name)
    except MeasurementUnit.DoesNotExist:
        # PRIORITY 4: Fallback to gram
        try:
            return MeasurementUnit.objects.get(name__iexact='gram')
        except:
            # PRIORITY 5: Last resort - any unit
            return MeasurementUnit.objects.first()

def get_conversion_cache():
    """
    Pre-load all unit conversions into memory for fast lookups.
    Returns a dictionary structure for O(1) lookup time.
    """
    cache = {
        'specific': {},  # ingredient_id -> {(from_id, to_id): multiplier}
        'generic': {}    # (from_id, to_id): multiplier
    }
    
    # Load all conversions in one query
    conversions = UnitConversion.objects.select_related('specific_ingredient').all()
    
    for conv in conversions:
        from_id = conv.from_unit_id
        to_id = conv.to_unit_id
        multiplier = conv.multiplier
        
        if conv.specific_ingredient_id:
            # Ingredient-specific conversion
            ing_id = conv.specific_ingredient_id
            if ing_id not in cache['specific']:
                cache['specific'][ing_id] = {}
            cache['specific'][ing_id][(from_id, to_id)] = multiplier
        else:
            # Generic conversion
            cache['generic'][(from_id, to_id)] = multiplier
    
    return cache

def convert_quantity(amount, from_unit, to_unit, ingredient=None, conversion_cache=None):
    """
    Convert a quantity from one unit to another using the UnitConversion table.
    Priority: 1) Ingredient-specific conversion, 2) Generic conversion
    Returns (converted_amount, multiplier) or (None, None) if no conversion exists.
    
    If conversion_cache is provided, uses in-memory lookups (much faster).
    Otherwise falls back to database queries.
    """
    if from_unit.measurement_unit_id == to_unit.measurement_unit_id:
        return (amount, Decimal('1'))
    
    from_id = from_unit.measurement_unit_id
    to_id = to_unit.measurement_unit_id
    
    # Use cache if provided (fast path)
    if conversion_cache:
        # PRIORITY 1: Check ingredient-specific conversion
        if ingredient and ingredient.ingredient_id in conversion_cache['specific']:
            specific_cache = conversion_cache['specific'][ingredient.ingredient_id]
            
            # Direct conversion
            if (from_id, to_id) in specific_cache:
                multiplier = specific_cache[(from_id, to_id)]
                converted = Decimal(amount) * multiplier
                return (converted, multiplier)
            
            # Reverse conversion
            if (to_id, from_id) in specific_cache:
                multiplier = specific_cache[(to_id, from_id)]
                converted = Decimal(amount) / multiplier
                return (converted, Decimal('1') / multiplier)
        
        # PRIORITY 2: Check generic conversion
        generic_cache = conversion_cache['generic']
        
        # Direct conversion
        if (from_id, to_id) in generic_cache:
            multiplier = generic_cache[(from_id, to_id)]
            converted = Decimal(amount) * multiplier
            return (converted, multiplier)
        
        # Reverse conversion
        if (to_id, from_id) in generic_cache:
            multiplier = generic_cache[(to_id, from_id)]
            converted = Decimal(amount) / multiplier
            return (converted, Decimal('1') / multiplier)
        
        # No conversion found
        return (None, None)
    
    # Fallback to database queries (slow path - for backwards compatibility)
    # PRIORITY 1: Check for INGREDIENT-SPECIFIC conversion (if ingredient provided)
    if ingredient:
        # Try direct conversion
        try:
            conversion = UnitConversion.objects.get(
                from_unit=from_unit, 
                to_unit=to_unit,
                specific_ingredient=ingredient
            )
            converted = Decimal(amount) * conversion.multiplier
            return (converted, conversion.multiplier)
        except UnitConversion.DoesNotExist:
            pass
        
        # Try reverse conversion
        try:
            conversion = UnitConversion.objects.get(
                from_unit=to_unit, 
                to_unit=from_unit,
                specific_ingredient=ingredient
            )
            converted = Decimal(amount) / conversion.multiplier
            return (converted, Decimal('1') / conversion.multiplier)
        except UnitConversion.DoesNotExist:
            pass
    
    # PRIORITY 2: Check for GENERIC conversion (specific_ingredient IS NULL)
    # Try direct conversion
    try:
        conversion = UnitConversion.objects.get(
            from_unit=from_unit, 
            to_unit=to_unit,
            specific_ingredient=None
        )
        converted = Decimal(amount) * conversion.multiplier
        return (converted, conversion.multiplier)
    except UnitConversion.DoesNotExist:
        pass
    
    # Try reverse conversion
    try:
        conversion = UnitConversion.objects.get(
            from_unit=to_unit, 
            to_unit=from_unit,
            specific_ingredient=None
        )
        converted = Decimal(amount) / conversion.multiplier
        return (converted, Decimal('1') / conversion.multiplier)
    except UnitConversion.DoesNotExist:
        pass
    
    # No conversion found (neither ingredient-specific nor generic)
    return (None, None)

@login_required
@require_POST
def save_unit_conversion(request):
    """Save a new unit conversion"""
    if not request.user.is_superuser:
        return JsonResponse({'success': False, 'error': 'Permission denied'}, status=403)
    
    try:
        data = json.loads(request.body)

        from_unit_id = data.get('from_unit_id')
        to_unit_id = data.get('to_unit_id')
        multiplier = data.get('multiplier')
        ingredient_name = data.get('ingredient_name')  # ← NEW: Get ingredient name if provided
        
        if not all([from_unit_id, to_unit_id, multiplier]):
            return JsonResponse({'success': False, 'error': 'Missing required fields'}, status=400)
        
        from_unit = MeasurementUnit.objects.get(measurement_unit_id=from_unit_id)
        to_unit = MeasurementUnit.objects.get(measurement_unit_id=to_unit_id)

        # Get specific ingredient if provided
        specific_ingredient = None
        if ingredient_name:
            try:
                specific_ingredient = Ingredient.objects.get(name__iexact=ingredient_name)
            except Ingredient.DoesNotExist:
                return JsonResponse({'success': False, 'error': f'Ingredient "{ingredient_name}" not found'}, status=404)

        try:
            mult = Decimal(multiplier)
            if mult <= 0:
                return JsonResponse({'success': False, 'error': 'Multiplier must be positive'}, status=400)
        except:
            return JsonResponse({'success': False, 'error': 'Invalid multiplier'}, status=400)
        
        # Check for existing conversion (with the same specific_ingredient setting)
        existing = UnitConversion.objects.filter(
            from_unit=from_unit, 
            to_unit=to_unit,
            specific_ingredient=specific_ingredient
        ).first()

        if existing:
            existing.multiplier = mult
            existing.save()
        else:
            # Create conversion (generic or ingredient-specific based on user choice)
            UnitConversion.objects.create(
                from_unit=from_unit,
                to_unit=to_unit,
                specific_ingredient=specific_ingredient,
                multiplier=mult
            )
        
        return JsonResponse({
            'success': True,
            'message': f'Conversion saved: 1 {from_unit.name} = {mult} {to_unit.name}'
        })
        
    except MeasurementUnit.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Unit not found'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

# ============================================
# UNIT CONVERSION MANAGEMENT
# ============================================

def scan_for_missing_conversions():
    """
    Scan all recipe ingredients and find missing unit conversions.
    Returns list of missing conversions needed.
    """
    from collections import defaultdict
    
    missing = []
    seen = set()  # To avoid duplicates
    
    # Get conversion cache for fast lookups
    conversion_cache = get_conversion_cache()
    
    # Get all recipe ingredients with their units
    recipe_ingredients = RecipeIngredient.objects.select_related(
        'ingredient',
        'ingredient__default_unit',
        'unit',
        'recipe'
    ).exclude(
        ingredient__default_unit__isnull=True  # Skip ingredients without shopping units
    ).exclude(
        unit__isnull=True  # Skip ingredients without recipe units
    )
    
    for ri in recipe_ingredients:
        ingredient = ri.ingredient
        from_unit = ri.unit
        to_unit = ingredient.default_unit
        
        # Skip if same unit
        if from_unit.measurement_unit_id == to_unit.measurement_unit_id:
            continue
        
        # Create unique key to avoid duplicates
        cache_key = f"{from_unit.measurement_unit_id}-{to_unit.measurement_unit_id}-{ingredient.ingredient_id}"
        
        if cache_key in seen:
            continue
        
        # Check if conversion exists
        converted_qty, multiplier = convert_quantity(
            Decimal('1'), 
            from_unit, 
            to_unit, 
            ingredient, 
            conversion_cache
        )
        
        if converted_qty is None:
            # Missing conversion found
            seen.add(cache_key)
            missing.append({
                'ingredient': ingredient.name,
                'ingredient_id': ingredient.ingredient_id,
                'from_unit': from_unit.name,
                'from_unit_id': from_unit.measurement_unit_id,
                'to_unit': to_unit.name,
                'to_unit_id': to_unit.measurement_unit_id,
                'recipe_example': ri.recipe.recipe_name
            })
    
    return missing

@login_required
def add_conversion(request):
    """Add a new unit conversion via AJAX"""
    if not request.user.is_superuser:
        return JsonResponse({'success': False, 'error': 'Permission denied'})
    
    if request.method == 'POST':
        try:
            from_unit_id = request.POST.get('from_unit')
            to_unit_id = request.POST.get('to_unit')
            specific_ingredient_id = request.POST.get('specific_ingredient')
            multiplier = request.POST.get('multiplier')
            
            if not all([from_unit_id, to_unit_id, multiplier]):
                return JsonResponse({'success': False, 'error': 'Missing required fields'})
            
            from_unit = MeasurementUnit.objects.get(measurement_unit_id=from_unit_id)
            to_unit = MeasurementUnit.objects.get(measurement_unit_id=to_unit_id)
            
            # Handle specific ingredient (can be null for generic conversions)
            specific_ingredient = None
            if specific_ingredient_id and specific_ingredient_id != 'null':
                specific_ingredient = Ingredient.objects.get(ingredient_id=specific_ingredient_id)
            
            # Check if conversion already exists
            existing = UnitConversion.objects.filter(
                from_unit=from_unit,
                to_unit=to_unit,
                specific_ingredient=specific_ingredient
            ).first()
            
            if existing:
                # Update existing
                existing.multiplier = multiplier
                existing.save()
            else:
                # Create new
                UnitConversion.objects.create(
                    from_unit=from_unit,
                    to_unit=to_unit,
                    specific_ingredient=specific_ingredient,
                    multiplier=multiplier
                )
            
            return JsonResponse({'success': True})
            
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'success': False, 'error': 'Invalid request method'})

@login_required
def unit_conversions_management(request):
    """Manage unit conversions"""
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to access this page.')
        return redirect('home')
    
    conversions = UnitConversion.objects.all().select_related(
        'from_unit', 
        'to_unit', 
        'specific_ingredient'
    ).order_by('from_unit__name', 'to_unit__name')
    
    all_units = MeasurementUnit.objects.all().order_by('name')
    all_ingredients = Ingredient.objects.all().order_by('name')
    
    # Scan for missing conversions across all recipes
    missing_conversions = scan_for_missing_conversions()
    
    context = {
        'conversions': conversions,
        'all_units': all_units,
        'all_ingredients': all_ingredients,
        'missing_conversions': json.dumps(missing_conversions) if missing_conversions else '[]',
        'missing_count': len(missing_conversions),
    }
    
    return render(request, 'unit_conversions_management.html', context)

@login_required
@require_POST
def add_unit_conversion_manual(request):
    """Add a new unit conversion"""
    if not request.user.is_superuser:
        return JsonResponse({'success': False, 'error': 'Permission denied'}, status=403)
    
    try:
        data = json.loads(request.body)
        
        from_unit_id = data.get('from_unit_id')
        to_unit_id = data.get('to_unit_id')
        specific_ingredient_id = data.get('specific_ingredient_id')  # ← NEW
        multiplier = data.get('multiplier')
        
        if not all([from_unit_id, to_unit_id, multiplier]):
            return JsonResponse({'success': False, 'error': 'All required fields must be filled'}, status=400)
        
        # Check if same unit
        if from_unit_id == to_unit_id:
            return JsonResponse({'success': False, 'error': 'Cannot convert a unit to itself'}, status=400)
        
        from_unit = MeasurementUnit.objects.get(measurement_unit_id=from_unit_id)
        to_unit = MeasurementUnit.objects.get(measurement_unit_id=to_unit_id)
        
        # Get specific ingredient if provided
        specific_ingredient = None
        if specific_ingredient_id:
            try:
                specific_ingredient = Ingredient.objects.get(ingredient_id=specific_ingredient_id)
            except Ingredient.DoesNotExist:
                return JsonResponse({'success': False, 'error': 'Ingredient not found'}, status=404)
        
        mult = Decimal(multiplier)
        if mult <= 0:
            return JsonResponse({'success': False, 'error': 'Multiplier must be positive'}, status=400)
        
        # Check if already exists (now includes specific_ingredient in the check)
        if UnitConversion.objects.filter(
            from_unit=from_unit, 
            to_unit=to_unit,
            specific_ingredient=specific_ingredient
        ).exists():
            if specific_ingredient:
                return JsonResponse({'success': False, 'error': f'This conversion already exists for {specific_ingredient.name}'}, status=400)
            else:
                return JsonResponse({'success': False, 'error': 'This generic conversion already exists'}, status=400)
        
        # Create conversion
        conversion = UnitConversion.objects.create(
            from_unit=from_unit,
            to_unit=to_unit,
            specific_ingredient=specific_ingredient,  # ← NEW
            multiplier=mult
        )
        
        return JsonResponse({
            'success': True,
            'conversion': {
                'id': conversion.unit_conversion_id,
                'from_unit': from_unit.name,
                'to_unit': to_unit.name,
                'specific_ingredient': specific_ingredient.name if specific_ingredient else 'Generic',
                'multiplier': str(mult)
            }
        })
        
    except MeasurementUnit.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Unit not found'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@login_required
@require_POST
def edit_unit_conversion(request):
    """Edit an existing unit conversion"""
    if not request.user.is_superuser:
        return JsonResponse({'success': False, 'error': 'Permission denied'}, status=403)
    
    try:
        data = json.loads(request.body)
        
        conversion_id = data.get('conversion_id')
        multiplier = data.get('multiplier')
        specific_ingredient_id = data.get('specific_ingredient_id')  # ← NEW
        
        if not all([conversion_id, multiplier]):
            return JsonResponse({'success': False, 'error': 'Missing required fields'}, status=400)
        
        conversion = UnitConversion.objects.get(unit_conversion_id=conversion_id)
        
        mult = Decimal(multiplier)
        if mult <= 0:
            return JsonResponse({'success': False, 'error': 'Multiplier must be positive'}, status=400)
        
        # Get specific ingredient if provided
        specific_ingredient = None
        if specific_ingredient_id:
            try:
                specific_ingredient = Ingredient.objects.get(ingredient_id=specific_ingredient_id)
            except Ingredient.DoesNotExist:
                return JsonResponse({'success': False, 'error': 'Ingredient not found'}, status=404)
        
        # Check if changing specific_ingredient would create a duplicate
        duplicate = UnitConversion.objects.filter(
            from_unit=conversion.from_unit,
            to_unit=conversion.to_unit,
            specific_ingredient=specific_ingredient
        ).exclude(unit_conversion_id=conversion_id).exists()
        
        if duplicate:
            if specific_ingredient:
                return JsonResponse({'success': False, 'error': f'A conversion for {specific_ingredient.name} already exists'}, status=400)
            else:
                return JsonResponse({'success': False, 'error': 'A generic conversion already exists'}, status=400)
        
        # Update the conversion
        conversion.multiplier = mult
        conversion.specific_ingredient = specific_ingredient
        conversion.save()
        
        applies_to = specific_ingredient.name if specific_ingredient else 'all ingredients'
        
        return JsonResponse({
            'success': True,
            'message': f'Updated: 1 {conversion.from_unit.name} = {mult} {conversion.to_unit.name} (applies to {applies_to})'
        })
        
    except UnitConversion.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Conversion not found'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@login_required
@require_POST
def delete_unit_conversion(request):
    """Delete a unit conversion"""
    if not request.user.is_superuser:
        return JsonResponse({'success': False, 'error': 'Permission denied'}, status=403)
    
    try:
        data = json.loads(request.body)
        conversion_id = data.get('conversion_id')
        
        if not conversion_id:
            return JsonResponse({'success': False, 'error': 'Conversion ID required'}, status=400)
        
        conversion = UnitConversion.objects.get(unit_conversion_id=conversion_id)
        conversion_text = f"1 {conversion.from_unit.name} = {conversion.multiplier} {conversion.to_unit.name}"
        conversion.delete()
        
        return JsonResponse({
            'success': True,
            'message': f'Deleted conversion: {conversion_text}'
        })
        
    except UnitConversion.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Conversion not found'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


# ============================================
# INGREDIENT BASE UNIT MANAGEMENT
# ============================================

@login_required
def ingredient_base_units_management(request):
    """Manage ingredient shopping units"""
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to access this page.')
        return redirect('home')
    
    # Get filter parameters
    search_query = request.GET.get('search', '')
    category_filter = request.GET.get('category', '')
    
    # Query ingredients
    ingredients = Ingredient.objects.select_related('category', 'default_unit').all()
    
    if search_query:
        ingredients = ingredients.filter(name__icontains=search_query)
    
    if category_filter:
        ingredients = ingredients.filter(category__ingredient_category_id=category_filter)
    
    ingredients = ingredients.order_by('name')
    
    # Get all units and categories
    all_units = MeasurementUnit.objects.all().order_by('name')
    categories = IngredientCategory.objects.all().order_by('name')
    
    # Simple list - no auto-calculation
    ingredients_with_base = [{'ingredient': ing} for ing in ingredients]
    
    context = {
        'ingredients_with_base': ingredients_with_base,
        'categories': categories,
        'all_units': all_units,
        'search_query': search_query,
        'category_filter': category_filter,
    }
    
    return render(request, 'ingredient_base_units_management.html', context)
    
    # Helper function to calculate auto unit
    def calculate_auto_unit(ingredient):
        """Calculate what the auto unit would be based on category"""
        if ingredient.category:
            preferred_unit_name = base_unit_map.get(ingredient.category.name, 'gram')
        else:
            preferred_unit_name = 'gram'
        
        # Look up from our pre-loaded units
        auto_unit = unit_lookup.get(preferred_unit_name.lower())
        if not auto_unit:
            # Fallback
            auto_unit = unit_lookup.get('gram') or all_units[0] if all_units else None
        return auto_unit
    
    # Compute the effective base unit for each ingredient efficiently
    ingredients_with_base = []
    for ingredient in ingredients:
        # Always calculate what the auto unit would be
        auto_unit = calculate_auto_unit(ingredient)
        
        # Determine effective unit (what's currently being used)
        if ingredient.default_unit:
            effective_base_unit = ingredient.default_unit
        else:
            effective_base_unit = auto_unit
        
        ingredients_with_base.append({
            'ingredient': ingredient,
            'effective_base_unit': effective_base_unit,
            'auto_base_unit': auto_unit,  # NEW: Always include the auto-calculated unit
            'is_manual': ingredient.default_unit is not None
        })
    
    # Get all categories for filters
    categories = IngredientCategory.objects.all().order_by('name')
    
    context = {
        'ingredients_with_base': ingredients_with_base,
        'categories': categories,
        'all_units': all_units,
        'search_query': search_query,
        'category_filter': category_filter,
    }
    
    return render(request, 'ingredient_base_units_management.html', context)

@login_required
@require_POST
def update_ingredient_base_unit(request):
    """Update an ingredient's default unit"""
    if not request.user.is_superuser:
        return JsonResponse({'success': False, 'error': 'Permission denied'}, status=403)
    
    try:
        data = json.loads(request.body)
        
        ingredient_id = data.get('ingredient_id')
        base_unit_id = data.get('base_unit_id')
        
        if not ingredient_id:
            return JsonResponse({'success': False, 'error': 'Ingredient ID required'}, status=400)
        
        ingredient = Ingredient.objects.get(ingredient_id=ingredient_id)
        
        # If base_unit_id is None or empty, clear the override
        if not base_unit_id:
            ingredient.default_unit = None
            ingredient.save()
            return JsonResponse({
                'success': True,
                'message': f'Cleared override for {ingredient.name} - now using automatic unit'
            })
        
        base_unit = MeasurementUnit.objects.get(measurement_unit_id=base_unit_id)
        ingredient.default_unit = base_unit
        ingredient.save()
        
        return JsonResponse({
            'success': True,
            'message': f'Updated {ingredient.name} default unit to {base_unit.name}'
        })
        
    except Ingredient.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Ingredient not found'}, status=404)
    except MeasurementUnit.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Unit not found'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@login_required
@require_POST
def send_meal_plan_shopping_list(request):
    """Send meal plan shopping list via email"""
    if not request.user.is_superuser:
        return JsonResponse({'success': False, 'error': 'Permission denied'}, status=403)
    
    try:
        data = json.loads(request.body)
        
        meal_plan_name = data.get('meal_plan_name')
        date_range = data.get('date_range')
        total_recipes = data.get('total_recipes')
        email = data.get('email')
        ingredients_by_category = data.get('ingredients', {})
        
        # Validate
        if not email:
            return JsonResponse({'success': False, 'error': 'Email address is required'}, status=400)
        
        if not ingredients_by_category:
            return JsonResponse({'success': False, 'error': 'No ingredients to send'}, status=400)
        
        # Build email content
        subject = f'🛒 Shopping List for {meal_plan_name}'
        
        # Plain text version
        text_content = f"""Shopping List for {meal_plan_name}

{date_range}
{total_recipes} recipes

Items to Buy:
"""
        
        for category in sorted(ingredients_by_category.keys()):
            text_content += f"\n{category}:\n"
            for item in ingredients_by_category[category]:
                qty = item['quantity']
                # Format quantity nicely
                if qty % 1 == 0:
                    qty_str = f"{int(qty)}"
                else:
                    qty_str = f"{qty:.2f}".rstrip('0').rstrip('.')
                text_content += f"☐ {qty_str} {item['unit']} {item['ingredient']}\n"
        
        text_content += "\n---\nGenerated by ALIVENTE ONLINE - Recipe Management"
        
        # HTML version
        html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <style>
        body {{
            font-family: Arial, sans-serif;
            max-width: 600px;
            margin: 0 auto;
            padding: 20px;
            color: #2c3e50;
        }}
        .header {{
            background: linear-gradient(135deg, #28a745 0%, #20c997 100%);
            color: white;
            padding: 30px;
            border-radius: 12px;
            margin-bottom: 30px;
        }}
        .header h1 {{
            margin: 0 0 10px 0;
            font-size: 28px;
        }}
        .header p {{
            margin: 5px 0;
            opacity: 0.95;
        }}
        .category {{
            margin-bottom: 25px;
        }}
        .category-header {{
            color: #28a745;
            font-size: 20px;
            font-weight: 600;
            padding-bottom: 10px;
            border-bottom: 2px solid #e9ecef;
            margin-bottom: 15px;
        }}
        ul {{
            list-style: none;
            padding-left: 0;
        }}
        li {{
            padding: 12px;
            border-bottom: 1px solid #e9ecef;
            display: flex;
            align-items: center;
            gap: 10px;
        }}
        li:hover {{
            background: #f8f9fa;
        }}
        li:last-child {{
            border-bottom: none;
        }}
        .checkbox {{
            width: 18px;
            height: 18px;
            border: 2px solid #28a745;
            border-radius: 3px;
            flex-shrink: 0;
        }}
        .quantity {{
            font-weight: 600;
            color: #28a745;
            margin-right: 5px;
        }}
        .footer {{
            margin-top: 30px;
            padding-top: 20px;
            border-top: 2px solid #e9ecef;
            text-align: center;
            color: #6c757d;
            font-size: 14px;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>🛒 Shopping List</h1>
        <p><strong>{meal_plan_name}</strong></p>
        <p>{date_range}</p>
        <p>{total_recipes} recipes</p>
    </div>
"""
        
        # Add categories
        for category in sorted(ingredients_by_category.keys()):
            html_content += f"""
    <div class="category">
        <div class="category-header">📌 {category}</div>
        <ul>
"""
            for item in ingredients_by_category[category]:
                qty = item['quantity']
                if qty % 1 == 0:
                    qty_str = f"{int(qty)}"
                else:
                    qty_str = f"{qty:.2f}".rstrip('0').rstrip('.')
                
                html_content += f"""
            <li>
                <div class="checkbox"></div>
                <span><span class="quantity">{qty_str} {item['unit']}</span> {item['ingredient']}</span>
            </li>
"""
            html_content += """
        </ul>
    </div>
"""
        
        html_content += f"""
    <div class="footer">
        <p>Generated by <strong>ALIVENTE ONLINE</strong> - Recipe Management</p>
        <p>{datetime.now().strftime("%B %d, %Y at %I:%M %p")}</p>
    </div>
</body>
</html>
"""
        
        # Send email using same method as recipe shopping list
        msg = EmailMultiAlternatives(
            subject,
            text_content,
            settings.DEFAULT_FROM_EMAIL,
            [email]
        )
        msg.attach_alternative(html_content, "text/html")
        msg.send()
        
        return JsonResponse({
            'success': True,
            'message': f'Shopping list sent to {email}'
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

@login_required
def check_ingredient_usage(request):
    """Check if ingredient is used in any recipes"""
    if request.method == 'POST':
        data = json.loads(request.body)
        ingredient_id = data.get('ingredient_id')
        
        try:
            ingredient = Ingredient.objects.get(ingredient_id=ingredient_id)
            
            # Count recipes using this ingredient
            usage_count = RecipeIngredient.objects.filter(ingredient=ingredient).count()
            
            # Get recipe names (limit to 5 for display)
            recipes = RecipeIngredient.objects.filter(ingredient=ingredient).select_related('recipe')[:5]
            recipe_names = [ri.recipe.recipe_name for ri in recipes]
            
            return JsonResponse({
                'success': True,
                'usage_count': usage_count,
                'recipe_names': recipe_names,
                'can_delete': usage_count == 0
            })
            
        except Ingredient.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Ingredient not found'})
    
    return JsonResponse({'success': False, 'error': 'Invalid request'})

@login_required
def delete_ingredient(request):
    """Delete ingredient if not used in any recipes"""
    if request.method == 'POST':
        data = json.loads(request.body)
        ingredient_id = data.get('ingredient_id')
        
        try:
            ingredient = Ingredient.objects.get(ingredient_id=ingredient_id)
            
            # Check if ingredient is used in any recipes
            usage_count = RecipeIngredient.objects.filter(ingredient=ingredient).count()
            
            if usage_count > 0:
                return JsonResponse({
                    'success': False,
                    'error': f'Cannot delete - ingredient is used in {usage_count} recipe(s)'
                })
            
            # Safe to delete
            ingredient_name = ingredient.name
            ingredient.delete()
            
            return JsonResponse({
                'success': True,
                'message': f'Successfully deleted ingredient: {ingredient_name}'
            })
            
        except Ingredient.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Ingredient not found'})
    
    return JsonResponse({'success': False, 'error': 'Invalid request'})

@login_required
def update_ingredient_full(request):
    """Update ingredient name, category, and shopping unit"""
    if request.method == 'POST':
        data = json.loads(request.body)
        ingredient_id = data.get('ingredient_id')
        new_name = data.get('name', '').strip()
        category_id = data.get('category_id')
        unit_id = data.get('unit_id')
        
        try:
            ingredient = Ingredient.objects.get(ingredient_id=ingredient_id)
            
            # Validate name is not empty
            if not new_name:
                return JsonResponse({'success': False, 'error': 'Ingredient name cannot be empty'})
            
            # Check for duplicate names (case-insensitive, excluding current ingredient)
            duplicate = Ingredient.objects.filter(name__iexact=new_name).exclude(ingredient_id=ingredient_id).first()
            if duplicate:
                return JsonResponse({
                    'success': False,
                    'error': f'An ingredient named "{new_name}" already exists'
                })
            
            # Validate category
            if not category_id:
                return JsonResponse({'success': False, 'error': 'Category must be selected'})
            
            try:
                category = IngredientCategory.objects.get(ingredient_category_id=category_id)
            except IngredientCategory.DoesNotExist:
                return JsonResponse({'success': False, 'error': 'Invalid category'})
            
            # Validate unit (optional - can be None)
            unit = None
            if unit_id:
                try:
                    unit = MeasurementUnit.objects.get(measurement_unit_id=unit_id)
                except MeasurementUnit.DoesNotExist:
                    return JsonResponse({'success': False, 'error': 'Invalid unit'})
            
            # Update ingredient
            ingredient.name = new_name
            ingredient.category = category
            ingredient.default_unit = unit
            ingredient.save()
            
            return JsonResponse({
                'success': True,
                'message': 'Ingredient updated successfully',
                'ingredient': {
                    'name': ingredient.name,
                    'category_id': ingredient.category.ingredient_category_id,
                    'category_name': ingredient.category.name,
                    'unit_id': ingredient.default_unit.measurement_unit_id if ingredient.default_unit else None,
                    'unit_name': ingredient.default_unit.name if ingredient.default_unit else None
                }
            })
            
        except Ingredient.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Ingredient not found'})
    
    return JsonResponse({'success': False, 'error': 'Invalid request'})

@login_required
def categories_management(request):
    """Manage ingredient categories"""
    # Get all categories sorted alphabetically
    categories = IngredientCategory.objects.all().order_by('name')
    
    # Get ingredient count and names for each category
    categories_with_count = []
    for category in categories:
        ingredients = Ingredient.objects.filter(category=category).order_by('name')
        ingredient_names = list(ingredients.values_list('name', flat=True))
        
        categories_with_count.append({
            'category': category,
            'ingredient_count': ingredients.count(),
            'ingredients': ingredient_names,  # List of ingredient names for tooltip
        })
    
    context = {
        'categories_with_count': categories_with_count
    }
    
    return render(request, 'categories_management.html', context)

@login_required
def add_category(request):
    """Add a new ingredient category"""
    if request.method == 'POST':
        data = json.loads(request.body)
        category_name = data.get('name', '').strip()
        
        # Validate name is not empty
        if not category_name:
            return JsonResponse({'success': False, 'error': 'Category name cannot be empty'})
        
        # Check for duplicate names (case-insensitive)
        duplicate = IngredientCategory.objects.filter(name__iexact=category_name).first()
        if duplicate:
            return JsonResponse({
                'success': False,
                'error': f'A category named "{category_name}" already exists'
            })
        
        # Create new category
        category = IngredientCategory.objects.create(name=category_name)
        
        return JsonResponse({
            'success': True,
            'message': 'Category added successfully',
            'category': {
                'id': category.ingredient_category_id,
                'name': category.name
            }
        })
   
    return JsonResponse({'success': False, 'error': 'Invalid request'})

@login_required
def update_category(request):
    """Update category name"""
    if request.method == 'POST':
        data = json.loads(request.body)
        category_id = data.get('category_id')
        new_name = data.get('name', '').strip()
        
        try:
            category = IngredientCategory.objects.get(ingredient_category_id=category_id)
            
            # Validate name is not empty
            if not new_name:
                return JsonResponse({'success': False, 'error': 'Category name cannot be empty'})
            
            # Check for duplicate names (case-insensitive, excluding current category)
            duplicate = IngredientCategory.objects.filter(name__iexact=new_name).exclude(ingredient_category_id=category_id).first()
            if duplicate:
                return JsonResponse({
                    'success': False,
                    'error': f'A category named "{new_name}" already exists'
                })
            
            # Update category
            category.name = new_name
            category.save()
            
            return JsonResponse({
                'success': True,
                'message': 'Category updated successfully',
                'category': {
                    'id': category.ingredient_category_id,
                    'name': category.name
                }
            })
            
        except IngredientCategory.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Category not found'})
    
    return JsonResponse({'success': False, 'error': 'Invalid request'})

@login_required
def check_category_usage(request):
    """Check if category is used by any ingredients"""
    if request.method == 'POST':
        data = json.loads(request.body)
        category_id = data.get('category_id')
        
        try:
            category = IngredientCategory.objects.get(ingredient_category_id=category_id)
            
            # Count ingredients using this category
            ingredient_count = Ingredient.objects.filter(category=category).count()
            
            # Get ingredient names (limit to 5 for display)
            ingredients = Ingredient.objects.filter(category=category)[:5]
            ingredient_names = [ing.name for ing in ingredients]
            
            return JsonResponse({
                'success': True,
                'usage_count': ingredient_count,
                'ingredient_names': ingredient_names,
                'can_delete': ingredient_count == 0
            })
            
        except IngredientCategory.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Category not found'})
    
    return JsonResponse({'success': False, 'error': 'Invalid request'})

@login_required
def delete_category(request):
    """Delete category if not used by any ingredients"""
    if request.method == 'POST':
        data = json.loads(request.body)
        category_id = data.get('category_id')
        
        try:
            category = IngredientCategory.objects.get(ingredient_category_id=category_id)
            
            # Check if category is used by any ingredients
            ingredient_count = Ingredient.objects.filter(category=category).count()
            
            if ingredient_count > 0:
                return JsonResponse({
                    'success': False,
                    'error': f'Cannot delete - category is used by {ingredient_count} ingredient(s)'
                })
            
            # Safe to delete
            category_name = category.name
            category.delete()
            
            return JsonResponse({
                'success': True,
                'message': f'Successfully deleted category: {category_name}'
            })
            
        except IngredientCategory.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Category not found'})
    
    return JsonResponse({'success': False, 'error': 'Invalid request'})

@login_required
def measurement_units_management(request):
    """Manage measurement units - with usage details for popups - OPTIMIZED"""
    from django.db.models import Count
    from collections import defaultdict
    
    # Get all units with usage counts
    units = MeasurementUnit.objects.annotate(
        recipe_usage=Count('recipeingredient', distinct=True),
        ingredient_usage=Count('ingredient', distinct=True),
        from_conversion_usage=Count('conversions_from', distinct=True),
        to_conversion_usage=Count('conversions_to', distinct=True)
    ).order_by('name')
    
    # Pre-fetch all recipe names by unit (single query)
    recipe_names_by_unit = defaultdict(list)
    recipe_data = (
        RecipeIngredient.objects
        .select_related('recipe', 'unit')
        .values('unit_id', 'recipe__recipe_name')
        .distinct()
        .order_by('recipe__recipe_name')
    )
    for item in recipe_data:
        recipe_names_by_unit[item['unit_id']].append(item['recipe__recipe_name'])
    
    # Pre-fetch all ingredient names by unit (single query)
    ingredient_names_by_unit = defaultdict(list)
    ingredient_data = (
        Ingredient.objects
        .filter(default_unit__isnull=False)
        .values('default_unit_id', 'name')
        .order_by('name')
    )
    for item in ingredient_data:
        ingredient_names_by_unit[item['default_unit_id']].append(item['name'])
    
    # Pre-fetch all conversions (two queries total)
    conversions_by_unit = defaultdict(list)
    
    # Conversions FROM each unit
    from_conversions = (
        UnitConversion.objects
        .select_related('from_unit', 'to_unit')
        .all()
    )
    for conv in from_conversions:
        from_abbr = conv.from_unit.abbreviation or conv.from_unit.name
        to_abbr = conv.to_unit.abbreviation or conv.to_unit.name
        conversions_by_unit[conv.from_unit_id].append(f"{from_abbr} → {to_abbr} (×{conv.multiplier})")
        conversions_by_unit[conv.to_unit_id].append(f"{from_abbr} → {to_abbr} (×{conv.multiplier})")
    
    # Build the final list
    units_with_count = []
    for unit in units:
        conversion_count = unit.from_conversion_usage + unit.to_conversion_usage
        total_usage = unit.recipe_usage + unit.ingredient_usage + conversion_count
        
        units_with_count.append({
            'unit': unit,
            'recipe_count': unit.recipe_usage,
            'ingredient_count': unit.ingredient_usage,
            'conversion_count': conversion_count,
            'total_usage': total_usage,
            'recipes': recipe_names_by_unit.get(unit.measurement_unit_id, []),
            'ingredients': ingredient_names_by_unit.get(unit.measurement_unit_id, []),
            'conversions': conversions_by_unit.get(unit.measurement_unit_id, []),
        })
    
    # Get unit type choices for the modal
    unit_types = MeasurementUnit.UNIT_TYPE_CHOICES
    
    context = {
        'units_with_count': units_with_count,
        'unit_types': unit_types
    }
    
    return render(request, 'measurement_units_management.html', context)

@login_required
def add_measurement_unit(request):
    """Add a new measurement unit"""
    if request.method == 'POST':
        data = json.loads(request.body)
        unit_name = data.get('name', '').strip()
        name_plural = data.get('name_plural', '').strip()
        abbreviation = data.get('abbreviation', '').strip()
        abbreviation_plural = data.get('abbreviation_plural', '').strip()
        unit_type = data.get('unit_type', 'other')
        
        # Validate name is not empty
        if not unit_name:
            return JsonResponse({'success': False, 'error': 'Unit name cannot be empty'})
        
        # Check for duplicate names (case-insensitive)
        duplicate = MeasurementUnit.objects.filter(name__iexact=unit_name).first()
        if duplicate:
            return JsonResponse({
                'success': False,
                'error': f'A unit named "{unit_name}" already exists'
            })
        
        # Check for duplicate abbreviations if provided (case-insensitive)
        if abbreviation:
            duplicate_abbr = MeasurementUnit.objects.filter(abbreviation__iexact=abbreviation).first()
            if duplicate_abbr:
                return JsonResponse({
                    'success': False,
                    'error': f'A unit with abbreviation "{abbreviation}" already exists'
                })
        
        # Create new unit with plural fields
        unit = MeasurementUnit.objects.create(
            name=unit_name,
            name_plural=name_plural if name_plural else None,
            abbreviation=abbreviation if abbreviation else None,
            abbreviation_plural=abbreviation_plural if abbreviation_plural else None,
            unit_type=unit_type
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Measurement unit added successfully',
            'unit': {
                'id': unit.measurement_unit_id,
                'name': unit.name,
                'name_plural': unit.name_plural,
                'abbreviation': unit.abbreviation,
                'abbreviation_plural': unit.abbreviation_plural,
                'unit_type': unit.unit_type,
                'unit_type_display': unit.get_unit_type_display()
            }
        })
    
    return JsonResponse({'success': False, 'error': 'Invalid request'})

@login_required
def update_measurement_unit(request):
    """Update measurement unit name, abbreviation, and type"""
    if request.method == 'POST':
        data = json.loads(request.body)
        unit_id = data.get('unit_id')
        new_name = data.get('name', '').strip()
        new_name_plural = data.get('name_plural', '').strip()
        new_abbreviation = data.get('abbreviation', '').strip()
        new_abbreviation_plural = data.get('abbreviation_plural', '').strip()
        new_unit_type = data.get('unit_type', 'other')
        
        try:
            unit = MeasurementUnit.objects.get(measurement_unit_id=unit_id)
            
            # Validate name is not empty
            if not new_name:
                return JsonResponse({'success': False, 'error': 'Unit name cannot be empty'})
            
            # Check for duplicate names (case-insensitive, excluding current unit)
            duplicate = MeasurementUnit.objects.filter(name__iexact=new_name).exclude(measurement_unit_id=unit_id).first()
            if duplicate:
                return JsonResponse({
                    'success': False,
                    'error': f'A unit named "{new_name}" already exists'
                })
            
            # Check for duplicate abbreviations if provided (case-insensitive, excluding current unit)
            if new_abbreviation:
                duplicate_abbr = MeasurementUnit.objects.filter(abbreviation__iexact=new_abbreviation).exclude(measurement_unit_id=unit_id).first()
                if duplicate_abbr:
                    return JsonResponse({
                        'success': False,
                        'error': f'A unit with abbreviation "{new_abbreviation}" already exists'
                    })
            
            # Update unit with plural fields
            unit.name = new_name
            unit.name_plural = new_name_plural if new_name_plural else None
            unit.abbreviation = new_abbreviation if new_abbreviation else None
            unit.abbreviation_plural = new_abbreviation_plural if new_abbreviation_plural else None
            unit.unit_type = new_unit_type
            unit.save()
            
            return JsonResponse({
                'success': True,
                'message': 'Measurement unit updated successfully',
                'unit': {
                    'id': unit.measurement_unit_id,
                    'name': unit.name,
                    'name_plural': unit.name_plural,
                    'abbreviation': unit.abbreviation,
                    'abbreviation_plural': unit.abbreviation_plural,
                    'unit_type': unit.unit_type,
                    'unit_type_display': unit.get_unit_type_display()
                }
            })
            
        except MeasurementUnit.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Measurement unit not found'})
    
    return JsonResponse({'success': False, 'error': 'Invalid request'})

@login_required
def check_unit_usage(request):
    """Check if measurement unit is used in recipes, ingredients, or conversions"""
    if request.method == 'POST':
        data = json.loads(request.body)
        unit_id = data.get('unit_id')
        
        try:
            unit = MeasurementUnit.objects.get(measurement_unit_id=unit_id)
            
            # Count usage in recipes
            recipe_count = RecipeIngredient.objects.filter(unit=unit).count()
            
            # Count usage in ingredients (as default unit)
            ingredient_count = Ingredient.objects.filter(default_unit=unit).count()
            
            # Count usage in unit conversions
            conversion_count = UnitConversion.objects.filter(
                Q(from_unit=unit) | Q(to_unit=unit)
            ).count()
            
            total_usage = recipe_count + ingredient_count + conversion_count
            
            # Get some example names (limit to 5 total)
            usage_examples = []
            
            if recipe_count > 0:
                recipes = RecipeIngredient.objects.filter(unit=unit).select_related('recipe')[:3]
                for ri in recipes:
                    usage_examples.append(f"Recipe: {ri.recipe.recipe_name}")
            
            if ingredient_count > 0 and len(usage_examples) < 5:
                ingredients = Ingredient.objects.filter(default_unit=unit)[:2]
                for ing in ingredients:
                    usage_examples.append(f"Ingredient: {ing.name}")
            
            return JsonResponse({
                'success': True,
                'total_usage': total_usage,
                'recipe_count': recipe_count,
                'ingredient_count': ingredient_count,
                'conversion_count': conversion_count,
                'usage_examples': usage_examples,
                'can_delete': total_usage == 0
            })
            
        except MeasurementUnit.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Measurement unit not found'})
    
    return JsonResponse({'success': False, 'error': 'Invalid request'})

@login_required
def delete_measurement_unit(request):
    """Delete measurement unit if not used anywhere"""
    if request.method == 'POST':
        data = json.loads(request.body)
        unit_id = data.get('unit_id')
        
        try:
            unit = MeasurementUnit.objects.get(measurement_unit_id=unit_id)
            
            # Check if unit is used anywhere
            recipe_count = RecipeIngredient.objects.filter(unit=unit).count()
            ingredient_count = Ingredient.objects.filter(default_unit=unit).count()
            conversion_count = UnitConversion.objects.filter(
                Q(from_unit=unit) | Q(to_unit=unit)
            ).count()
            
            total_usage = recipe_count + ingredient_count + conversion_count
            
            if total_usage > 0:
                usage_details = []
                if recipe_count > 0:
                    usage_details.append(f'{recipe_count} recipe(s)')
                if ingredient_count > 0:
                    usage_details.append(f'{ingredient_count} ingredient(s)')
                if conversion_count > 0:
                    usage_details.append(f'{conversion_count} conversion(s)')
                
                return JsonResponse({
                    'success': False,
                    'error': f'Cannot delete - unit is used in {", ".join(usage_details)}'
                })
            
            # Safe to delete
            unit_name = unit.name
            unit.delete()
            
            return JsonResponse({
                'success': True,
                'message': f'Successfully deleted measurement unit: {unit_name}'
            })
            
        except MeasurementUnit.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Measurement unit not found'})
    
    return JsonResponse({'success': False, 'error': 'Invalid request'})

# ============================================
# FILE EXTRACTION FUNCTIONS
# ============================================

def extract_text_from_pdf(file):
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        raise Exception(f"Error reading PDF: {str(e)}")


def extract_text_from_docx(file):
    """Extract text from Word document"""
    try:
        doc = Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        raise Exception(f"Error reading Word document: {str(e)}")


def extract_text_from_image(file):
    """For images, we'll pass directly to Claude's vision API"""
    # Convert to base64 for Claude API
    try:
        image = Image.open(file)
        buffered = BytesIO()
        image.save(buffered, format="PNG")
        img_base64 = base64.b64encode(buffered.getvalue()).decode()
        return img_base64
    except Exception as e:
        raise Exception(f"Error processing image: {str(e)}")


# ============================================
# AI EXTRACTION FUNCTION
# ============================================

# Replace the extract_recipe_with_ai function in your views.py

def extract_recipe_with_ai(content, file_type):
    """Use Claude AI to extract recipe data with structured ingredients"""
    
    # Get API key from settings
    api_key = getattr(settings, 'ANTHROPIC_API_KEY', None)
    if not api_key:
        raise Exception("ANTHROPIC_API_KEY not found in settings")
    
    client = anthropic.Anthropic(api_key=api_key)
    
    # Updated prompt for structured ingredient extraction
    system_prompt = """You are a recipe extraction expert. Extract recipe information from the provided content and return it in JSON format.

Extract the following fields:
- recipe_name: The name of the recipe
- description: A brief description (if available)
- prep_time: Preparation time in minutes (number only)
- cook_time: Cooking time in minutes (number only)
- total_time: Total time in minutes (number only)
- servings: Number of servings (number only)
- ingredients: Array of ingredient objects with these fields:
  * quantity: The amount (e.g., "2", "1/4", "1.5") - extract the number only
  * measurement: The unit (e.g., "cups", "tablespoons", "teaspoons", "packets", "cloves") - use singular lowercase
  * ingredient: The ingredient name (e.g., "flour", "olive oil", "frozen artichokes")
  * preparation: Any preparation notes (e.g., "chopped", "diced", "minced", "grated") - empty string if none
- instructions: Array of instruction strings (step by step)

For ingredients, parse each one carefully:
Example: "2 packets Frozen Artichokes" should be:
  {"quantity": "2", "measurement": "packets", "ingredient": "Frozen Artichokes", "preparation": ""}

Example: "1/4 teaspoon salt" should be:
  {"quantity": "1/4", "measurement": "teaspoon", "ingredient": "salt", "preparation": ""}

Example: "2 tablespoons olive oil, extra virgin" should be:
  {"quantity": "2", "measurement": "tablespoons", "ingredient": "olive oil", "preparation": "extra virgin"}

Example: "1 teaspoon minced fresh garlic" should be:
  {"quantity": "1", "measurement": "teaspoon", "ingredient": "fresh garlic", "preparation": "minced"}

Return ONLY valid JSON with these fields. If a field is not found, use null for numbers or empty string/array for text."""
    
    try:
        if file_type in ['jpg', 'jpeg', 'png']:
            # Use vision API for images
            message = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=4096,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": f"image/{file_type}",
                                    "data": content,
                                },
                            },
                            {
                                "type": "text",
                                "text": "Extract the recipe information from this image and return it in the JSON format specified."
                            }
                        ],
                    }
                ],
                system=system_prompt
            )
        else:
            # Use text API for PDF/Word
            message = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=4096,
                messages=[
                    {
                        "role": "user",
                        "content": f"Extract the recipe information from this text and return it in the JSON format specified:\n\n{content}"
                    }
                ],
                system=system_prompt
            )
        
        # Parse the response
        response_text = message.content[0].text
        
        # Extract JSON from response (Claude might wrap it in markdown)
        if "```json" in response_text:
            json_start = response_text.find("```json") + 7
            json_end = response_text.find("```", json_start)
            response_text = response_text[json_start:json_end].strip()
        elif "```" in response_text:
            json_start = response_text.find("```") + 3
            json_end = response_text.find("```", json_start)
            response_text = response_text[json_start:json_end].strip()
        
        recipe_data = json.loads(response_text)
        
        # Validate and set defaults
        recipe_data.setdefault('recipe_name', 'Imported Recipe')
        recipe_data.setdefault('description', '')
        recipe_data.setdefault('prep_time', None)
        recipe_data.setdefault('cook_time', None)
        recipe_data.setdefault('total_time', None)
        recipe_data.setdefault('servings', 4)
        recipe_data.setdefault('ingredients', [])
        recipe_data.setdefault('instructions', [])
        
        # Ensure ingredients have all required fields
        for ing in recipe_data['ingredients']:
            ing.setdefault('quantity', '')
            ing.setdefault('measurement', '')
            ing.setdefault('ingredient', '')
            ing.setdefault('preparation', '')
        
        return recipe_data
        
    except Exception as e:
        print(f"AI Extraction Error: {str(e)}")
        return None

# AJAX endpoint to add new measurement
@login_required
@require_POST
def add_measurement_ajax(request):
    """Add new measurement unit via AJAX"""
    try:
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        name_plural = data.get('name_plural', '').strip()
        abbreviation = data.get('abbreviation', '').strip()
        abbreviation_plural = data.get('abbreviation_plural', '').strip()
        unit_type = data.get('unit_type', 'other')
        
        if not name:
            return JsonResponse({'success': False, 'message': 'Name is required'})
        
        # Check if already exists
        if MeasurementUnit.objects.filter(name__iexact=name).exists():
            return JsonResponse({'success': False, 'message': 'Measurement already exists'})
        
        # Check for duplicate abbreviations if provided
        if abbreviation:
            if MeasurementUnit.objects.filter(abbreviation__iexact=abbreviation).exists():
                return JsonResponse({'success': False, 'message': f'A unit with abbreviation "{abbreviation}" already exists'})
        
        # Create new measurement with all fields
        measurement = MeasurementUnit.objects.create(
            name=name,
            name_plural=name_plural if name_plural else None,
            abbreviation=abbreviation if abbreviation else None,
            abbreviation_plural=abbreviation_plural if abbreviation_plural else None,
            unit_type=unit_type
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Measurement added successfully',
            'measurement_id': measurement.measurement_unit_id,
            'name': measurement.name
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

# AJAX endpoint to add new ingredient
@login_required
@require_POST
def add_ingredient_ajax(request):
    try:
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        category_id = data.get('category_id')
        shopping_unit_id = data.get('shopping_unit_id')  # NEW
        
        if not name:
            return JsonResponse({'success': False, 'message': 'Ingredient name is required'})
        
        if not shopping_unit_id:
            return JsonResponse({'success': False, 'message': 'Shopping unit is required'})
        
        # Check if already exists
        if Ingredient.objects.filter(name__iexact=name).exists():
            return JsonResponse({'success': False, 'message': 'Ingredient already exists'})
        
        # Create ingredient with category AND shopping unit
        ingredient = Ingredient.objects.create(
            name=name,
            category_id=category_id if category_id else None,
            default_unit_id=shopping_unit_id  # NEW: Set shopping unit
        )
        
        return JsonResponse({
            'success': True,
            'ingredient_id': ingredient.ingredient_id,
            'name': ingredient.name
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

@login_required
@require_POST
def add_preparation_ajax(request):
    """Add new preparation method via AJAX"""
    try:
        data = json.loads(request.body)
        name = data.get('name', '').strip()
        
        if not name:
            return JsonResponse({'success': False, 'message': 'Name is required'})
        
        # Lowercase for consistency
        name = name.lower()
        
        # Check if already exists
        if PreparationMethod.objects.filter(name__iexact=name).exists():
            return JsonResponse({'success': False, 'message': 'Preparation method already exists'})
        
        # Create new preparation method
        preparation = PreparationMethod.objects.create(name=name)
        
        return JsonResponse({
            'success': True,
            'message': 'Preparation method added successfully',
            'preparation_id': preparation.preparation_method_id,
            'name': preparation.name
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

@login_required
@require_POST
def add_preparation_ajax(request):
    """Add new preparation method via AJAX"""
    try:
        data = json.loads(request.body)
        name = data.get('name', '').strip().lower()
        
        if not name:
            return JsonResponse({'success': False, 'message': 'Name is required'})
        
        # Check if already exists
        if PreparationMethod.objects.filter(name__iexact=name).exists():
            return JsonResponse({'success': False, 'message': 'Preparation method already exists'})
        
        # Create new preparation method
        preparation = PreparationMethod.objects.create(name=name)
        
        return JsonResponse({
            'success': True,
            'message': 'Preparation method added successfully',
            'preparation_id': preparation.preparation_method_id,
            'name': preparation.name
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

@login_required
@require_POST
def find_matching_recipes(request):
    """
    Find recipes that match selected ingredients with smart weighting.
    
    Logic:
    1. Auto-include herbs/spices, oil, water
    2. Calculate match % for each recipe
    3. Weight protein ingredients heavily
    4. Return top matches sorted by %
    """
    try:
        # Get selected ingredient IDs from frontend
        data = json.loads(request.body)
        selected_ingredient_ids = data.get('ingredient_ids', [])
        
        # Convert to integers
        selected_ingredient_ids = [int(id) for id in selected_ingredient_ids]
        
        # ========== AUTO-INCLUDE COMMON INGREDIENTS ==========
        # Get all ingredients in "Herbs & Spices" category
        herbs_spices_category = IngredientCategory.objects.filter(
            name='Herbs & Spices'
        ).first()
        
        if herbs_spices_category:
            herbs_spices_ids = list(
                Ingredient.objects.filter(
                    category=herbs_spices_category
                ).values_list('ingredient_id', flat=True)
            )
            selected_ingredient_ids.extend(herbs_spices_ids)
        
        # Auto-include Oil and Water (by name search)
        common_items = Ingredient.objects.filter(
            name__in=['Oil', 'Olive Oil', 'Vegetable Oil', 'Water']
        ).values_list('ingredient_id', flat=True)
        selected_ingredient_ids.extend(list(common_items))
        
        # Remove duplicates
        selected_ingredient_ids = list(set(selected_ingredient_ids))
        
        # ========== CHECK IF USER SELECTED ANY PROTEIN ==========
        # Get all proteins from "Meat", "Poultry", "Fish & Seafood" categories
        protein_categories = IngredientCategory.objects.filter(
            name__in=['Meat', 'Poultry', 'Fish & Seafood']
        )
        protein_ingredient_ids = list(
            Ingredient.objects.filter(
                category__in=protein_categories
            ).values_list('ingredient_id', flat=True)
        )
        
        # Check if user selected any protein
        user_selected_protein = any(
            ing_id in selected_ingredient_ids 
            for ing_id in protein_ingredient_ids
        )
        
        # ========== FIND MATCHING RECIPES ==========
        recipes = Recipe.objects.prefetch_related(
            'recipe_ingredients__ingredient',
            'recipe_ingredients__ingredient__category'
        ).all()
        
        results = []
        
        for recipe in recipes:
            required_ingredients = recipe.recipe_ingredients.all()
            total_ingredients = required_ingredients.count()
            
            if total_ingredients == 0:
                continue  # Skip recipes with no ingredients
            
            matched_count = 0
            weighted_matched = 0
            weighted_total = 0
            missing_ingredients = []
            
            for req_ing in required_ingredients:
                ing_id = req_ing.ingredient.ingredient_id
                is_protein = ing_id in protein_ingredient_ids
                
                # Assign weights
                if user_selected_protein and is_protein:
                    weight = 3.0  # Protein is 3x more important
                else:
                    weight = 1.0
                
                weighted_total += weight
                
                # Check if user has this ingredient
                if ing_id in selected_ingredient_ids:
                    matched_count += 1
                    weighted_matched += weight
                else:
                    missing_ingredients.append({
                        'name': req_ing.ingredient.name,
                        'is_protein': is_protein
                    })
            
            # Calculate weighted match percentage
            match_percentage = (weighted_matched / weighted_total * 100) if weighted_total > 0 else 0

            # Store ALL recipes with their match percentages (no filtering yet)
            results.append({
                'recipe_id': recipe.recipe_id,
                'recipe_name': recipe.recipe_name,
                'recipe_image': recipe.recipe_image.url if recipe.recipe_image else None,
                'match_percentage': round(match_percentage, 1),
                'matched_count': matched_count,
                'total_ingredients': total_ingredients,
                'missing_ingredients': missing_ingredients,
                'prep_time': recipe.prep_time,
                'cook_time': recipe.cook_time,
                'difficulty_level': recipe.difficulty_level
            })

            # Sort by match percentage (highest first)
            results.sort(key=lambda x: x['match_percentage'], reverse=True)

            # SMART FILTERING LOGIC
            # If we have 5+ recipes with ≥50% match, use 50% threshold
            # Otherwise, take top 5 recipes regardless of percentage
            recipes_above_50 = [r for r in results if r['match_percentage'] >= 50]

            if len(recipes_above_50) >= 5:
                # Use standard 50% threshold
                results = recipes_above_50
            else:
                # Take top 5 recipes (or all if less than 5 total)
                results = results[:5]
        
        # Return top 10 for initial display
        # Frontend can request more with pagination
        return JsonResponse({
            'success': True,
            'total_results': len(results),
            'results': results[:10],  # Top 10
            'all_results': results  # All results for "Show More"
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=400)

### CELEBRATION MANAGEMENT VIEWS ###
@login_required
def celebration_management(request):
    """Main celebration management page"""
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to access this page.')
        return redirect('home')
    
    # Handle Contact CRUD
    if request.method == 'POST':
        action = request.POST.get('action')
        
        if action == 'add_contact':
            name = request.POST.get('name')
            relationship = request.POST.get('relationship')
            email = request.POST.get('email')
            phone = request.POST.get('phone')
            notes = request.POST.get('notes')
            
            if name:
                Contact.objects.create(
                    name=name,
                    relationship=relationship,
                    email=email,
                    phone=phone,
                    notes=notes,
                    created_by=request.user
                )
                messages.success(request, f'Contact "{name}" added successfully!')
            else:
                messages.error(request, 'Name is required.')
        
        elif action == 'edit_contact':
            contact_id = request.POST.get('contact_id')
            try:
                contact = Contact.objects.get(id=contact_id, created_by=request.user)
                contact.name = request.POST.get('name')
                contact.relationship = request.POST.get('relationship')
                contact.email = request.POST.get('email')
                contact.phone = request.POST.get('phone')
                contact.notes = request.POST.get('notes')
                contact.save()
                messages.success(request, f'Contact "{contact.name}" updated successfully!')
            except Contact.DoesNotExist:
                messages.error(request, 'Contact not found.')
        
        elif action == 'delete_contact':
            contact_id = request.POST.get('contact_id')
            try:
                contact = Contact.objects.get(id=contact_id, created_by=request.user)
                name = contact.name
                contact.delete()
                messages.success(request, f'Contact "{name}" deleted successfully!')
            except Contact.DoesNotExist:
                messages.error(request, 'Contact not found.')
        
        elif action == 'add_event':
            contact_id = request.POST.get('contact_id')
            try:
                contact = Contact.objects.get(id=contact_id, created_by=request.user)
                event_type = request.POST.get('event_type')
                event_date_str = request.POST.get('event_date')
                priority = request.POST.get('priority', 'normal')
                notes = request.POST.get('event_notes')
                
                # Notification settings
                notify_one_week = request.POST.get('notify_one_week') == 'on'
                notify_one_day = request.POST.get('notify_one_day') == 'on'
                notify_same_day = request.POST.get('notify_same_day') == 'on'
                
                if event_date_str:
                    # Parse the date
                    from datetime import datetime
                    event_date = datetime.strptime(event_date_str, '%Y-%m-%d').date()
                    
                    # For namedays, set year to 1900 (placeholder year)
                    if event_type == 'nameday':
                        event_date = event_date.replace(year=1900)
                    
                    CelebrationEvent.objects.create(
                        contact=contact,
                        event_type=event_type,
                        event_date=event_date,
                        is_recurring=True,  # Always recurring for now
                        priority=priority,
                        notes=notes,
                        notify_one_week=notify_one_week,
                        notify_one_day=notify_one_day,
                        notify_same_day=notify_same_day,
                        notify_demetri=request.POST.get('notify_demetri', 'on') == 'on',
                        notify_angy=request.POST.get('notify_angy', 'on') == 'on',
                        notify_erene=request.POST.get('notify_erene', 'on') == 'on',
                        notify_alexandra=request.POST.get('notify_alexandra', 'on') == 'on',
                        created_by=request.user
                    )
                    messages.success(request, f'{event_type.title()} event added for {contact.name}!')
                else:
                    messages.error(request, 'Event date is required.')
            except Contact.DoesNotExist:
                messages.error(request, 'Contact not found.')
            except ValueError:
                messages.error(request, 'Invalid date format.')
        
        elif action == 'edit_event':
            event_id = request.POST.get('event_id')
            try:
                event = CelebrationEvent.objects.get(id=event_id, created_by=request.user)
                event.event_type = request.POST.get('event_type')
                event_date_str = request.POST.get('event_date')
                event.priority = request.POST.get('priority', 'normal')
                event.notes = request.POST.get('event_notes')
                event.notify_one_week = request.POST.get('notify_one_week') == 'on'
                event.notify_one_day = request.POST.get('notify_one_day') == 'on'
                event.notify_same_day = request.POST.get('notify_same_day') == 'on'
                event.notify_demetri = request.POST.get('notify_demetri', 'on') == 'on'
                event.notify_angy = request.POST.get('notify_angy', 'on') == 'on'
                event.notify_erene = request.POST.get('notify_erene', 'on') == 'on'
                event.notify_alexandra = request.POST.get('notify_alexandra', 'on') == 'on'
                
                # Handle date update
                if event_date_str:
                    from datetime import datetime
                    event_date = datetime.strptime(event_date_str, '%Y-%m-%d').date()
                    
                    # For namedays, set year to 1900 (placeholder year)
                    if event.event_type == 'nameday':
                        event_date = event_date.replace(year=1900)
                    
                    event.event_date = event_date
                
                event.save()
                messages.success(request, 'Event updated successfully!')
            except CelebrationEvent.DoesNotExist:
                messages.error(request, 'Event not found.')
            except ValueError:
                messages.error(request, 'Invalid date format.')
        
        elif action == 'delete_event':
            event_id = request.POST.get('event_id')
            try:
                event = CelebrationEvent.objects.get(id=event_id, created_by=request.user)
                event.delete()
                messages.success(request, 'Event deleted successfully!')
            except CelebrationEvent.DoesNotExist:
                messages.error(request, 'Event not found.')
        
        return redirect('celebration_management')
    
    # Get all contacts with their events
    contacts = Contact.objects.filter(created_by=request.user).prefetch_related('celebration_events')
    
    # Get upcoming events for dashboard
    today = timezone.now().date()
    all_events = []
    
    for contact in contacts:
        for event in contact.celebration_events.all():
            next_date = event.get_next_occurrence()
            if next_date:
                days_until = (next_date - today).days
                if days_until <= 90:  # Show events in next 90 days
                    all_events.append({
                        'contact': contact,
                        'event': event,
                        'next_date': next_date,
                        'days_until': days_until
                    })
    
    # Sort by days until
    all_events.sort(key=lambda x: x['days_until'])
    
    return render(request, 'celebration_management.html', {
        'contacts': contacts,
        'upcoming_events': all_events[:10],  # Top 10 upcoming
        'relationship_choices': Contact.RELATIONSHIP_CHOICES,
        'event_type_choices': CelebrationEvent.EVENT_TYPE_CHOICES,
        'priority_choices': CelebrationEvent.PRIORITY_CHOICES,
    })

@login_required
def celebration_calendar(request):
    """Calendar view of all celebrations"""
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to access this page.')
        return redirect('home')
    
    # Get today's date
    today = timezone.now().date()
    
    # Get all events for this user
    events = CelebrationEvent.objects.filter(created_by=request.user).select_related('contact')
    
    # Find first month with events (from today forward)
    first_event_date = None
    
    for event in events:
        next_occurrence = event.get_next_occurrence()
        if next_occurrence and next_occurrence >= today:
            if first_event_date is None or next_occurrence < first_event_date:
                first_event_date = next_occurrence
    
    # Get month and year from request, or use first event month, or default to current
    if 'month' in request.GET and 'year' in request.GET:
        month = int(request.GET.get('month'))
        year = int(request.GET.get('year'))
    elif first_event_date:
        month = first_event_date.month
        year = first_event_date.year
    else:
        month = today.month
        year = today.year
    
    # Build calendar
    cal = monthcalendar(year, month)
    
    # Map events to calendar days
    events_by_day = {}
    for event in events:
        next_occurrence = event.get_next_occurrence()
        if next_occurrence and next_occurrence.month == month and next_occurrence.year == year:
            day = next_occurrence.day
            if day not in events_by_day:
                events_by_day[day] = []
            events_by_day[day].append(event)
    
    # Get all upcoming events for timeline view (next 365 days)
    all_events = []
    contacts = Contact.objects.filter(created_by=request.user).prefetch_related('celebration_events')
    
    for contact in contacts:
        for event in contact.celebration_events.all():
            next_date = event.get_next_occurrence()
            if next_date:
                days_until = (next_date - today).days
                if days_until <= 365:  # Show events in next year
                    all_events.append({
                        'contact': contact,
                        'event': event,
                        'next_date': next_date,
                        'days_until': days_until
                    })
    
    # Sort by next occurrence date
    all_events.sort(key=lambda x: x['next_date'])
    
    # Previous and next month/year
    if month == 1:
        prev_month, prev_year = 12, year - 1
    else:
        prev_month, prev_year = month - 1, year
    
    if month == 12:
        next_month, next_year = 1, year + 1
    else:
        next_month, next_year = month + 1, year
    
    return render(request, 'celebration_calendar.html', {
        'calendar': cal,
        'month': month,
        'year': year,
        'month_name': month_name[month],
        'events_by_day': events_by_day,
        'all_events': all_events,
        'prev_month': prev_month,
        'prev_year': prev_year,
        'next_month': next_month,
        'next_year': next_year,
        'today': today,
        'default_view': request.GET.get('view', 'calendar'),
    })

@login_required
@require_POST
def import_celebrations(request):
    """Import contacts and events from Excel file"""
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to access this page.')
        return redirect('home')
    
    if 'excel_file' not in request.FILES:
        messages.error(request, 'No file uploaded.')
        return redirect('celebration_management')
    
    excel_file = request.FILES['excel_file']
    skip_duplicates = request.POST.get('skip_duplicates') == 'on'
    
    try:
        import pandas as pd
        
        # Read Excel file
        df = pd.read_excel(excel_file)
        
        # Validate columns (flexible - support both "EVENT" and "EVENT TYPE")
        event_column = 'EVENT TYPE' if 'EVENT TYPE' in df.columns else 'EVENT'
        required_columns = ['NAME', 'RELATIONSHIP', event_column, 'DATE']
        if not all(col in df.columns for col in required_columns):
            messages.error(request, f'Excel file must have columns: NAME, RELATIONSHIP, EVENT/EVENT TYPE, DATE')
            return redirect('celebration_management')
        
        # Map relationship values to model choices
        relationship_map = {
            'family': 'family',
            'friend': 'friend',
            'colleague': 'colleague',
            'other': 'other',
        }
        
        # Map event types
        event_type_map = {
            'birthday': 'birthday',
            'nameday': 'nameday',
            'anniversary': 'anniversary',
            'custom': 'custom',
        }
        
        # Map priority values
        priority_map = {
            'high': 'high',
            'normal': 'normal',
            'low': 'low',
        }
        
        contacts_created = 0
        contacts_skipped = 0
        events_created = 0
        
        # Group by NAME to create contacts
        for name, group in df.groupby('NAME'):
            name = str(name).strip()
            
            if not name:
                continue
            
            # Check if contact exists
            if skip_duplicates and Contact.objects.filter(created_by=request.user, name__iexact=name).exists():
                contacts_skipped += 1
                continue
            
            # Get relationship from first row
            relationship_value = str(group.iloc[0]['RELATIONSHIP']).strip().lower()
            relationship = relationship_map.get(relationship_value, 'other')
            
            # Create contact
            contact = Contact.objects.create(
                name=name,
                relationship=relationship,
                created_by=request.user
            )
            contacts_created += 1
            
            # Create events for this contact
            for _, row in group.iterrows():
                event_type_value = str(row[event_column]).strip().lower()
                event_type = event_type_map.get(event_type_value, 'custom')
                
                # Parse date
                event_date = pd.to_datetime(row['DATE']).date()
                
                # For birthdays and namedays without birth year, use placeholder year 1900
                event_date = event_date.replace(year=1900)
                
                # Get priority (default to 'high' if column doesn't exist)
                if 'PRIORITY' in df.columns:
                    priority_value = str(row['PRIORITY']).strip().lower()
                    priority = priority_map.get(priority_value, 'high')
                else:
                    priority = 'normal'
                
                # Parse notification settings
                notify_one_week = False
                notify_one_day = False
                notify_same_day = False
                
                if 'Notification Settings' in df.columns:
                    notification_settings = str(row['Notification Settings']).lower()
                    
                    if 'one week' in notification_settings or '1 week' in notification_settings:
                        notify_one_week = True
                    if 'one day' in notification_settings or '1 day' in notification_settings:
                        notify_one_day = True
                    if 'same day' in notification_settings:
                        notify_same_day = True
                    
                    # If "all" is mentioned, enable all notifications
                    if 'all' in notification_settings:
                        notify_one_week = True
                        notify_one_day = True
                        notify_same_day = True
                else:
                    # Default: all notifications enabled
                    notify_one_week = True
                    notify_one_day = True
                    notify_same_day = True
                
                # Create event
                CelebrationEvent.objects.create(
                    contact=contact,
                    event_type=event_type,
                    event_date=event_date,
                    is_recurring=True,
                    priority=priority,
                    notify_one_week=notify_one_week,
                    notify_one_day=notify_one_day,
                    notify_same_day=notify_same_day,
                    created_by=request.user
                )
                events_created += 1
        
        # Success message
        msg = f'Successfully imported {contacts_created} contacts and {events_created} events.'
        if contacts_skipped > 0:
            msg += f' Skipped {contacts_skipped} duplicate contacts.'
        messages.success(request, msg)
        
    except Exception as e:
        messages.error(request, f'Error importing file: {str(e)}')
    
    return redirect('celebration_management')

@login_required
def celebration_dashboard(request):
    """Dashboard showing upcoming celebrations"""
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to access this page.')
        return redirect('home')
    
    # Get all contacts with their events
    today = timezone.now().date()
    contacts = Contact.objects.filter(created_by=request.user).prefetch_related('celebration_events')
    
    # Get upcoming events for dashboard (next 30 days)
    all_events = []
    
    for contact in contacts:
        for event in contact.celebration_events.all():
            next_date = event.get_next_occurrence()
            if next_date:
                days_until = (next_date - today).days
                if days_until <= 30:  # Show events in next 30 days only
                    all_events.append({
                        'contact': contact,
                        'event': event,
                        'next_date': next_date,
                        'days_until': days_until
                    })
    
    # Sort by days until
    all_events.sort(key=lambda x: x['days_until'])
    
    return render(request, 'celebration_dashboard.html', {
        'upcoming_events': all_events,
    })

@login_required
def notification_settings(request):
    """Manage email notification recipients for administration items"""
    if not request.user.is_superuser:
        messages.error(request, 'You do not have permission to access this page.')
        return redirect('home')
    
    # Handle form submission
    if request.method == 'POST':
        notification_type = request.POST.get('notification_type')
        to_addresses = request.POST.get('to_addresses', '')
        cc_addresses = request.POST.get('cc_addresses', '')
        
        recipient, created = NotificationRecipient.objects.get_or_create(
            notification_type=notification_type,
            defaults={'created_by': request.user}
        )
        recipient.to_addresses = to_addresses
        recipient.cc_addresses = cc_addresses
        recipient.save()
        
        messages.success(request, f'{recipient.get_notification_type_display()} email addresses updated successfully!')
        return redirect('notification_settings')
    
    # Get only administration notification types
    admin_types = [
        'daily_report', 
        'new_lease_upload',
        'expense_needs_approval',
        'expense_approved',
        'expense_paid',
        'friday_status_report_supervisor', 
        'friday_status_report_staff'
    ]
    
    notification_settings = {}
    
    for type_code, type_name in NotificationRecipient.NOTIFICATION_TYPES:
        if type_code in admin_types:
            try:
                recipient = NotificationRecipient.objects.get(notification_type=type_code)
                notification_settings[type_code] = {
                    'name': type_name,
                    'to_emails': recipient.to_addresses,
                    'cc_emails': recipient.cc_addresses
                }
            except NotificationRecipient.DoesNotExist:
                notification_settings[type_code] = {
                    'name': type_name,
                    'to_emails': '',
                    'cc_emails': ''
                }
    
    return render(request, 'notification_settings.html', {
        'notification_settings': notification_settings,
    })

@login_required
def personal_notification_settings(request):
    """Manage email notification recipients for personal items"""
    
    # Handle form submission
    if request.method == 'POST':
        notification_type = request.POST.get('notification_type')
        to_addresses = request.POST.get('to_addresses', '')
        cc_addresses = request.POST.get('cc_addresses', '')
        
        recipient, created = NotificationRecipient.objects.get_or_create(
            notification_type=notification_type,
            defaults={'created_by': request.user}
        )
        recipient.to_addresses = to_addresses
        recipient.cc_addresses = cc_addresses
        recipient.save()
        
        messages.success(request, f'{recipient.get_notification_type_display()} email addresses updated successfully!')
        return redirect('personal_notification_settings')
    
    # Get only personal notification types
    personal_types = ['celebration_reminder', 'document_expiry']
    
    notification_settings = {}
    
    for type_code, type_name in NotificationRecipient.NOTIFICATION_TYPES:
        if type_code in personal_types:
            try:
                recipient = NotificationRecipient.objects.get(notification_type=type_code)
                notification_settings[type_code] = {
                    'name': type_name,
                    'to_emails': recipient.to_addresses,
                    'cc_emails': recipient.cc_addresses
                }
            except NotificationRecipient.DoesNotExist:
                notification_settings[type_code] = {
                    'name': type_name,
                    'to_emails': '',
                    'cc_emails': ''
                }
    
    return render(request, 'personal_notification_settings.html', {
        'notification_settings': notification_settings,
    })

@login_required
@require_POST
def update_event_notifications(request, event_id):
    """Update notification preferences for a specific event via AJAX"""
    import json
    
    try:
        event = CelebrationEvent.objects.get(id=event_id)
        
        # Get the JSON data from request
        data = json.loads(request.body)
        
        # Update the notification preferences
        event.notify_demetri = data.get('notify_demetri', False)
        event.notify_angy = data.get('notify_angy', False)
        event.notify_erene = data.get('notify_erene', False)
        event.notify_alexandra = data.get('notify_alexandra', False)
        
        event.save()
        
        return JsonResponse({
            'success': True,
            'message': 'Notification preferences updated'
        })
    
    except CelebrationEvent.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'Event not found'
        }, status=404)
    
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

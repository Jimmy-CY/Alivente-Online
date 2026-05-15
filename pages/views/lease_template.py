"""
Lease template generator views and helpers.

Extracted from pages/views/main.py as part of the modular views migration
(section ### LEASE TEMPLATE GENERATOR ###).

Contains 9 functions, all lease-related:
    is_superuser(user)              - tiny User predicate helper.
                                      NOTE: appears unused (no callers
                                      found in views modules), kept
                                      verbatim in case templates or
                                      URL patterns reference it.
    generate_lease_agreement_view   - main lease generation form view
                                      and POST handler.
    get_ordinal_day                 - "1st", "2nd", "3rd" etc.
    number_to_words_greek           - integer -> Greek words.
    format_date_greek               - date -> "1η Ιανουαρίου 2026".
    translate_to_greek              - large English->Greek dict for
                                      property/furniture/keys.
    generate_lease_document         - DocxTemplate-based generator.
    prepare_lease_template_data     - large field-mapping helper.
    create_basic_lease_document     - python-docx fallback when no
                                      template file is present.

NOTE: a 10th function (spell_check_instructions) was originally extracted
here but was misplaced - it spell-checks recipe instructions, not lease
content. It was relocated to the ### RECIPE MANAGEMENT ### section in
main.py in a follow-up commit.

URL patterns remain registered in pages/urls.py.
The `logger` instance is set up at module level (was inherited from main.py).
"""

import json
import logging
import os
import tempfile
from datetime import datetime

from django.conf import settings
from django.contrib import messages
from django.contrib.auth.decorators import login_required, permission_required
from django.http import FileResponse, JsonResponse
from django.shortcuts import get_object_or_404, render
from django.views.decorators.http import require_POST

from docx import Document
from docxtpl import DocxTemplate

from ..models import props, tenant

logger = logging.getLogger(__name__)


def is_superuser(user):
    """Check if user is superuser"""
    return user.is_superuser

@login_required
@permission_required('auth.can_edit_tenants', raise_exception=True)
def generate_lease_agreement_view(request):
    """
    View to display the lease agreement generation form and handle document generation
    """
    # Get properties and tenants from your actual models
    properties = props.objects.filter(prop_available_for_rent='Yes').order_by('prop_name')
    tenants = tenant.objects.select_related('prop').all().order_by('tenant_name')
    
    context = {
        'properties': properties,
        'tenants': tenants,
    }
    
    if request.method == 'POST':
        try:
            # Extract form data
            country = request.POST.get('country')
            language = request.POST.get('language')
            property_id = request.POST.get('property')
            tenant_id = request.POST.get('tenant')
            
            # Get additional data from modal (JSON format)
            additional_data_json = request.POST.get('additional_data', '{}')
            
            # Validate additional_data is valid JSON
            try:
                additional_data = json.loads(additional_data_json)
            except json.JSONDecodeError:
                additional_data = {}
                logger.warning(f"Invalid JSON in additional_data: {additional_data_json}")
            
            # Validate required fields
            if not all([country, language, property_id, tenant_id]):
                messages.error(request, 'Please fill in all required fields.')
                return render(request, 'generate_lease_agreement.html', context)
            
            # Validate additional data was provided
            if not additional_data:
                messages.error(request, 'Please complete the additional data by clicking "Review & Complete Data".')
                return render(request, 'generate_lease_agreement.html', context)
            
            # Get property object from database
            try:
                property_obj = get_object_or_404(props, prop_id=property_id)
            except:
                messages.error(request, 'Selected property not found.')
                return render(request, 'generate_lease_agreement.html', context)
            
            # Handle tenant (existing or new)
            if tenant_id == 'new_tenant':
                # Create a mock tenant object with new tenant data
                new_tenant_data = additional_data.get('new_tenant_data', {})
                if not new_tenant_data.get('tenant_name'):
                    messages.error(request, 'Tenant name is required for new tenant.')
                    return render(request, 'generate_lease_agreement.html', context)
                
                tenant_obj = type('MockTenant', (), {
                    'tenant_name': new_tenant_data.get('tenant_name', ''),
                    'tenant_type': new_tenant_data.get('tenant_type', 'Individual'),
                    'tenant_passport_id': new_tenant_data.get('tenant_passport_id', ''),
                    'tenant_passport_country': new_tenant_data.get('tenant_passport_country', ''),
                    'tenant_email': new_tenant_data.get('tenant_email', ''),
                    'tenant_contact_number': new_tenant_data.get('tenant_contact_number', ''),
                    'tenant_address': new_tenant_data.get('tenant_address', ''),
                    'tenant_rent': None,
                    'tenant_deposit': None,
                    'tenant_lease_start_date': None,
                    'tenant_lease_end_date': None,
                    'tenant_levies': None,
                    'tenant_payment_terms': None,
                    'tenant_rental_type': None,
                    'tenant_renewal': None,
                    'tenant_renewal_period': None,
                })()
            else:
                # Get existing tenant
                try:
                    tenant_obj = get_object_or_404(tenant, tenant_id=tenant_id)
                except:
                    messages.error(request, 'Selected tenant not found.')
                    return render(request, 'generate_lease_agreement.html', context)
            
            # Generate the lease agreement
            document_path = generate_lease_document(
                country=country,
                language=language,
                property_obj=property_obj,
                tenant_obj=tenant_obj,
                additional_data=additional_data
            )
            
            if document_path:
                # Create filename
                filename = f'lease_agreement_{property_obj.prop_name}_{tenant_obj.tenant_name}_{country}_{language}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
                filename = filename.replace(' ', '_').replace('/', '_')  # Clean filename
                
                # Return the generated document
                try:
                    response = FileResponse(
                        open(document_path, 'rb'),
                        as_attachment=True,
                        filename=filename
                    )
                    # messages.success(request, 'Lease agreement generated successfully!')
                    return response
                except Exception as e:
                    logger.error(f"Error returning file: {str(e)}")
                    messages.error(request, 'Error downloading the generated document.')
            else:
                messages.error(request, 'Error generating lease agreement. Please check the template file exists.')
                
        except Exception as e:
            logger.error(f"Error generating lease agreement: {str(e)}")
            messages.error(request, f'An error occurred while generating the lease agreement: {str(e)}')
    
    return render(request, 'generate_lease_agreement.html', context)

def get_ordinal_day(day):
    """Convert day number to ordinal (1st, 2nd, 3rd, etc.)"""
    day = int(day)
    if 10 <= day % 100 <= 20:
        suffix = 'th'
    else:
        suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    return f"{day}{suffix}"

def number_to_words_greek(num):
    """Convert number to Greek words with proper capitalization"""
    if num == 0:
        return "Μηδέν"
    
    ones = ["", "Ένα", "Δύο", "Τρία", "Τέσσερα", "Πέντε", "Έξι", "Επτά", "Οκτώ", "Εννέα",
            "Δέκα", "Έντεκα", "Δώδεκα", "Δεκατρία", "Δεκατέσσερα", "Δεκαπέντε", 
            "Δεκαέξι", "Δεκαεπτά", "Δεκαοκτώ", "Δεκαεννέα"]
    
    tens = ["", "", "Είκοσι", "Τριάντα", "Σαράντα", "Πενήντα", "Εξήντα", "Εβδομήντα", "Ογδόντα", "Ενενήντα"]
    
    hundreds = ["", "Εκατό", "Διακόσια", "Τριακόσια", "Τετρακόσια", "Πεντακόσια", 
                "Εξακόσια", "Επτακόσια", "Οκτακόσια", "Εννιακόσια"]
    
    num = int(num)  # Remove decimals
    
    if num < 20:
        return ones[num]
    elif num < 100:
        return tens[num // 10] + ("" if num % 10 == 0 else " " + ones[num % 10])
    elif num < 1000:
        return hundreds[num // 100] + ("" if num % 100 == 0 else " " + number_to_words_greek(num % 100))
    elif num < 1000000:
        thousands = num // 1000
        remainder = num % 1000
        
        if thousands == 1:
            thousands_text = "Χίλια"
        elif thousands < 20:
            thousands_text = ones[thousands] + " Χιλιάδες"
        else:
            thousands_text = number_to_words_greek(thousands) + " Χιλιάδες"
            
        return thousands_text + ("" if remainder == 0 else " " + number_to_words_greek(remainder))
    else:
        return str(int(num))  # Fallback for very large numbers

def format_date_greek(date_obj):
    """Convert date object to Greek format with ordinal day and Greek month name"""
    if not date_obj:
        return 'N/A'
    
    # Greek month names (genitive case for dates)
    greek_months = {
        1: "Ιανουαρίου", 2: "Φεβρουαρίου", 3: "Μαρτίου", 4: "Απριλίου",
        5: "Μαΐου", 6: "Ιουνίου", 7: "Ιουλίου", 8: "Αυγούστου",
        9: "Σεπτεμβρίου", 10: "Οκτωβρίου", 11: "Νοεμβρίου", 12: "Δεκεμβρίου"
    }
    
    day = date_obj.day
    month = greek_months[date_obj.month]
    year = date_obj.year
    
    # Add Greek ordinal suffix to day
    ordinal_day = f"{day}η"
    
    return f"{ordinal_day} {month} {year}"

def translate_to_greek(text):
    """Translate common property and name terms to Greek"""
    if not text or text.strip() == '':
        return text
    
    # Complete translations dictionary
    translations = {
        # Common street types
        'Street': 'Οδός',
        'St': 'Οδός',
        'Avenue': 'Λεωφόρος',
        'Ave': 'Λεωφόρος',
        'Road': 'Δρόμος',
        'Rd': 'Δρόμος',
        'Lane': 'Δρομάκι',
        'Square': 'Πλατεία',
        'Plaza': 'Πλατεία',
        
        # Specific street names and areas
        'Eleftheroupoleos': 'Ελευθερουπόλεως',
        'Agias': 'Αγίας',
        'Annas': 'Άννας',
        'Dikaiosynis': 'Δικαιοσύνης',
        'Ionion': 'Ιόνιον',
        'Aristoteli': 'Αριστοτέλη',
        'Valaoriti': 'Βαλαωρίτη',
        'Pindarou': 'Πινδάρου',
        'Evagora': 'Ευαγόρα',
        'Palikaridi': 'Παλικαρίδη',
        'Agios': 'Άγιος',
        'Dometios': 'Δομέτιος',
        'Aristoteli Valaoriti': 'Αριστοτέλη Βαλαωρίτη',
        
        # Common area names in Cyprus
        'Nicosia': 'Λευκωσία',
        'Limassol': 'Λεμεσός',
        'Larnaca': 'Λάρνακα',
        'Paphos': 'Πάφος',
        'Famagusta': 'Αμμόχωστος',
        'Kyrenia': 'Κερύνεια',
        'Engomi': 'Έγκωμη',
        'Strovolos': 'Στρόβολος',
        'Lakatamia': 'Λακαταμια',
        'Latsia': 'Λατσιά',
        'Aglandjia': 'Αγλαντζιά',
        'Agia Thekla': 'Αγία Θέκλα',
        'Sotira': 'Σωτήρα',
        'Cyprus': 'Κύπρος',
        'Agios Dometios': 'Άγιος Δομέτιος',
        
        # Common building types
        'Flat': 'Διαμέρισμα',
        'Apartment': 'Διαμέρισμα',
        'House': 'Σπίτι',
        'Villa': 'Βίλα',
        'Villas': 'Βίλες',
        'Office': 'Γραφείο',
        'Building': 'Κτίριο',
        'Complex': 'Συγκρότημα',
        'Tower': 'Πύργος',
        'Center': 'Κέντρο',
        'Centre': 'Κέντρο',
        
        # FURNITURE AND APPLIANCES
        'Television': 'Τηλεόραση',
        'Televisions': 'Τηλεοράσεις',
        'TV': 'Τηλεόραση',
        'TVs': 'Τηλεοράσεις',
        'Washing Machine': 'Πλυντήριο',
        'Washing Machines': 'Πλυντήρια',
        'Dishwasher': 'Πλυντήριο Πιάτων',
        'Dishwashers': 'Πλυντήρια Πιάτων',
        'Oven': 'Φούρνος',
        'Ovens': 'Φούρνοι',
        'Stove': 'Ηλεκτρικές εστίες',
        'Stoves': 'Ηλεκτρικές εστίες',
        'Extractor Fan': 'Απορροφητήρας',
        'Extractor Fans': 'Απορροφητήρες',
        'Fridge Freezer': 'Ψυγειοκαταψύκτης',
        'Fridge Freezers': 'Ψυγειοκαταψύκτες',
        'Airconditioner': 'Κλιματιστικό',
        'Airconditioners': 'Κλιματιστικά',
        'Air Conditioner': 'Κλιματιστικό',
        'Air Conditioners': 'Κλιματιστικά',
        'Microwave': 'Φούρνος Μικροκυμάτων',
        'Microwaves': 'Φούρνοι Μικροκυμάτων',
        'Three Seater Couch': 'Τριθέσιος Καναπές',
        'Three Seater Couches': 'Τριθέσιοι Καναπέδες',
        'Two Seater Couch': 'Διθέσιος Καναπές',
        'Two Seater Couches': 'Διθέσιοι Καναπέδες',
        'Lounge Chair': 'Πολυθρόνα',
        'Lounge Chairs': 'Πολυθρόνες',
        'Coffee Tables': 'Τραπεζάκια Σαλονιού',
        'Coffee Table': 'Τραπεζάκι Σαλονιού',
        'Corner Tables': 'Γωνιακά Τραπέζια',
        'Corner Table': 'Γωνιακό Τραπέζι',
        'Standing Lamps': 'Φωτιστικά Δαπέδου',
        'Standing Lamp': 'Φωτιστικό Δαπέδου',
        'Table Lamps': 'Φωτιστικά Τραπεζιού',
        'Table Lamp': 'Φωτιστικό Τραπεζιού',
        'Dining Tables': 'Τραπεζαρίες',
        'Dining Table': 'Τραπεζαρία',
        'Dining Chairs': 'Καρέκλες Τραπεζαρίας',
        'Dining Chair': 'Καρέκλα Τραπεζαρίας',
        'Single Beds': 'Μονά Κρεβάτια',
        'Single Bed': 'Μονό Κρεβάτι',
        'Queen Beds': 'Διπλά Κρεβάτια',
        'Queen Bed': 'Διπλό Κρεβάτι',
        'Double Bed': 'Διπλό Κρεβάτι',
        'Double Beds': 'Διπλά Κρεβάτια',
        'Outside Tables': 'Εξωτερικά Τραπέζια',
        'Outside Table': 'Εξωτερικό Τραπέζι',
        'Outside Chairs': 'Εξωτερικές Καρέκλες',
        'Outside Chair': 'Εξωτερική Καρέκλα',
        'TV Table': 'Τραπέζι Tηλεόρασης',
        'TV Tables': 'Τραπέζια Tηλεόρασης',
        'Bed': 'Κρεβάτι',
        'Beds': 'Κρεβάτια',
        'Sofa': 'Καναπές',
        'Sofas': 'Καναπέδες',
        'Table': 'Τραπέζι',
        'Tables': 'Τραπέζια',
        'Chair': 'Καρέκλα',
        'Chairs': 'Καρέκλες',
        'Lamp': 'Λάμπα',
        'Lamps': 'Λάμπες',
        'Desk': 'Γραφείο',
        'Desks': 'Γραφεία',
        'Wardrobe': 'Ντουλάπα',
        'Wardrobes': 'Ντουλάπες',
        'Refrigerator': 'Ψυγείο',
        'Refrigerators': 'Ψυγεία',
        'Fridge': 'Ψυγείο',
        'Fridges': 'Ψυγεία',
        'Heater': 'Θερμάστρα',
        'Heaters': 'Θερμάστρες',
        'Curtain': 'Κουρτίνα',
        'Curtains': 'Κουρτίνες',
        'Mirror': 'Καθρέφτης',
        'Mirrors': 'Καθρέφτες',
        'Bookshelf': 'Βιβλιοθήκη',
        'Bookshelves': 'Βιβλιοθήκες',
        'AC': 'Κλιματιστικό',
        'ACs': 'Κλιματιστικά',
        'Bedside Table': 'Κομοδίνο',
        'Bedside Tables': 'Κομοδίνα',
        'Roller Blind': 'Ρολό',
        'Roller Blinds': 'Ρολά',
        'Stool': 'Σκαμπό',
        'Stools': 'Σκαμπό',
        
        # Keys - Complete singular and plural forms
        'Key': 'Κλειδί',
        'Keys': 'Κλειδιά',
        'Front Door': 'Κεντρική Πόρτα',
        'Front Door Key': 'Κλειδί Κεντρικής Πόρτας',
        'Front Door Keys': 'Κλειδιά Κεντρικής Πόρτας',
        'Main Door': 'Κεντρική Πόρτα',
        'Main Door Key': 'Κλειδί Κεντρικής Πόρτας',
        'Main Door Keys': 'Κλειδιά Κεντρικής Πόρτας',
        'Building Key': 'Κλειδί Κτιρίου',
        'Building Keys': 'Κλειδιά Κτιρίου',
        'Entrance': 'Είσοδος',
        'Entrance Key': 'Κλειδί Εισόδου',
        'Entrance Keys': 'Κλειδιά Εισόδου',
        'Mailbox': 'Γραμματοκιβώτιο',
        'Mailbox Key': 'Κλειδί Γραμματοκιβωτίου',
        'Mailbox Keys': 'Κλειδιά Γραμματοκιβωτίου',
        'Post Box Key': 'Κλειδί Γραμματοκιβωτίου',
        'Post Box Keys': 'Κλειδιά Γραμματοκιβωτίου',
        'Storeroom Key': 'Κλειδί Αποθήκης',
        'Storeroom Keys': 'Κλειδιά Αποθήκης',
        'Storage Key': 'Κλειδί Αποθήκης',
        'Storage Keys': 'Κλειδιά Αποθήκης',
        'Garage Key': 'Κλειδί Γκαράζ',
        'Garage Keys': 'Κλειδιά Γκαράζ',
        'Garage': 'Γκαράζ',
        'Storage': 'Αποθήκη',
        'Balcony Key': 'Κλειδί Μπαλκονιού',
        'Balcony Keys': 'Κλειδιά Μπαλκονιού',
        'Balcony': 'Μπαλκόνι',
        'Terrace Key': 'Κλειδί Βεράντας',
        'Terrace Keys': 'Κλειδιά Βεράντας',
        'Terrace': 'Βεράντα',
        
        # Common names - CORRECTED
        'Demetri': 'Δημήτρης',
        'Manias': 'Μανιάς',
        'Foti': 'Φώτι',  # CORRECTED - was Φώτης
        'Pitta': 'Πίττα',
        'George': 'Γεώργιος',
        'Maria': 'Μαρία',
        'Andreas': 'Ανδρέας',
        'Christina': 'Χριστίνα',
        'Kostas': 'Κώστας',
        'Nikos': 'Νίκος',
        
        # Company terms
        'Limited': 'Περιορισμένη',
        'Ltd': 'Περιορισμένη',
        'Company': 'Εταιρεία',
        'Co': 'Εταιρεία',
    }
    
    # Simple direct matching - no regex complications
    if text in translations:
        return translations[text]
    
    # Case insensitive fallback
    for eng, greek in translations.items():
        if text.lower() == eng.lower():
            return greek
    
    # Handle compound names and addresses by translating individual parts
    if ' ' in text:
        words = text.split()
        translated_words = []
        for word in words:
            # Clean punctuation
            clean_word = word.strip('.,;:!?()[]{}')
            punctuation = word[len(clean_word):]
            
            if clean_word in translations:
                translated_words.append(translations[clean_word] + punctuation)
            else:
                # Case insensitive check
                found = False
                for eng, greek in translations.items():
                    if clean_word.lower() == eng.lower():
                        translated_words.append(greek + punctuation)
                        found = True
                        break
                if not found:
                    translated_words.append(word)
        
        result = ' '.join(translated_words)
        if result != text:
            return result
    
    return text

def generate_lease_document(country, language, property_obj, tenant_obj, additional_data):
    """
    Generate the actual lease agreement document using Word templates
    """
    try:
        # Define template paths based on country and language
        template_mapping = {
            'cyprus': {
                'english': 'lease_templates/cyprus_english_lease_template.docx',
                'greek': 'lease_templates/cyprus_greek_lease_template.docx'
            },
            'greece': {
                'greek': 'lease_templates/greece_greek_lease_template.docx'
            },
            'spain': {
                'spanish': 'lease_templates/spain_spanish_lease_template.docx'
            }
        }
        
        # Get template path
        template_filename = template_mapping.get(country, {}).get(language)
        if not template_filename:
            logger.error(f"No template mapping found for country: {country}, language: {language}")
            return None
            
        template_path = os.path.join(settings.BASE_DIR, 'pages', 'templates', template_filename)

        # Check if template exists
        if not os.path.exists(template_path):
            logger.error(f"Template file not found: {template_path}")
            # Try fallback - create basic document
            return create_basic_lease_document(country, language, property_obj, tenant_obj, additional_data)
        
        # Load template
        doc = DocxTemplate(template_path)
        
        # Prepare data for template
        template_data = prepare_lease_template_data(
            country, language, property_obj, tenant_obj, additional_data
        )
        
        # Render document
        doc.render(template_data)
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        
        logger.info(f"Lease document generated successfully: {temp_file.name}")
        return temp_file.name
        
    except Exception as e:
        logger.error(f"Error in generate_lease_document: {str(e)}")
        return None

def prepare_lease_template_data(country, language, property_obj, tenant_obj, additional_data):
    """
    Prepare data dictionary for the lease agreement template using your models
    """
    # Get current date components
    current_date = datetime.now()
    
    # Helper function to convert numbers to words with proper capitalization
    def number_to_words(num):
        """Convert number to English words with proper capitalization"""
        if num == 0:
            return "Zero"
        
        ones = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine", 
                "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen", 
                "Seventeen", "Eighteen", "Nineteen"]
        
        tens = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
        
        num = int(num)  # Remove decimals
        
        if num < 20:
            return ones[num]
        elif num < 100:
            return tens[num // 10] + ("" if num % 10 == 0 else " " + ones[num % 10])
        elif num < 1000:
            return ones[num // 100] + " Hundred" + ("" if num % 100 == 0 else " " + number_to_words(num % 100))
        elif num < 1000000:
            return number_to_words(num // 1000) + " Thousand" + ("" if num % 1000 == 0 else " " + number_to_words(num % 1000))
        else:
            return str(int(num))  # Fallback for very large numbers
    
    # Helper function to pluralize items
    def pluralize_item(item_name, count):
        """Add 's' to item names when count > 1"""
        if count > 1:
            # Handle special cases - check the last word only for compound items
            words = item_name.split()
            last_word = words[-1]
            
            if last_word.endswith('y') and len(last_word) > 1 and last_word[-2] not in 'aeiou':
                # Only change y to ies for words ending in consonant + y (like "Key" but not "Boy")
                words[-1] = last_word[:-1] + 'ies'
                return ' '.join(words)
            elif last_word.endswith(('s', 'sh', 'ch', 'x', 'z')):
                words[-1] = last_word + 'es'
                return ' '.join(words)
            else:
                words[-1] = last_word + 's'
                return ' '.join(words)
        return item_name
    
    # Get dates from additional_data or tenant object
    start_date_str = additional_data.get('start_date') or (
        tenant_obj.tenant_lease_start_date.strftime('%Y-%m-%d') 
        if hasattr(tenant_obj, 'tenant_lease_start_date') and tenant_obj.tenant_lease_start_date else ''
    )
    end_date_str = additional_data.get('end_date') or (
        tenant_obj.tenant_lease_end_date.strftime('%Y-%m-%d') 
        if hasattr(tenant_obj, 'tenant_lease_end_date') and tenant_obj.tenant_lease_end_date else ''
    )
    
    # Convert date strings to datetime objects
    start_date_obj = None
    end_date_obj = None
    duration_days = 0
    duration_months = 0
    
    try:
        if start_date_str:
            start_date_obj = datetime.strptime(start_date_str, '%Y-%m-%d').date()
        if end_date_str:
            end_date_obj = datetime.strptime(end_date_str, '%Y-%m-%d').date()
        
        if start_date_obj and end_date_obj:
            duration_days = (end_date_obj - start_date_obj).days
            duration_months = max(1, round(duration_days / 30.44))  # More accurate month calculation
    except ValueError as e:
        logger.warning(f"Date parsing error: {str(e)}")
    
    # Get financial information - prefer additional_data over database (remove decimals)
    monthly_rent = int(float(additional_data.get('monthly_rent') or (
        tenant_obj.tenant_rent if hasattr(tenant_obj, 'tenant_rent') and tenant_obj.tenant_rent else 0
    )))
    security_deposit = int(float(additional_data.get('security_deposit') or (
        tenant_obj.tenant_deposit if hasattr(tenant_obj, 'tenant_deposit') and tenant_obj.tenant_deposit else 0
    )))
    communal_expenses = int(float(additional_data.get('communal_expenses', 0)))
    
    # Get extension information
    extension_period = int(additional_data.get('extension_period', 0))
    extension_date_str = additional_data.get('extension_date', '')
    extension_date_obj = None
    if extension_date_str:
        try:
            extension_date_obj = datetime.strptime(extension_date_str, '%Y-%m-%d').date()
        except ValueError:
            pass
    
    # Handle rental increase percentage (remove decimals if whole number)
    rental_increase = float(additional_data.get('rental_increase', 0))
    rental_increase_formatted = f"{int(rental_increase)}" if rental_increase == int(rental_increase) else f"{rental_increase}"
    
    # Get currency symbol based on country
    currency_symbols = {
        'cyprus': '€',
        'greece': '€',
        'spain': '€'
    }
    currency = currency_symbols.get(country, '€')
    
    # Build full property address
    address_parts = [
        property_obj.prop_address1,
        property_obj.prop_address2,
        property_obj.prop_suburb,
        property_obj.prop_city,
        property_obj.prop_province,
        property_obj.prop_country,
        property_obj.prop_pcode
    ]
    full_address = ', '.join([part for part in address_parts if part and part.strip()])
    
    # Get tenant information (handle both existing and new tenants)
    if additional_data.get('is_new_tenant') and additional_data.get('new_tenant_data'):
        new_tenant_data = additional_data['new_tenant_data']
        tenant_name = new_tenant_data.get('tenant_name', 'N/A')
        tenant_address = new_tenant_data.get('tenant_address', 'N/A')
        tenant_phone = new_tenant_data.get('tenant_contact_number', 'N/A')
        tenant_email = new_tenant_data.get('tenant_email', 'N/A')
        tenant_type = new_tenant_data.get('tenant_type', 'Individual')
        
        if tenant_type == 'Company':
            tenant_passport_id = 'N/A'
            tenant_passport_country = 'N/A'
            tenant_registration_number = new_tenant_data.get('tenant_registration_number', 'N/A')
            tenant_contact_person = new_tenant_data.get('tenant_contact_person', 'N/A')
        else:
            tenant_passport_id = new_tenant_data.get('tenant_passport_id', 'N/A')
            tenant_passport_country = new_tenant_data.get('tenant_passport_country', 'N/A')
            tenant_registration_number = 'N/A'
            tenant_contact_person = 'N/A'
    else:
        tenant_name = tenant_obj.tenant_name or 'N/A'
        # Use property address as tenant address for existing tenants
        tenant_address = full_address or 'N/A'
        tenant_phone = getattr(tenant_obj, 'tenant_contact_number', 'N/A') or 'N/A'
        tenant_email = getattr(tenant_obj, 'tenant_email', 'N/A') or 'N/A'
        tenant_type = getattr(tenant_obj, 'tenant_type', 'Individual') or 'Individual'
        tenant_contact_person = getattr(tenant_obj, 'tenant_contact_person', 'N/A') or 'N/A'

        # Get additional information for existing tenants from form data
        existing_tenant_data = additional_data.get('existing_tenant_data', {})
        if tenant_type == 'Company':
            tenant_passport_id = 'N/A'
            tenant_passport_country = 'N/A'
            tenant_registration_number = existing_tenant_data.get('tenant_registration_number', 'N/A')
        else:
            tenant_passport_id = existing_tenant_data.get('tenant_passport_id', 'N/A')
            tenant_passport_country = existing_tenant_data.get('tenant_passport_country', 'N/A')
            tenant_registration_number = 'N/A'
    
    # Second tenant information
    has_second_tenant = additional_data.get('has_second_tenant', False)
    second_tenant_data = additional_data.get('second_tenant_data', {})
    
    # Second tenant variables for template
    if has_second_tenant and second_tenant_data:
        second_tenant_name = second_tenant_data.get('tenant_name', 'N/A')
        second_tenant_passport_id = second_tenant_data.get('tenant_passport_id', 'N/A')
        second_tenant_passport_country = second_tenant_data.get('tenant_passport_country', 'N/A')
        second_tenant_address = second_tenant_data.get('tenant_address', 'N/A')
        second_tenant_contact_number = second_tenant_data.get('tenant_contact_number', 'N/A')
        second_tenant_email = second_tenant_data.get('tenant_email', 'N/A')
    else:
        second_tenant_name = ''
        second_tenant_passport_id = ''
        second_tenant_passport_country = ''
        second_tenant_address = ''
        second_tenant_contact_number = ''
        second_tenant_email = ''
    
    # Process furniture and appliances
    furniture_items = {}
    furniture_list = []
    furniture_list_greek = []
    for key, value in additional_data.items():
        if key.startswith('furniture_') and int(value or 0) > 0:
            count = int(value)
            item_name = key.replace('furniture_', '').replace('_', ' ').title()
            pluralized_name = pluralize_item(item_name, count)
            furniture_items[pluralized_name] = count
            furniture_list.append(f"{count} x {pluralized_name}")
            
            # Create Greek version - translate the pluralized name
            item_name_greek = translate_to_greek(pluralized_name)
            furniture_list_greek.append(f"{count} x {item_name_greek}")
            
    # Get fully furnished checkbox value
    fully_furnished = additional_data.get('fully_furnished') == 'true'
    
    # Get parking and storeroom checkbox values
    has_parking = additional_data.get('amenity_parking') == True  # Boolean, not string
    has_storeroom = additional_data.get('amenity_storeroom') == True  # Boolean, not string

    # Process keys
    keys_items = {}
    keys_list = []
    keys_list_greek = []
    for key, value in additional_data.items():
        if key.startswith('keys_') and int(value or 0) > 0:
            count = int(value)
            key_name = key.replace('keys_', '').replace('_', ' ').title() + ' Key'
            pluralized_key = pluralize_item(key_name, count)
            keys_items[pluralized_key] = count
            keys_list.append(f"{count} x {pluralized_key}")
            
            # Create Greek version - translate the pluralized key name
            key_name_greek = translate_to_greek(pluralized_key)
            keys_list_greek.append(f"{count} x {key_name_greek}")
            
    # Prepare template data
    template_data = {
        # Current date components with ordinal day
        'today_day': get_ordinal_day(current_date.day),
        'today_month': current_date.strftime('%B'),
        'today_year': current_date.strftime('%Y'),
        'today_date': current_date.strftime('%B %d, %Y'),
        
        # Greek date formats
        'today_date_greek': format_date_greek(current_date.date()),
        'lease_start_date_greek': format_date_greek(start_date_obj) if start_date_obj else 'N/A',
        'lease_end_date_greek': format_date_greek(end_date_obj) if end_date_obj else 'N/A',
        
        # Document information
        'document_date': current_date.strftime('%B %d, %Y'),
        'document_id': f'LA-{current_date.strftime("%Y%m%d-%H%M%S")}',
        'country': country.title(),
        'language': language.title(),
        
        # Landlord information
        'landlord_name': additional_data.get('landlord_name', 'Alivente Limited'),
        'landlord_contact_person': additional_data.get('landlord_contact_person', 'Demetri Manias'),
        'landlord_contact_person_greek': translate_to_greek(additional_data.get('landlord_contact_person', 'Demetri Manias')),
        'landlord_registration_number': additional_data.get('landlord_registration_number', 'HE123456'),
        'landlord_phone': additional_data.get('landlord_phone', '+357-96668557'),
        'landlord_phone_number': additional_data.get('landlord_phone', '+357-96668557'),
        'landlord_email': additional_data.get('landlord_email', 'demetri.manias@alivente.com'),
        'landlord_address': additional_data.get('landlord_address', 'Alivente House, Dikaiosynis 13A, Engomi, Nicosia, Cyprus, 2412'),
        
        # Tenant information
        'tenant_name': tenant_name,
        'tenant_full_name': tenant_name,
        'tenant_passport_id': tenant_passport_id,
        'tenant_passport_country': tenant_passport_country,
        'tenant_address': tenant_address,
        'tenant_phone': tenant_phone,
        'tenant_email': tenant_email,
        'tenant_type': tenant_type,
        'tenant_registration_number': tenant_registration_number,
        'tenant_contact_person': tenant_contact_person,
        
        # Second tenant information
        'has_second_tenant': has_second_tenant,
        'second_tenant_name': second_tenant_name,
        'second_tenant_passport_id': second_tenant_passport_id,
        'second_tenant_passport_country': second_tenant_passport_country,
        'second_tenant_address': second_tenant_address,
        'second_tenant_contact_number': second_tenant_contact_number,
        'second_tenant_email': second_tenant_email,
        
        # Property information
        'property_name': property_obj.prop_name or 'N/A',
        'property_address': full_address or 'N/A',
        'property_address1': property_obj.prop_address1 or '',
        'property_address2': property_obj.prop_address2 or '',
        'property_suburb': property_obj.prop_suburb or '',
        'property_city': property_obj.prop_city or '',
        'property_province': property_obj.prop_province or '',
        'property_country': property_obj.prop_country or '',
        'property_pcode': property_obj.prop_pcode or '',
        'property_floor_area': property_obj.prop_floor_area or 'N/A',
        'property_year_built': property_obj.prop_year_built or 'N/A',
        
        # Property information - Greek translations
        'property_address1_greek': translate_to_greek(property_obj.prop_address1 or ''),
        'property_address2_greek': translate_to_greek(property_obj.prop_address2 or ''),
        'property_suburb_greek': translate_to_greek(property_obj.prop_suburb or ''),
        'property_city_greek': translate_to_greek(property_obj.prop_city or ''),
        'property_province_greek': translate_to_greek(property_obj.prop_province or ''),
        
        # Lease terms
        'lease_start_date': start_date_obj.strftime('%B %d, %Y') if start_date_obj else 'N/A',
        'lease_end_date': end_date_obj.strftime('%B %d, %Y') if end_date_obj else 'N/A',
        'lease_start_date_short': start_date_obj.strftime('%d/%m/%Y') if start_date_obj else 'N/A',
        'lease_end_date_short': end_date_obj.strftime('%d/%m/%Y') if end_date_obj else 'N/A',
        'duration_months': duration_months,
        'duration_months_words': number_to_words(duration_months),
        'duration_days': duration_days,
        
        # Financial terms - numbers (without decimals)
        'rental_amount': monthly_rent,
        'communal_expenses': communal_expenses,
        'security_deposit': security_deposit,
        'deposit_amount': security_deposit,
        
        # Financial terms - words (with proper capitalization)
        'rental_amount_words': number_to_words(monthly_rent),
        'communal_expenses_words': number_to_words(communal_expenses),
        'deposit_amount_words': number_to_words(security_deposit),
        
        # Greek financial terms in words
        'rental_amount_words_greek': number_to_words_greek(monthly_rent),
        'deposit_amount_words_greek': number_to_words_greek(security_deposit),
        'communal_expenses_words_greek': number_to_words_greek(communal_expenses),
        
        # Financial terms - with thousand separators (the template uses the raw numbers)
        'rental_amount': f'{monthly_rent:,}',
        'communal_expenses': f'{communal_expenses:,}',
        'security_deposit': f'{security_deposit:,}',
        
        # Additional formatted versions for compatibility
        'rental_amount_formatted': f'{currency}{monthly_rent:,}',
        'communal_expenses_formatted': f'{currency}{communal_expenses:,}',
        'deposit_amount_formatted': f'{currency}{security_deposit:,}',
        'security_deposit_formatted': f'{currency}{security_deposit:,}',
        'total_rent_formatted': f'{currency}{(monthly_rent * duration_months):,}' if duration_months else f'{currency}0',
        'monthly_rent_formatted': f'{currency}{monthly_rent:,}',
        
        # Extension and increase terms
        'lease_extension_period': extension_period,
        'lease_extension_period_words': number_to_words(extension_period),
        'lease_extension_period_words_greek': number_to_words_greek(extension_period),
        'lease_extension_date': extension_date_obj.strftime('%B %d, %Y') if extension_date_obj else 'N/A',
        'lease_extension_date_greek': format_date_greek(extension_date_obj) if extension_date_obj else 'N/A',
        'percentage_rental_increase': rental_increase_formatted,
        
        # Currency
        'currency': currency,
        
        # Furniture and appliances (with proper pluralization and comma-separated lists)
        'furniture_items': furniture_items,
        'furniture_list': ', '.join(furniture_list) if furniture_list else 'None',
        'furniture_list_greek': ', '.join(furniture_list_greek) if furniture_list_greek else 'Κανένα',
        'fully_furnished': fully_furnished,
        
        # Parking and Storeroom
        'has_parking': has_parking,
        'has_storeroom': has_storeroom,
        
        # Keys (with proper pluralization and comma-separated lists)
        'keys_items': keys_items,
        'keys_list': ', '.join(keys_list) if keys_list else 'None',
        'keys_list_greek': ', '.join(keys_list_greek) if keys_list_greek else 'Κανένα',
        
        # Legacy fields for compatibility
        'start_date': start_date_obj.strftime('%B %d, %Y') if start_date_obj else 'N/A',
        'end_date': end_date_obj.strftime('%B %d, %Y') if end_date_obj else 'N/A',
        'monthly_rent': monthly_rent,
        
        # Generation metadata
        'generation_timestamp': current_date.strftime('%Y-%m-%d %H:%M:%S'),
    }
    
    return template_data

def create_basic_lease_document(country, language, property_obj, tenant_obj, additional_data):
    """
    Create a basic lease document when no template is available
    This is a fallback method for demonstration purposes
    """
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Helper function to pluralize items (for fallback function)
        def pluralize_item(item_name, count):
            """Add 's' to item names when count > 1"""
            if count > 1:
                # Handle special cases
                if item_name.endswith('y'):
                    return item_name[:-1] + 'ies'  # e.g., "Key" -> "Keys"
                elif item_name.endswith(('s', 'sh', 'ch', 'x', 'z')):
                    return item_name + 'es'
                else:
                    return item_name + 's'
            return item_name
        
        # Create a new document
        document = Document()
        
        # Add title
        title = document.add_heading('LEASE AGREEMENT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = document.add_paragraph(f'{country.title()} - {language.title()}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        document.add_paragraph('')  # Empty line
        
        # Document info
        document.add_paragraph(f'Document ID: LA-{datetime.now().strftime("%Y%m%d-%H%M%S")}')
        document.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
        document.add_paragraph('')
        
        # Property information
        document.add_heading('Property Information', level=1)
        document.add_paragraph(f'Property Name: {property_obj.prop_name or "N/A"}')
        
        address_parts = [
            property_obj.prop_address1,
            property_obj.prop_address2,
            property_obj.prop_suburb,
            property_obj.prop_city,
            property_obj.prop_province,
            property_obj.prop_country,
            property_obj.prop_pcode
        ]
        full_address = ', '.join([part for part in address_parts if part])
        document.add_paragraph(f'Address: {full_address or "N/A"}')
        
        if property_obj.prop_floor_area:
            document.add_paragraph(f'Floor Area: {property_obj.prop_floor_area} m²')
        if property_obj.prop_year_built:
            document.add_paragraph(f'Year Built: {property_obj.prop_year_built}')
        
        document.add_paragraph('')
        
        # Tenant information
        document.add_heading('Tenant Information', level=1)
        if additional_data.get('is_new_tenant') and additional_data.get('new_tenant_data'):
            new_tenant_data = additional_data['new_tenant_data']
            document.add_paragraph(f'Tenant Name: {new_tenant_data.get("tenant_name", "N/A")}')
            document.add_paragraph(f'Passport/ID: {new_tenant_data.get("tenant_passport_id", "N/A")}')
            document.add_paragraph(f'Email: {new_tenant_data.get("tenant_email", "N/A")}')
            document.add_paragraph(f'Phone: {new_tenant_data.get("tenant_contact_number", "N/A")}')
            document.add_paragraph(f'Address: {new_tenant_data.get("tenant_address", "N/A")}')
        else:
            document.add_paragraph(f'Tenant Name: {tenant_obj.tenant_name or "N/A"}')
            document.add_paragraph(f'Email: {getattr(tenant_obj, "tenant_email", "N/A") or "N/A"}')
            document.add_paragraph(f'Phone: {getattr(tenant_obj, "tenant_contact_number", "N/A") or "N/A"}')
        
        document.add_paragraph('')
        
        # Lease terms
        document.add_heading('Lease Terms', level=1)
        
        # Get dates
        start_date = additional_data.get('start_date') or (
            tenant_obj.tenant_lease_start_date.strftime('%Y-%m-%d') 
            if hasattr(tenant_obj, 'tenant_lease_start_date') and tenant_obj.tenant_lease_start_date else 'N/A'
        )
        end_date = additional_data.get('end_date') or (
            tenant_obj.tenant_lease_end_date.strftime('%Y-%m-%d') 
            if hasattr(tenant_obj, 'tenant_lease_end_date') and tenant_obj.tenant_lease_end_date else 'N/A'
        )
        
        document.add_paragraph(f'Start Date: {start_date}')
        document.add_paragraph(f'End Date: {end_date}')
        
        # Financial terms (without decimals, with thousand separators)
        monthly_rent = int(float(additional_data.get('monthly_rent') or (
            tenant_obj.tenant_rent if hasattr(tenant_obj, 'tenant_rent') and tenant_obj.tenant_rent else 0
        )))
        security_deposit = int(float(additional_data.get('security_deposit') or (
            tenant_obj.tenant_deposit if hasattr(tenant_obj, 'tenant_deposit') and tenant_obj.tenant_deposit else 0
        )))
        communal_expenses = int(float(additional_data.get('communal_expenses', 0)))
        
        document.add_paragraph(f'Monthly Rent: €{monthly_rent:,}')
        document.add_paragraph(f'Security Deposit: €{security_deposit:,}')
        if communal_expenses > 0:
            document.add_paragraph(f'Communal Expenses: €{communal_expenses:,}')
        
        if hasattr(tenant_obj, 'tenant_levies') and tenant_obj.tenant_levies:
            document.add_paragraph(f'Levies: €{int(float(tenant_obj.tenant_levies)):,}')
        
        # Furniture and appliances (with pluralization)
        furniture_list = []
        for key, value in additional_data.items():
            if key.startswith('furniture_') and int(value or 0) > 0:
                count = int(value)
                item_name = key.replace('furniture_', '').replace('_', ' ').title()
                pluralized_name = pluralize_item(item_name, count)
                furniture_list.append(f"{count} x {pluralized_name}")
        
        if furniture_list:
            document.add_paragraph('')
            document.add_heading('Furniture and Appliances', level=1)
            document.add_paragraph(', '.join(furniture_list))
        
        # Keys (with pluralization)
        keys_list = []
        for key, value in additional_data.items():
            if key.startswith('keys_') and int(value or 0) > 0:
                count = int(value)
                key_name = key.replace('keys_', '').replace('_', ' ').title() + ' Key'
                pluralized_key = pluralize_item(key_name, count)
                keys_list.append(f"{count} x {pluralized_key}")
        
        if keys_list:
            document.add_paragraph('')
            document.add_heading('Keys Provided', level=1)
            document.add_paragraph(', '.join(keys_list))
        
        # Signatures
        document.add_paragraph('')
        document.add_heading('Signatures', level=1)
        document.add_paragraph('Landlord: _________________________    Date: _________')
        document.add_paragraph('Tenant: _________________________     Date: _________')
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        document.save(temp_file.name)
        
        logger.info(f"Basic lease document created: {temp_file.name}")
        return temp_file.name
        
    except Exception as e:
        logger.error(f"Error creating basic lease document: {str(e)}")
        return None